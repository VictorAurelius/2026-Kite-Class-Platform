#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script tạo báo cáo thực tập tốt nghiệp dạng Word (.docx)
Format theo mẫu "Huong dan trinh bay bao cao TTTN.pdf" - ĐH GTVT

Cấu trúc báo cáo:
1. Bìa chính (có bảng thông tin sinh viên, logo, chữ CỬ NHÂN màu vàng)
2. Bìa phụ (thêm trường Đơn vị thực tập)
3. Bản nhận xét của cơ sở thực tập
4. Lời cảm ơn
5. Mục lục + Danh mục hình vẽ + Danh mục bảng biểu
6. Danh mục từ viết tắt
7. 4 Chương nội dung chính
8. Tài liệu tham khảo (IEEE)
9. Phụ lục
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

# ============== THÔNG TIN SINH VIÊN ==============
STUDENT_INFO = {
    "name": "Nguyễn Văn Kiệt",
    "student_id": "221230890",
    "class": "CNTT1-K63",
    "course": "63",
    "major": "Công nghệ thông tin",
    "department": "Công nghệ thông tin",
    "degree": "Cử nhân",
    "university": "Đại học Giao thông Vận tải",
}

INTERNSHIP_INFO = {
    "company": "SY PARTNERS., JSC",
    "address": "Tầng 3, Tòa nhà Luxury, Số 99 Võ Chí Công, Quận Tây Hồ, Hà Nội",
    "position": "Software Engineer",
    "advisor": "TS. Nguyễn Đức Dư",
    "company_mentor": "Trịnh Công Vượng (Project Manager)",
    "start_date": "01/12/2025",
    "end_date": "01/03/2026",
}

# ============== CONSTANTS ==============
FONT_NAME = 'Times New Roman'
FONT_SIZE_NORMAL = Pt(13)  # Đoạn văn
FONT_SIZE_CHAPTER = Pt(18)  # Chương - Theo quy định UTC
FONT_SIZE_SECTION = Pt(16)  # Mục 1.1 - Theo quy định UTC
FONT_SIZE_SUBSECTION = Pt(14)  # Tiểu mục 1.1.1 - Theo quy định UTC
FONT_SIZE_TABLE = Pt(12)
FONT_SIZE_CAPTION = Pt(13)

LINE_SPACING = 1.2  # Theo quy định UTC
FIRST_LINE_INDENT = Cm(1.0)  # Theo quy định UTC

MARGIN_LEFT = Cm(3.0)
MARGIN_RIGHT = Cm(2.0)
MARGIN_TOP = Cm(2.5)  # Theo quy định UTC
MARGIN_BOTTOM = Cm(2.5)  # Theo quy định UTC


def set_document_margins(doc):
    """Thiết lập căn lề cho toàn bộ document"""
    for section in doc.sections:
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT


def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def remove_cell_borders(cell):
    """Remove all borders from a cell"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_page_border(section):
    """Thêm khung viền cho trang bìa"""
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'text')

    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '24')  # Border width (1/8 point)
        border.set(qn('w:space'), '24')  # Space from text (points)
        border.set(qn('w:color'), '000000')  # Black
        pgBorders.append(border)

    sectPr.append(pgBorders)


def remove_page_border(section):
    """Xóa khung viền khỏi section (nếu có)"""
    sectPr = section._sectPr
    # Tìm và xóa tất cả pgBorders elements
    pgBorders_list = sectPr.findall(qn('w:pgBorders'))
    for pgBorders in pgBorders_list:
        sectPr.remove(pgBorders)


def add_page_number_header(doc):
    """
    Thêm số trang ở giữa phía TRÊN đầu trang (header)
    Chỉ thêm số trang từ Mục lục trở đi (section 2+)
    Không thêm số trang cho 2 trang bìa (section 0, 1)
    """
    for i, section in enumerate(doc.sections):
        # Bỏ qua 2 trang bìa (section 0 và 1)
        if i < 2:
            continue

        header = section.header
        header.is_linked_to_previous = False

        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        for r in p.runs:
            r.font.name = FONT_NAME
            r.font.size = FONT_SIZE_NORMAL

        # Reset số trang về 1 cho section đầu tiên có số trang (Mục lục)
        if i == 2:
            section.start_type = 0  # Continuous
            sectPr = section._sectPr
            pgNumType = sectPr.find(qn('w:pgNumType'))
            if pgNumType is None:
                pgNumType = OxmlElement('w:pgNumType')
                sectPr.append(pgNumType)
            pgNumType.set(qn('w:start'), '1')


def set_heading_font(style, font_name, font_size, bold=True, italic=False):
    """
    Thiết lập font cho Heading style một cách đầy đủ.

    Word Heading styles mặc định sử dụng theme fonts (Calibri Headings)
    và theme colors (màu xanh). Cần:
    - XÓA các thuộc tính theme (asciiTheme, hAnsiTheme, themeColor)
    - Thiết lập font và color trực tiếp
    """
    # Đảm bảo rPr element tồn tại
    rPr = style._element.get_or_add_rPr()

    # Xóa rFonts cũ (có thể chứa theme fonts) và tạo mới
    old_rFonts = rPr.find(qn('w:rFonts'))
    if old_rFonts is not None:
        rPr.remove(old_rFonts)

    # Tạo rFonts mới với font name trực tiếp (không dùng theme)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    # KHÔNG đặt asciiTheme, hAnsiTheme để tránh bị override bởi theme
    rPr.insert(0, rFonts)

    # Thiết lập font size
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rPr.append(sz)
    sz.set(qn('w:val'), str(int(font_size.pt * 2)))  # Word uses half-points

    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = OxmlElement('w:szCs')
        rPr.append(szCs)
    szCs.set(qn('w:val'), str(int(font_size.pt * 2)))

    # Thiết lập bold
    b = rPr.find(qn('w:b'))
    if bold:
        if b is None:
            b = OxmlElement('w:b')
            rPr.append(b)
    else:
        if b is not None:
            rPr.remove(b)

    # Thiết lập italic
    i = rPr.find(qn('w:i'))
    if italic:
        if i is None:
            i = OxmlElement('w:i')
            rPr.append(i)
    else:
        if i is not None:
            rPr.remove(i)

    # Xóa color cũ (có thể chứa themeColor) và tạo mới
    old_color = rPr.find(qn('w:color'))
    if old_color is not None:
        rPr.remove(old_color)

    # Tạo color mới với màu đen trực tiếp (không dùng theme)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '000000')
    # KHÔNG đặt themeColor để tránh bị override
    rPr.append(color)


def setup_styles(doc):
    """Thiết lập các style chuẩn cho document bao gồm Heading styles"""
    # Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE_NORMAL

    pf = style.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Thiết lập font cho Normal style qua XML
    rPr = style._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)

    # Heading 1 style (Chương) - 14pt Bold, Left
    h1_style = doc.styles['Heading 1']
    # Thiết lập qua XML để xóa theme
    set_heading_font(h1_style, FONT_NAME, Pt(14), bold=True, italic=False)
    # Thiết lập qua API để đảm bảo
    h1_style.font.name = FONT_NAME
    h1_style.font.size = Pt(14)
    h1_style.font.bold = True
    h1_style.font.italic = False
    h1_style.font.color.rgb = RGBColor(0, 0, 0)
    h1_style.paragraph_format.space_before = Pt(12)
    h1_style.paragraph_format.space_after = Pt(6)
    h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Heading 2 style (Mục 1.1, 1.2) - 13pt Bold, Left
    h2_style = doc.styles['Heading 2']
    set_heading_font(h2_style, FONT_NAME, Pt(13), bold=True, italic=False)
    h2_style.font.name = FONT_NAME
    h2_style.font.size = Pt(13)
    h2_style.font.bold = True
    h2_style.font.italic = False
    h2_style.font.color.rgb = RGBColor(0, 0, 0)
    h2_style.paragraph_format.space_before = Pt(12)
    h2_style.paragraph_format.space_after = Pt(6)
    h2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Heading 3 style (Tiểu mục 1.1.1) - 13pt Bold + Italic, Left
    h3_style = doc.styles['Heading 3']
    set_heading_font(h3_style, FONT_NAME, Pt(13), bold=True, italic=True)
    h3_style.font.name = FONT_NAME
    h3_style.font.size = Pt(13)
    h3_style.font.bold = True
    h3_style.font.italic = True
    h3_style.font.color.rgb = RGBColor(0, 0, 0)
    h3_style.paragraph_format.space_before = Pt(6)
    h3_style.paragraph_format.space_after = Pt(6)
    h3_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Caption style cho bảng và hình
    try:
        caption_style = doc.styles['Caption']
    except KeyError:
        caption_style = doc.styles.add_style('Caption', 1)  # 1 = paragraph style
    set_heading_font(caption_style, FONT_NAME, Pt(12), bold=True, italic=True)
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_style.paragraph_format.space_before = Pt(6)
    caption_style.paragraph_format.space_after = Pt(6)


def set_font(run, size=FONT_SIZE_NORMAL, bold=False, italic=False, color=None):
    """Helper to set font properties"""
    run.font.name = FONT_NAME
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)


def add_chapter_title(doc, number, text, add_page_break=True):
    """
    Tiêu đề chương: Sử dụng Heading 1 style để tạo mục lục tự động
    Format: "1. GIỚI THIỆU CHUNG VỀ ĐƠN VỊ THỰC TẬP"
    - Heading 1: 18pt Bold, Center, Times New Roman (theo quy định UTC)
    """
    if add_page_break:
        doc.add_page_break()

    # Sử dụng Heading 1 style để Word có thể tạo mục lục tự động
    p = doc.add_paragraph(style='Heading 1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Căn giữa theo quy định UTC

    run = p.add_run(f"{number}. {text.upper()}")

    # Thiết lập font trực tiếp cho run để đảm bảo không bị theme override
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = FONT_SIZE_CHAPTER  # 18pt theo quy định UTC
    run.font.bold = True
    run.font.italic = False
    run.font.color.rgb = RGBColor(0, 0, 0)

    return p


def add_section_title(doc, text):
    """
    Tiêu đề mục (1.1, 1.2): Sử dụng Heading 2 style để tạo mục lục tự động
    - Heading 2: 16pt Bold, Left, Times New Roman (theo quy định UTC)
    """
    p = doc.add_paragraph(style='Heading 2')
    run = p.add_run(text)

    # Thiết lập font trực tiếp cho run
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = FONT_SIZE_SECTION  # 16pt theo quy định UTC
    run.font.bold = True
    run.font.italic = False
    run.font.color.rgb = RGBColor(0, 0, 0)

    return p


def add_subsection_title(doc, text):
    """
    Tiêu đề tiểu mục (1.1.1, 1.1.2): Sử dụng Heading 3 style để tạo mục lục tự động
    - Heading 3: 14pt Bold, Left, Times New Roman (theo quy định UTC)
    """
    p = doc.add_paragraph(style='Heading 3')
    run = p.add_run(text)

    # Thiết lập font trực tiếp cho run
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = FONT_SIZE_SUBSECTION  # 14pt theo quy định UTC
    run.font.bold = True
    run.font.italic = False  # Không italic theo quy định UTC
    run.font.color.rgb = RGBColor(0, 0, 0)

    return p


def add_seq_field(paragraph, seq_name, prefix="", reset_chapter=False):
    """
    Thêm SEQ field vào paragraph để đánh số tự động
    SEQ field giống như khi dùng Insert > Caption trong Word

    Args:
        paragraph: Paragraph object
        seq_name: Tên sequence (ví dụ: "Table", "Figure") - KHÔNG thêm số chapter
        prefix: Text đứng trước số (ví dụ: "Bảng ", "Hình ")
        reset_chapter: True nếu cần reset về 1 ở đầu chapter mới
    """
    if prefix:
        run = paragraph.add_run(prefix)
        set_font(run, Pt(12), bold=True, italic=True)

    # Tạo SEQ field
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    # Thêm \r để reset numbering nếu cần (ở đầu chapter mới)
    if reset_chapter:
        instrText.text = f' SEQ {seq_name} \\* ARABIC \\r 1 '
    else:
        instrText.text = f' SEQ {seq_name} \\* ARABIC '

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    # Placeholder text (sẽ được cập nhật khi F9 trong Word)
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    run2 = paragraph.add_run("1")  # Placeholder number
    set_font(run2, Pt(12), bold=True, italic=True)

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run2._r.append(fldChar3)

    return paragraph


def add_table_caption(doc, chapter_num, caption_text, reset=False):
    """
    Thêm caption cho bảng với SEQ field tự động đánh số
    Format: "Bảng X.Y. Caption text" (với X là số chương, Y là SEQ tự động)

    Args:
        doc: Document object
        chapter_num: Số chương (1, 2, 3, 4)
        caption_text: Nội dung caption
        reset: True nếu đây là bảng đầu tiên trong chapter (reset về 1)

    Lưu ý: Trong Word, bấm Ctrl+A rồi F9 để cập nhật tất cả fields
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    # Thêm "Bảng X." với X là số chương
    run = p.add_run(f"Bảng {chapter_num}.")
    set_font(run, Pt(12), bold=True, italic=True)

    # Thêm SEQ field cho số thứ tự - QUAN TRỌNG: Dùng "Table" chung cho tất cả chapter
    # Để TOC \c "Table" có thể tìm thấy
    add_seq_field(p, "Table", "", reset_chapter=reset)

    # Thêm phần caption text
    run = p.add_run(f". {caption_text}")
    set_font(run, Pt(12), bold=True, italic=True)

    return p


def add_figure_caption(doc, chapter_num, caption_text, reset=False):
    """
    Thêm caption cho hình với SEQ field tự động đánh số
    Format: "Hình X.Y. Caption text" (với X là số chương, Y là SEQ tự động)

    Args:
        doc: Document object
        chapter_num: Số chương (1, 2, 3, 4)
        caption_text: Nội dung caption
        reset: True nếu đây là hình đầu tiên trong chapter (reset về 1)

    Lưu ý: Trong Word, bấm Ctrl+A rồi F9 để cập nhật tất cả fields
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)

    # Thêm "Hình X." với X là số chương
    run = p.add_run(f"Hình {chapter_num}.")
    set_font(run, Pt(12), bold=True, italic=True)

    # Thêm SEQ field cho số thứ tự - QUAN TRỌNG: Dùng "Figure" chung cho tất cả chapter
    # Để TOC \c "Figure" có thể tìm thấy
    add_seq_field(p, "Figure", "", reset_chapter=reset)

    # Thêm phần caption text
    run = p.add_run(f". {caption_text}")
    set_font(run, Pt(12), bold=True, italic=True)

    return p


def add_paragraph_text(doc, text, first_line_indent=True):
    """
    Đoạn văn: 13pt, Justify, thụt đầu dòng 1.27cm, giãn dòng 1.5
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = LINE_SPACING

    if first_line_indent:
        p.paragraph_format.first_line_indent = FIRST_LINE_INDENT

    run = p.add_run(text)
    set_font(run, FONT_SIZE_NORMAL)

    return p


def add_bullet_list(doc, items):
    """
    Thêm danh sách bullet

    Format: Bullet nhô ra ngoài, nội dung thẳng hàng khi xuống dòng
    - left_indent: Vị trí nội dung text (1.5cm)
    - first_line_indent: Bullet nhô ra (-0.5cm)
    """
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        # Text nằm ở 1.5cm từ lề trái
        p.paragraph_format.left_indent = Cm(1.5)
        # Bullet nhô ra ngoài 0.5cm (hanging indent)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.line_spacing = LINE_SPACING

        run = p.add_run(item)
        set_font(run, FONT_SIZE_NORMAL)


def add_table_with_caption(doc, chapter_num, caption_text, headers, rows, col_widths=None, reset=False):
    """
    Thêm bảng với caption sử dụng SEQ field để đánh số tự động
    Caption ở PHÍA TRÊN bảng theo đúng chuẩn Word

    Args:
        doc: Document object
        chapter_num: Số chương (1, 2, 3, 4)
        caption_text: Nội dung caption
        headers: List các header cột
        rows: List các hàng dữ liệu
        col_widths: List độ rộng cột (cm)
        reset: True nếu đây là bảng đầu tiên trong chapter
    """
    # Thêm caption với SEQ field
    add_table_caption(doc, chapter_num, caption_text, reset=reset)

    # Tạo bảng
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        if col_widths and i < len(col_widths):
            header_cells[i].width = Cm(col_widths[i])
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_font(run, FONT_SIZE_TABLE, bold=True)
        set_cell_shading(header_cells[i], 'D9E2F3')

    for row_data in rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            row.cells[i].text = str(cell_text)
            if col_widths and i < len(col_widths):
                row.cells[i].width = Cm(col_widths[i])
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    set_font(run, FONT_SIZE_TABLE)

    doc.add_paragraph()
    return table


def add_figure_placeholder(doc, chapter_num, caption_text):
    """
    Thêm placeholder cho hình vẽ với caption sử dụng SEQ field
    Caption ở PHÍA DƯỚI hình theo đúng chuẩn Word

    Args:
        doc: Document object
        chapter_num: Số chương (1, 2, 3, 4)
        caption_text: Nội dung caption
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)

    run = p.add_run("[Chèn hình vẽ tại đây]")
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_NORMAL
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Thêm caption với SEQ field
    add_figure_caption(doc, chapter_num, caption_text)


# ============== TRANG BÌA CHÍNH ==============
def add_cover_page(doc):
    """Tạo trang bìa chính theo mẫu PDF (trang 1)"""
    import os

    # TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI")
    set_font(run, Pt(14), bold=False)

    # KHOA CÔNG NGHỆ THÔNG TIN (đậm, gạch chân)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("KHOA CÔNG NGHỆ THÔNG TIN")
    set_font(run, Pt(14), bold=True)
    run.font.underline = True

    # Logo
    logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'templates', 'logo_utc.png')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(24)

    if os.path.exists(logo_path):
        run = p.add_run()
        run.add_picture(logo_path, width=Cm(3.5))
    else:
        run = p.add_run("[LOGO TRƯỜNG]")
        run.font.color.rgb = RGBColor(128, 128, 128)
        set_font(run, Pt(12))

    # BÁO CÁO (gạch chân)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("BÁO CÁO")
    set_font(run, Pt(14), bold=True)
    run.font.underline = True

    # THỰC TẬP TỐT NGHIỆP
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("THỰC TẬP TỐT NGHIỆP")
    set_font(run, Pt(22), bold=True)

    # CỬ NHÂN (không màu vàng, không gạch chân)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("CỬ NHÂN")
    set_font(run, Pt(22), bold=True)

    # Bảng thông tin sinh viên (CÓ viền)
    table = doc.add_table(rows=9, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'  # Thêm style có viền

    info_rows = [
        ("Sinh viên thực hiện", STUDENT_INFO["name"]),
        ("Mã sinh viên", STUDENT_INFO["student_id"]),
        ("Lớp", STUDENT_INFO["class"]),
        ("Khóa", STUDENT_INFO["course"]),
        ("Ngành đào tạo", STUDENT_INFO["major"]),
        ("Đơn vị thực tập", INTERNSHIP_INFO["company"]),
        ("Giảng viên hướng dẫn", INTERNSHIP_INFO["advisor"]),
        ("CBHD tại đơn vị TT", INTERNSHIP_INFO["company_mentor"]),
        ("Thời gian thực tập", f"Từ ngày {INTERNSHIP_INFO['start_date']} đến ngày {INTERNSHIP_INFO['end_date']}"),
    ]

    for i, (label, value) in enumerate(info_rows):
        row = table.rows[i]
        # Label cell
        row.cells[0].text = label
        row.cells[0].width = Cm(5.0)
        for paragraph in row.cells[0].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_font(run, Pt(13))

        # Value cell (không có dấu ":")
        row.cells[1].text = value
        row.cells[1].width = Cm(9.0)
        for paragraph in row.cells[1].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_font(run, Pt(13))

    # Khoảng trống
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # Hà Nội – 2026
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Hà Nội – 2026")
    set_font(run, Pt(14), bold=True, italic=True)

    # QUAN TRỌNG: Tạo section break để tách bìa chính thành section riêng
    doc.add_section(WD_SECTION.NEW_PAGE)


# ============== TRANG BÌA PHỤ ==============
def add_secondary_cover_page(doc):
    """Tạo trang bìa phụ theo mẫu PDF (trang 2) - thêm Đơn vị thực tập

    Trang này sẽ nằm trong section 1 (nhờ section break ở cuối add_cover_page)
    """
    import os

    # Section break đã được tạo trong add_cover_page, nội dung này sẽ ở section 1

    # TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI")
    set_font(run, Pt(14), bold=False)

    # KHOA CÔNG NGHỆ THÔNG TIN (đậm, gạch chân)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("KHOA CÔNG NGHỆ THÔNG TIN")
    set_font(run, Pt(14), bold=True)
    run.font.underline = True

    # Khoảng trống thay logo (giống bìa chính)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(24)

    # BÁO CÁO (gạch chân)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("BÁO CÁO")
    set_font(run, Pt(14), bold=True)
    run.font.underline = True

    # THỰC TẬP TỐT NGHIỆP
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("THỰC TẬP TỐT NGHIỆP")
    set_font(run, Pt(22), bold=True)

    # CỬ NHÂN (không màu vàng, không gạch chân)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("CỬ NHÂN")
    set_font(run, Pt(22), bold=True)

    # Bảng thông tin sinh viên (CÓ viền, có thêm Đơn vị thực tập)
    table = doc.add_table(rows=9, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'  # Thêm style có viền

    info_rows = [
        ("Sinh viên thực hiện", STUDENT_INFO["name"]),
        ("Mã sinh viên", STUDENT_INFO["student_id"]),
        ("Lớp", STUDENT_INFO["class"]),
        ("Khóa", STUDENT_INFO["course"]),
        ("Ngành đào tạo", STUDENT_INFO["major"]),
        ("Đơn vị thực tập", INTERNSHIP_INFO["company"]),
        ("Giảng viên hướng dẫn", INTERNSHIP_INFO["advisor"]),
        ("CBHD tại đơn vị TT", INTERNSHIP_INFO["company_mentor"]),
        ("Thời gian thực tập", f"Từ ngày {INTERNSHIP_INFO['start_date']} đến ngày {INTERNSHIP_INFO['end_date']}"),
    ]

    for i, (label, value) in enumerate(info_rows):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].width = Cm(5.0)
        for paragraph in row.cells[0].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_font(run, Pt(13))

        row.cells[1].text = value
        row.cells[1].width = Cm(9.0)
        for paragraph in row.cells[1].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_font(run, Pt(13))

    # Khoảng trống (3 paragraphs giống bìa chính)
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # Hà Nội – 2026
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Hà Nội – 2026")
    set_font(run, Pt(14), bold=True, italic=True)

    # QUAN TRỌNG: Tạo section break để tách bìa phụ thành section riêng
    # Sau section break này, các trang còn lại sẽ KHÔNG có border
    doc.add_section(WD_SECTION.NEW_PAGE)


# ============== BẢN NHẬN XÉT CỦA CƠ SỞ THỰC TẬP ==============
def add_company_review_page(doc):
    """Tạo trang Bản nhận xét của cơ sở thực tập (trang 3)"""
    doc.add_page_break()

    # CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    set_font(run, Pt(13), bold=True)

    # Độc lập - Tự do - Hạnh phúc
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Độc lập - Tự do - Hạnh phúc")
    set_font(run, Pt(13), bold=True)
    run.font.underline = True

    # Ngày tháng năm
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Hà Nội, ngày …. tháng … năm 2026")
    set_font(run, Pt(13), italic=True)

    # Tiêu đề
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("BẢN NHẬN XÉT CỦA CƠ SỞ THỰC TẬP")
    set_font(run, Pt(14), bold=True)

    # Kính gửi
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Kính gửi: ")
    set_font(run, Pt(13), bold=True, italic=True)
    run = p.add_run("Khoa Công nghệ thông tin, Trường Đại học Giao thông vận tải")
    set_font(run, Pt(13), bold=True, italic=True)

    # Thông tin cơ sở thực tập
    info_lines = [
        f"Cơ sở thực tập: {INTERNSHIP_INFO['company']}",
        "Người đại diện: ",
        "Chức vụ: ",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_font(run, Pt(13))

    # Xác nhận sinh viên
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Xác nhận sinh viên:")
    set_font(run, Pt(13), bold=True)

    # Thông tin sinh viên
    p = doc.add_paragraph()
    run = p.add_run(f"Họ tên: {STUDENT_INFO['name']}")
    set_font(run, Pt(13))
    run = p.add_run(f"\t\t\tMã sinh viên: {STUDENT_INFO['student_id']}")
    set_font(run, Pt(13))

    p = doc.add_paragraph()
    run = p.add_run(f"Lớp: {STUDENT_INFO['class']}")
    set_font(run, Pt(13))

    p = doc.add_paragraph()
    run = p.add_run(f"Đã thực tập tốt nghiệp tại cơ sở trong thời gian từ: {INTERNSHIP_INFO['start_date']} đến: {INTERNSHIP_INFO['end_date']}")
    set_font(run, Pt(13))

    # Các mục nhận xét
    review_items = [
        "Nội dung thực tập:",
        "Về tinh thần, ý thức, thái độ đối với công việc được giao:",
        "Về trình độ, kỹ năng làm việc/ khả năng thực hành:",
        "Ưu điểm nổi bật:",
        "Hạn chế cần khắc phục:",
        "Các nhận xét khác (nếu có):",
    ]

    for item in review_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run(f"- {item}")
        set_font(run, Pt(13))
        run.font.underline = True

    # Điểm thực tập
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run("Điểm thực tập (thang điểm 10): …… điểm.")
    set_font(run, Pt(13))

    # Chữ ký
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("ĐẠI DIỆN CƠ SỞ THỰC TẬP")
    set_font(run, Pt(13), bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("(Họ tên, chữ ký và đóng dấu nếu có)")
    set_font(run, Pt(13), italic=True)


# ============== LỜI CẢM ƠN ==============
def add_acknowledgment_page(doc):
    """Tạo trang Lời cảm ơn (trang 4) - theo mẫu tham khảo"""
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("LỜI CẢM ƠN")
    set_font(run, Pt(14), bold=True, italic=True)

    # Nội dung lời cảm ơn - chi tiết hơn theo mẫu tham khảo
    add_paragraph_text(doc,
        "Trong suốt quá trình thực tập và hoàn thành báo cáo này, em đã nhận được sự "
        "quan tâm, hướng dẫn và giúp đỡ quý báu từ nhiều tập thể và cá nhân. Đây là nguồn động "
        "viên to lớn, giúp em có thêm động lực và kiến thức để hoàn thành tốt đợt thực tập của mình.")

    add_paragraph_text(doc,
        f"Trước hết, em xin bày tỏ lòng biết ơn sâu sắc đến {INTERNSHIP_INFO['advisor']}, giảng "
        "viên hướng dẫn thuộc Trường Đại học Giao thông Vận tải. Thầy đã tận tình hướng dẫn, "
        "định hướng nội dung thực tập, đóng góp nhiều ý kiến chuyên môn quan trọng và luôn "
        "theo sát, hỗ trợ em trong suốt quá trình thực hiện đề tài. Những kiến thức chuyên sâu, "
        "kinh nghiệm thực tiễn cũng như sự nghiêm túc trong học thuật mà thầy truyền đạt đã "
        "giúp em nâng cao tư duy chuyên môn và hoàn thiện báo cáo một cách tốt hơn.")

    add_paragraph_text(doc,
        "Em xin chân thành cảm ơn Khoa Công nghệ thông tin, Trường Đại học Giao "
        "thông Vận tải đã tạo điều kiện thuận lợi để em được tham gia thực tập, tiếp cận với môi "
        "trường làm việc thực tế và vận dụng những kiến thức đã học vào thực tiễn. Sự hỗ trợ "
        "của Khoa là nền tảng quan trọng giúp sinh viên có cơ hội học hỏi, rèn luyện kỹ năng và "
        "tích lũy kinh nghiệm thực tế.")

    add_paragraph_text(doc,
        f"Bên cạnh đó, em xin gửi lời cảm ơn chân thành đến {INTERNSHIP_INFO['company']} cùng các "
        "anh/chị trong đơn vị đã nhiệt tình hướng dẫn, chia sẻ kinh nghiệm chuyên môn, tạo điều "
        "kiện thuận lợi để em được tham gia vào các công việc thực tế, qua đó giúp em hiểu rõ "
        "hơn về quy trình làm việc cũng như yêu cầu của môi trường nghề nghiệp sau này.")

    add_paragraph_text(doc,
        "Cuối cùng, em xin cảm ơn gia đình, bạn bè và những người thân đã luôn quan "
        "tâm, động viên, hỗ trợ em cả về tinh thần lẫn vật chất trong suốt thời gian thực tập và học tập.")

    add_paragraph_text(doc,
        "Mặc dù đã rất cố gắng, song do thời gian thực tập và kinh nghiệm thực tiễn còn "
        "hạn chế, báo cáo không tránh khỏi những thiếu sót. Em rất mong nhận được sự đóng "
        "góp ý kiến từ quý thầy cô để báo cáo được hoàn thiện hơn.")

    add_paragraph_text(doc, "Em xin chân thành cảm ơn!")

    # Chữ ký
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Hà Nội, ngày ... tháng 01 năm 2026")
    set_font(run, Pt(13), italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Sinh viên thực hiện")
    set_font(run, Pt(13), bold=True)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(STUDENT_INFO["name"])
    set_font(run, Pt(13))


# ============== MỤC LỤC ==============
def add_toc_page(doc):
    """Thêm trang Mục lục"""
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("MỤC LỤC")
    set_font(run, Pt(14), bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Tạo mục lục tự động: References → Table of Contents]")
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(128, 128, 128)


def add_list_of_figures(doc):
    """
    Thêm Danh mục hình vẽ với TOC field tự động

    Sau khi mở file Word:
    1. Bấm Ctrl+A (chọn tất cả)
    2. Bấm F9 (cập nhật tất cả fields)
    3. Danh mục hình vẽ sẽ tự động hiển thị
    """
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("DANH MỤC HÌNH VẼ")
    set_font(run, Pt(14), bold=True)

    # Thêm TOC field tự động cho hình vẽ
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_toc_field(p, toc_type="Figure")

    # Thêm hướng dẫn
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("(Bấm Ctrl+A rồi F9 trong Word để cập nhật danh mục)")
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(128, 128, 128)


def add_toc_field(paragraph, toc_type="Table"):
    """
    Thêm TOC field cho danh mục bảng biểu hoặc hình vẽ

    Args:
        paragraph: Paragraph object
        toc_type: "Table" hoặc "Figure"
    """
    run = paragraph.add_run()

    # Begin field character
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    # Instruction text
    # TOC \h \z \c "Table" - creates table of figures for "Table" label
    # \h: Include hyperlinks
    # \z: Hide page numbers in web layout view
    # \c "Table": Only include items with "Table" label
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f'TOC \\h \\z \\c "{toc_type}"'

    # End field character
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    # Add all elements to run
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    return run


def add_list_of_tables(doc):
    """
    Thêm Danh mục bảng biểu với TOC field tự động

    Sau khi mở file Word:
    1. Bấm Ctrl+A (chọn tất cả)
    2. Bấm F9 (cập nhật tất cả fields)
    3. Danh mục bảng biểu sẽ tự động hiển thị
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("DANH MỤC BẢNG BIỂU")
    set_font(run, Pt(14), bold=True)

    # Thêm TOC field tự động cho bảng biểu
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_toc_field(p, toc_type="Table")

    # Thêm hướng dẫn
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("(Bấm Ctrl+A rồi F9 trong Word để cập nhật danh mục)")
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(128, 128, 128)


def add_abbreviations(doc):
    """Thêm Danh mục thuật ngữ và từ viết tắt - Phân biệt rõ ràng giữa thuật ngữ và từ viết tắt"""
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("DANH MỤC THUẬT NGỮ VÀ TỪ VIẾT TẮT")
    set_font(run, Pt(14), bold=True)

    # ============ Phần 1: Thuật ngữ (Terms/Terminology) ============
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("1. THUẬT NGỮ")
    set_font(run, FONT_SIZE_NORMAL, bold=True)

    # 3 cột: Thuật ngữ | Tiếng Anh | Giải thích
    terms = [
        ("Batch Processing", "Batch Processing", "Xử lý hàng loạt dữ liệu theo lịch trình định sẵn"),
        ("Microservices", "Microservices", "Kiến trúc phần mềm chia nhỏ ứng dụng thành các dịch vụ độc lập"),
        ("Multi-tenant", "Multi-tenancy", "Kiến trúc cho phép nhiều tổ chức dùng chung một hệ thống"),
        ("RESTful API", "RESTful API", "Kiến trúc thiết kế API theo nguyên tắc REST"),
        ("Shiteki", "指摘 (tiếng Nhật)", "Phản hồi, góp ý từ quá trình review code/thiết kế"),
        ("Offshore Development", "Offshore Development", "Mô hình phát triển phần mềm thuê ngoài quốc tế"),
    ]

    table1 = doc.add_table(rows=1, cols=3)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers1 = ["Thuật ngữ", "Tiếng Anh", "Giải thích"]
    col_widths1 = [Cm(4.0), Cm(4.5), Cm(7.5)]
    header_cells1 = table1.rows[0].cells
    for i, (header, width) in enumerate(zip(headers1, col_widths1)):
        header_cells1[i].text = header
        header_cells1[i].width = width
        for paragraph in header_cells1[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_font(run, FONT_SIZE_TABLE, bold=True)
        set_cell_shading(header_cells1[i], 'D9E2F3')

    for term, english, explanation in terms:
        row = table1.add_row()
        row.cells[0].text = term
        row.cells[0].width = col_widths1[0]
        row.cells[1].text = english
        row.cells[1].width = col_widths1[1]
        row.cells[2].text = explanation
        row.cells[2].width = col_widths1[2]
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font(run, FONT_SIZE_TABLE)

    # ============ Phần 2: Từ viết tắt (Abbreviations/Acronyms) ============
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("2. TỪ VIẾT TẮT")
    set_font(run, FONT_SIZE_NORMAL, bold=True)

    # 3 cột: Từ viết tắt | Tiếng Anh đầy đủ | Nghĩa tiếng Việt
    abbreviations = [
        ("AI", "Artificial Intelligence", "Trí tuệ nhân tạo"),
        ("API", "Application Programming Interface", "Giao diện lập trình ứng dụng"),
        ("BrSE", "Bridge System Engineer", "Kỹ sư cầu nối (Việt-Nhật)"),
        ("CSDL", "Cơ sở dữ liệu", "Database"),
        ("DB", "Database", "Cơ sở dữ liệu"),
        ("DevOps", "Development and Operations", "Phát triển và vận hành"),
        ("IDE", "Integrated Development Environment", "Môi trường phát triển tích hợp"),
        ("MVP", "Minimum Viable Product", "Sản phẩm khả thi tối thiểu"),
        ("QA", "Quality Assurance", "Đảm bảo chất lượng"),
        ("SaaS", "Software as a Service", "Phần mềm như một dịch vụ"),
        ("SYP", "SY PARTNERS., JSC", "Công ty SY PARTNERS"),
    ]

    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers2 = ["Từ viết tắt", "Tiếng Anh đầy đủ", "Nghĩa tiếng Việt"]
    col_widths2 = [Cm(3.0), Cm(6.0), Cm(7.0)]
    header_cells2 = table2.rows[0].cells
    for i, (header, width) in enumerate(zip(headers2, col_widths2)):
        header_cells2[i].text = header
        header_cells2[i].width = width
        for paragraph in header_cells2[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_font(run, FONT_SIZE_TABLE, bold=True)
        set_cell_shading(header_cells2[i], 'D9E2F3')

    for abbr, full_name, meaning in abbreviations:
        row = table2.add_row()
        row.cells[0].text = abbr
        row.cells[0].width = col_widths2[0]
        row.cells[1].text = full_name
        row.cells[1].width = col_widths2[1]
        row.cells[2].text = meaning
        row.cells[2].width = col_widths2[2]
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font(run, FONT_SIZE_TABLE)


# ============== CHƯƠNG 1: GIỚI THIỆU CHUNG VỀ ĐƠN VỊ THỰC TẬP ==============
def add_chapter1(doc):
    """Chương 1: Giới thiệu chung về đơn vị thực tập - theo mẫu tham khảo"""
    add_chapter_title(doc, "1", "GIỚI THIỆU CHUNG VỀ ĐƠN VỊ THỰC TẬP")

    # 1.1 Thông tin chung về đơn vị thực tập
    add_section_title(doc, "1.1. Thông tin chung về đơn vị thực tập")

    add_paragraph_text(doc,
        f"Tên đơn vị: {INTERNSHIP_INFO['company']} (viết tắt: SYP)")

    add_paragraph_text(doc, "Lĩnh vực hoạt động:")

    add_paragraph_text(doc,
        f"{INTERNSHIP_INFO['company']} là công ty công nghệ thông tin chuyên về "
        "phát triển phần mềm gia công (offshore development) cho thị trường Nhật Bản và quốc tế. "
        "Công ty được thành lập tại Hà Nội, Việt Nam vào năm 2022, với đội ngũ lãnh đạo có hơn 20 năm "
        "kinh nghiệm trong lĩnh vực phát triển phần mềm offshore trên toàn cầu. Công ty tập trung vào "
        "các lĩnh vực:")

    add_bullet_list(doc, [
        "Phát triển phần mềm gia công cho khách hàng Nhật Bản",
        "Thiết kế hệ thống và cơ sở dữ liệu",
        "Phát triển ứng dụng web và mobile",
        "Tích hợp hệ thống và DevOps",
        "Tư vấn và chuyển giao công nghệ",
    ])

    add_paragraph_text(doc, "Thông tin liên hệ:")

    add_bullet_list(doc, [
        f"Địa chỉ: {INTERNSHIP_INFO['address']}",
        "Website: https://syp.vn",
        "Quy mô: Hơn 95 nhân viên (tính đến tháng 6/2024)",
    ])

    add_paragraph_text(doc, "Thành tựu đã đạt được:")

    add_paragraph_text(doc,
        "Từ khi thành lập năm 2022 đến nay, công ty đã đạt được nhiều thành tựu đáng kể trong "
        "lĩnh vực phát triển phần mềm offshore cho thị trường Nhật Bản:")

    add_bullet_list(doc, [
        "Phát triển thành công hơn 30 dự án phần mềm cho các khách hàng Nhật Bản thuộc nhiều lĩnh vực: "
        "tài chính, bán lẻ, y tế, giáo dục",
        "Phục vụ hơn 15 khách hàng doanh nghiệp lớn tại Nhật Bản với tỷ lệ hài lòng cao (trên 90%)",
        "Tăng trưởng quy mô từ 20 nhân viên ban đầu lên hơn 95 nhân viên sau 3 năm hoạt động",
        "Xây dựng quy trình phát triển phần mềm chuyên nghiệp theo chuẩn quốc tế với tỷ lệ "
        "dự án giao đúng hạn đạt 95%",
        "Áp dụng thành công các công nghệ hiện đại: Java Spring Boot, Oracle Database, Spring Batch, "
        "RESTful API, Git/Maven, AI Integration",
        "Đào tạo và phát triển nguồn nhân lực chất lượng cao với hơn 80% kỹ sư có kinh nghiệm 3+ năm",
        "Xây dựng văn hóa doanh nghiệp lấy con người làm trung tâm với tỷ lệ giữ chân nhân tài cao",
    ])

    add_paragraph_text(doc, "Cơ cấu tổ chức:")

    add_paragraph_text(doc,
        "Công ty được tổ chức theo mô hình chuẩn của doanh nghiệp phần mềm, với các bộ phận chính:")

    add_bullet_list(doc, [
        "Ban Giám đốc: Điều hành và quản lý chiến lược công ty",
        "Phòng Phát triển phần mềm (Development): Thiết kế và lập trình các dự án",
        "Phòng Kiểm thử (QA/QC): Đảm bảo chất lượng sản phẩm",
        "Phòng BrSE (Bridge System Engineer): Kết nối với khách hàng Nhật Bản",
        "Phòng Nhân sự và Hành chính: Quản lý nguồn nhân lực",
    ])

    add_paragraph_text(doc, "Quy mô và định hướng phát triển:")

    add_paragraph_text(doc,
        "Công ty hướng tới trở thành đối tác chiến lược tin cậy cho các doanh nghiệp Nhật Bản "
        "trong lĩnh vực phát triển phần mềm. Các mục tiêu chính bao gồm:")

    add_bullet_list(doc, [
        "Mở rộng quy mô lên 200+ nhân viên vào năm 2027",
        "Tăng cường hợp tác với khách hàng Nhật Bản lớn",
        "Phát triển năng lực AI và Machine Learning trong các dự án",
        "Xây dựng môi trường làm việc đạt chuẩn quốc tế",
    ])

    # 1.2 Chức năng, nhiệm vụ của bộ phận thực tập
    add_section_title(doc, "1.2. Chức năng, nhiệm vụ của bộ phận thực tập")

    add_paragraph_text(doc,
        "Sinh viên được phân công vào bộ phận Development với vai trò Software Engineer. "
        "Bộ phận thực tập tại công ty có các chức năng và nhiệm vụ chính sau đây:")

    add_paragraph_text(doc, "Tổ chức và quản lý chương trình thực tập:")

    add_paragraph_text(doc,
        "Phối hợp với giảng viên hướng dẫn để lập kế hoạch, phân công nhiệm vụ thực tập "
        "cho sinh viên. Đảm bảo thực tập phù hợp với chương trình đào tạo, đáp ứng yêu cầu "
        "về thời lượng, nội dung chuyên môn và mục tiêu phát triển kỹ năng thực tiễn.")

    add_paragraph_text(doc, "Hướng dẫn và giám sát sinh viên thực tập:")

    add_paragraph_text(doc,
        "Hướng dẫn sinh viên thực hiện các công việc thiết kế hệ thống, viết tài liệu kỹ thuật, "
        "tham gia vào quy trình review và cải tiến chất lượng. Tổ chức các buổi training định kỳ, "
        "kiểm tra tiến độ và góp ý chuyên môn để đảm bảo chất lượng công việc.")

    add_paragraph_text(doc, "Cung cấp điều kiện cơ sở vật chất và tài nguyên hỗ trợ:")

    add_paragraph_text(doc,
        "Cung cấp máy tính, phần mềm chuyên dụng (IntelliJ IDEA, VS Code, Oracle Database, "
        "các công cụ thiết kế), tài liệu tham khảo và môi trường phát triển. Hỗ trợ sinh viên "
        "tiếp cận các dự án thực tế của công ty.")

    add_paragraph_text(doc, "Đánh giá và phản hồi:")

    add_paragraph_text(doc,
        "Tham gia đánh giá kết quả thực tập dựa trên các tiêu chí: chất lượng thiết kế, "
        "tinh thần học hỏi, khả năng làm việc nhóm và thái độ làm việc. Phối hợp với "
        "giảng viên hướng dẫn để đưa ra nhận xét cuối cùng.")

    # 1.3 Môi trường làm việc và quy trình công tác
    add_section_title(doc, "1.3. Môi trường làm việc và quy trình công tác")

    add_paragraph_text(doc, "Môi trường làm việc:")

    add_paragraph_text(doc,
        f"Công ty có văn phòng hiện đại tại {INTERNSHIP_INFO['address']}. "
        "Môi trường làm việc được trang bị đầy đủ:")

    add_bullet_list(doc, [
        "Hệ thống máy tính cấu hình cao, phòng họp với thiết bị video conference",
        "Không gian làm việc thoáng đãng, có Wi-Fi tốc độ cao",
        "Môi trường thân thiện, chuyên nghiệp với đội ngũ nhân viên trẻ trung",
        "Văn hóa chia sẻ kiến thức qua các buổi training nội bộ hàng tuần",
    ])

    add_paragraph_text(doc, "Quy trình công tác:")

    add_paragraph_text(doc,
        "Công ty áp dụng mô hình làm việc hybrid, kết hợp giữa làm việc tại văn phòng "
        "và làm việc từ xa. Quy trình phát triển phần mềm theo mô hình Agile/Scrum "
        "với các đặc điểm:")

    add_bullet_list(doc, [
        "Sprint 2 tuần với daily standup meeting hàng ngày",
        "Quy trình review nhiều cấp: Leader Review → Customer Review → End-user Review",
        "Hệ thống quản lý task bằng Backlog (công cụ quản lý dự án của Nhật)",
        "Sử dụng Git cho quản lý mã nguồn và tài liệu",
        "Thời gian làm việc linh hoạt, tập trung vào kết quả",
    ])


# ============== CHƯƠNG 2: NỘI DUNG THỰC TẬP ==============
def add_chapter2(doc):
    """Chương 2: Nội dung thực tập - theo mẫu tham khảo"""
    add_chapter_title(doc, "2", "NỘI DUNG THỰC TẬP")

    # 2.1 Mục tiêu và yêu cầu của đợt thực tập
    add_section_title(doc, "2.1. Mục tiêu và yêu cầu của đợt thực tập")

    add_paragraph_text(doc,
        "Mục tiêu của đợt thực tập nhằm giúp sinh viên củng cố và vận dụng những kiến "
        "thức lý thuyết đã học vào môi trường làm việc thực tế, qua đó nâng cao năng lực chuyên "
        "môn và kỹ năng nghề nghiệp trong lĩnh vực công nghệ thông tin.")

    add_paragraph_text(doc, "Cụ thể, đợt thực tập hướng tới các mục tiêu sau:")

    add_bullet_list(doc, [
        "Giúp sinh viên hiểu rõ hơn về môi trường làm việc, quy trình công tác và yêu cầu "
        "chuyên môn trong lĩnh vực phát triển phần mềm offshore",
        "Rèn luyện kỹ năng thiết kế hệ thống: cơ sở dữ liệu, màn hình, API và batch processing",
        "Nâng cao kỹ năng làm việc độc lập, làm việc nhóm, quản lý thời gian và báo cáo công việc",
        "Tạo điều kiện cho sinh viên tiếp cận với các công nghệ, công cụ và phương pháp "
        "thiết kế phần mềm hiện đại theo tiêu chuẩn Nhật Bản",
        "Học cách sử dụng AI (Claude AI) [4] hỗ trợ kiểm tra chất lượng thiết kế",
        "Chuẩn bị nền tảng kiến thức và kỹ năng cần thiết cho đồ án tốt nghiệp cũng như "
        "công việc sau khi ra trường",
    ])

    add_paragraph_text(doc,
        "Yêu cầu của đợt thực tập là sinh viên phải thực hiện đầy đủ kế hoạch thực tập đã đề ra, "
        "chấp hành nghiêm túc nội quy của công ty, hoàn thành các nhiệm vụ được giao, "
        "báo cáo tiến độ đúng thời hạn và tổng hợp kết quả thực tập thành báo cáo theo đúng "
        "quy định của Nhà trường.")

    # 2.2 Kế hoạch thực tập
    add_section_title(doc, "2.2. Kế hoạch thực tập")

    add_paragraph_text(doc,
        "Bảng kế hoạch thực tập được xây dựng theo từng tuần nhằm đảm bảo tiến độ và "
        "chất lượng công việc trong suốt thời gian thực tập.")

    # Bảng đầu tiên trong Chapter 2 - reset=True để đánh số từ 1
    add_table_with_caption(doc, 2, "Kế hoạch thực tập chi tiết",
        ["Tuần", "Thời gian", "Nội dung công việc chính"],
        [
            ("1", "01/12 – 07/12", "Làm quen môi trường, tìm hiểu dự án, ôn tập kiến thức"),
            ("2", "08/12 – 14/12", "Training thiết kế cơ sở dữ liệu (DB Design)"),
            ("3", "15/12 – 21/12", "Thực hành thiết kế bảng, index, constraints"),
            ("4", "22/12 – 28/12", "Training và thực hành thiết kế màn hình (Screen Design)"),
            ("5", "29/12 – 04/01", "Hoàn thiện Screen Design, bắt đầu API Design"),
            ("6", "05/01 – 11/01", "Training và thực hành thiết kế API RESTful"),
            ("7", "12/01 – 18/01", "Giới thiệu AI Checker, training thiết kế Batch"),
            ("8", "19/01 – 25/01", "Thực hành thiết kế Batch Processing"),
            ("9", "26/01 – 01/02", "Thiết kế độc lập, xử lý Shiteki (feedback) đợt 1"),
            ("10", "02/02 – 08/02", "Xử lý Shiteki đợt 2, hoàn thiện thiết kế"),
            ("11", "09/02 – 15/02", "Tổng hợp sản phẩm, kiểm tra chất lượng"),
            ("12", "16/02 – 22/02", "Viết báo cáo thực tập, chuẩn bị tài liệu"),
            ("13", "23/02 – 01/03", "Hoàn thiện báo cáo, nộp sản phẩm cuối cùng"),
        ],
        col_widths=[2.0, 3.5, 10.5],
        reset=True
    )

    # 2.3 Các công việc đã thực hiện
    add_section_title(doc, "2.3. Các công việc đã thực hiện")

    add_paragraph_text(doc,
        "Trong thời gian thực tập tại công ty, em đã thực hiện các công việc "
        "theo sự phân công và hướng dẫn trực tiếp của cán bộ hướng dẫn tại đơn vị.")

    add_paragraph_text(doc, "Mô tả các công việc đã thực hiện:")

    add_bullet_list(doc, [
        "Tìm hiểu đề tài thực tập và các yêu cầu kỹ thuật do cán bộ hướng dẫn giao",
        "Nghiên cứu tài liệu liên quan đến thiết kế hệ thống phần mềm theo chuẩn Nhật Bản",
        "Thực hiện thiết kế cơ sở dữ liệu, màn hình, API và batch processing",
        "Ghi chép tiến độ công việc và báo cáo kết quả định kỳ cho cán bộ hướng dẫn",
        "Hoàn thiện sản phẩm và tổng hợp nội dung báo cáo thực tập",
    ])

    add_paragraph_text(doc, "Vai trò và trách nhiệm của sinh viên:")

    add_paragraph_text(doc,
        "Sinh viên là người trực tiếp thực hiện toàn bộ các công việc được giao, chủ động "
        "nghiên cứu, học hỏi và áp dụng kiến thức chuyên môn để hoàn thành nhiệm vụ. "
        "Đồng thời, sinh viên có trách nhiệm tuân thủ kế hoạch thực tập, đảm bảo tiến độ "
        "công việc và chất lượng kết quả theo yêu cầu của cán bộ hướng dẫn.")

    add_subsection_title(doc, "2.3.1. Thiết kế cơ sở dữ liệu")

    add_paragraph_text(doc,
        "Trong quá trình thực tập, sinh viên được giao nhiệm vụ thiết kế cơ sở dữ liệu "
        "cho các module của hệ thống. Công việc bắt đầu với việc nghiên cứu cấu trúc "
        "database hiện có bao gồm: Entity Info (thông tin cơ bản về entity như tên logic, "
        "tên vật lý, hệ thống), Column Info (chi tiết các cột bao gồm tên logic, tên vật lý, "
        "kiểu dữ liệu, ràng buộc), và Index Info (thông tin về Primary Key, Foreign Key, các index).")

    add_paragraph_text(doc,
        "Các kiến thức về Oracle Database [1] được áp dụng trong quá trình thiết kế bao gồm:")

    add_bullet_list(doc, [
        "Kiểu dữ liệu: CHAR, VARCHAR2, NUMBER, DATE, CLOB, BLOB, JSON",
        "Constraints: PRIMARY KEY, FOREIGN KEY, NOT NULL, UNIQUE, CHECK",
        "Index: B-tree index, Bitmap index, Function-based index",
        "Naming convention: Quy tắc đặt tên chuẩn cho table, column, index",
    ])

    add_paragraph_text(doc,
        "Sinh viên học được cách định nghĩa bảng theo chuẩn với các thành phần: "
        "tên vật lý (sử dụng PascalCase hoặc snake_case), kiểu dữ liệu theo chuẩn Oracle, "
        "các ràng buộc (NOT NULL, UNIQUE, DEFAULT), và Primary Key (định danh duy nhất).")

    add_subsection_title(doc, "2.3.2. Thiết kế màn hình")

    add_paragraph_text(doc,
        "Thiết kế màn hình là quá trình xác định giao diện người dùng, bao gồm bố cục, "
        "các thành phần tương tác và luồng xử lý. Đây là loại thiết kế phức tạp nhất "
        "vì cần tham chiếu đến nhiều tài liệu liên quan.")

    add_paragraph_text(doc, "Các thành phần chính trong thiết kế màn hình bao gồm:")

    add_bullet_list(doc, [
        "Layout: Bố cục tổng thể của màn hình",
        "Items: Định nghĩa từng phần tử trên màn hình (ID, tên, loại, I/O, bắt buộc)",
        "Validation: Quy tắc kiểm tra dữ liệu nhập (đơn lẻ và tương quan)",
        "Messages: Các thông báo hiển thị cho người dùng",
        "Item Control: Điều khiển trạng thái các phần tử (hiển thị/ẩn, enable/disable)",
    ])

    add_paragraph_text(doc,
        "Validation được chia thành hai loại: Validate đơn lẻ (kiểm tra từng trường độc lập "
        "như bắt buộc nhập, độ dài, format, kiểu dữ liệu) và Validate tương quan (kiểm tra "
        "logic phụ thuộc giữa các trường như ngày kết thúc phải sau ngày bắt đầu).")

    add_subsection_title(doc, "2.3.3. Thiết kế API RESTful")

    add_paragraph_text(doc,
        "Thiết kế API là quá trình định nghĩa các endpoint để frontend và các hệ thống "
        "khác giao tiếp với backend. API trong dự án tuân theo kiến trúc RESTful [3], [5] với "
        "các nguyên tắc: Resource-based URL, sử dụng HTTP Methods chuẩn, Status Codes "
        "rõ ràng, và Content-Type là application/json.")

    add_paragraph_text(doc, "Cấu trúc thiết kế API bao gồm các thành phần:")

    add_bullet_list(doc, [
        "API ID: Mã định danh duy nhất cho mỗi API",
        "HTTP Method: GET (đọc), POST (tạo), PUT (cập nhật toàn bộ), PATCH (cập nhật một phần), DELETE (xóa)",
        "Endpoint: URL path của API",
        "Request Parameters và Request Body: Các tham số đầu vào",
        "Response: Định dạng phản hồi với status, data, message",
        "Error Handling: Xử lý các trường hợp lỗi với status code phù hợp",
    ])

    add_subsection_title(doc, "2.3.4. Thiết kế Batch Processing")

    add_paragraph_text(doc,
        "Batch processing (xử lý hàng loạt) là phương pháp xử lý khối lượng lớn dữ liệu "
        "theo lịch trình định sẵn, không cần tương tác người dùng. Công việc thiết kế batch "
        "áp dụng kiến trúc Spring Batch [2] với các thành phần: Job (đơn vị công việc cao nhất), "
        "Step (các bước trong một Job), ItemReader, ItemProcessor, và ItemWriter.")

    add_paragraph_text(doc, "Một thiết kế batch tiêu chuẩn bao gồm các phần:")

    add_bullet_list(doc, [
        "Tổng quan chức năng: Mục đích, đối tượng batch, cách khởi động",
        "Shell Script: Định nghĩa tham số đầu vào, mã kết thúc",
        "Xử lý Java: Logic xử lý chính với 5 block cơ bản (Chuẩn bị, Khởi tạo, Kiểm tra, Xử lý chính, Kết thúc)",
        "Yêu cầu tìm kiếm: Các câu SQL SELECT",
        "Yêu cầu cập nhật: Các câu SQL INSERT/UPDATE/DELETE",
    ])

    add_subsection_title(doc, "2.3.5. Quy trình xử lý Review (Shiteki)")

    add_paragraph_text(doc,
        "Sau khi hoàn thành thiết kế, sản phẩm được đưa qua quy trình review nhiều cấp: "
        "Leader Review → Customer Review → End-user Review. Mỗi cấp review sẽ đưa ra "
        "các phản hồi (shiteki) cần được xử lý trước khi chuyển sang cấp tiếp theo.")

    add_paragraph_text(doc, "Các loại shiteki thường gặp và cách xử lý:")

    add_bullet_list(doc, [
        "Lỗi chính tả (sai tên bảng, cột, API): Kiểm tra lại tài liệu tham khảo",
        "Lỗi logic (sai điều kiện, thiếu trường hợp): Phân tích lại nghiệp vụ",
        "Thiếu thông tin (thiếu validation, message): Bổ sung theo yêu cầu",
        "Không nhất quán (khác biệt giữa các phần): Đồng bộ toàn bộ thiết kế",
    ])

    # 2.4 Công nghệ, công cụ và kỹ thuật sử dụng
    add_section_title(doc, "2.4. Công nghệ, công cụ và kỹ thuật sử dụng")

    add_paragraph_text(doc,
        "Trong quá trình thực tập tại dự án SORA STEP4, em đã được tiếp cận và nghiên cứu nhiều "
        "công nghệ hiện đại được sử dụng rộng rãi trong phát triển phần mềm offshore tại Nhật Bản. "
        "Dưới đây là mô tả chi tiết về các công nghệ chính mà em đã học hỏi và áp dụng trong thực tế:")

    add_subsection_title(doc, "2.4.1. Oracle Database")

    add_paragraph_text(doc,
        "Oracle Database [1] là hệ quản trị cơ sở dữ liệu quan hệ hàng đầu thế giới, được sử dụng "
        "rộng rãi trong các doanh nghiệp lớn và dự án offshore tại Nhật Bản nhờ tính ổn định, bảo mật "
        "và hiệu năng cao. Oracle Database tuân thủ đầy đủ các thuộc tính ACID (Atomicity, Consistency, "
        "Isolation, Durability), đảm bảo tính toàn vẹn dữ liệu trong mọi tình huống.")

    add_paragraph_text(doc, "Các tính năng nổi bật của Oracle Database:")

    add_bullet_list(doc, [
        "ACID Compliance: Đảm bảo transaction được thực hiện toàn vẹn hoặc rollback hoàn toàn",
        "PL/SQL: Ngôn ngữ lập trình mở rộng của SQL, cho phép viết stored procedures, functions, triggers",
        "Constraints: Hỗ trợ đầy đủ PRIMARY KEY, FOREIGN KEY, NOT NULL, UNIQUE, CHECK constraints",
        "Advanced Indexing: B-tree index, Bitmap index, Function-based index để tối ưu truy vấn",
        "Data Types: Hỗ trợ đa dạng kiểu dữ liệu như VARCHAR2, NUMBER, DATE, CLOB, BLOB, JSON",
    ])

    add_paragraph_text(doc,
        "Trong dự án SORA STEP4, Oracle Database được sử dụng làm hệ quản trị cơ sở dữ liệu chính. "
        "Em đã học cách thiết kế database schema với naming convention chuẩn Nhật Bản, định nghĩa "
        "constraints để đảm bảo tính toàn vẹn dữ liệu, và tạo index để tối ưu hóa hiệu năng truy vấn. "
        "Kiến thức về Oracle data types giúp em chọn đúng kiểu dữ liệu phù hợp cho từng column, đảm bảo "
        "hiệu quả lưu trữ và xử lý.")

    add_subsection_title(doc, "2.4.2. Java Spring Boot")

    add_paragraph_text(doc,
        "Spring Boot [8] là framework phát triển ứng dụng Java phổ biến nhất hiện nay, "
        "được xây dựng dựa trên Spring Framework với mục tiêu đơn giản hóa việc cấu hình và "
        "triển khai ứng dụng. Spring Boot theo triết lý \"Convention over Configuration\", "
        "giúp developer tập trung vào logic nghiệp vụ thay vì cấu hình phức tạp.")

    add_paragraph_text(doc, "Các tính năng nổi bật của Spring Boot:")

    add_bullet_list(doc, [
        "Auto-configuration: Tự động cấu hình các thành phần dựa trên dependencies có trong classpath",
        "Dependency Injection: Quản lý dependencies giữa các component thông qua IoC container",
        "Spring Data JPA: Đơn giản hóa thao tác với database thông qua ORM, tự động sinh SQL",
        "Embedded Server: Tích hợp sẵn Tomcat/Jetty, không cần deploy file WAR riêng",
        "Layered Architecture: Tách biệt rõ ràng Controller - Service - Repository giúp code dễ bảo trì",
    ])

    add_paragraph_text(doc,
        "Trong dự án SORA STEP4, Spring Boot được sử dụng làm nền tảng backend để xây dựng các "
        "RESTful API phục vụ frontend. Em đã học cách áp dụng kiến trúc layered của Spring Boot: "
        "Controller xử lý HTTP requests, Service chứa business logic, Repository tương tác với database. "
        "Spring Data JPA giúp em thao tác với Oracle Database một cách dễ dàng thông qua các Entity và "
        "Repository interface.")

    add_subsection_title(doc, "2.4.3. Spring Batch")

    add_paragraph_text(doc,
        "Spring Batch [2] là framework mạnh mẽ cho việc xử lý batch processing trong Java, "
        "được thiết kế để xử lý khối lượng lớn dữ liệu một cách hiệu quả và đáng tin cậy. Spring Batch "
        "cung cấp các tính năng như transaction management, chunk processing, restart/retry mechanism, "
        "và job scheduling.")

    add_paragraph_text(doc, "Các thành phần chính của Spring Batch:")

    add_bullet_list(doc, [
        "Job: Đơn vị công việc cao nhất, đại diện cho một batch process hoàn chỉnh",
        "Step: Các bước xử lý trong một Job, có thể chạy tuần tự hoặc song song",
        "Chunk Processing: Xử lý dữ liệu theo từng chunk (ví dụ 1000 records/chunk) để tối ưu memory",
        "ItemReader: Đọc dữ liệu từ nguồn (database, file, API)",
        "ItemProcessor: Xử lý/biến đổi dữ liệu đã đọc",
        "ItemWriter: Ghi dữ liệu đã xử lý vào đích (database, file)",
    ])

    add_paragraph_text(doc,
        "Trong dự án SORA STEP4, Spring Batch được sử dụng để xử lý các tác vụ định kỳ như import "
        "dữ liệu hàng loạt, generate reports, synchronize data giữa các hệ thống. Em đã học cách thiết kế "
        "batch job với chunk processing để xử lý hiệu quả hàng triệu records, cấu hình retry logic khi gặp lỗi, "
        "và monitor batch execution status. Kiến trúc Job-Step-Chunk của Spring Batch giúp code dễ mở rộng "
        "và maintain.")

    add_subsection_title(doc, "2.4.4. RESTful API")

    add_paragraph_text(doc,
        "RESTful API [3], [5] là kiến trúc thiết kế API dựa trên các nguyên tắc của REST "
        "(Representational State Transfer), sử dụng HTTP protocol để giao tiếp giữa client và server. "
        "RESTful API được ưa chuộng nhờ tính đơn giản, dễ hiểu, và khả năng mở rộng tốt.")

    add_paragraph_text(doc, "Các nguyên tắc thiết kế RESTful API:")

    add_bullet_list(doc, [
        "Resource-based: Mọi thứ đều là resource, được định danh bằng URI (ví dụ: /users, /orders)",
        "HTTP Methods: Sử dụng GET (đọc), POST (tạo), PUT (cập nhật toàn bộ), PATCH (cập nhật từng phần), DELETE (xóa)",
        "Stateless: Mỗi request độc lập, server không lưu trữ session state của client",
        "Status Codes: Sử dụng HTTP status codes chuẩn (200 OK, 201 Created, 400 Bad Request, 404 Not Found, 500 Internal Server Error)",
        "HATEOAS: Response chứa links để client navigate đến các resource liên quan (optional)",
    ])

    add_paragraph_text(doc,
        "Trong dự án SORA STEP4, tất cả các API được thiết kế theo chuẩn RESTful. Em đã học cách "
        "định nghĩa endpoint theo resource-based URL, chọn HTTP method phù hợp cho từng thao tác, "
        "thiết kế request/response body với JSON format, và xử lý error với status code rõ ràng. "
        "Kỹ năng thiết kế RESTful API giúp em xây dựng các API dễ sử dụng, dễ document, và dễ maintain.")

    add_subsection_title(doc, "2.4.5. Git và Maven")

    add_paragraph_text(doc,
        "Git là hệ thống quản lý phiên bản phân tán (distributed version control system) phổ biến nhất "
        "hiện nay, cho phép nhiều developer làm việc đồng thời trên cùng một codebase. Maven là công cụ "
        "build automation và dependency management cho Java projects, giúp quản lý thư viện, compile code, "
        "chạy tests và package ứng dụng.")

    add_paragraph_text(doc, "Các tính năng chính:")

    add_bullet_list(doc, [
        "Git: Version control, branching/merging, conflict resolution, commit history, remote repository (GitHub/GitLab)",
        "Maven: Dependency management với pom.xml, build lifecycle (compile, test, package, install), plugin ecosystem",
        "Git Workflow: Feature branching, pull request, code review, merge to main branch",
        "Maven Convention: Cấu trúc thư mục chuẩn (src/main/java, src/test/java), naming convention",
        "CI/CD Integration: Tích hợp với Jenkins/GitLab CI để tự động build và deploy",
    ])

    add_paragraph_text(doc,
        "Trong dự án SORA STEP4, Git được sử dụng để quản lý mã nguồn thiết kế và tài liệu. Em đã học "
        "cách commit code với message rõ ràng, tạo branch cho từng feature, merge code sau khi review, "
        "và resolve conflicts khi có. Maven được sử dụng để quản lý dependencies của Spring Boot project, "
        "compile code, chạy tests, và package thành JAR file. IntelliJ IDEA được sử dụng làm IDE chính, "
        "tích hợp sẵn Git và Maven để thuận tiện trong development.")

    add_subsection_title(doc, "2.4.6. Tổng hợp công nghệ sử dụng")

    add_table_with_caption(doc, 2, "Công nghệ và công cụ sử dụng trong thực tập SORA STEP4",
        ["Loại", "Tên", "Mục đích"],
        [
            ("Database", "Oracle Database [1]", "Lưu trữ dữ liệu quan hệ, đảm bảo ACID"),
            ("Backend Framework", "Java Spring Boot [8]", "Xây dựng RESTful API và business logic"),
            ("Batch Processing", "Spring Batch [2]", "Xử lý dữ liệu hàng loạt, scheduled jobs"),
            ("API Architecture", "RESTful API [3], [5]", "Thiết kế API theo chuẩn REST"),
            ("Version Control", "Git, GitHub", "Quản lý mã nguồn và collaboration"),
            ("Build Tool", "Maven", "Dependency management và build automation"),
            ("IDE", "IntelliJ IDEA", "Môi trường phát triển Java"),
            ("AI Tools", "Claude AI [4]", "Hỗ trợ kiểm tra chất lượng thiết kế"),
        ],
        col_widths=[4.0, 5.5, 6.5]
    )


# ============== CHƯƠNG 3: KẾT QUẢ VÀ ĐÁNH GIÁ ==============
def add_chapter3(doc):
    """Chương 3: Kết quả và đánh giá - theo mẫu tham khảo"""
    add_chapter_title(doc, "3", "KẾT QUẢ VÀ ĐÁNH GIÁ")

    # 3.1 Kết quả đạt được
    add_section_title(doc, "3.1. Kết quả đạt được trong quá trình thực tập")

    add_paragraph_text(doc,
        "Trong suốt thời gian thực tập tại công ty, em đã hoàn thành "
        "đầy đủ các nội dung và nhiệm vụ theo kế hoạch thực tập đã đề ra dưới sự hướng dẫn của "
        "cán bộ phụ trách. Các công việc được giao đều được thực hiện nghiêm túc, đúng "
        "tiến độ và đảm bảo yêu cầu về chất lượng.")

    add_paragraph_text(doc,
        "Thông qua quá trình thực tập, em đã từng bước tiếp cận với các công việc chuyên "
        "môn trong lĩnh vực thiết kế hệ thống phần mềm, từ việc nghiên cứu tài liệu, phân tích yêu cầu "
        "đến triển khai và hoàn thiện các nội dung liên quan đến đề tài thực tập. Kết quả đạt được "
        "không chỉ thể hiện qua sản phẩm hoặc nội dung công việc đã hoàn thành mà còn ở sự "
        "tiến bộ rõ rệt về tư duy, kỹ năng và thái độ làm việc.")

    add_paragraph_text(doc, "Kết quả đạt được:")

    add_bullet_list(doc, [
        "Hoàn thành các nội dung công việc theo kế hoạch đề ra",
        "Hoàn thành thiết kế cơ sở dữ liệu cho các module được giao đạt tiêu chuẩn",
        "Hoàn thành thiết kế màn hình với đầy đủ validation và message",
        "Hoàn thành thiết kế API RESTful theo chuẩn công ty",
        "Tham gia sử dụng và góp ý cải tiến hệ thống AI Checker",
        "Nâng cao kỹ năng thiết kế, tư duy logic và khả năng giải quyết vấn đề",
        "Hoàn thành báo cáo thực tập đúng quy định, phản ánh trung thực quá trình thực hiện",
    ])

    # 3.2 Kiến thức và kỹ năng tích lũy được
    add_section_title(doc, "3.2. Kiến thức và kỹ năng tích lũy được")

    add_subsection_title(doc, "3.2.1. Kiến thức chuyên môn")

    add_paragraph_text(doc,
        "Trong quá trình thực tập, em đã củng cố và mở rộng các kiến thức chuyên môn đã "
        "được học trên giảng đường. Đặc biệt là kiến thức về thiết kế hệ thống, phân tích yêu cầu, "
        "cơ sở dữ liệu và quy trình phát triển phần mềm. Việc áp dụng lý thuyết vào các "
        "bài toán thực tế giúp em hiểu sâu hơn bản chất vấn đề, đồng thời nâng cao khả năng vận "
        "dụng kiến thức vào thực tiễn.")

    add_bullet_list(doc, [
        "Nắm vững quy trình thiết kế hệ thống phần mềm chuyên nghiệp theo chuẩn Nhật Bản",
        "Hiểu sâu về thiết kế cơ sở dữ liệu với Oracle Database [1]",
        "Biết cách thiết kế API RESTful [3], [5] theo chuẩn",
        "Hiểu về kiến trúc Spring Batch [2] và thiết kế batch processing",
        "Biết cách sử dụng AI (Claude) [4] hỗ trợ kiểm tra chất lượng thiết kế",
    ])

    add_subsection_title(doc, "3.2.2. Kỹ năng làm việc nhóm")

    add_paragraph_text(doc,
        "Quá trình thực tập giúp em rèn luyện kỹ năng làm việc nhóm thông qua việc trao "
        "đổi, thảo luận và phối hợp với cán bộ hướng dẫn và các đồng nghiệp. Em học được "
        "cách lắng nghe ý kiến đóng góp, chia sẻ công việc hợp lý, hỗ trợ lẫn nhau để hoàn thành "
        "nhiệm vụ chung.")

    add_bullet_list(doc, [
        "Kỹ năng giao tiếp và trao đổi với team member trong môi trường chuyên nghiệp",
        "Kỹ năng review thiết kế và tiếp nhận feedback (shiteki)",
        "Kỹ năng làm việc với khách hàng nước ngoài (thông qua BrSE)",
        "Kỹ năng giao tiếp trong môi trường học thuật và chuyên môn",
    ])

    add_subsection_title(doc, "3.2.3. Kỹ năng phân tích và giải quyết vấn đề")

    add_paragraph_text(doc,
        "Thông qua các công việc được giao, em đã rèn luyện khả năng phân tích yêu cầu, "
        "xác định vấn đề và đề xuất hướng giải quyết phù hợp. Khi gặp khó khăn trong quá trình "
        "thực hiện, em học được cách chủ động tìm kiếm tài liệu, tham khảo ý kiến mentor "
        "và tự đánh giá, điều chỉnh phương án thực hiện.")

    add_bullet_list(doc, [
        "Kỹ năng đọc hiểu và phân tích yêu cầu nghiệp vụ từ tài liệu khách hàng",
        "Kỹ năng viết QA (Question & Answer) để xác nhận yêu cầu không rõ ràng",
        "Kỹ năng tư duy logic, xử lý tình huống và khắc phục lỗi trong thiết kế",
        "Kỹ năng sử dụng AI hỗ trợ công việc một cách hiệu quả",
    ])

    # 3.3 Thuận lợi và khó khăn
    add_section_title(doc, "3.3. Thuận lợi và khó khăn")

    add_paragraph_text(doc, "Thuận lợi:")

    add_bullet_list(doc, [
        "Nhận được sự quan tâm, hướng dẫn tận tình của cán bộ hướng dẫn trong suốt quá trình thực tập",
        "Môi trường học tập và làm việc tại công ty thân thiện, chuyên nghiệp, "
        "tạo điều kiện thuận lợi cho việc học hỏi và nghiên cứu",
        "Được tiếp cận với các tài liệu chuyên môn, cơ sở vật chất và công cụ hỗ trợ "
        "phục vụ cho quá trình thực tập",
        "Kiến thức nền tảng đã được trang bị trong quá trình học tập tại trường giúp "
        "sinh viên dễ dàng tiếp cận nội dung thực tập",
    ])

    add_paragraph_text(doc, "Khó khăn:")

    add_bullet_list(doc, [
        "Một số kiến thức và công nghệ còn mới (Oracle Database [1], Spring Batch [2], Java SE [6]), "
        "đòi hỏi em phải tự nghiên cứu và học hỏi thêm trong thời gian ngắn",
        "Kinh nghiệm thực tiễn còn hạn chế nên trong giai đoạn đầu gặp một số khó khăn "
        "khi triển khai công việc",
        "Thời gian thực tập có hạn, trong khi khối lượng công việc và yêu cầu chuyên môn "
        "tương đối nhiều",
        "Kiến thức tiếng Nhật còn hạn chế trong việc đọc hiểu một số tài liệu",
    ])

    add_paragraph_text(doc,
        "Tuy nhiên, nhờ sự hướng dẫn của cán bộ hướng dẫn và sự nỗ lực của bản thân, em đã từng bước "
        "khắc phục được những khó khăn trên và hoàn thành tốt đợt thực tập.")


# ============== CHƯƠNG 4: NHẬN XÉT VÀ ĐỊNH HƯỚNG ==============
def add_chapter4(doc):
    """Chương 4: Nhận xét và định hướng - theo mẫu tham khảo"""
    add_chapter_title(doc, "4", "NHẬN XÉT VÀ ĐỊNH HƯỚNG")

    # 4.1 Nhận xét chung về đợt thực tập
    add_section_title(doc, "4.1. Nhận xét chung về đợt thực tập")

    add_paragraph_text(doc,
        f"Đợt thực tập tại {INTERNSHIP_INFO['company']} "
        "là một trải nghiệm học tập có ý nghĩa và mang lại nhiều giá trị thiết thực đối với em.")

    add_paragraph_text(doc,
        "Thông qua quá trình thực tập, em đã có cơ hội tiếp cận với môi trường làm việc mang "
        "tính chuyên nghiệp cao, từ đó hiểu rõ hơn về yêu cầu và tính chất công việc "
        "trong lĩnh vực phát triển phần mềm offshore cho thị trường Nhật Bản.")

    add_paragraph_text(doc,
        "Các nội dung thực tập được xây dựng phù hợp với chương trình đào tạo, gắn liền "
        "giữa lý thuyết và thực tiễn, giúp em từng bước làm quen với quy trình làm việc, phương "
        "pháp thiết kế và triển khai các nhiệm vụ chuyên môn. Sự hướng dẫn tận tình của "
        "cán bộ hướng dẫn đã giúp em định hướng đúng đắn, kịp thời khắc phục những hạn chế trong "
        "quá trình thực hiện.")

    add_paragraph_text(doc,
        "Nhìn chung, đợt thực tập đã đạt được các mục tiêu đề ra, góp phần nâng cao kiến "
        "thức, kỹ năng và ý thức nghề nghiệp của em, đồng thời tạo nền tảng quan trọng cho quá trình "
        "học tập và làm việc sau này.")

    # 4.2 Bài học kinh nghiệm rút ra
    add_section_title(doc, "4.2. Bài học kinh nghiệm rút ra")

    add_paragraph_text(doc,
        "Từ quá trình thực tập, em đã rút ra được nhiều bài học kinh nghiệm quý báu. "
        "Trước hết là bài học về tinh thần tự giác và chủ động học tập. Trong môi trường thực "
        "tế, việc tự tìm hiểu tài liệu, chủ động đặt câu hỏi và đề xuất giải pháp là yếu tố quan "
        "trọng giúp nâng cao hiệu quả công việc.")

    add_paragraph_text(doc,
        "Bên cạnh đó, em nhận thức rõ hơn về tầm quan trọng của việc nắm vững kiến "
        "thức nền tảng và khả năng vận dụng linh hoạt kiến thức đã học vào các tình huống cụ "
        "thể. Việc làm việc theo kế hoạch, tuân thủ quy trình và đảm bảo tiến độ cũng là những "
        "kinh nghiệm cần thiết được rút ra trong suốt quá trình thực tập.")

    add_paragraph_text(doc, "Các bài học kinh nghiệm cụ thể:")

    add_bullet_list(doc, [
        "Cần chủ động trong việc học hỏi và đặt câu hỏi (viết QA) khi gặp khó khăn",
        "Tầm quan trọng của việc đọc kỹ tài liệu và hiểu rõ yêu cầu trước khi thực hiện",
        "Cần kiểm tra kỹ lưỡng (self-review) trước khi gửi sản phẩm cho leader review",
        "Kỹ năng mềm (giao tiếp, làm việc nhóm) quan trọng không kém kỹ năng chuyên môn",
        "Học cách tiếp nhận ý kiến góp ý (shiteki) một cách nghiêm túc và cải thiện",
    ])

    # 4.3 Định hướng nghề nghiệp và học tập sau thực tập
    add_section_title(doc, "4.3. Định hướng nghề nghiệp và học tập sau thực tập")

    add_paragraph_text(doc,
        "Sau đợt thực tập, em đã có cái nhìn rõ ràng hơn về định hướng nghề nghiệp trong "
        "tương lai. Trên cơ sở những kiến thức và kỹ năng đã tích lũy được, em định hướng tiếp "
        "tục nâng cao trình độ chuyên môn trong lĩnh vực công nghệ thông tin, đặc biệt là các "
        "mảng liên quan đến thiết kế hệ thống và phát triển phần mềm.")

    add_paragraph_text(doc, "Kết nối với đồ án tốt nghiệp:")

    add_paragraph_text(doc,
        "Đợt thực tập này có ý nghĩa quan trọng như một bước chuẩn bị nền tảng cho đồ án "
        "tốt nghiệp của em với đề tài \"XÂY DỰNG HỆ THỐNG SAAS CUNG CẤP DỊCH VỤ ĐÀO TẠO\". "
        "Các kiến thức và kỹ năng tích lũy được trong quá trình thực tập có liên hệ trực tiếp "
        "với đồ án tốt nghiệp:")

    add_bullet_list(doc, [
        "Kinh nghiệm thiết kế cơ sở dữ liệu Oracle sẽ được áp dụng để xây dựng database schema "
        "cho hệ thống SaaS đào tạo với các entity như Organization, User, Course, Enrollment, Assessment",
        "Kiến thức về kiến trúc Multi-tenant học được từ các dự án offshore sẽ được vận dụng để "
        "thiết kế hệ thống cho phép nhiều tổ chức đào tạo sử dụng chung nền tảng SaaS",
        "Kỹ năng thiết kế API RESTful sẽ được áp dụng để xây dựng backend API cho hệ thống SaaS "
        "sử dụng Java Spring Boot và PostgreSQL",
        "Kinh nghiệm thiết kế batch processing với Spring Batch sẽ giúp em triển khai các tính năng "
        "như import dữ liệu hàng loạt, generate báo cáo định kỳ, đồng bộ dữ liệu giữa các module",
        "Quy trình làm việc chuyên nghiệp học được từ công ty sẽ được áp dụng vào quản lý "
        "dự án đồ án tốt nghiệp: phân tích yêu cầu, thiết kế hệ thống, review thiết kế, testing",
    ])

    add_paragraph_text(doc,
        "Trong thời gian tới, em sẽ tập trung củng cố kiến thức chuyên ngành, học hỏi "
        "thêm các công nghệ mới, nâng cao kỹ năng lập trình, kỹ năng làm việc nhóm và kỹ năng "
        "nghiên cứu. Đồng thời, em cũng sẽ chuẩn bị tốt cho đồ án tốt nghiệp và sẵn sàng tham "
        "gia vào môi trường làm việc chuyên nghiệp sau khi ra trường.")

    add_paragraph_text(doc, "Các mục tiêu cụ thể:")

    add_bullet_list(doc, [
        "Hoàn thành tốt đồ án tốt nghiệp về hệ thống SaaS đào tạo - áp dụng kiến thức đã tích lũy",
        "Tiếp tục học hỏi và nâng cao kỹ năng lập trình (Java, Spring Boot, PostgreSQL)",
        "Tìm hiểu sâu hơn về AI và ứng dụng trong phát triển phần mềm",
        "Cải thiện kỹ năng ngoại ngữ (tiếng Anh, tiếng Nhật) để làm việc trong môi trường quốc tế",
        "Xây dựng portfolio cá nhân với các dự án thực tế",
    ])

    add_paragraph_text(doc,
        "Đợt thực tập là bước đệm quan trọng giúp em xác định rõ mục tiêu học tập và "
        "nghề nghiệp, tạo động lực để không ngừng rèn luyện và phát triển bản thân trong tương lai.")

    # 4.4 Những đóng góp của đề tài
    add_section_title(doc, "4.4. Những đóng góp của đề tài")

    add_subsection_title(doc, "4.4.1. Đóng góp về mặt sản phẩm")

    add_bullet_list(doc, [
        "Hoàn thành các thiết kế cơ sở dữ liệu, màn hình, API theo chuẩn doanh nghiệp Nhật Bản",
        "Tham gia vào quy trình thiết kế Batch Processing theo kiến trúc Spring Batch",
        "Đóng góp vào việc cải tiến chất lượng thiết kế thông qua quy trình review nhiều cấp",
        "Xây dựng sản phẩm hoặc mô hình phục vụ học tập và nghiên cứu",
    ])

    add_subsection_title(doc, "4.4.2. Đóng góp về mặt quy trình")

    add_bullet_list(doc, [
        "Hiểu và áp dụng quy trình làm việc thực tế trong doanh nghiệp offshore Nhật Bản",
        "Nắm vững quy trình xử lý Shiteki (review feedback) từ nhiều cấp độ: "
        "Leader → Customer → End-user",
        "Tích lũy kinh nghiệm sử dụng AI (Claude) hỗ trợ kiểm tra chất lượng thiết kế",
        "Hình thành phong cách làm việc chủ động và khoa học",
    ])

    add_subsection_title(doc, "4.4.3. Đóng góp về mặt kiến thức")

    add_bullet_list(doc, [
        "Tổng hợp kiến thức về thiết kế hệ thống phần mềm theo chuẩn quốc tế",
        "Chuẩn bị nền tảng kiến thức vững chắc cho đồ án tốt nghiệp về hệ thống SaaS đào tạo",
        "Tài liệu báo cáo có thể làm tham khảo cho các sinh viên khóa sau về quy trình "
        "thực tập tại doanh nghiệp offshore",
    ])


# ============== TÀI LIỆU THAM KHẢO ==============
def add_ieee_reference(doc, ref_num, author, title, source_type, year, url=None, accessed=None, publisher=None, pages=None):
    """
    Thêm một tài liệu tham khảo theo chuẩn IEEE Citation Style

    IEEE Standard Format:
    - Online: [1] A. Author, "Title," Website/Source, Month Day, Year. [Online]. Available: URL. (accessed Month Day, Year).
    - Book: [1] A. Author, Title. City: Publisher, Year.
    - Journal: [1] A. Author, "Article title," Journal Name, vol. X, no. Y, pp. Z, Year.

    Args:
        doc: Document object
        ref_num: Số thứ tự tài liệu [1], [2]...
        author: Tên tác giả (e.g., "J. Smith" hoặc "Smith, J.")
        title: Tiêu đề tài liệu
        source_type: Loại nguồn ("online", "book", "journal", "conference")
        year: Năm xuất bản (e.g., "2024")
        url: URL (cho tài liệu online)
        accessed: Ngày truy cập (e.g., "Jan. 15, 2026")
        publisher: Nhà xuất bản (cho sách)
        pages: Số trang (cho journal/conference)
    """
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.63)
    p.paragraph_format.first_line_indent = Cm(-0.63)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # [1]
    run = p.add_run(f"[{ref_num}] ")
    set_font(run, FONT_SIZE_NORMAL)

    # Author,
    run = p.add_run(f"{author}, ")
    set_font(run, FONT_SIZE_NORMAL)

    # "Title" (in nghiêng cho article/online, không nghiêng cho book)
    if source_type.lower() in ["online", "journal", "conference"]:
        run = p.add_run(f'"{title}," ')
        set_font(run, FONT_SIZE_NORMAL, italic=True)
    else:  # book
        run = p.add_run(f"{title}. ")
        set_font(run, FONT_SIZE_NORMAL, italic=True)

    # Format theo loại nguồn
    if source_type.lower() == "online" and url:
        # Online format: Website, Year. [Online]. Available: URL. (accessed Date).
        run = p.add_run(f"{year}. ")
        set_font(run, FONT_SIZE_NORMAL)

        run = p.add_run("[Online]. Available: ")
        set_font(run, FONT_SIZE_NORMAL)

        run = p.add_run(url)
        set_font(run, FONT_SIZE_NORMAL)
        run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color for URL
        run.font.underline = True

        if accessed:
            run = p.add_run(f" (accessed {accessed})")
            set_font(run, FONT_SIZE_NORMAL)

        run = p.add_run(".")
        set_font(run, FONT_SIZE_NORMAL)

    elif source_type.lower() == "book":
        # Book format: City: Publisher, Year.
        if publisher:
            run = p.add_run(f"{publisher}, ")
            set_font(run, FONT_SIZE_NORMAL)
        run = p.add_run(f"{year}.")
        set_font(run, FONT_SIZE_NORMAL)

    elif source_type.lower() == "journal":
        # Journal format: Journal Name, vol. X, no. Y, pp. Z, Year.
        if pages:
            run = p.add_run(f"pp. {pages}, ")
            set_font(run, FONT_SIZE_NORMAL)
        run = p.add_run(f"{year}.")
        set_font(run, FONT_SIZE_NORMAL)

    else:
        # Generic format
        run = p.add_run(f"{year}.")
        set_font(run, FONT_SIZE_NORMAL)


def add_references(doc):
    """
    Tài liệu tham khảo theo chuẩn IEEE
    Sử dụng Heading 1 để có thể thêm vào mục lục tự động
    """
    doc.add_page_break()

    # Sử dụng Heading 1 để có thể tạo mục lục
    p = doc.add_paragraph(style='Heading 1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Căn giữa theo quy định UTC

    run = p.add_run("TÀI LIỆU THAM KHẢO")
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = FONT_SIZE_CHAPTER  # 18pt theo quy định UTC
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    # IEEE format references - Ví dụ chuẩn IEEE Citation Style

    # [1] Online resource - Database documentation
    add_ieee_reference(doc,
        ref_num=1,
        author="Oracle Corporation",
        title="Oracle Database 19c Documentation",
        source_type="online",
        year="2024",
        url="https://docs.oracle.com/en/database/oracle/oracle-database/19/",
        accessed="Jan. 15, 2026"
    )

    # [2] Online resource - Framework documentation
    add_ieee_reference(doc,
        ref_num=2,
        author="VMware Inc.",
        title="Spring Batch Reference Documentation v5.0",
        source_type="online",
        year="2024",
        url="https://docs.spring.io/spring-batch/docs/current/reference/html/",
        accessed="Jan. 18, 2026"
    )

    # [3] Online resource - API design guide
    add_ieee_reference(doc,
        ref_num=3,
        author="M. Masse",
        title="REST API Design Rulebook",
        source_type="book",
        year="2011",
        publisher="O'Reilly Media"
    )

    # [4] Online resource - API documentation
    add_ieee_reference(doc,
        ref_num=4,
        author="Anthropic PBC",
        title="Claude API Reference Documentation",
        source_type="online",
        year="2025",
        url="https://docs.anthropic.com/en/api/",
        accessed="Jan. 20, 2026"
    )

    # [5] Technical standard
    add_ieee_reference(doc,
        ref_num=5,
        author="R. T. Fielding",
        title="Architectural Styles and the Design of Network-based Software Architectures",
        source_type="online",
        year="2000",
        url="https://www.ics.uci.edu/~fielding/pubs/dissertation/top.htm",
        accessed="Jan. 10, 2026"
    )

    # [6] Java documentation
    add_ieee_reference(doc,
        ref_num=6,
        author="Oracle Corporation",
        title="Java SE 17 Documentation",
        source_type="online",
        year="2024",
        url="https://docs.oracle.com/en/java/javase/17/",
        accessed="Jan. 12, 2026"
    )

    # [7] Tài liệu nội bộ
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.63)
    p.paragraph_format.first_line_indent = Cm(-0.63)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("[7] ")
    set_font(run, FONT_SIZE_NORMAL)
    run = p.add_run("Tài liệu thiết kế nội bộ dự án SORA STEP4, SY PARTNERS., JSC (không công khai), 2025.")
    set_font(run, FONT_SIZE_NORMAL)

    # [8] Spring Boot documentation
    add_ieee_reference(doc,
        ref_num=8,
        author="VMware Inc.",
        title="Spring Boot Reference Documentation v3.2",
        source_type="online",
        year="2024",
        url="https://docs.spring.io/spring-boot/docs/current/reference/html/",
        accessed="Jan. 25, 2026"
    )

    # [9] Maven documentation
    add_ieee_reference(doc,
        ref_num=9,
        author="Apache Software Foundation",
        title="Maven Documentation - Welcome to Apache Maven",
        source_type="online",
        year="2025",
        url="https://maven.apache.org/guides/",
        accessed="Jan. 28, 2026"
    )

    # [10] Git documentation
    add_ieee_reference(doc,
        ref_num=10,
        author="Software Freedom Conservancy",
        title="Git Documentation - Reference Manual",
        source_type="online",
        year="2025",
        url="https://git-scm.com/doc",
        accessed="Jan. 22, 2026"
    )

    # [11] IntelliJ IDEA documentation
    add_ieee_reference(doc,
        ref_num=11,
        author="JetBrains s.r.o.",
        title="IntelliJ IDEA Documentation - The Java IDE",
        source_type="online",
        year="2025",
        url="https://www.jetbrains.com/idea/documentation/",
        accessed="Jan. 30, 2026"
    )

    # [12] Database normalization concepts
    add_ieee_reference(doc,
        ref_num=12,
        author="E. F. Codd",
        title="Further Normalization of the Data Base Relational Model",
        source_type="book",
        year="1972",
        publisher="IBM Research Report RJ909"
    )

    # [13] Microservices Architecture Book
    add_ieee_reference(doc,
        ref_num=13,
        author="C. Richardson",
        title="Microservices Patterns: With Examples in Java",
        source_type="book",
        year="2018",
        publisher="Manning Publications"
    )


# ============== PHỤ LỤC ==============
def add_appendix(doc):
    """Phụ lục - sử dụng Heading 1 để có thể thêm vào mục lục"""
    doc.add_page_break()

    # Sử dụng Heading 1 để có thể tạo mục lục
    p = doc.add_paragraph(style='Heading 1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Căn giữa theo quy định UTC

    run = p.add_run("PHỤ LỤC")
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = FONT_SIZE_CHAPTER  # 18pt theo quy định UTC
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Phụ lục A: Nhật ký thực tập chi tiết (theo mẫu báo cáo tham khảo)
    add_section_title(doc, "Phụ lục A: Nhật ký thực tập chi tiết")

    add_paragraph_text(doc,
        f"Thời gian thực tập: Từ ngày {INTERNSHIP_INFO['start_date']} đến ngày {INTERNSHIP_INFO['end_date']}",
        first_line_indent=False)

    # Create diary table
    diary_data = [
        ("1", "01/12 – 07/12/2025",
         "Onboarding, tìm hiểu dự án SORA STEP4, quy trình làm việc, công cụ (Git, IntelliJ, Oracle SQL Developer), ôn tập Java/Spring Boot/Oracle Database, đọc Business Requirement Document",
         "Hiểu quy trình phát triển offshore theo chuẩn Nhật Bản, cấu trúc tổ chức dự án, công cụ quản lý. Hoàn thành onboarding, nắm overview dự án"),

        ("2", "08/12 – 14/12/2025",
         "Training Database Design, học Oracle Database (data types, constraints, indexes), naming convention, nghiên cứu cấu trúc DB SORA, đọc ERD",
         "Nắm vững kiểu dữ liệu Oracle (VARCHAR2, NUMBER, DATE, CLOB, BLOB), constraints (PK, FK, NOT NULL, UNIQUE, CHECK), indexes (B-tree, Bitmap, Function-based). Hoàn thành training DB Design"),

        ("3", "15/12 – 21/12/2025",
         "Thiết kế bảng cho module User Management, định nghĩa columns (data type, length, constraints), tạo Entity Info và Column Info documents, review với mentor và xử lý feedback đợt 1",
         "Phân tích yêu cầu nghiệp vụ xác định entities/relationships, áp dụng normalization (1NF, 2NF, 3NF), viết tài liệu thiết kế DB theo chuẩn. Hoàn thành thiết kế 5 bảng, pass review sau 1 lần sửa"),

        ("4", "22/12 – 28/12/2025",
         "Thiết kế indexes, tạo Index Info document (loại index, columns, mục đích), học query optimization và EXPLAIN PLAN, training Screen Design, tìm hiểu wireframe và UI/UX principles",
         "Hiểu khi nào tạo index và loại phù hợp (B-tree vs Bitmap), analyze query performance, quy trình thiết kế màn hình từ wireframe đến detailed design. Hoàn thành Index Info cho tất cả bảng"),

        ("5", "29/12 – 04/01/2026",
         "Thiết kế màn hình User List, định nghĩa layout (header, search bar, table, pagination, buttons), tạo Items document (ID, tên, loại, I/O type, required flag), validation rules, messages",
         "Phân tích màn hình thành components, hiểu validation (client-side vs server-side), viết validation rules đầy đủ edge cases. Hoàn thành thiết kế User List với items, validation, messages"),

        ("6", "05/01 – 11/01/2026",
         "Thiết kế màn hình User Detail, correlation validation (logic phụ thuộc giữa trường), Item Control document (hiển thị/ẩn, enable/disable fields), training RESTful API (HTTP methods, status codes, request/response)",
         "Thiết kế validation phức tạp multiple conditions, state management (view vs edit mode), nguyên tắc RESTful API. Hoàn thành thiết kế User Detail, sẵn sàng API Design"),

        ("7", "12/01 – 18/01/2026",
         "Thiết kế API cho User Management (GET/POST/PUT/DELETE /users), Request Parameters (query/path params, body), Response structure (status, data, message, pagination), Error Handling (400/401/403/404/500), giới thiệu Claude AI [4] để review thiết kế",
         "Thiết kế RESTful API resource-based, hiểu idempotency (GET, PUT, DELETE), sử dụng AI tools review và cải thiện. Hoàn thành 5 APIs chính với đầy đủ request/response specs"),

        ("8", "19/01 – 25/01/2026",
         "Thiết kế APIs cho Course Management, sử dụng Claude AI kiểm tra chất lượng, fix issues (naming inconsistency, missing error cases, incomplete docs), training Spring Batch (Job, Step, ItemReader, ItemProcessor, ItemWriter)",
         "Viết API documentation rõ ràng đầy đủ, hiểu chunk processing, integrate AI vào workflow nâng cao chất lượng. Hoàn thành APIs cho Course Management, hiểu cơ bản Spring Batch"),

        ("9", "26/01 – 01/02/2026",
         "Thiết kế Batch Job Import User từ CSV, Job structure 5 blocks (Chuẩn bị, Khởi tạo, Kiểm tra, Xử lý, Kết thúc), Shell Script (input params, return codes), pseudo code ItemReader/Processor/Writer, SQL queries (SELECT, INSERT, UPDATE)",
         "Phân tích batch requirements chia steps, transaction management trong batch, xử lý error và rollback chunk processing. Hoàn thành thiết kế Batch Job đầu tiên với 5 blocks và SQL"),

        ("10", "02/02 – 08/02/2026",
         "Thiết kế Batch Job Generate Monthly Report, áp dụng chunk processing (size=1000), xử lý feedback (Shiteki) đợt 1 từ Leader Review, fix lỗi (logic sai, thiếu validation, inconsistent naming), refactor đảm bảo nhất quán DB/Screen/API/Batch",
         "Xử lý review feedback có hệ thống, hiểu tầm quan trọng consistency, tự review trước submit. Hoàn thành Batch Job thứ 2, xử lý 80% feedback đợt 1"),

        ("11", "09/02 – 15/02/2026",
         "Hoàn thành feedback đợt 1 submit lại, thiết kế độc lập module Enrollment Management, áp dụng tất cả kiến thức (DB/Screen/API/Batch), sử dụng Claude AI tự review, tổng hợp tài liệu theo chuẩn template",
         "Làm việc độc lập từ phân tích đến hoàn thiện thiết kế, tự kiểm tra chất lượng và phát hiện lỗi, time management đúng deadline. Hoàn thành Enrollment Management, pass review ngay lần đầu"),

        ("12", "16/02 – 22/02/2026",
         "Xử lý feedback đợt 2 từ Customer Review, tinh chỉnh theo góp ý khách hàng, kiểm tra cross-reference (DB↔API, Screen↔API), viết báo cáo thực tập, chuẩn bị tài liệu tổng hợp (ERD, API spec, Batch flow diagram)",
         "Communication với khách hàng qua tài liệu, tầm quan trọng cross-reference checking, viết báo cáo kỹ thuật chuẩn học thuật. Hoàn thành xử lý feedback đợt 2, bắt đầu báo cáo"),

        ("13", "23/02 – 01/03/2026",
         "Hoàn thiện tất cả tài liệu (DB/Screen/API/Batch), tạo tổng hợp (System Architecture, Database ERD, API Documentation), viết hoàn chỉnh báo cáo theo chuẩn UTC, chuẩn bị slide thuyết trình, nộp sản phẩm và nhận feedback tích cực",
         "Tổng hợp và trình bày kết quả, viết báo cáo kỹ thuật chuyên nghiệp, kỹ năng thuyết trình và communication. Hoàn thành đầy đủ sản phẩm, đánh giá cao từ mentor và công ty"),
    ]

    # Create table (14 rows = 1 header + 13 data rows)
    table = doc.add_table(rows=14, cols=4)
    table.style = 'Table Grid'

    # Set column widths
    table.columns[0].width = Cm(1.2)  # Tuần
    table.columns[1].width = Cm(3.5)  # Thời gian
    table.columns[2].width = Cm(6.5)  # Nội dung công việc
    table.columns[3].width = Cm(5.3)  # Kỹ năng & Kết quả

    # Header row
    header_cells = table.rows[0].cells
    headers = ["Tuần", "Thời gian", "Nội dung công việc", "Kỹ năng & Kết quả"]
    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        cell.text = header_text

        # Format header cell
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_TABLE
        run.font.bold = True

        # Cell shading (light gray)
        shading = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
        cell._element.get_or_add_tcPr().append(shading)

        # Vertical alignment
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for row_idx, (week, period, work, skills) in enumerate(diary_data, start=1):
        row_cells = table.rows[row_idx].cells

        # Week number (centered)
        row_cells[0].text = week
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Period (centered)
        row_cells[1].text = period
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Work content (justified)
        row_cells[2].text = work
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        row_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # Skills & Results (justified)
        row_cells[3].text = skills
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        row_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # Format all cells
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = FONT_NAME
                    run.font.size = FONT_SIZE_TABLE
                # Set spacing
                paragraph.space_before = Pt(0)
                paragraph.space_after = Pt(0)
                paragraph_format = paragraph.paragraph_format
                paragraph_format.line_spacing = 1.15

    doc.add_paragraph()  # Spacing

    # Phụ lục B: Hình ảnh, tài liệu minh chứng
    add_section_title(doc, "Phụ lục B: Hình ảnh, tài liệu minh chứng")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Đính kèm hình ảnh minh chứng quá trình thực tập]")
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_NORMAL
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Phụ lục C: Sản phẩm thực tập
    add_section_title(doc, "Phụ lục C: Sản phẩm thực tập")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Mô tả hoặc đính kèm các sản phẩm thiết kế đã hoàn thành]")
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_NORMAL
    run.font.color.rgb = RGBColor(128, 128, 128)


# ============== HÀM CHÍNH ==============
def create_report():
    """Hàm chính tạo báo cáo"""
    print("Đang tạo báo cáo thực tập theo mẫu mới...")

    doc = Document()

    # Thiết lập document
    # Lưu ý: margins sẽ được apply lại sau khi tạo xong tất cả sections
    setup_styles(doc)

    # 1. Trang bìa chính
    add_cover_page(doc)

    # 2. Trang bìa phụ
    add_secondary_cover_page(doc)

    # 3. Bản nhận xét của cơ sở thực tập - BỎ (sẽ xin vật lý)
    # add_company_review_page(doc)

    # 4. Lời cảm ơn
    add_acknowledgment_page(doc)

    # 5. Mục lục + Danh mục bảng biểu (KHÔNG có danh mục hình vẽ vì không có hình)
    add_toc_page(doc)
    add_list_of_tables(doc)

    # 6. Danh mục từ viết tắt
    add_abbreviations(doc)

    # 7. Nội dung chính - 4 chương
    add_chapter1(doc)  # Giới thiệu chung về đơn vị thực tập
    add_chapter2(doc)  # Nội dung thực tập
    add_chapter3(doc)  # Kết quả và đánh giá
    add_chapter4(doc)  # Nhận xét và định hướng

    # 8. Tài liệu tham khảo
    add_references(doc)

    # 9. Phụ lục
    add_appendix(doc)

    # DEBUG: Kiểm tra số sections
    print(f"DEBUG: Total sections = {len(doc.sections)}")

    # QUAN TRỌNG: Thêm khung viền TRƯỚC KHI thêm số trang và apply margins
    # để tránh border properties bị propagate
    if len(doc.sections) >= 3:
        # Bước 1: Thêm border cho sections 0 và 1
        add_page_border(doc.sections[0])  # Trang bìa chính
        add_page_border(doc.sections[1])  # Trang bìa phụ

        # Bước 2: Xóa EXPLICITLY borders từ section 2 trở đi
        for i in range(2, len(doc.sections)):
            remove_page_border(doc.sections[i])

        print(f"DEBUG: Applied border to sections 0 and 1")
        print(f"DEBUG: Explicitly removed borders from sections 2-{len(doc.sections)-1}")
    else:
        print(f"WARNING: Không đủ 3 sections! Chỉ có {len(doc.sections)} sections")

    # Apply margins cho TẤT CẢ sections SAU KHI đã apply/remove borders
    set_document_margins(doc)

    # Thêm số trang SAU CÙNG
    add_page_number_header(doc)

    # Lưu file
    output_path = "BAO_CAO_THUC_TAP.docx"
    doc.save(output_path)

    print(f"Da tao file: {output_path}")
    print(f"Cau truc bao cao:")
    print(f"  1. Bia chinh (co khung vien, bang thong tin SV co vien, logo UTC)")
    print(f"  2. Bia phu")
    print(f"  3. Loi cam on")
    print(f"  4. Muc luc + Danh muc bang bieu")
    print(f"  5. Danh muc tu viet tat")
    print(f"  6. 4 Chuong noi dung chinh")
    print(f"  7. Tai lieu tham khao (IEEE citations)")
    print(f"  8. Phu luc")
    print(f"")
    print(f"Luu y: Ban nhan xet cua co so thuc tap se xin vat ly (khong co trong file Word)")

    return output_path


if __name__ == "__main__":
    create_report()
