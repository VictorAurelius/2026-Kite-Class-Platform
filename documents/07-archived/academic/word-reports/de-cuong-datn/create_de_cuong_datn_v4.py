#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script tạo Đề cương Đồ án Tốt nghiệp Cử nhân dạng Word (.docx)
Format theo mẫu "Mau-Decuong DATN-Cử nhân.pdf" - ĐH GTVT

Cấu trúc đề cương:
1. Header: Trường + Khoa | Quốc hiệu
2. Ngày tháng
3. Tiêu đề: ĐỀ CƯƠNG ĐỒ ÁN TỐT NGHIỆP CỬ NHÂN
4. Thông tin sinh viên
5. Thông tin giảng viên hướng dẫn
6. Tên đề tài
7. 4 mục nội dung chính
8. Chữ ký các bên
9. Logo UTC
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============== THÔNG TIN SINH VIÊN ==============
STUDENT_INFO = {
    "name": "Nguyễn Văn Kiệt",
    "student_id": "221230890",
    "class": "CNTT1-K63",
    "course": "63",
    "phone": "...",
    "email": "...",
    "major": "Công nghệ thông tin",
    "system": "Chính quy",
}

# ============== THÔNG TIN GIẢNG VIÊN ==============
ADVISOR_INFO = {
    "name": "TS. Nguyễn Đức Dư",
    "department": "Khoa Công nghệ thông tin - Trường ĐH GTVT",
    "phone": "...",
    "email": "...",
}

# ============== THÔNG TIN ĐỀ TÀI ==============
THESIS_INFO = {
    "title": "XÂY DỰNG HỆ THỐNG SAAS CUNG CẤP DỊCH VỤ ĐÀO TẠO",
}

# ============== CONSTANTS ==============
FONT_NAME = 'Times New Roman'
FONT_SIZE_NORMAL = Pt(13)
FONT_SIZE_TITLE = Pt(14)

MARGIN_LEFT = Cm(2.5)
MARGIN_RIGHT = Cm(2.0)
MARGIN_TOP = Cm(2.0)
MARGIN_BOTTOM = Cm(2.0)


def set_document_margins(doc):
    """Thiết lập căn lề cho toàn bộ document"""
    for section in doc.sections:
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT


def set_font(run, size=FONT_SIZE_NORMAL, bold=False, italic=False, underline=False, color=None):
    """Helper to set font properties"""
    run.font.name = FONT_NAME
    run.font.size = size
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if color:
        run.font.color.rgb = color
    # Đảm bảo font cho tiếng Việt
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)


def set_cell_borders(cell, border_size=4):
    """Set cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(border_size))
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def remove_cell_borders(cell):
    """Remove all borders from a cell"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_header_section(doc):
    """Thêm phần header: Trường + Khoa | Quốc hiệu"""
    # Tạo bảng 2 cột cho header
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Cột trái: Trường + Khoa
    cell_left_1 = table.rows[0].cells[0]
    cell_left_1.width = Cm(8)
    p = cell_left_1.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TRƯỜNG ĐẠI HỌC GTVT")
    set_font(run, Pt(12), bold=True)
    remove_cell_borders(cell_left_1)

    cell_left_2 = table.rows[1].cells[0]
    p = cell_left_2.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("KHOA CÔNG NGHỆ THÔNG TIN")
    set_font(run, Pt(12), bold=True)
    remove_cell_borders(cell_left_2)

    # Cột phải: Quốc hiệu
    cell_right_1 = table.rows[0].cells[1]
    cell_right_1.width = Cm(9)
    p = cell_right_1.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    set_font(run, Pt(12), bold=True)
    remove_cell_borders(cell_right_1)

    cell_right_2 = table.rows[1].cells[1]
    p = cell_right_2.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Độc lập - Tự do - Hạnh phúc")
    set_font(run, Pt(12), bold=True, underline=True)
    remove_cell_borders(cell_right_2)

    # Ngày tháng
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Hà Nội, ngày ..... tháng ..... năm 2026")
    set_font(run, Pt(13), italic=True)


def add_title(doc):
    """Thêm tiêu đề: ĐỀ CƯƠNG ĐỒ ÁN TỐT NGHIỆP CỬ NHÂN"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)

    run = p.add_run("ĐỀ CƯƠNG ĐỒ ÁN TỐT NGHIỆP ")
    set_font(run, Pt(14), bold=True)

    # CỬ NHÂN với highlight màu vàng
    run = p.add_run("CỬ NHÂN")
    set_font(run, Pt(14), bold=True, underline=True)
    # Thêm highlight màu vàng
    rPr = run._element.get_or_add_rPr()
    highlight = OxmlElement('w:highlight')
    highlight.set(qn('w:val'), 'yellow')
    rPr.append(highlight)


def add_student_info(doc):
    """Thêm thông tin sinh viên"""
    # Họ và tên sinh viên
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Họ và tên sinh viên: ")
    set_font(run, Pt(13), bold=True)
    run = p.add_run(STUDENT_INFO["name"])
    set_font(run, Pt(13))

    # Mã SV, Lớp, Khóa (trên cùng một dòng)
    p = doc.add_paragraph()
    run = p.add_run("Mã SV")
    set_font(run, Pt(13))
    run = p.add_run(f"\t\t: {STUDENT_INFO['student_id']}")
    set_font(run, Pt(13))
    run = p.add_run(f"\t\tLớp: {STUDENT_INFO['class']}")
    set_font(run, Pt(13))
    run = p.add_run(f"\t\tKhóa: {STUDENT_INFO['course']}")
    set_font(run, Pt(13))

    # Số điện thoại, Email
    p = doc.add_paragraph()
    run = p.add_run("Số điện thoại")
    set_font(run, Pt(13))
    run = p.add_run(f"\t: {STUDENT_INFO['phone']}")
    set_font(run, Pt(13))
    run = p.add_run(f"\t\tEmail: {STUDENT_INFO['email']}")
    set_font(run, Pt(13))

    # Ngành
    p = doc.add_paragraph()
    run = p.add_run("Ngành")
    set_font(run, Pt(13))
    run = p.add_run(f"\t\t: ")
    set_font(run, Pt(13))
    run = p.add_run("Công nghệ thông tin")
    set_font(run, Pt(13), underline=True)
    run = p.add_run("/Khoa học máy tính")
    set_font(run, Pt(13))

    # Hệ
    p = doc.add_paragraph()
    run = p.add_run(f"Hệ: {STUDENT_INFO['system']}")
    set_font(run, Pt(13))


def add_advisor_info(doc):
    """Thêm thông tin giảng viên hướng dẫn"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Giảng viên (cán bộ) hướng dẫn: ")
    set_font(run, Pt(13), bold=True)
    run = p.add_run(ADVISOR_INFO["name"])
    set_font(run, Pt(13))

    p = doc.add_paragraph()
    run = p.add_run("Đơn vị công tác")
    set_font(run, Pt(13))
    run = p.add_run(f"\t: {ADVISOR_INFO['department']}")
    set_font(run, Pt(13))

    p = doc.add_paragraph()
    run = p.add_run("Số điện thoại")
    set_font(run, Pt(13))
    run = p.add_run(f"\t: {ADVISOR_INFO['phone']}")
    set_font(run, Pt(13))
    run = p.add_run(f"\t\tEmail: {ADVISOR_INFO['email']}")
    set_font(run, Pt(13))


def add_thesis_title(doc):
    """Thêm tên đề tài"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Tên đề tài: ")
    set_font(run, Pt(13), bold=True)
    run = p.add_run(THESIS_INFO["title"])
    set_font(run, Pt(13))


def add_section_title(doc, number, title):
    """Thêm tiêu đề mục"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"{number}. {title}")
    set_font(run, Pt(13), bold=True)


def add_section_content(doc, text, indent=True):
    """Thêm nội dung mục"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_font(run, Pt(13))


def add_bullet_item(doc, text):
    """Thêm bullet item"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(f"- {text}")
    set_font(run, Pt(13))


def add_subsection_title(doc, title):
    """Thêm tiêu đề tiểu mục (2.1, 2.2...)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    set_font(run, Pt(13), bold=True)


def add_content_sections(doc):
    """Thêm 4 mục nội dung chính - theo mẫu tham khảo"""

    # 1. Nội dung, phạm vi của đề tài
    add_section_title(doc, "1", "Nội dung, phạm vi của đề tài")

    # Nội dung của đề tài
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Nội dung của đề tài")
    set_font(run, Pt(13), bold=True)

    add_section_content(doc,
        "Đề tài xây dựng nền tảng multi-tenant SaaS cho phép các tổ chức giáo dục triển khai hệ thống quản lý lớp học "
        "với branding cá nhân hóa tự động bởi AI. Kiến trúc hybrid: KiteHub (Modular Monolith) quản lý trung tâm, "
        "KiteClass (Microservices) cho từng instance khách hàng.")

    add_bullet_item(doc, "Phân tích yêu cầu và thiết kế kiến trúc hybrid (KiteHub + KiteClass)")
    add_bullet_item(doc, "Xây dựng KiteClass Core Services (Gateway, Core, Frontend) và KiteHub Platform")
    add_bullet_item(doc, "Tích hợp AI Agent (GPT-4, DALL-E 3) tự động tạo branding")
    add_bullet_item(doc, "Triển khai thử nghiệm trên AWS EKS và đánh giá kết quả")

    # Phạm vi của đề tài
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Phạm vi của đề tài")
    set_font(run, Pt(13), bold=True)

    add_section_content(doc,
        "Hệ thống được xây dựng ở mức MVP, tập trung chức năng cốt lõi cho các tổ chức giáo dục nhỏ và vừa tại Việt Nam.")

    add_bullet_item(doc, "Chức năng: Quản lý student/teacher/class, attendance, assignment, grading, billing (thanh toán học phí), "
        "AI branding, auto-provisioning")
    add_bullet_item(doc, "Công nghệ: Java Spring Boot, Next.js, PostgreSQL, Redis, AWS EKS, OpenAI GPT-4/DALL-E 3")
    add_bullet_item(doc, "Giới hạn: MVP scope, chưa triển khai mobile app, chưa tối ưu cho >1000 concurrent users/instance")

    # 2. Công nghệ, công cụ và ngôn ngữ lập trình
    add_section_title(doc, "2", "Công nghệ, công cụ và ngôn ngữ lập trình")

    add_section_content(doc,
        "Đề tài sử dụng bộ công nghệ full-stack hiện đại phù hợp với xu hướng 2025-2026, "
        "đảm bảo tính scalable, dễ bảo trì và có tiềm năng ứng dụng thực tế:")

    add_bullet_item(doc, "Backend: Java 21 LTS, Spring Boot 3.2, Spring Security, PostgreSQL 15, Redis 7.x")
    add_bullet_item(doc, "Frontend: Next.js 14 (App Router), React, TypeScript")
    add_bullet_item(doc, "AI Services: OpenAI GPT-4 (text generation), DALL-E 3 (image generation)")
    add_bullet_item(doc, "DevOps: Docker, Kubernetes (AWS EKS), GitHub Actions, Terraform")
    add_bullet_item(doc, "Tools: Git/GitHub, IntelliJ IDEA, VS Code, Prometheus/Grafana")

    # 3. Các kết quả chính dự kiến đạt được
    add_section_title(doc, "3", "Các kết quả chính dự kiến đạt được")

    add_section_content(doc,
        "Qua quá trình nghiên cứu và thực hiện đề tài, các kết quả dự kiến đạt được bao gồm:")

    add_bullet_item(doc, "Nắm vững kiến thức về kiến trúc Microservices, multi-tenant SaaS, kiến trúc hybrid")
    add_bullet_item(doc, "Có kỹ năng phát triển full-stack với Java Spring Boot và Next.js")
    add_bullet_item(doc, "Có kinh nghiệm triển khai hệ thống trên cloud với Kubernetes")
    add_bullet_item(doc, "Có khả năng tích hợp AI vào sản phẩm thực tế")
    add_bullet_item(doc, "Báo cáo đồ án tốt nghiệp hoàn chỉnh (UML diagrams, ERD, API docs, slides, video demo)")
    add_bullet_item(doc, "Chương trình hoàn chỉnh: KiteClass Platform MVP (KiteHub Platform + KiteClass Instance)")
    add_bullet_item(doc, "AI Agent tự động tạo branding từ ảnh upload với chi phí thấp và thời gian nhanh, "
        "auto-provisioning tự động setup infrastructure")
    add_bullet_item(doc, "Hệ thống đạt thời gian phản hồi nhanh, hỗ trợ nhiều người dùng đồng thời, độ sẵn sàng cao")
    add_bullet_item(doc, "Mã nguồn trên GitHub với tài liệu hướng dẫn đầy đủ")
    add_bullet_item(doc, "Giá trị ứng dụng: Giải quyết bài toán thực tế cho các trung tâm giáo dục nhỏ tại Việt Nam")

    # 4. Kế hoạch thực hiện đề tài
    add_section_title(doc, "4", "Kế hoạch thực hiện đề tài")

    add_section_content(doc,
        "Kế hoạch thực hiện đề tài với các công việc chi tiết (4 tháng từ 02/2026 đến 05/2026). "
        "Mỗi công việc kéo dài khoảng 1 tuần, các công việc được thực hiện đồng thời để tối ưu thời gian.")

    # Bảng kế hoạch chi tiết - mỗi task ~1 tuần (1 header + 24 data rows)
    table = doc.add_table(rows=25, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    headers = ["STT", "Nội dung công việc", "Thời gian dự kiến", "Ghi chú"]
    header_widths = [Cm(1.2), Cm(7.5), Cm(4.5), Cm(3)]

    for i, (header, width) in enumerate(zip(headers, header_widths)):
        cell = table.rows[0].cells[i]
        cell.width = width
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_font(run, Pt(12), bold=True)
        set_cell_shading(cell, 'D9E2F3')

    # Data rows - Mỗi task ~1 tuần, phát triển KiteClass trước KiteHub
    plan_data = [
        # Phase 1: Research & Analysis (3 tuần: 01/02 - 21/02)
        ("1", "Nghiên cứu công nghệ, phân tích đối thủ cạnh tranh", "01/02 – 14/02/2026", "Phân tích BeeClass, Udemy"),
        ("2", "Phân tích yêu cầu nghiệp vụ, use case modeling", "08/02 – 21/02/2026", "Mô hình hóa use cases"),

        # Phase 2: System Design (2 tuần: 22/02 - 07/03)
        ("3", "Thiết kế kiến trúc hệ thống (Hybrid Architecture)", "22/02 – 28/02/2026", "Sơ đồ PlantUML"),
        ("4", "Thiết kế database schema & API specification", "01/03 – 07/03/2026", "ERD, tài liệu API"),

        # Phase 3: KiteClass Development - Backend & Frontend song song (8 tuần: 08/03 - 02/05)
        ("5", "KiteClass Gateway: Auth, routing + Unit tests", "08/03 – 14/03/2026", "Xác thực, định tuyến"),
        ("6", "KiteClass Frontend: Base setup + Auth pages + Unit tests", "08/03 – 14/03/2026", "Cơ sở giao diện"),

        ("7", "KiteClass Core: Student & Teacher modules + Unit tests", "15/03 – 28/03/2026", "Quản lý SV, GV"),
        ("8", "KiteClass Frontend: Student & Teacher pages + Unit tests", "15/03 – 28/03/2026", "Giao diện SV, GV"),

        ("9", "KiteClass Core: Class & Attendance modules + Unit tests", "29/03 – 11/04/2026", "Lớp học, điểm danh"),
        ("10", "KiteClass Frontend: Class & Attendance pages + Unit tests", "29/03 – 11/04/2026", "Giao diện lớp, điểm danh"),

        ("11", "KiteClass Core: Assignment & Grading modules + Unit tests", "12/04 – 18/04/2026", "Bài tập, chấm điểm"),
        ("12", "KiteClass Frontend: Assignment pages + Unit tests", "12/04 – 18/04/2026", "Giao diện bài tập"),

        ("13", "KiteClass Core: Billing & Invoice module + Unit tests", "19/04 – 25/04/2026", "Thanh toán, hóa đơn"),
        ("14", "KiteClass Frontend: Billing & Invoice pages + Unit tests", "19/04 – 25/04/2026", "Giao diện thanh toán"),

        ("15", "Integration testing KiteClass (cross-service)", "26/04 – 02/05/2026", "Kiểm thử liên kết"),

        # Phase 4: KiteHub Platform Development (3 tuần: 03/05 - 23/05)
        ("16", "KiteHub: Auth & Tenant Management + Unit tests", "03/05 – 09/05/2026", "Xác thực, quản lý tenant"),
        ("17", "KiteHub: Billing System + Unit tests", "03/05 – 09/05/2026", "Thanh toán điện tử"),

        ("18", "KiteHub: Admin Dashboard + Unit tests", "10/05 – 16/05/2026", "Giao diện quản trị"),
        ("19", "KiteHub: Auto-provisioning + Unit tests", "10/05 – 16/05/2026", "Tự động triển khai"),

        ("20", "KiteHub: AI Agent + Unit tests", "17/05 – 23/05/2026", "Tự động tạo branding"),
        ("21", "Integration testing KiteHub", "17/05 – 23/05/2026", "Kiểm thử platform"),

        # Phase 5: System Testing & Deployment (1 tuần: 24/05 - 25/05)
        ("22", "Load testing & performance tuning", "24/05 – 25/05/2026", "Kiểm thử tải, tối ưu"),
        ("23", "AWS EKS production deployment", "24/05 – 25/05/2026", "Triển khai production"),

        # Phase 6: Documentation (1 tuần: 26/05 - 31/05)
        ("24", "Hoàn thiện thesis report, slides, demo video", "26/05 – 31/05/2026", "Chuẩn bị bảo vệ"),
    ]

    for row_idx, (stt, content, time, note) in enumerate(plan_data):
        row = table.rows[row_idx + 1]

        # STT
        cell = row.cells[0]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(stt)
        set_font(run, Pt(12))

        # Nội dung
        cell = row.cells[1]
        p = cell.paragraphs[0]
        run = p.add_run(content)
        set_font(run, Pt(12))

        # Thời gian
        cell = row.cells[2]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(time)
        set_font(run, Pt(12))

        # Ghi chú
        cell = row.cells[3]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(note)
        set_font(run, Pt(12))


def add_signatures(doc):
    """Thêm phần chữ ký"""
    # Khoảng trống
    for _ in range(2):
        doc.add_paragraph()

    # Bảng chữ ký 4 cột
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    titles = ["Trưởng Khoa", "Trưởng Bộ môn", "Giảng viên hướng dẫn", "Sinh viên thực hiện"]
    subtitles = ["(Ký và ghi rõ họ tên)", "(Ký và ghi rõ họ tên)", "(Ký và ghi rõ họ tên)", "(Ký và ghi rõ họ tên)"]

    for i, (title, subtitle) in enumerate(zip(titles, subtitles)):
        # Title
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        set_font(run, Pt(12), bold=True)
        remove_cell_borders(cell)

        # Subtitle
        cell = table.rows[1].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        set_font(run, Pt(11), italic=True)
        remove_cell_borders(cell)


def add_logo(doc):
    """Thêm logo UTC ở cuối trang"""
    import os

    # Khoảng trống cho chữ ký
    for _ in range(4):
        doc.add_paragraph()

    logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'logo_utc.png')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if os.path.exists(logo_path):
        run = p.add_run()
        run.add_picture(logo_path, width=Cm(3.0))
    else:
        run = p.add_run("[LOGO UTC]")
        run.font.color.rgb = RGBColor(128, 128, 128)
        set_font(run, Pt(12))


def create_de_cuong():
    """Hàm chính tạo đề cương"""
    print("Đang tạo Đề cương Đồ án Tốt nghiệp...")

    doc = Document()

    # Thiết lập document
    set_document_margins(doc)

    # 1. Header
    add_header_section(doc)

    # 2. Tiêu đề
    add_title(doc)

    # 3. Thông tin sinh viên
    add_student_info(doc)

    # 4. Thông tin giảng viên
    add_advisor_info(doc)

    # 5. Tên đề tài
    add_thesis_title(doc)

    # 6. Nội dung chính (4 mục)
    add_content_sections(doc)

    # 7. Chữ ký
    add_signatures(doc)

    # 8. Logo
    add_logo(doc)

    # Lưu file
    output_path = "DE_CUONG_DATN.docx"
    doc.save(output_path)

    print(f"Đã tạo file: {output_path}")
    print(f"Cấu trúc đề cương:")
    print(f"  - Thông tin sinh viên: {STUDENT_INFO['name']} - {STUDENT_INFO['student_id']}")
    print(f"  - Giảng viên hướng dẫn: {ADVISOR_INFO['name']}")
    print(f"  - Tên đề tài: {THESIS_INFO['title'][:50]}...")
    print(f"  - 4 mục nội dung chính + Bảng kế hoạch")
    print(f"  - Chữ ký 4 bên + Logo UTC")

    return output_path


if __name__ == "__main__":
    create_de_cuong()
