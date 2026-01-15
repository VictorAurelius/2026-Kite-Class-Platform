#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script tạo báo cáo khảo sát KiteClass Platform dạng Word (.docx)
Format theo quy định trình bày đồ án tốt nghiệp - ĐH GTVT

Quy chuẩn áp dụng:
- Căn lề: trên 2.5cm, dưới 2.5cm, trái 3cm, phải 2cm
- Số trang: giữa, phía trên đầu trang
- Chương: Times New Roman 18pt, Bold, căn giữa
- Mục (1.1): Times New Roman 16pt, Bold, căn trái
- Tiểu mục (1.1.1): Times New Roman 14pt, Bold, căn trái
- Đoạn văn: Times New Roman 13pt, Justify, thụt đầu dòng 1cm, giãn dòng 1.2
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============== CONSTANTS theo quy định ==============
FONT_NAME = 'Times New Roman'
FONT_SIZE_NORMAL = Pt(13)
FONT_SIZE_CHAPTER = Pt(18)
FONT_SIZE_SECTION = Pt(16)
FONT_SIZE_SUBSECTION = Pt(14)
FONT_SIZE_TABLE = Pt(12)
FONT_SIZE_CAPTION = Pt(13)

LINE_SPACING = 1.2
FIRST_LINE_INDENT = Cm(1.0)

MARGIN_LEFT = Cm(3.0)
MARGIN_RIGHT = Cm(2.0)
MARGIN_TOP = Cm(2.5)
MARGIN_BOTTOM = Cm(2.5)


def set_document_margins(doc):
    """Thiết lập căn lề cho toàn bộ document"""
    for section in doc.sections:
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT


def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_page_number_header(doc):
    """Thêm số trang ở giữa phía TRÊN đầu trang (header)"""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        for r in p.runs:
            r.font.name = FONT_NAME
            r.font.size = FONT_SIZE_NORMAL


def setup_styles(doc):
    """Thiết lập các style chuẩn cho document"""
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE_NORMAL
    pf = style.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)


def set_font(run, size=FONT_SIZE_NORMAL, bold=False, italic=False):
    """Helper to set font properties"""
    run.font.name = FONT_NAME
    run.font.size = size
    run.bold = bold
    run.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)


def add_chapter_title(doc, text, add_page_break=True):
    """Tiêu đề chương: 18pt, Bold, căn giữa"""
    if add_page_break:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text.upper())
    set_font(run, FONT_SIZE_CHAPTER, bold=True)
    return p


def add_section_title(doc, text):
    """Tiêu đề mục (1.1, 1.2): 16pt, Bold, căn trái"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_font(run, FONT_SIZE_SECTION, bold=True)
    return p


def add_subsection_title(doc, text):
    """Tiêu đề tiểu mục (1.1.1): 14pt, Bold, căn trái"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_font(run, FONT_SIZE_SUBSECTION, bold=True)
    return p


def add_paragraph_text(doc, text, first_line_indent=True):
    """Đoạn văn: 13pt, Justify, thụt đầu dòng 1cm, giãn dòng 1.2"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = LINE_SPACING
    if first_line_indent:
        p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    run = p.add_run(text)
    set_font(run, FONT_SIZE_NORMAL)
    return p


def add_bullet_list(doc, items):
    """Thêm danh sách bullet"""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.line_spacing = LINE_SPACING
        run = p.add_run(item)
        set_font(run, FONT_SIZE_NORMAL)


def add_numbered_list(doc, items, start=1):
    """Thêm danh sách đánh số"""
    for i, item in enumerate(items, start):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.line_spacing = LINE_SPACING
        run = p.add_run(f"{i}. {item}")
        set_font(run, FONT_SIZE_NORMAL)


def add_table_with_caption(doc, caption, headers, rows):
    """Thêm bảng với tiêu đề (caption) ở PHÍA TRÊN bảng"""
    p_caption = doc.add_paragraph()
    p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_caption.paragraph_format.space_before = Pt(12)
    p_caption.paragraph_format.space_after = Pt(6)
    run = p_caption.add_run(caption)
    set_font(run, FONT_SIZE_CAPTION, bold=True)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_font(run, FONT_SIZE_TABLE, bold=True)
        set_cell_shading(header_cells[i], 'D9E2F3')

    for row_data in rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            row.cells[i].text = str(cell_text)
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    set_font(run, FONT_SIZE_TABLE)

    doc.add_paragraph()
    return table


def add_figure_placeholder(doc, caption):
    """Thêm placeholder cho hình vẽ với caption ở PHÍA DƯỚI hình"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("[Chèn biểu đồ tại đây]")
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_NORMAL
    run.font.color.rgb = RGBColor(128, 128, 128)

    p_caption = doc.add_paragraph()
    p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_caption.paragraph_format.space_after = Pt(12)
    run = p_caption.add_run(caption)
    set_font(run, FONT_SIZE_CAPTION, bold=True)


def add_title_page(doc):
    """Tạo trang bìa theo mẫu quy định ĐH GTVT"""
    import os

    # 1. TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI")
    set_font(run, Pt(14), bold=False)

    # 2. KHOA CÔNG NGHỆ THÔNG TIN
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("KHOA CÔNG NGHỆ THÔNG TIN")
    set_font(run, Pt(14), bold=True)
    run.font.underline = True

    # 3. LOGO
    logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'logo_utc.png')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(24)
    if os.path.exists(logo_path):
        run = p.add_run()
        run.add_picture(logo_path, width=Cm(3.5))
    else:
        run = p.add_run("[LOGO TRƯỜNG]")
        run.font.color.rgb = RGBColor(128, 128, 128)
        set_font(run, Pt(12))

    # 4. BÁO CÁO KHẢO SÁT
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("BÁO CÁO KHẢO SÁT")
    set_font(run, Pt(26), bold=True)

    # 5. ĐỀ TÀI
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("ĐỀ TÀI")
    set_font(run, Pt(14), bold=False)

    # 6. TÊN ĐỀ TÀI
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(36)
    run = p.add_run("KHẢO SÁT NHU CẦU NGƯỜI DÙNG\nCHO NỀN TẢNG QUẢN LÝ LỚP HỌC KITECLASS")
    set_font(run, Pt(16), bold=True)

    # 7. Thông tin sinh viên
    info = [
        ("Giảng viên hướng dẫn", "ThS. Nguyễn Đức Dư"),
        ("Sinh viên thực hiện", "Nguyễn Văn Kiệt"),
        ("Lớp", "[Tên lớp]"),
        ("Mã sinh viên", "[MSSV]"),
    ]

    for label, value in info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(3)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run1 = p.add_run(f"{label}")
        set_font(run1, Pt(14), bold=False)
        run2 = p.add_run(f"\t: {value}")
        set_font(run2, Pt(14), bold=False)

    # 8. Hà Nội – 2026
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Hà Nội – 2026")
    set_font(run, Pt(14), bold=True)


def add_toc_page(doc):
    """Thêm trang Mục lục"""
    add_chapter_title(doc, "MỤC LỤC")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Tạo mục lục tự động: References → Table of Contents]")
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(128, 128, 128)


def add_list_of_tables(doc):
    """Thêm Danh mục bảng biểu"""
    add_chapter_title(doc, "DANH MỤC BẢNG BIỂU")
    tables = [
        "Bảng 1.1. Thông tin các actors trong hệ thống",
        "Bảng 1.2. Timeline khảo sát",
        "Bảng 2.1. Phân bố mẫu khảo sát theo đối tượng",
        "Bảng 2.2. Công cụ quản lý hiện tại của chủ trung tâm",
        "Bảng 2.3. Mức độ khó khăn các công việc quản lý",
        "Bảng 2.4. Đánh giá mức độ quan trọng các tính năng",
        "Bảng 2.5. Mức giá chấp nhận được của chủ trung tâm",
        "Bảng 2.6. Kênh thông báo ưa thích của phụ huynh",
        "Bảng 3.1. Ma trận ưu tiên tính năng",
        "Bảng 3.2. Đề xuất gói dịch vụ",
    ]
    for item in tables:
        p = doc.add_paragraph()
        run = p.add_run(item)
        set_font(run, FONT_SIZE_NORMAL)


def add_list_of_figures(doc):
    """Thêm Danh mục hình vẽ"""
    add_chapter_title(doc, "DANH MỤC HÌNH VẼ")
    figures = [
        "Hình 1.1. Quy trình khảo sát 3 giai đoạn",
        "Hình 1.2. Ma trận actors trong hệ thống KiteClass",
        "Hình 2.1. Biểu đồ phân bố quy mô trung tâm",
        "Hình 2.2. Biểu đồ pain points của chủ trung tâm",
        "Hình 2.3. Biểu đồ nhu cầu tính năng gamification",
        "Hình 2.4. Biểu đồ mức độ sẵn sàng cài app của phụ huynh",
        "Hình 3.1. Feature Prioritization Matrix",
        "Hình 3.2. Persona chủ trung tâm điển hình",
    ]
    for item in figures:
        p = doc.add_paragraph()
        run = p.add_run(item)
        set_font(run, FONT_SIZE_NORMAL)


def add_abbreviations(doc):
    """Thêm Danh mục từ viết tắt"""
    add_chapter_title(doc, "DANH MỤC TỪ VIẾT TẮT")
    abbreviations = [
        ("AI", "Artificial Intelligence - Trí tuệ nhân tạo"),
        ("API", "Application Programming Interface - Giao diện lập trình ứng dụng"),
        ("GVHD", "Giảng viên hướng dẫn"),
        ("HV", "Học viên"),
        ("LMS", "Learning Management System - Hệ thống quản lý học tập"),
        ("MVP", "Minimum Viable Product - Sản phẩm khả dụng tối thiểu"),
        ("QR", "Quick Response - Mã phản hồi nhanh"),
        ("SaaS", "Software as a Service - Phần mềm dạng dịch vụ"),
        ("TT", "Trung tâm"),
        ("UI/UX", "User Interface/User Experience - Giao diện/Trải nghiệm người dùng"),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Từ viết tắt", "Giải thích"]
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_font(run, FONT_SIZE_TABLE, bold=True)
        set_cell_shading(header_cells[i], 'D9E2F3')

    for abbr, meaning in abbreviations:
        row = table.add_row()
        row.cells[0].text = abbr
        row.cells[1].text = meaning
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font(run, FONT_SIZE_TABLE)

    doc.add_paragraph()


def add_introduction(doc):
    """MỞ ĐẦU"""
    add_chapter_title(doc, "MỞ ĐẦU")

    add_section_title(doc, "1. Giới thiệu về nội dung khảo sát")

    add_paragraph_text(doc,
        "Báo cáo này trình bày kết quả khảo sát nhu cầu người dùng cho dự án "
        "KiteClass Platform - nền tảng quản lý lớp học trực tuyến theo kiến trúc "
        "Microservices. Khảo sát được thực hiện nhằm thu thập thông tin từ các "
        "đối tượng người dùng tiềm năng, bao gồm: chủ trung tâm, quản trị viên, "
        "giáo viên, học viên và phụ huynh.")

    add_paragraph_text(doc,
        "Kết quả khảo sát sẽ là cơ sở quan trọng để xác định các tính năng ưu tiên, "
        "thiết kế giao diện người dùng và xây dựng chiến lược phát triển sản phẩm "
        "phù hợp với nhu cầu thực tế của thị trường giáo dục Việt Nam.")

    add_section_title(doc, "2. Mục đích khảo sát")

    add_paragraph_text(doc, "Khảo sát được thực hiện với các mục đích sau:")
    add_bullet_list(doc, [
        "Xác định pain points (vấn đề đau đầu) của người dùng hiện tại",
        "Đánh giá mức độ quan trọng của các tính năng đề xuất",
        "Hiểu workflow làm việc hàng ngày của từng đối tượng",
        "Thu thập yêu cầu mới từ người dùng",
        "Đánh giá mức độ sẵn sàng chi trả cho giải pháp"
    ])

    add_section_title(doc, "3. Phương pháp khảo sát")

    add_paragraph_text(doc,
        "Khảo sát được thực hiện theo phương pháp hỗn hợp (mixed methods), "
        "kết hợp cả định lượng (quantitative) và định tính (qualitative):")

    add_bullet_list(doc, [
        "Phase 1: Khảo sát online qua Google Forms (2 tuần)",
        "Phase 2: Khảo sát chuyên sâu với logic branching (1 tuần)",
        "Phase 3: Phỏng vấn trực tiếp qua video call (2 tuần)"
    ])

    add_section_title(doc, "4. Cấu trúc báo cáo")

    add_paragraph_text(doc,
        "Ngoài phần Mở đầu và Kết luận, báo cáo được tổ chức thành 3 nội dung chính:")

    add_paragraph_text(doc,
        "Nội dung 1: Phương pháp và quy trình khảo sát - Trình bày chi tiết về "
        "đối tượng khảo sát, công cụ sử dụng và quy trình thực hiện.")

    add_paragraph_text(doc,
        "Nội dung 2: Kết quả khảo sát - Phân tích chi tiết kết quả thu được từ "
        "từng nhóm đối tượng, bao gồm cả dữ liệu định lượng và định tính.")

    add_paragraph_text(doc,
        "Nội dung 3: Phân tích và đề xuất - Tổng hợp insights, đề xuất tính năng "
        "ưu tiên và chiến lược phát triển sản phẩm.")


def add_content1(doc):
    """NỘI DUNG 1: Phương pháp và quy trình khảo sát"""
    add_chapter_title(doc, "NỘI DUNG 1\nPHƯƠNG PHÁP VÀ QUY TRÌNH KHẢO SÁT")

    add_section_title(doc, "1.1. Đối tượng khảo sát")

    add_subsection_title(doc, "1.1.1. Các actors trong hệ thống")

    add_paragraph_text(doc,
        "Hệ thống KiteClass Platform được thiết kế phục vụ nhiều nhóm người dùng "
        "khác nhau, được phân thành 2 cấp: KiteHub (nền tảng trung tâm) và "
        "KiteClass Instance (mỗi trung tâm giáo dục).")

    add_table_with_caption(doc,
        "Bảng 1.1. Thông tin các actors trong hệ thống",
        ["Actor", "Mô tả", "Số lượng ước tính", "Độ ưu tiên"],
        [
            ("CENTER_OWNER", "Chủ trung tâm, người ra quyết định mua", "50-100", "P0 - Cao nhất"),
            ("CENTER_ADMIN", "Quản trị viên vận hành hàng ngày", "100-200", "P0 - Cao nhất"),
            ("TEACHER", "Giáo viên trực tiếp giảng dạy", "500-1000", "P1 - Cao"),
            ("STUDENT", "Học viên sử dụng hệ thống", "5000-10000", "P1 - Cao"),
            ("PARENT", "Phụ huynh theo dõi con em", "3000-6000", "P1 - Cao"),
        ]
    )

    add_subsection_title(doc, "1.1.2. Tiêu chí chọn mẫu")

    add_paragraph_text(doc,
        "Để đảm bảo tính đại diện, mẫu khảo sát được chọn theo các tiêu chí sau:")

    add_bullet_list(doc, [
        "Quy mô trung tâm: Đa dạng từ nhỏ (<50 HV), vừa (50-200), lớn (>200)",
        "Lĩnh vực: Ngoại ngữ, Toán/Lý/Hóa, Kỹ năng mềm, Âm nhạc",
        "Vị trí địa lý: TP.HCM, Hà Nội, các tỉnh",
        "Kinh nghiệm CNTT: Từ thấp đến cao"
    ])

    add_section_title(doc, "1.2. Công cụ và kênh khảo sát")

    add_subsection_title(doc, "1.2.1. Công cụ khảo sát")

    add_bullet_list(doc, [
        "Google Forms: Khảo sát online Phase 1 và 2",
        "Zoom/Google Meet: Phỏng vấn trực tiếp Phase 3",
        "Google Sheets: Tổng hợp và phân tích dữ liệu",
        "Miro: Affinity mapping cho phân tích định tính"
    ])

    add_subsection_title(doc, "1.2.2. Kênh tiếp cận")

    add_paragraph_text(doc, "Các kênh tiếp cận được sử dụng cho từng đối tượng:")

    add_bullet_list(doc, [
        "CENTER_OWNER: Email trực tiếp, LinkedIn, Zalo",
        "CENTER_ADMIN: Email qua Owner, Zalo Group",
        "TEACHER: Email từ Admin, Facebook Group",
        "STUDENT: In-app survey, Zalo, Email",
        "PARENT: Zalo notification, Email"
    ])

    add_section_title(doc, "1.3. Quy trình thực hiện")

    add_subsection_title(doc, "1.3.1. Timeline khảo sát")

    add_table_with_caption(doc,
        "Bảng 1.2. Timeline khảo sát",
        ["Tuần", "Hoạt động", "Output"],
        [
            ("1", "Thiết kế bảng khảo sát, pilot test", "Bảng câu hỏi hoàn chỉnh"),
            ("2-3", "Phát khảo sát online Phase 1", "215 responses"),
            ("4", "Phân tích sơ bộ, thiết kế Phase 2", "Báo cáo sơ bộ"),
            ("5", "Khảo sát chuyên sâu Phase 2", "58 responses chi tiết"),
            ("6-7", "Phỏng vấn trực tiếp", "18 interview transcripts"),
            ("8", "Tổng hợp và báo cáo", "Báo cáo khảo sát hoàn chỉnh"),
        ]
    )

    add_figure_placeholder(doc, "Hình 1.1. Quy trình khảo sát 3 giai đoạn")


def add_content2(doc):
    """NỘI DUNG 2: Kết quả khảo sát"""
    add_chapter_title(doc, "NỘI DUNG 2\nKẾT QUẢ KHẢO SÁT")

    add_section_title(doc, "2.1. Tổng quan mẫu khảo sát")

    add_paragraph_text(doc,
        "Tổng số responses thu được: 215 phiếu khảo sát online và 18 cuộc phỏng vấn "
        "trực tiếp. Phân bố mẫu theo đối tượng như sau:")

    add_table_with_caption(doc,
        "Bảng 2.1. Phân bố mẫu khảo sát theo đối tượng",
        ["Đối tượng", "Số lượng", "Tỷ lệ (%)", "Phỏng vấn"],
        [
            ("Chủ trung tâm", "32", "14.9%", "8"),
            ("Quản trị viên", "28", "13.0%", "5"),
            ("Giáo viên", "45", "20.9%", "3"),
            ("Học viên", "68", "31.6%", "1"),
            ("Phụ huynh", "42", "19.5%", "1"),
            ("Tổng", "215", "100%", "18"),
        ]
    )

    add_section_title(doc, "2.2. Kết quả từ Chủ trung tâm (CENTER_OWNER)")

    add_subsection_title(doc, "2.2.1. Quy mô và lĩnh vực")

    add_paragraph_text(doc,
        "Phân bố quy mô trung tâm tham gia khảo sát cho thấy đa số là các trung tâm "
        "quy mô nhỏ và vừa, phù hợp với đối tượng mục tiêu của KiteClass Platform.")

    add_figure_placeholder(doc, "Hình 2.1. Biểu đồ phân bố quy mô trung tâm")

    add_paragraph_text(doc, "Phân bố theo lĩnh vực giảng dạy:")
    add_bullet_list(doc, [
        "Ngoại ngữ (Anh, Trung, Hàn, Nhật): 56.3%",
        "Toán - Lý - Hóa: 25.0%",
        "Tin học / Lập trình: 9.4%",
        "Kỹ năng mềm: 6.3%",
        "Âm nhạc / Mỹ thuật: 3.0%"
    ])

    add_subsection_title(doc, "2.2.2. Công cụ quản lý hiện tại")

    add_table_with_caption(doc,
        "Bảng 2.2. Công cụ quản lý hiện tại của chủ trung tâm",
        ["Công cụ", "Tỷ lệ sử dụng", "Ghi chú"],
        [
            ("Excel / Google Sheets", "78.1%", "Phổ biến nhất"),
            ("Sổ sách giấy", "34.4%", "Dùng song song"),
            ("Zalo / Facebook Group", "71.9%", "Liên lạc"),
            ("Phần mềm quản lý", "15.6%", "BeeClass, EduSoft"),
            ("Không dùng công cụ nào", "6.3%", "Quy mô rất nhỏ"),
        ]
    )

    add_subsection_title(doc, "2.2.3. Pain Points chính")

    add_paragraph_text(doc,
        "Kết quả khảo sát cho thấy các vấn đề lớn nhất mà chủ trung tâm đang gặp phải:")

    add_table_with_caption(doc,
        "Bảng 2.3. Mức độ khó khăn các công việc quản lý (thang điểm 1-5)",
        ["Công việc", "Điểm TB", "Độ lệch chuẩn", "Xếp hạng"],
        [
            ("Quản lý học phí, công nợ", "4.2", "0.8", "1"),
            ("Trao đổi với phụ huynh", "3.9", "0.9", "2"),
            ("Báo cáo, thống kê", "3.8", "1.0", "3"),
            ("Điểm danh, theo dõi học viên", "3.5", "0.9", "4"),
            ("Quản lý lịch học, phòng học", "3.2", "1.1", "5"),
        ]
    )

    add_figure_placeholder(doc, "Hình 2.2. Biểu đồ pain points của chủ trung tâm")

    add_paragraph_text(doc, "Trích dẫn từ phỏng vấn:")

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    run = p.add_run(
        '"Mỗi tháng tôi mất ít nhất 2 ngày chỉ để đối chiếu học phí, nhắc nhở phụ huynh '
        'đóng tiền. Nhiều khi nhắc nhiều quá cũng ngại." - Chủ TT Anh ngữ, Hà Nội')
    set_font(run, FONT_SIZE_NORMAL, italic=True)

    add_subsection_title(doc, "2.2.4. Nhu cầu tính năng")

    add_table_with_caption(doc,
        "Bảng 2.4. Đánh giá mức độ quan trọng các tính năng (thang điểm 1-5)",
        ["Tính năng", "Điểm TB", "% Rất cần (4-5)", "Xếp hạng"],
        [
            ("Quản lý học phí tự động", "4.6", "87.5%", "1"),
            ("Thông báo tự động cho phụ huynh", "4.4", "81.3%", "2"),
            ("Báo cáo doanh thu tự động", "4.3", "78.1%", "3"),
            ("Thanh toán online (QR, MoMo)", "4.1", "71.9%", "4"),
            ("Cổng thông tin phụ huynh", "3.9", "65.6%", "5"),
            ("Video bài giảng online", "3.5", "53.1%", "6"),
            ("Hệ thống điểm thưởng (Gamification)", "3.2", "43.8%", "7"),
        ]
    )

    add_subsection_title(doc, "2.2.5. Mức giá chấp nhận")

    add_table_with_caption(doc,
        "Bảng 2.5. Mức giá chấp nhận được của chủ trung tâm",
        ["Mức giá/tháng", "Tỷ lệ (%)", "Ghi chú"],
        [
            ("Dưới 500.000 VND", "31.3%", "Quy mô nhỏ"),
            ("500.000 - 1.000.000 VND", "43.8%", "Phổ biến nhất"),
            ("1.000.000 - 2.000.000 VND", "18.8%", "Quy mô vừa"),
            ("Trên 2.000.000 VND", "6.3%", "Quy mô lớn"),
        ]
    )

    add_section_title(doc, "2.3. Kết quả từ Giáo viên (TEACHER)")

    add_subsection_title(doc, "2.3.1. Thói quen giảng dạy")

    add_paragraph_text(doc, "Kết quả khảo sát giáo viên cho thấy:")
    add_bullet_list(doc, [
        "Số lớp trung bình/tuần: 8.5 lớp",
        "Số học viên trung bình/lớp: 12 học viên",
        "86.7% sử dụng công nghệ trong giảng dạy (slide, video)",
        "71.1% giao bài tập qua Zalo/Facebook"
    ])

    add_subsection_title(doc, "2.3.2. Điểm danh và theo dõi")

    add_paragraph_text(doc, "Phương thức điểm danh hiện tại:")
    add_bullet_list(doc, [
        "Gọi tên, ghi sổ tay: 62.2%",
        "Ghi vào Excel sau buổi học: 24.4%",
        "Dùng app điểm danh: 8.9%",
        "Không điểm danh thường xuyên: 4.4%"
    ])

    add_paragraph_text(doc,
        "84.4% giáo viên cho biết muốn có hệ thống điểm danh nhanh (QR code hoặc 1-click) "
        "và tự động thông báo cho phụ huynh.")

    add_section_title(doc, "2.4. Kết quả từ Học viên (STUDENT)")

    add_subsection_title(doc, "2.4.1. Thiết bị và thói quen")

    add_paragraph_text(doc, "Thông tin về thiết bị và thói quen học tập:")
    add_bullet_list(doc, [
        "Thiết bị chính: Điện thoại (72.1%), Laptop (25.0%), Tablet (2.9%)",
        "Thời gian tự học/ngày: 1-2 giờ (57.4%), 2-3 giờ (29.4%)",
        "73.5% thích học có video minh họa"
    ])

    add_subsection_title(doc, "2.4.2. Gamification")

    add_paragraph_text(doc,
        "Kết quả cho thấy học viên rất hứng thú với các tính năng gamification:")

    add_figure_placeholder(doc, "Hình 2.3. Biểu đồ nhu cầu tính năng gamification")

    add_bullet_list(doc, [
        "82.4% thích được thưởng điểm khi học tốt",
        "76.5% muốn có huy hiệu thành tích",
        "67.6% quan tâm đến bảng xếp hạng lớp",
        "Loại phần thưởng hấp dẫn nhất: Giảm học phí (45.6%), Quà tặng (32.4%)"
    ])

    add_section_title(doc, "2.5. Kết quả từ Phụ huynh (PARENT)")

    add_subsection_title(doc, "2.5.1. Nhu cầu theo dõi con")

    add_paragraph_text(doc, "Phụ huynh mong muốn được thông tin về:")
    add_bullet_list(doc, [
        "Con vắng học: 95.2% (quan trọng nhất)",
        "Điểm kiểm tra: 88.1%",
        "Tiến độ học tập: 78.6%",
        "Nhận xét của giáo viên: 71.4%"
    ])

    add_subsection_title(doc, "2.5.2. Kênh thông báo")

    add_table_with_caption(doc,
        "Bảng 2.6. Kênh thông báo ưa thích của phụ huynh",
        ["Kênh", "Xếp hạng 1", "Xếp hạng 2", "Xếp hạng 3"],
        [
            ("Zalo", "64.3%", "21.4%", "9.5%"),
            ("App riêng", "21.4%", "35.7%", "28.6%"),
            ("SMS", "9.5%", "28.6%", "33.3%"),
            ("Email", "4.8%", "14.3%", "28.6%"),
        ]
    )

    add_subsection_title(doc, "2.5.3. Cổng thông tin phụ huynh")

    add_paragraph_text(doc,
        "88.1% phụ huynh cho biết sẵn sàng cài app mới để theo dõi con, với điều kiện "
        "app dễ sử dụng và cung cấp thông tin kịp thời.")

    add_figure_placeholder(doc, "Hình 2.4. Biểu đồ mức độ sẵn sàng cài app của phụ huynh")


def add_content3(doc):
    """NỘI DUNG 3: Phân tích và đề xuất"""
    add_chapter_title(doc, "NỘI DUNG 3\nPHÂN TÍCH VÀ ĐỀ XUẤT")

    add_section_title(doc, "3.1. Tổng hợp Key Insights")

    add_subsection_title(doc, "3.1.1. Pain Points chính")

    add_numbered_list(doc, [
        "Quản lý học phí và công nợ là vấn đề đau đầu nhất (87.5% chủ TT)",
        "Thiếu kênh liên lạc hiệu quả với phụ huynh (81.3%)",
        "Tốn thời gian làm báo cáo thủ công (78.1%)",
        "Điểm danh thủ công, khó theo dõi (71.9%)",
        "Khó quản lý khi quy mô tăng lên (65.6%)"
    ])

    add_subsection_title(doc, "3.1.2. Cơ hội phát triển")

    add_bullet_list(doc, [
        "78.1% chưa dùng phần mềm quản lý chuyên dụng",
        "74.2% sẵn sàng trả 500k-1tr/tháng cho giải pháp tốt",
        "88.1% phụ huynh muốn có app theo dõi con",
        "82.4% học viên thích gamification"
    ])

    add_section_title(doc, "3.2. Feature Prioritization")

    add_paragraph_text(doc,
        "Dựa trên kết quả khảo sát, các tính năng được phân loại theo ma trận "
        "quan trọng - khả thi như sau:")

    add_table_with_caption(doc,
        "Bảng 3.1. Ma trận ưu tiên tính năng",
        ["Ưu tiên", "Tính năng", "Quan trọng (1-5)", "Ghi chú"],
        [
            ("P0 - Bắt buộc", "Quản lý học phí tự động", "4.6", "MVP"),
            ("P0 - Bắt buộc", "Điểm danh nhanh (QR/1-click)", "4.4", "MVP"),
            ("P0 - Bắt buộc", "Thông báo tự động", "4.4", "MVP"),
            ("P1 - Quan trọng", "Thanh toán VietQR", "4.1", "Phase 2"),
            ("P1 - Quan trọng", "Cổng phụ huynh", "3.9", "Phase 2"),
            ("P1 - Quan trọng", "Báo cáo tự động", "4.3", "Phase 2"),
            ("P2 - Bổ sung", "Video bài giảng", "3.5", "Phase 3"),
            ("P2 - Bổ sung", "Gamification", "3.2", "Phase 3"),
        ]
    )

    add_figure_placeholder(doc, "Hình 3.1. Feature Prioritization Matrix")

    add_section_title(doc, "3.3. User Personas")

    add_subsection_title(doc, "3.3.1. Persona 1: Chủ trung tâm điển hình")

    add_paragraph_text(doc,
        'Cô Hương, 35 tuổi, chủ trung tâm tiếng Anh "Bright English" với 120 học viên '
        'và 8 giáo viên. Cô mở trung tâm được 3 năm, đang dùng Excel để quản lý. '
        'Pain point lớn nhất là theo dõi học phí và liên lạc với phụ huynh.')

    add_paragraph_text(doc, "Đặc điểm:")
    add_bullet_list(doc, [
        "Giỏi chuyên môn nhưng không rành công nghệ",
        "Làm việc 10-12 tiếng/ngày, 2-3 tiếng cho hành chính",
        "Sẵn sàng trả 800k-1tr/tháng cho giải pháp tiết kiệm thời gian",
        "Ưu tiên: Dễ sử dụng > Nhiều tính năng"
    ])

    add_figure_placeholder(doc, "Hình 3.2. Persona chủ trung tâm điển hình")

    add_section_title(doc, "3.4. Đề xuất chiến lược")

    add_subsection_title(doc, "3.4.1. Đề xuất gói dịch vụ")

    add_table_with_caption(doc,
        "Bảng 3.2. Đề xuất gói dịch vụ",
        ["Gói", "Giá/tháng", "Đối tượng", "Tính năng chính"],
        [
            ("Basic", "290.000 VND", "<50 HV", "Quản lý cơ bản, điểm danh, thông báo"),
            ("Standard", "490.000 VND", "50-200 HV", "Basic + VietQR, báo cáo, Parent Portal"),
            ("Premium", "990.000 VND", ">200 HV", "Standard + Video, Gamification, API"),
        ]
    )

    add_subsection_title(doc, "3.4.2. Chiến lược tiếp cận thị trường")

    add_bullet_list(doc, [
        "Focus vào trung tâm ngoại ngữ (56.3% thị trường)",
        "Demo miễn phí 1 tháng để vượt qua rào cản tâm lý",
        "Tích hợp Zalo để leverage thói quen hiện có của phụ huynh",
        "Thiết kế UI đơn giản, ít click, phù hợp người không rành CNTT"
    ])

    add_subsection_title(doc, "3.4.3. Roadmap đề xuất")

    add_numbered_list(doc, [
        "MVP (3 tháng): Core features - Quản lý HV, học phí, điểm danh, thông báo",
        "Phase 2 (2 tháng): VietQR payment, Parent Portal, báo cáo nâng cao",
        "Phase 3 (2 tháng): Video VOD, Gamification, Live streaming (optional)"
    ])


def add_conclusion(doc):
    """KẾT LUẬN"""
    add_chapter_title(doc, "KẾT LUẬN")

    add_section_title(doc, "1. Tổng kết kết quả khảo sát")

    add_paragraph_text(doc,
        "Qua quá trình khảo sát với 215 respondents và 18 cuộc phỏng vấn, báo cáo "
        "đã thu thập được những insights quan trọng về nhu cầu của người dùng tiềm năng "
        "cho KiteClass Platform.")

    add_paragraph_text(doc, "Những kết luận chính:")
    add_bullet_list(doc, [
        "Thị trường có nhu cầu thực sự về giải pháp quản lý trung tâm giáo dục",
        "78.1% chưa dùng phần mềm chuyên dụng, đây là cơ hội lớn",
        "Pain point lớn nhất là quản lý học phí và liên lạc phụ huynh",
        "Mức giá 500k-1tr/tháng được chấp nhận bởi đa số (74.2%)",
        "Phụ huynh sẵn sàng cài app mới (88.1%)",
        "Gamification được học viên hưởng ứng tích cực (82.4%)"
    ])

    add_section_title(doc, "2. Đóng góp của khảo sát")

    add_bullet_list(doc, [
        "Xác định được 5 tính năng ưu tiên cao nhất cho MVP",
        "Hiểu rõ workflow và pain points của từng đối tượng",
        "Có cơ sở để thiết kế gói dịch vụ phù hợp",
        "Xây dựng được user personas để định hướng thiết kế UI/UX",
        "Xác định được kênh tiếp cận hiệu quả (Zalo, Facebook)"
    ])

    add_section_title(doc, "3. Hướng phát triển tiếp theo")

    add_bullet_list(doc, [
        "Thiết kế prototype dựa trên insights từ khảo sát",
        "Thực hiện Usability Testing với 10 users đại diện",
        "Hoàn thiện thiết kế UI/UX trước khi phát triển",
        "Pilot với 5-10 trung tâm trong 2 tháng đầu"
    ])


def add_references(doc):
    """TÀI LIỆU THAM KHẢO"""
    add_chapter_title(doc, "TÀI LIỆU THAM KHẢO")

    references = [
        "[1] Portigal, S. (2013). Interviewing Users: How to Uncover Compelling Insights. "
        "Rosenfeld Media.",
        "[2] Kuniavsky, M. (2012). Observing the User Experience: A Practitioner's Guide "
        "to User Research. Morgan Kaufmann.",
        "[3] Goodman, E., Kuniavsky, M., & Moed, A. (2012). Observing the User Experience: "
        "A Practitioner's Guide to User Research. Morgan Kaufmann.",
        "[4] Kế hoạch khảo sát và phỏng vấn KiteClass Platform V3.1. Tài liệu nội bộ.",
        "[5] BeeClass Platform Analysis Report. Tài liệu nội bộ.",
    ]

    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        run = p.add_run(ref)
        set_font(run, FONT_SIZE_NORMAL)


def add_appendix(doc):
    """PHỤ LỤC"""
    add_chapter_title(doc, "PHỤ LỤC")

    add_section_title(doc, "Phụ lục A: Mẫu phiếu khảo sát CENTER_OWNER")

    add_paragraph_text(doc,
        "Phiếu khảo sát gồm 4 phần với tổng cộng 11 câu hỏi, thời gian hoàn thành "
        "trung bình 5-7 phút. Chi tiết phiếu khảo sát được đính kèm trong file riêng.")

    add_section_title(doc, "Phụ lục B: Kịch bản phỏng vấn")

    add_paragraph_text(doc,
        "Kịch bản phỏng vấn Discovery Interview (30 phút) cho chủ trung tâm được "
        "thiết kế với 10 câu hỏi mở, tập trung vào context, workflow và pain points.")

    add_section_title(doc, "Phụ lục C: Danh sách người tham gia phỏng vấn")

    interview_list = [
        ("1", "Cô Hương", "Chủ TT Anh ngữ", "Hà Nội", "45 phút"),
        ("2", "Anh Minh", "Chủ TT Toán", "TP.HCM", "35 phút"),
        ("3", "Cô Lan", "Quản trị viên", "Đà Nẵng", "30 phút"),
        ("4", "Chị Thu", "Phụ huynh", "Hà Nội", "25 phút"),
        ("...", "...", "...", "...", "..."),
    ]

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["#", "Tên", "Vai trò", "Địa điểm", "Thời lượng"]
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_font(run, FONT_SIZE_TABLE, bold=True)
        set_cell_shading(header_cells[i], 'D9E2F3')

    for row_data in interview_list:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            row.cells[i].text = str(cell_text)
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    set_font(run, FONT_SIZE_TABLE)

    doc.add_paragraph()

    # Chữ ký
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Hà Nội, ngày ... tháng 01 năm 2026")
    set_font(run, FONT_SIZE_NORMAL, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Sinh viên thực hiện")
    set_font(run, FONT_SIZE_NORMAL, bold=True)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Nguyễn Văn Kiệt")
    set_font(run, FONT_SIZE_NORMAL)


def create_report():
    """Hàm chính tạo báo cáo khảo sát"""
    print("Đang tạo báo cáo khảo sát KiteClass theo quy định ĐH GTVT...")

    doc = Document()

    # Thiết lập document
    set_document_margins(doc)
    setup_styles(doc)

    # Trang bìa
    add_title_page(doc)

    # Các phần đầu
    add_toc_page(doc)
    add_list_of_tables(doc)
    add_list_of_figures(doc)
    add_abbreviations(doc)

    # Nội dung chính
    add_introduction(doc)
    add_content1(doc)
    add_content2(doc)
    add_content3(doc)
    add_conclusion(doc)
    add_references(doc)
    add_appendix(doc)

    # Thêm số trang
    add_page_number_header(doc)

    # Lưu file
    output_path = "BAO_CAO_KHAO_SAT_KITECLASS.docx"
    doc.save(output_path)

    print(f"✓ Đã tạo file: {output_path}")
    print(f"✓ Căn lề: Trái 3cm, Phải 2cm, Trên 2.5cm, Dưới 2.5cm")
    print(f"✓ Số trang: Giữa, phía trên đầu trang")
    print(f"✓ Chương: 18pt Bold, Mục: 16pt Bold, Tiểu mục: 14pt Bold")
    print(f"✓ Nội dung: 13pt, giãn dòng 1.2, thụt đầu dòng 1cm")

    return output_path


if __name__ == "__main__":
    create_report()
