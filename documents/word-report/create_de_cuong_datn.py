#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script tạo Đề cương Đồ án Tốt nghiệp Cử nhân dạng Word (.docx)
Format theo mẫu "Mau-Decuong DATN-Cử nhân.pdf" - ĐH GTVT

Cấu trúc đề cương:
1. Header: Trường + Khoa | Quốc hiệu
2. Ngày tháng
3. Tiêu đề: ĐỀ CƯƠNG ĐỒ ÁN TỐT NGHIỆP CỬ NHÂN
4. Thông tin sinh viên
5. Thông tin giảng viên hướng dẫn
6. Tên đề tài
7. 4 mục nội dung chính
8. Chữ ký các bên
9. Logo UTC
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============== THÔNG TIN SINH VIÊN ==============
STUDENT_INFO = {
    "name": "Nguyễn Văn Kiệt",
    "student_id": "221230890",
    "class": "CNTT1-K63",
    "course": "63",
    "phone": "...",
    "email": "...",
    "major": "Công nghệ thông tin",
    "system": "Chính quy",
}

# ============== THÔNG TIN GIẢNG VIÊN ==============
ADVISOR_INFO = {
    "name": "ThS. Nguyễn Đức Dư",
    "department": "Khoa Công nghệ thông tin - Trường ĐH GTVT",
    "phone": "...",
    "email": "...",
}

# ============== THÔNG TIN ĐỀ TÀI ==============
THESIS_INFO = {
    "title": "Xây dựng hệ thống quản lý lớp học trực tuyến theo kiến trúc Microservices - KiteClass Platform",
}

# ============== CONSTANTS ==============
FONT_NAME = 'Times New Roman'
FONT_SIZE_NORMAL = Pt(13)
FONT_SIZE_TITLE = Pt(14)

MARGIN_LEFT = Cm(2.5)
MARGIN_RIGHT = Cm(2.0)
MARGIN_TOP = Cm(2.0)
MARGIN_BOTTOM = Cm(2.0)


def set_document_margins(doc):
    """Thiết lập căn lề cho toàn bộ document"""
    for section in doc.sections:
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT


def set_font(run, size=FONT_SIZE_NORMAL, bold=False, italic=False, underline=False, color=None):
    """Helper to set font properties"""
    run.font.name = FONT_NAME
    run.font.size = size
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if color:
        run.font.color.rgb = color
    # Đảm bảo font cho tiếng Việt
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)


def set_cell_borders(cell, border_size=4):
    """Set cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(border_size))
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def remove_cell_borders(cell):
    """Remove all borders from a cell"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_header_section(doc):
    """Thêm phần header: Trường + Khoa | Quốc hiệu"""
    # Tạo bảng 2 cột cho header
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Cột trái: Trường + Khoa
    cell_left_1 = table.rows[0].cells[0]
    cell_left_1.width = Cm(8)
    p = cell_left_1.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TRƯỜNG ĐẠI HỌC GTVT")
    set_font(run, Pt(12), bold=True)
    remove_cell_borders(cell_left_1)

    cell_left_2 = table.rows[1].cells[0]
    p = cell_left_2.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("KHOA CÔNG NGHỆ THÔNG TIN")
    set_font(run, Pt(12), bold=True)
    remove_cell_borders(cell_left_2)

    # Cột phải: Quốc hiệu
    cell_right_1 = table.rows[0].cells[1]
    cell_right_1.width = Cm(9)
    p = cell_right_1.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    set_font(run, Pt(12), bold=True)
    remove_cell_borders(cell_right_1)

    cell_right_2 = table.rows[1].cells[1]
    p = cell_right_2.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Độc lập - Tự do - Hạnh phúc")
    set_font(run, Pt(12), bold=True, underline=True)
    remove_cell_borders(cell_right_2)

    # Ngày tháng
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Hà Nội, ngày ..... tháng ..... năm 2026")
    set_font(run, Pt(13), italic=True)


def add_title(doc):
    """Thêm tiêu đề: ĐỀ CƯƠNG ĐỒ ÁN TỐT NGHIỆP CỬ NHÂN"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)

    run = p.add_run("ĐỀ CƯƠNG ĐỒ ÁN TỐT NGHIỆP ")
    set_font(run, Pt(14), bold=True)

    # CỬ NHÂN với highlight màu vàng
    run = p.add_run("CỬ NHÂN")
    set_font(run, Pt(14), bold=True, underline=True)
    # Thêm highlight màu vàng
    rPr = run._element.get_or_add_rPr()
    highlight = OxmlElement('w:highlight')
    highlight.set(qn('w:val'), 'yellow')
    rPr.append(highlight)


def add_student_info(doc):
    """Thêm thông tin sinh viên"""
    # Họ và tên sinh viên
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Họ và tên sinh viên: ")
    set_font(run, Pt(13), bold=True)
    run = p.add_run(STUDENT_INFO["name"])
    set_font(run, Pt(13))

    # Mã SV, Lớp, Khóa (trên cùng một dòng)
    p = doc.add_paragraph()
    run = p.add_run("Mã SV")
    set_font(run, Pt(13))
    run = p.add_run(f"\t\t: {STUDENT_INFO['student_id']}")
    set_font(run, Pt(13))
    run = p.add_run(f"\t\tLớp: {STUDENT_INFO['class']}")
    set_font(run, Pt(13))
    run = p.add_run(f"\t\tKhóa: {STUDENT_INFO['course']}")
    set_font(run, Pt(13))

    # Số điện thoại, Email
    p = doc.add_paragraph()
    run = p.add_run("Số điện thoại")
    set_font(run, Pt(13))
    run = p.add_run(f"\t: {STUDENT_INFO['phone']}")
    set_font(run, Pt(13))
    run = p.add_run(f"\t\tEmail: {STUDENT_INFO['email']}")
    set_font(run, Pt(13))

    # Ngành
    p = doc.add_paragraph()
    run = p.add_run("Ngành")
    set_font(run, Pt(13))
    run = p.add_run(f"\t\t: ")
    set_font(run, Pt(13))
    run = p.add_run("Công nghệ thông tin")
    set_font(run, Pt(13), underline=True)
    run = p.add_run("/Khoa học máy tính")
    set_font(run, Pt(13))

    # Hệ
    p = doc.add_paragraph()
    run = p.add_run(f"Hệ: {STUDENT_INFO['system']}")
    set_font(run, Pt(13))


def add_advisor_info(doc):
    """Thêm thông tin giảng viên hướng dẫn"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Giảng viên (cán bộ) hướng dẫn: ")
    set_font(run, Pt(13), bold=True)
    run = p.add_run(ADVISOR_INFO["name"])
    set_font(run, Pt(13))

    p = doc.add_paragraph()
    run = p.add_run("Đơn vị công tác")
    set_font(run, Pt(13))
    run = p.add_run(f"\t: {ADVISOR_INFO['department']}")
    set_font(run, Pt(13))

    p = doc.add_paragraph()
    run = p.add_run("Số điện thoại")
    set_font(run, Pt(13))
    run = p.add_run(f"\t: {ADVISOR_INFO['phone']}")
    set_font(run, Pt(13))
    run = p.add_run(f"\t\tEmail: {ADVISOR_INFO['email']}")
    set_font(run, Pt(13))


def add_thesis_title(doc):
    """Thêm tên đề tài"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Tên đề tài: ")
    set_font(run, Pt(13), bold=True)
    run = p.add_run(THESIS_INFO["title"])
    set_font(run, Pt(13))


def add_section_title(doc, number, title):
    """Thêm tiêu đề mục"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"{number}. {title}")
    set_font(run, Pt(13), bold=True)


def add_section_content(doc, text, indent=True):
    """Thêm nội dung mục"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_font(run, Pt(13))


def add_bullet_item(doc, text):
    """Thêm bullet item"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(f"- {text}")
    set_font(run, Pt(13))


def add_content_sections(doc):
    """Thêm 4 mục nội dung chính"""

    # 1. Nội dung, phạm vi của đề tài
    add_section_title(doc, "1", "Nội dung, phạm vi của đề tài")

    add_section_content(doc,
        "Trong bối cảnh chuyển đổi số giáo dục đang diễn ra mạnh mẽ, nhu cầu về các nền tảng "
        "học trực tuyến linh hoạt, có khả năng mở rộng và tùy chỉnh cao ngày càng tăng. "
        "Đề tài hướng đến xây dựng một nền tảng quản lý lớp học trực tuyến hiện đại, "
        "áp dụng kiến trúc microservices.")

    add_section_content(doc, "Phạm vi nghiên cứu:", indent=True)
    add_bullet_item(doc, "Đối tượng: Các tổ chức giáo dục, doanh nghiệp đào tạo, giảng viên độc lập")
    add_bullet_item(doc, "Phạm vi chức năng: Quản lý lớp học, người dùng, nội dung học tập và vận hành nền tảng")
    add_bullet_item(doc, "Kiến trúc: Microservices với KiteClass Core Services + Expand Services + KiteHub Platform")

    # 2. Công nghệ, công cụ và ngôn ngữ lập trình
    add_section_title(doc, "2", "Công nghệ, công cụ và ngôn ngữ lập trình")

    add_bullet_item(doc, "Backend: Node.js / NestJS, RESTful API, GraphQL")
    add_bullet_item(doc, "Frontend: React.js / Next.js")
    add_bullet_item(doc, "Database: PostgreSQL, MongoDB, Redis")
    add_bullet_item(doc, "Message Queue: RabbitMQ / Apache Kafka")
    add_bullet_item(doc, "Container & Orchestration: Docker, Kubernetes")
    add_bullet_item(doc, "CI/CD: GitHub Actions, Jenkins")
    add_bullet_item(doc, "Cloud Platform: AWS / GCP / Azure")

    # 3. Các kết quả chính dự kiến đạt được
    add_section_title(doc, "3", "Các kết quả chính dự kiến đạt được")

    add_bullet_item(doc, "Hệ thống Core Services (Main Class, User, CMC) hoạt động ổn định")
    add_bullet_item(doc, "Nền tảng KiteHub Platform với Sale, Message, Maintaining Services")
    add_bullet_item(doc, "API Gateway tích hợp và điều phối các microservices")
    add_bullet_item(doc, "Dashboard quản trị trực quan, realtime monitoring")
    add_bullet_item(doc, "Tài liệu thiết kế hệ thống và hướng dẫn sử dụng đầy đủ")
    add_bullet_item(doc, "Báo cáo đồ án tốt nghiệp hoàn chỉnh")

    # 4. Kế hoạch thực hiện đề tài
    add_section_title(doc, "4", "Kế hoạch thực hiện đề tài")

    # Bảng kế hoạch
    table = doc.add_table(rows=8, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    headers = ["STT", "Nội dung công việc", "Thời gian dự kiến", "Ghi chú"]
    header_widths = [Cm(1.5), Cm(8), Cm(4), Cm(3)]

    for i, (header, width) in enumerate(zip(headers, header_widths)):
        cell = table.rows[0].cells[i]
        cell.width = width
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_font(run, Pt(12), bold=True)
        set_cell_shading(cell, 'D9E2F3')

    # Data rows
    plan_data = [
        ("1", "Nghiên cứu tài liệu, phân tích yêu cầu", "Tuần 1-2", ""),
        ("2", "Thiết kế kiến trúc tổng thể hệ thống", "Tuần 3-4", ""),
        ("3", "Phát triển Core Services (Main Class, User, CMC)", "Tuần 5-8", "Phase 1"),
        ("4", "Phát triển KiteHub Platform", "Tuần 9-12", "Phase 2"),
        ("5", "Tích hợp, kiểm thử hệ thống", "Tuần 13-14", ""),
        ("6", "Viết báo cáo, hoàn thiện tài liệu", "Tuần 15-16", ""),
        ("7", "Bảo vệ đồ án tốt nghiệp", "Tuần 17", ""),
    ]

    for row_idx, (stt, content, time, note) in enumerate(plan_data):
        row = table.rows[row_idx + 1]

        # STT
        cell = row.cells[0]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(stt)
        set_font(run, Pt(12))

        # Nội dung
        cell = row.cells[1]
        p = cell.paragraphs[0]
        run = p.add_run(content)
        set_font(run, Pt(12))

        # Thời gian
        cell = row.cells[2]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(time)
        set_font(run, Pt(12))

        # Ghi chú
        cell = row.cells[3]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(note)
        set_font(run, Pt(12))


def add_signatures(doc):
    """Thêm phần chữ ký"""
    # Khoảng trống
    for _ in range(2):
        doc.add_paragraph()

    # Bảng chữ ký 4 cột
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    titles = ["Trưởng Khoa", "Trưởng Bộ môn", "Giảng viên hướng dẫn", "Sinh viên thực hiện"]
    subtitles = ["(Ký và ghi rõ họ tên)", "(Ký và ghi rõ họ tên)", "(Ký và ghi rõ họ tên)", "(Ký và ghi rõ họ tên)"]

    for i, (title, subtitle) in enumerate(zip(titles, subtitles)):
        # Title
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        set_font(run, Pt(12), bold=True)
        remove_cell_borders(cell)

        # Subtitle
        cell = table.rows[1].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        set_font(run, Pt(11), italic=True)
        remove_cell_borders(cell)


def add_logo(doc):
    """Thêm logo UTC ở cuối trang"""
    import os

    # Khoảng trống cho chữ ký
    for _ in range(4):
        doc.add_paragraph()

    logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'logo_utc.png')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if os.path.exists(logo_path):
        run = p.add_run()
        run.add_picture(logo_path, width=Cm(3.0))
    else:
        run = p.add_run("[LOGO UTC]")
        run.font.color.rgb = RGBColor(128, 128, 128)
        set_font(run, Pt(12))


def create_de_cuong():
    """Hàm chính tạo đề cương"""
    print("Đang tạo Đề cương Đồ án Tốt nghiệp...")

    doc = Document()

    # Thiết lập document
    set_document_margins(doc)

    # 1. Header
    add_header_section(doc)

    # 2. Tiêu đề
    add_title(doc)

    # 3. Thông tin sinh viên
    add_student_info(doc)

    # 4. Thông tin giảng viên
    add_advisor_info(doc)

    # 5. Tên đề tài
    add_thesis_title(doc)

    # 6. Nội dung chính (4 mục)
    add_content_sections(doc)

    # 7. Chữ ký
    add_signatures(doc)

    # 8. Logo
    add_logo(doc)

    # Lưu file
    output_path = "DE_CUONG_DATN.docx"
    doc.save(output_path)

    print(f"Đã tạo file: {output_path}")
    print(f"Cấu trúc đề cương:")
    print(f"  - Thông tin sinh viên: {STUDENT_INFO['name']} - {STUDENT_INFO['student_id']}")
    print(f"  - Giảng viên hướng dẫn: {ADVISOR_INFO['name']}")
    print(f"  - Tên đề tài: {THESIS_INFO['title'][:50]}...")
    print(f"  - 4 mục nội dung chính + Bảng kế hoạch")
    print(f"  - Chữ ký 4 bên + Logo UTC")

    return output_path


if __name__ == "__main__":
    create_de_cuong()
