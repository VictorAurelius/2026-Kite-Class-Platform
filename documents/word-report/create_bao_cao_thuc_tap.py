#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script tạo báo cáo thực tập tốt nghiệp dạng Word (.docx)
Format theo mẫu "Huong dan trinh bay bao cao TTTN.pdf" - ĐH GTVT

Cấu trúc báo cáo:
1. Bìa chính (có bảng thông tin sinh viên, logo, chữ CỬ NHÂN màu vàng)
2. Bìa phụ (thêm trường Đơn vị thực tập)
3. Bản nhận xét của cơ sở thực tập
4. Lời cảm ơn
5. Mục lục + Danh mục hình vẽ + Danh mục bảng biểu
6. Danh mục từ viết tắt
7. 4 Chương nội dung chính
8. Tài liệu tham khảo (IEEE)
9. Phụ lục
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============== THÔNG TIN SINH VIÊN ==============
STUDENT_INFO = {
    "name": "Nguyễn Văn Kiệt",
    "student_id": "221230890",
    "class": "CNTT1-K63",
    "course": "63",
    "major": "Công nghệ thông tin",
    "department": "Công nghệ thông tin",
    "degree": "Cử nhân",
    "university": "Đại học Giao thông Vận tải",
}

INTERNSHIP_INFO = {
    "company": "SY PARTNERS., JSC",
    "address": "Tầng 3, Tòa nhà Luxury, Số 99 Võ Chí Công, Quận Tây Hồ, Hà Nội",
    "position": "Software Engineer",
    "advisor": "ThS. Nguyễn Đức Dư",
    "company_mentor": "[Tên CBHD]",
    "start_date": "...",
    "end_date": "...",
}

# ============== CONSTANTS ==============
FONT_NAME = 'Times New Roman'
FONT_SIZE_NORMAL = Pt(13)
FONT_SIZE_CHAPTER = Pt(14)
FONT_SIZE_SECTION = Pt(13)
FONT_SIZE_SUBSECTION = Pt(13)
FONT_SIZE_TABLE = Pt(12)
FONT_SIZE_CAPTION = Pt(13)

LINE_SPACING = 1.5
FIRST_LINE_INDENT = Cm(1.27)

MARGIN_LEFT = Cm(3.0)
MARGIN_RIGHT = Cm(2.0)
MARGIN_TOP = Cm(2.0)
MARGIN_BOTTOM = Cm(2.0)


def set_document_margins(doc):
    """Thiết lập căn lề cho toàn bộ document"""
    for section in doc.sections:
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT


def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def remove_cell_borders(cell):
    """Remove all borders from a cell"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_page_number_header(doc):
    """Thêm số trang ở giữa phía TRÊN đầu trang (header)"""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False

        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        for r in p.runs:
            r.font.name = FONT_NAME
            r.font.size = FONT_SIZE_NORMAL


def set_heading_font(style, font_name, font_size, bold=True, italic=False):
    """
    Thiết lập font cho Heading style một cách đầy đủ.

    Word Heading styles mặc định sử dụng theme fonts (Calibri Headings)
    và theme colors (màu xanh). Cần:
    - XÓA các thuộc tính theme (asciiTheme, hAnsiTheme, themeColor)
    - Thiết lập font và color trực tiếp
    """
    # Đảm bảo rPr element tồn tại
    rPr = style._element.get_or_add_rPr()

    # Xóa rFonts cũ (có thể chứa theme fonts) và tạo mới
    old_rFonts = rPr.find(qn('w:rFonts'))
    if old_rFonts is not None:
        rPr.remove(old_rFonts)

    # Tạo rFonts mới với font name trực tiếp (không dùng theme)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    # KHÔNG đặt asciiTheme, hAnsiTheme để tránh bị override bởi theme
    rPr.insert(0, rFonts)

    # Thiết lập font size
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rPr.append(sz)
    sz.set(qn('w:val'), str(int(font_size.pt * 2)))  # Word uses half-points

    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = OxmlElement('w:szCs')
        rPr.append(szCs)
    szCs.set(qn('w:val'), str(int(font_size.pt * 2)))

    # Thiết lập bold
    b = rPr.find(qn('w:b'))
    if bold:
        if b is None:
            b = OxmlElement('w:b')
            rPr.append(b)
    else:
        if b is not None:
            rPr.remove(b)

    # Thiết lập italic
    i = rPr.find(qn('w:i'))
    if italic:
        if i is None:
            i = OxmlElement('w:i')
            rPr.append(i)
    else:
        if i is not None:
            rPr.remove(i)

    # Xóa color cũ (có thể chứa themeColor) và tạo mới
    old_color = rPr.find(qn('w:color'))
    if old_color is not None:
        rPr.remove(old_color)

    # Tạo color mới với màu đen trực tiếp (không dùng theme)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '000000')
    # KHÔNG đặt themeColor để tránh bị override
    rPr.append(color)


def setup_styles(doc):
    """Thiết lập các style chuẩn cho document bao gồm Heading styles"""
    # Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE_NORMAL

    pf = style.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Thiết lập font cho Normal style qua XML
    rPr = style._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)

    # Heading 1 style (Chương) - 14pt Bold, Left
    h1_style = doc.styles['Heading 1']
    # Thiết lập qua XML để xóa theme
    set_heading_font(h1_style, FONT_NAME, Pt(14), bold=True, italic=False)
    # Thiết lập qua API để đảm bảo
    h1_style.font.name = FONT_NAME
    h1_style.font.size = Pt(14)
    h1_style.font.bold = True
    h1_style.font.italic = False
    h1_style.font.color.rgb = RGBColor(0, 0, 0)
    h1_style.paragraph_format.space_before = Pt(12)
    h1_style.paragraph_format.space_after = Pt(6)
    h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Heading 2 style (Mục 1.1, 1.2) - 13pt Bold, Left
    h2_style = doc.styles['Heading 2']
    set_heading_font(h2_style, FONT_NAME, Pt(13), bold=True, italic=False)
    h2_style.font.name = FONT_NAME
    h2_style.font.size = Pt(13)
    h2_style.font.bold = True
    h2_style.font.italic = False
    h2_style.font.color.rgb = RGBColor(0, 0, 0)
    h2_style.paragraph_format.space_before = Pt(12)
    h2_style.paragraph_format.space_after = Pt(6)
    h2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Heading 3 style (Tiểu mục 1.1.1) - 13pt Bold + Italic, Left
    h3_style = doc.styles['Heading 3']
    set_heading_font(h3_style, FONT_NAME, Pt(13), bold=True, italic=True)
    h3_style.font.name = FONT_NAME
    h3_style.font.size = Pt(13)
    h3_style.font.bold = True
    h3_style.font.italic = True
    h3_style.font.color.rgb = RGBColor(0, 0, 0)
    h3_style.paragraph_format.space_before = Pt(6)
    h3_style.paragraph_format.space_after = Pt(6)
    h3_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Caption style cho bảng và hình
    try:
        caption_style = doc.styles['Caption']
    except KeyError:
        caption_style = doc.styles.add_style('Caption', 1)  # 1 = paragraph style
    set_heading_font(caption_style, FONT_NAME, Pt(12), bold=True, italic=True)
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_style.paragraph_format.space_before = Pt(6)
    caption_style.paragraph_format.space_after = Pt(6)


def set_font(run, size=FONT_SIZE_NORMAL, bold=False, italic=False, color=None):
    """Helper to set font properties"""
    run.font.name = FONT_NAME
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)


def add_chapter_title(doc, number, text, add_page_break=True):
    """
    Tiêu đề chương: Sử dụng Heading 1 style để tạo mục lục tự động
    Format: "1. GIỚI THIỆU CHUNG VỀ ĐƠN VỊ THỰC TẬP"
    - Heading 1: 14pt Bold, Times New Roman, màu đen
    """
    if add_page_break:
        doc.add_page_break()

    # Sử dụng Heading 1 style để Word có thể tạo mục lục tự động
    p = doc.add_paragraph(style='Heading 1')
    run = p.add_run(f"{number}. {text.upper()}")

    # Thiết lập font trực tiếp cho run để đảm bảo không bị theme override
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.italic = False
    run.font.color.rgb = RGBColor(0, 0, 0)

    return p


def add_section_title(doc, text):
    """
    Tiêu đề mục (1.1, 1.2): Sử dụng Heading 2 style để tạo mục lục tự động
    - Heading 2: 13pt Bold, Times New Roman, màu đen
    """
    p = doc.add_paragraph(style='Heading 2')
    run = p.add_run(text)

    # Thiết lập font trực tiếp cho run
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.italic = False
    run.font.color.rgb = RGBColor(0, 0, 0)

    return p


def add_subsection_title(doc, text):
    """
    Tiêu đề tiểu mục (1.1.1, 1.1.2): Sử dụng Heading 3 style để tạo mục lục tự động
    - Heading 3: 13pt Bold + Italic, Times New Roman, màu đen
    """
    p = doc.add_paragraph(style='Heading 3')
    run = p.add_run(text)

    # Thiết lập font trực tiếp cho run
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.italic = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    return p


def add_seq_field(paragraph, seq_name, prefix=""):
    """
    Thêm SEQ field vào paragraph để đánh số tự động
    SEQ field giống như khi dùng Insert > Caption trong Word

    Args:
        paragraph: Paragraph object
        seq_name: Tên sequence (ví dụ: "Table", "Figure")
        prefix: Text đứng trước số (ví dụ: "Bảng ", "Hình ")
    """
    run = paragraph.add_run(prefix)
    set_font(run, Pt(12), bold=True, italic=True)

    # Tạo SEQ field
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' SEQ {seq_name} \\* ARABIC '

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    # Placeholder text (sẽ được cập nhật khi F9 trong Word)
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    run2 = paragraph.add_run("1")  # Placeholder number
    set_font(run2, Pt(12), bold=True, italic=True)

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run2._r.append(fldChar3)

    return paragraph


def add_table_caption(doc, chapter_num, caption_text):
    """
    Thêm caption cho bảng với SEQ field tự động đánh số
    Format: "Bảng X.Y. Caption text" (với X là số chương, Y là SEQ tự động)

    Lưu ý: Trong Word, bấm Ctrl+A rồi F9 để cập nhật tất cả fields
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    # Thêm "Bảng X." với X là số chương
    run = p.add_run(f"Bảng {chapter_num}.")
    set_font(run, Pt(12), bold=True, italic=True)

    # Thêm SEQ field cho số thứ tự trong chương
    add_seq_field(p, f"Table{chapter_num}", "")

    # Thêm phần caption text
    run = p.add_run(f". {caption_text}")
    set_font(run, Pt(12), bold=True, italic=True)

    return p


def add_figure_caption(doc, chapter_num, caption_text):
    """
    Thêm caption cho hình với SEQ field tự động đánh số
    Format: "Hình X.Y. Caption text" (với X là số chương, Y là SEQ tự động)

    Lưu ý: Trong Word, bấm Ctrl+A rồi F9 để cập nhật tất cả fields
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)

    # Thêm "Hình X." với X là số chương
    run = p.add_run(f"Hình {chapter_num}.")
    set_font(run, Pt(12), bold=True, italic=True)

    # Thêm SEQ field cho số thứ tự trong chương
    add_seq_field(p, f"Figure{chapter_num}", "")

    # Thêm phần caption text
    run = p.add_run(f". {caption_text}")
    set_font(run, Pt(12), bold=True, italic=True)

    return p


def add_paragraph_text(doc, text, first_line_indent=True):
    """
    Đoạn văn: 13pt, Justify, thụt đầu dòng 1.27cm, giãn dòng 1.5
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = LINE_SPACING

    if first_line_indent:
        p.paragraph_format.first_line_indent = FIRST_LINE_INDENT

    run = p.add_run(text)
    set_font(run, FONT_SIZE_NORMAL)

    return p


def add_bullet_list(doc, items):
    """Thêm danh sách bullet"""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.line_spacing = LINE_SPACING

        run = p.add_run(item)
        set_font(run, FONT_SIZE_NORMAL)


def add_table_with_caption(doc, chapter_num, caption_text, headers, rows, col_widths=None):
    """
    Thêm bảng với caption sử dụng SEQ field để đánh số tự động
    Caption ở PHÍA TRÊN bảng theo đúng chuẩn Word

    Args:
        doc: Document object
        chapter_num: Số chương (1, 2, 3, 4)
        caption_text: Nội dung caption
        headers: List các header cột
        rows: List các hàng dữ liệu
        col_widths: List độ rộng cột (cm)
    """
    # Thêm caption với SEQ field
    add_table_caption(doc, chapter_num, caption_text)

    # Tạo bảng
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        if col_widths and i < len(col_widths):
            header_cells[i].width = Cm(col_widths[i])
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_font(run, FONT_SIZE_TABLE, bold=True)
        set_cell_shading(header_cells[i], 'D9E2F3')

    for row_data in rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            row.cells[i].text = str(cell_text)
            if col_widths and i < len(col_widths):
                row.cells[i].width = Cm(col_widths[i])
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    set_font(run, FONT_SIZE_TABLE)

    doc.add_paragraph()
    return table


def add_figure_placeholder(doc, chapter_num, caption_text):
    """
    Thêm placeholder cho hình vẽ với caption sử dụng SEQ field
    Caption ở PHÍA DƯỚI hình theo đúng chuẩn Word

    Args:
        doc: Document object
        chapter_num: Số chương (1, 2, 3, 4)
        caption_text: Nội dung caption
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)

    run = p.add_run("[Chèn hình vẽ tại đây]")
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_NORMAL
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Thêm caption với SEQ field
    add_figure_caption(doc, chapter_num, caption_text)


# ============== TRANG BÌA CHÍNH ==============
def add_cover_page(doc):
    """Tạo trang bìa chính theo mẫu PDF (trang 1)"""
    import os

    # TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI")
    set_font(run, Pt(14), bold=False)

    # KHOA CÔNG NGHỆ THÔNG TIN (đậm, gạch chân)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("KHOA CÔNG NGHỆ THÔNG TIN")
    set_font(run, Pt(14), bold=True)
    run.font.underline = True

    # Logo
    logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'logo_utc.png')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(24)

    if os.path.exists(logo_path):
        run = p.add_run()
        run.add_picture(logo_path, width=Cm(3.5))
    else:
        run = p.add_run("[LOGO TRƯỜNG]")
        run.font.color.rgb = RGBColor(128, 128, 128)
        set_font(run, Pt(12))

    # BÁO CÁO (gạch chân)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("BÁO CÁO")
    set_font(run, Pt(14), bold=True)
    run.font.underline = True

    # THỰC TẬP TỐT NGHIỆP
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("THỰC TẬP TỐT NGHIỆP")
    set_font(run, Pt(22), bold=True)

    # CỬ NHÂN (màu vàng, gạch chân)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("CỬ NHÂN")
    set_font(run, Pt(22), bold=True, color=RGBColor(255, 192, 0))  # Màu vàng
    run.font.underline = True

    # Bảng thông tin sinh viên (không có viền)
    table = doc.add_table(rows=8, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_rows = [
        ("Sinh viên thực hiện", STUDENT_INFO["name"]),
        ("Mã sinh viên", STUDENT_INFO["student_id"]),
        ("Lớp", STUDENT_INFO["class"]),
        ("Khóa", STUDENT_INFO["course"]),
        ("Ngành đào tạo", STUDENT_INFO["major"]),
        ("Giảng viên hướng dẫn", INTERNSHIP_INFO["advisor"]),
        ("CBHD tại đơn vị TT", INTERNSHIP_INFO["company_mentor"]),
        ("Thời gian thực tập", f"Từ ngày {INTERNSHIP_INFO['start_date']} đến ngày {INTERNSHIP_INFO['end_date']}"),
    ]

    for i, (label, value) in enumerate(info_rows):
        row = table.rows[i]
        # Label cell
        row.cells[0].text = label
        row.cells[0].width = Cm(5.0)
        for paragraph in row.cells[0].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_font(run, Pt(13))
        remove_cell_borders(row.cells[0])

        # Value cell
        row.cells[1].text = f": {value}"
        row.cells[1].width = Cm(9.0)
        for paragraph in row.cells[1].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_font(run, Pt(13))
        remove_cell_borders(row.cells[1])

    # Khoảng trống
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # Hà Nội – 2026
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Hà Nội – 2026")
    set_font(run, Pt(14), bold=True, italic=True)


# ============== TRANG BÌA PHỤ ==============
def add_secondary_cover_page(doc):
    """Tạo trang bìa phụ theo mẫu PDF (trang 2) - thêm Đơn vị thực tập"""
    import os

    doc.add_page_break()

    # TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI")
    set_font(run, Pt(14), bold=False)

    # KHOA CÔNG NGHỆ THÔNG TIN (đậm, gạch chân)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(36)
    run = p.add_run("KHOA CÔNG NGHỆ THÔNG TIN")
    set_font(run, Pt(14), bold=True)
    run.font.underline = True

    # BÁO CÁO (gạch chân)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("BÁO CÁO")
    set_font(run, Pt(14), bold=True)
    run.font.underline = True

    # THỰC TẬP TỐT NGHIỆP
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("THỰC TẬP TỐT NGHIỆP")
    set_font(run, Pt(22), bold=True)

    # CỬ NHÂN (màu vàng, gạch chân)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("CỬ NHÂN")
    set_font(run, Pt(22), bold=True, color=RGBColor(255, 192, 0))
    run.font.underline = True

    # Bảng thông tin sinh viên (có thêm Đơn vị thực tập)
    table = doc.add_table(rows=9, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_rows = [
        ("Sinh viên thực hiện", STUDENT_INFO["name"]),
        ("Mã sinh viên", STUDENT_INFO["student_id"]),
        ("Lớp", STUDENT_INFO["class"]),
        ("Khóa", STUDENT_INFO["course"]),
        ("Ngành đào tạo", STUDENT_INFO["major"]),
        ("Giảng viên hướng dẫn", INTERNSHIP_INFO["advisor"]),
        ("CBHD tại đơn vị TT", INTERNSHIP_INFO["company_mentor"]),
        ("Đơn vị thực tập", INTERNSHIP_INFO["company"]),
        ("Thời gian thực tập", f"Từ ngày {INTERNSHIP_INFO['start_date']} đến ngày {INTERNSHIP_INFO['end_date']}"),
    ]

    for i, (label, value) in enumerate(info_rows):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].width = Cm(5.0)
        for paragraph in row.cells[0].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_font(run, Pt(13))
        remove_cell_borders(row.cells[0])

        row.cells[1].text = f": {value}"
        row.cells[1].width = Cm(9.0)
        for paragraph in row.cells[1].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_font(run, Pt(13))
        remove_cell_borders(row.cells[1])

    # Khoảng trống
    for _ in range(2):
        p = doc.add_paragraph()

    # Hà Nội – 2026
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Hà Nội – 2026")
    set_font(run, Pt(14), bold=True, italic=True)


# ============== BẢN NHẬN XÉT CỦA CƠ SỞ THỰC TẬP ==============
def add_company_review_page(doc):
    """Tạo trang Bản nhận xét của cơ sở thực tập (trang 3)"""
    doc.add_page_break()

    # CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    set_font(run, Pt(13), bold=True)

    # Độc lập - Tự do - Hạnh phúc
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Độc lập - Tự do - Hạnh phúc")
    set_font(run, Pt(13), bold=True)
    run.font.underline = True

    # Ngày tháng năm
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Hà Nội, ngày …. tháng … năm 2026")
    set_font(run, Pt(13), italic=True)

    # Tiêu đề
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("BẢN NHẬN XÉT CỦA CƠ SỞ THỰC TẬP")
    set_font(run, Pt(14), bold=True)

    # Kính gửi
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Kính gửi: ")
    set_font(run, Pt(13), bold=True, italic=True)
    run = p.add_run("Khoa Công nghệ thông tin, Trường Đại học Giao thông vận tải")
    set_font(run, Pt(13), bold=True, italic=True)

    # Thông tin cơ sở thực tập
    info_lines = [
        f"Cơ sở thực tập: {INTERNSHIP_INFO['company']}",
        "Người đại diện: ",
        "Chức vụ: ",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_font(run, Pt(13))

    # Xác nhận sinh viên
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Xác nhận sinh viên:")
    set_font(run, Pt(13), bold=True)

    # Thông tin sinh viên
    p = doc.add_paragraph()
    run = p.add_run(f"Họ tên: {STUDENT_INFO['name']}")
    set_font(run, Pt(13))
    run = p.add_run(f"\t\t\tMã sinh viên: {STUDENT_INFO['student_id']}")
    set_font(run, Pt(13))

    p = doc.add_paragraph()
    run = p.add_run(f"Lớp: {STUDENT_INFO['class']}")
    set_font(run, Pt(13))

    p = doc.add_paragraph()
    run = p.add_run(f"Đã thực tập tốt nghiệp tại cơ sở trong thời gian từ: {INTERNSHIP_INFO['start_date']} đến: {INTERNSHIP_INFO['end_date']}")
    set_font(run, Pt(13))

    # Các mục nhận xét
    review_items = [
        "Nội dung thực tập:",
        "Về tinh thần, ý thức, thái độ đối với công việc được giao:",
        "Về trình độ, kỹ năng làm việc/ khả năng thực hành:",
        "Ưu điểm nổi bật:",
        "Hạn chế cần khắc phục:",
        "Các nhận xét khác (nếu có):",
    ]

    for item in review_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run(f"- {item}")
        set_font(run, Pt(13))
        run.font.underline = True

    # Điểm thực tập
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run("Điểm thực tập (thang điểm 10): …… điểm.")
    set_font(run, Pt(13))

    # Chữ ký
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("ĐẠI DIỆN CƠ SỞ THỰC TẬP")
    set_font(run, Pt(13), bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("(Họ tên, chữ ký và đóng dấu nếu có)")
    set_font(run, Pt(13), italic=True)


# ============== LỜI CẢM ƠN ==============
def add_acknowledgment_page(doc):
    """Tạo trang Lời cảm ơn (trang 4)"""
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("LỜI CẢM ƠN")
    set_font(run, Pt(14), bold=True, italic=True)

    # Nội dung lời cảm ơn
    add_paragraph_text(doc,
        "Trong quá trình thực tập tốt nghiệp, em đã nhận được sự hướng dẫn và giúp đỡ "
        "tận tình từ nhiều cá nhân và tổ chức. Em xin được gửi lời cảm ơn chân thành đến:")

    add_paragraph_text(doc,
        f"Thầy {INTERNSHIP_INFO['advisor']} - Giảng viên hướng dẫn tại trường, đã tận tình "
        "hướng dẫn, góp ý và định hướng cho em trong suốt quá trình thực tập và viết báo cáo.")

    add_paragraph_text(doc,
        f"Ban lãnh đạo và các anh chị tại {INTERNSHIP_INFO['company']} đã tạo điều kiện thuận lợi, "
        "hướng dẫn và chia sẻ kinh nghiệm quý báu trong thời gian em thực tập tại công ty.")

    add_paragraph_text(doc,
        "Khoa Công nghệ thông tin, Trường Đại học Giao thông vận tải đã tạo điều kiện "
        "cho em được thực tập tại doanh nghiệp để có cơ hội học hỏi và phát triển.")

    add_paragraph_text(doc,
        "Mặc dù đã cố gắng hoàn thành báo cáo một cách tốt nhất, nhưng không thể tránh "
        "khỏi những thiếu sót. Em rất mong nhận được sự góp ý của các thầy cô để báo cáo "
        "được hoàn thiện hơn.")

    add_paragraph_text(doc, "Em xin chân thành cảm ơn!")

    # Chữ ký
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Hà Nội, ngày ... tháng 01 năm 2026")
    set_font(run, Pt(13), italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Sinh viên thực hiện")
    set_font(run, Pt(13), bold=True)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(STUDENT_INFO["name"])
    set_font(run, Pt(13))


# ============== MỤC LỤC ==============
def add_toc_page(doc):
    """Thêm trang Mục lục"""
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("MỤC LỤC")
    set_font(run, Pt(14), bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Tạo mục lục tự động: References → Table of Contents]")
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(128, 128, 128)


def add_list_of_figures(doc):
    """Thêm Danh mục hình vẽ"""
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("DANH MỤC HÌNH VẼ")
    set_font(run, Pt(14), bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Tạo danh mục hình vẽ tự động hoặc liệt kê thủ công]")
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(128, 128, 128)


def add_list_of_tables(doc):
    """Thêm Danh mục bảng biểu"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("DANH MỤC BẢNG BIỂU")
    set_font(run, Pt(14), bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Tạo danh mục bảng biểu tự động hoặc liệt kê thủ công]")
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(128, 128, 128)


def add_abbreviations(doc):
    """Thêm Danh mục từ viết tắt"""
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("DANH MỤC TỪ VIẾT TẮT")
    set_font(run, Pt(14), bold=True)

    abbreviations = [
        ("AI", "Artificial Intelligence - Trí tuệ nhân tạo"),
        ("API", "Application Programming Interface - Giao diện lập trình ứng dụng"),
        ("CSDL", "Cơ sở dữ liệu"),
        ("IDE", "Integrated Development Environment - Môi trường phát triển tích hợp"),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Từ viết tắt", "Giải thích"]
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_font(run, FONT_SIZE_TABLE, bold=True)
        set_cell_shading(header_cells[i], 'D9E2F3')

    for abbr, meaning in abbreviations:
        row = table.add_row()
        row.cells[0].text = abbr
        row.cells[0].width = Cm(3.0)
        row.cells[1].text = meaning
        row.cells[1].width = Cm(13.0)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font(run, FONT_SIZE_TABLE)


# ============== CHƯƠNG 1: GIỚI THIỆU CHUNG VỀ ĐƠN VỊ THỰC TẬP ==============
def add_chapter1(doc):
    """Chương 1: Giới thiệu chung về đơn vị thực tập"""
    add_chapter_title(doc, "1", "GIỚI THIỆU CHUNG VỀ ĐƠN VỊ THỰC TẬP")

    # 1.1 Thông tin chung về đơn vị thực tập
    add_section_title(doc, "1.1. Thông tin chung về đơn vị thực tập")

    add_paragraph_text(doc,
        f"{INTERNSHIP_INFO['company']} (viết tắt: SYP) là công ty công nghệ thông tin chuyên về "
        "phát triển phần mềm gia công (offshore development). Công ty được thành lập "
        "tại Hà Nội, Việt Nam vào năm 2022, với đội ngũ lãnh đạo có hơn 20 năm kinh nghiệm "
        "trong lĩnh vực phát triển phần mềm offshore trên toàn cầu.")

    add_bullet_list(doc, [
        f"Tên công ty: {INTERNSHIP_INFO['company']}",
        f"Địa chỉ: {INTERNSHIP_INFO['address']}",
        "Website: https://syp.vn",
        "Lĩnh vực: Phát triển phần mềm gia công (Offshore Development)",
        "Quy mô: Hơn 95 nhân viên (tính đến tháng 6/2024)",
    ])

    # 1.2 Chức năng, nhiệm vụ của bộ phận thực tập
    add_section_title(doc, "1.2. Chức năng, nhiệm vụ của bộ phận thực tập")

    add_paragraph_text(doc,
        "Sinh viên được phân công vào bộ phận Development với vai trò Software Engineer. "
        "Nhiệm vụ chính của bộ phận bao gồm:")

    add_bullet_list(doc, [
        "Thiết kế hệ thống phần mềm theo yêu cầu khách hàng",
        "Phát triển và bảo trì các ứng dụng web, mobile",
        "Kiểm thử và đảm bảo chất lượng sản phẩm",
        "Tích hợp và triển khai hệ thống",
    ])

    # 1.3 Môi trường làm việc và quy trình công tác
    add_section_title(doc, "1.3. Môi trường làm việc và quy trình công tác")

    add_paragraph_text(doc,
        "Công ty áp dụng mô hình làm việc hybrid, kết hợp giữa làm việc tại văn phòng "
        "và làm việc từ xa. Quy trình phát triển phần mềm theo mô hình Agile/Scrum "
        "với các sprint kéo dài 2 tuần.")

    add_paragraph_text(doc,
        "Môi trường làm việc chuyên nghiệp, năng động với đội ngũ nhân viên trẻ trung. "
        "Công ty có văn hóa chia sẻ kiến thức thông qua các buổi training nội bộ "
        "và hỗ trợ nhân viên phát triển kỹ năng chuyên môn.")


# ============== CHƯƠNG 2: NỘI DUNG THỰC TẬP ==============
def add_chapter2(doc):
    """Chương 2: Nội dung thực tập"""
    add_chapter_title(doc, "2", "NỘI DUNG THỰC TẬP")

    # 2.1 Mục tiêu và yêu cầu của đợt thực tập
    add_section_title(doc, "2.1. Mục tiêu và yêu cầu của đợt thực tập")

    add_paragraph_text(doc, "Mục tiêu của đợt thực tập:")

    add_bullet_list(doc, [
        "Tiếp cận môi trường doanh nghiệp và quy trình phát triển phần mềm chuyên nghiệp",
        "Áp dụng kiến thức đã học vào thực tế công việc thiết kế hệ thống",
        "Rèn luyện kỹ năng thiết kế cơ sở dữ liệu, màn hình, API và batch processing",
        "Học cách sử dụng AI trong quy trình kiểm tra chất lượng",
        "Chuẩn bị nền tảng kiến thức cho đồ án tốt nghiệp",
    ])

    # 2.2 Kế hoạch thực tập
    add_section_title(doc, "2.2. Kế hoạch thực tập")

    add_table_with_caption(doc, 2, "Kế hoạch thực tập theo tuần",
        ["Tuần", "Nội dung công việc"],
        [
            ("1-2", "Làm quen môi trường, tìm hiểu dự án, training thiết kế DB"),
            ("3-4", "Thực hành thiết kế cơ sở dữ liệu"),
            ("5-6", "Training và thực hành thiết kế màn hình"),
            ("7-8", "Training và thực hành thiết kế API RESTful"),
            ("9-10", "Giới thiệu AI Checker, training thiết kế Batch"),
            ("11-12", "Thiết kế độc lập, hoàn thành báo cáo"),
        ],
        col_widths=[3.0, 13.0]
    )

    # 2.3 Các công việc đã thực hiện
    add_section_title(doc, "2.3. Các công việc đã thực hiện")

    add_subsection_title(doc, "2.3.1. Thiết kế cơ sở dữ liệu")

    add_paragraph_text(doc,
        "Trong quá trình thực tập, sinh viên được giao nhiệm vụ thiết kế cơ sở dữ liệu "
        "cho các module của hệ thống. Công việc bắt đầu với việc nghiên cứu cấu trúc "
        "database hiện có bao gồm: Entity Info (thông tin cơ bản về entity như tên logic, "
        "tên vật lý, hệ thống), Column Info (chi tiết các cột bao gồm tên logic, tên vật lý, "
        "kiểu dữ liệu, ràng buộc), và Index Info (thông tin về Primary Key, Foreign Key, các index).")

    add_paragraph_text(doc,
        "Các kiến thức về Oracle Database được áp dụng trong quá trình thiết kế bao gồm:")

    add_bullet_list(doc, [
        "Kiểu dữ liệu: CHAR, VARCHAR2, NUMBER, DATE, CLOB, BLOB, JSON",
        "Constraints: PRIMARY KEY, FOREIGN KEY, NOT NULL, UNIQUE, CHECK",
        "Index: B-tree index, Bitmap index, Function-based index",
        "Naming convention: Quy tắc đặt tên chuẩn cho table, column, index",
    ])

    add_paragraph_text(doc,
        "Sinh viên học được cách định nghĩa bảng theo chuẩn với các thành phần: "
        "tên vật lý (sử dụng PascalCase hoặc snake_case), kiểu dữ liệu theo chuẩn Oracle, "
        "các ràng buộc (NOT NULL, UNIQUE, DEFAULT), và Primary Key (định danh duy nhất).")

    add_subsection_title(doc, "2.3.2. Thiết kế màn hình")

    add_paragraph_text(doc,
        "Thiết kế màn hình là quá trình xác định giao diện người dùng, bao gồm bố cục, "
        "các thành phần tương tác và luồng xử lý. Đây là loại thiết kế phức tạp nhất "
        "vì cần tham chiếu đến nhiều tài liệu liên quan.")

    add_paragraph_text(doc, "Các thành phần chính trong thiết kế màn hình bao gồm:")

    add_bullet_list(doc, [
        "Layout: Bố cục tổng thể của màn hình",
        "Items: Định nghĩa từng phần tử trên màn hình (ID, tên, loại, I/O, bắt buộc)",
        "Validation: Quy tắc kiểm tra dữ liệu nhập (đơn lẻ và tương quan)",
        "Messages: Các thông báo hiển thị cho người dùng",
        "Item Control: Điều khiển trạng thái các phần tử (hiển thị/ẩn, enable/disable)",
    ])

    add_paragraph_text(doc,
        "Validation được chia thành hai loại: Validate đơn lẻ (kiểm tra từng trường độc lập "
        "như bắt buộc nhập, độ dài, format, kiểu dữ liệu) và Validate tương quan (kiểm tra "
        "logic phụ thuộc giữa các trường như ngày kết thúc phải sau ngày bắt đầu).")

    add_subsection_title(doc, "2.3.3. Thiết kế API RESTful")

    add_paragraph_text(doc,
        "Thiết kế API là quá trình định nghĩa các endpoint để frontend và các hệ thống "
        "khác giao tiếp với backend. API trong dự án tuân theo kiến trúc RESTful với "
        "các nguyên tắc: Resource-based URL, sử dụng HTTP Methods chuẩn, Status Codes "
        "rõ ràng, và Content-Type là application/json.")

    add_paragraph_text(doc, "Cấu trúc thiết kế API bao gồm các thành phần:")

    add_bullet_list(doc, [
        "API ID: Mã định danh duy nhất cho mỗi API",
        "HTTP Method: GET (đọc), POST (tạo), PUT (cập nhật toàn bộ), PATCH (cập nhật một phần), DELETE (xóa)",
        "Endpoint: URL path của API",
        "Request Parameters và Request Body: Các tham số đầu vào",
        "Response: Định dạng phản hồi với status, data, message",
        "Error Handling: Xử lý các trường hợp lỗi với status code phù hợp",
    ])

    add_subsection_title(doc, "2.3.4. Thiết kế Batch Processing")

    add_paragraph_text(doc,
        "Batch processing (xử lý hàng loạt) là phương pháp xử lý khối lượng lớn dữ liệu "
        "theo lịch trình định sẵn, không cần tương tác người dùng. Công việc thiết kế batch "
        "áp dụng kiến trúc Spring Batch với các thành phần: Job (đơn vị công việc cao nhất), "
        "Step (các bước trong một Job), ItemReader, ItemProcessor, và ItemWriter.")

    add_paragraph_text(doc, "Một thiết kế batch tiêu chuẩn bao gồm các phần:")

    add_bullet_list(doc, [
        "Tổng quan chức năng: Mục đích, đối tượng batch, cách khởi động",
        "Shell Script: Định nghĩa tham số đầu vào, mã kết thúc",
        "Xử lý Java: Logic xử lý chính với 5 block cơ bản (Chuẩn bị, Khởi tạo, Kiểm tra, Xử lý chính, Kết thúc)",
        "Yêu cầu tìm kiếm: Các câu SQL SELECT",
        "Yêu cầu cập nhật: Các câu SQL INSERT/UPDATE/DELETE",
    ])

    add_subsection_title(doc, "2.3.5. Quy trình xử lý Review (Shiteki)")

    add_paragraph_text(doc,
        "Sau khi hoàn thành thiết kế, sản phẩm được đưa qua quy trình review nhiều cấp: "
        "Leader Review → Customer Review → End-user Review. Mỗi cấp review sẽ đưa ra "
        "các phản hồi (shiteki) cần được xử lý trước khi chuyển sang cấp tiếp theo.")

    add_paragraph_text(doc, "Các loại shiteki thường gặp và cách xử lý:")

    add_bullet_list(doc, [
        "Lỗi chính tả (sai tên bảng, cột, API): Kiểm tra lại tài liệu tham khảo",
        "Lỗi logic (sai điều kiện, thiếu trường hợp): Phân tích lại nghiệp vụ",
        "Thiếu thông tin (thiếu validation, message): Bổ sung theo yêu cầu",
        "Không nhất quán (khác biệt giữa các phần): Đồng bộ toàn bộ thiết kế",
    ])

    # 2.4 Công nghệ, công cụ và kỹ thuật sử dụng
    add_section_title(doc, "2.4. Công nghệ, công cụ và kỹ thuật sử dụng")

    add_table_with_caption(doc, 2, "Công nghệ và công cụ sử dụng",
        ["Loại", "Tên", "Mục đích"],
        [
            ("Ngôn ngữ", "Java, SQL", "Lập trình backend, truy vấn CSDL"),
            ("Framework", "Spring Boot, Spring Batch", "Phát triển ứng dụng"),
            ("CSDL", "Oracle Database", "Lưu trữ dữ liệu"),
            ("IDE", "IntelliJ IDEA, VS Code", "Viết code"),
            ("Quản lý mã nguồn", "Git, GitHub", "Version control"),
            ("AI Tools", "Claude AI", "Kiểm tra chất lượng thiết kế"),
        ],
        col_widths=[4.0, 5.0, 7.0]
    )


# ============== CHƯƠNG 3: KẾT QUẢ VÀ ĐÁNH GIÁ ==============
def add_chapter3(doc):
    """Chương 3: Kết quả và đánh giá"""
    add_chapter_title(doc, "3", "KẾT QUẢ VÀ ĐÁNH GIÁ")

    # 3.1 Kết quả đạt được
    add_section_title(doc, "3.1. Kết quả đạt được trong quá trình thực tập")

    add_paragraph_text(doc,
        "Các thiết kế được hoàn thành đạt tiêu chuẩn của doanh nghiệp và khách hàng:")

    add_bullet_list(doc, [
        "Hoàn thành thiết kế cơ sở dữ liệu cho các module được giao",
        "Hoàn thành thiết kế màn hình theo yêu cầu",
        "Hoàn thành thiết kế API RESTful",
        "Tham gia sử dụng và cải tiến hệ thống AI Checker",
    ])

    # 3.2 Kiến thức và kỹ năng tích lũy được
    add_section_title(doc, "3.2. Kiến thức và kỹ năng tích lũy được")

    add_subsection_title(doc, "3.2.1. Kiến thức chuyên môn")

    add_bullet_list(doc, [
        "Nắm vững quy trình thiết kế hệ thống phần mềm chuyên nghiệp",
        "Hiểu sâu về thiết kế cơ sở dữ liệu với Oracle",
        "Biết cách thiết kế API RESTful theo chuẩn",
        "Hiểu về kiến trúc hệ thống nhiều lớp",
    ])

    add_subsection_title(doc, "3.2.2. Kỹ năng làm việc nhóm")

    add_bullet_list(doc, [
        "Kỹ năng giao tiếp và trao đổi với team member",
        "Kỹ năng review code và thiết kế",
        "Kỹ năng làm việc với khách hàng nước ngoài (thông qua BrSE)",
    ])

    add_subsection_title(doc, "3.2.3. Kỹ năng phân tích, giải quyết vấn đề")

    add_bullet_list(doc, [
        "Kỹ năng đọc hiểu và phân tích yêu cầu nghiệp vụ",
        "Kỹ năng viết QA để xác nhận yêu cầu",
        "Kỹ năng sử dụng AI hỗ trợ công việc",
    ])

    # 3.3 Thuận lợi và khó khăn
    add_section_title(doc, "3.3. Thuận lợi và khó khăn")

    add_subsection_title(doc, "3.3.1. Thuận lợi")

    add_bullet_list(doc, [
        "Được hướng dẫn tận tình từ mentor tại công ty",
        "Môi trường làm việc chuyên nghiệp, thân thiện",
        "Có cơ hội tiếp cận công nghệ và quy trình mới",
    ])

    add_subsection_title(doc, "3.3.2. Khó khăn")

    add_bullet_list(doc, [
        "Ban đầu còn bỡ ngỡ với quy trình làm việc doanh nghiệp",
        "Cần thời gian để làm quen với các tài liệu kỹ thuật",
        "Kiến thức tiếng Nhật còn hạn chế trong giao tiếp",
    ])


# ============== CHƯƠNG 4: NHẬN XÉT VÀ ĐỊNH HƯỚNG ==============
def add_chapter4(doc):
    """Chương 4: Nhận xét và định hướng"""
    add_chapter_title(doc, "4", "NHẬN XÉT VÀ ĐỊNH HƯỚNG")

    # 4.1 Nhận xét chung về đợt thực tập
    add_section_title(doc, "4.1. Nhận xét chung về đợt thực tập")

    add_paragraph_text(doc,
        "Qua quá trình thực tập tại công ty, sinh viên đã có cơ hội được tiếp cận "
        "với môi trường làm việc chuyên nghiệp, học hỏi quy trình phát triển phần mềm "
        "theo chuẩn quốc tế và tích lũy được nhiều kiến thức, kỹ năng quý báu.")

    add_paragraph_text(doc,
        "Đợt thực tập đã giúp sinh viên hiểu rõ hơn về vai trò của Software Engineer "
        "trong quy trình phát triển phần mềm và chuẩn bị tốt hơn cho công việc sau khi tốt nghiệp.")

    # 4.2 Bài học kinh nghiệm rút ra
    add_section_title(doc, "4.2. Bài học kinh nghiệm rút ra")

    add_bullet_list(doc, [
        "Cần chủ động trong việc học hỏi và đặt câu hỏi khi gặp khó khăn",
        "Tầm quan trọng của việc đọc kỹ tài liệu trước khi thực hiện",
        "Cần kiểm tra kỹ lưỡng trước khi gửi sản phẩm review",
        "Kỹ năng mềm quan trọng không kém kỹ năng chuyên môn",
    ])

    # 4.3 Định hướng nghề nghiệp và học tập sau thực tập
    add_section_title(doc, "4.3. Định hướng nghề nghiệp và học tập sau thực tập")

    add_paragraph_text(doc,
        "Sau đợt thực tập, sinh viên định hướng tiếp tục phát triển theo hướng "
        "Software Engineer, đặc biệt là trong lĩnh vực thiết kế và phát triển hệ thống.")

    add_paragraph_text(doc, "Các mục tiêu cụ thể:")

    add_bullet_list(doc, [
        "Hoàn thành tốt đồ án tốt nghiệp với kiến thức tích lũy được",
        "Tiếp tục học hỏi và nâng cao kỹ năng lập trình",
        "Tìm hiểu sâu hơn về AI và ứng dụng trong phát triển phần mềm",
        "Cải thiện kỹ năng ngoại ngữ (tiếng Anh, tiếng Nhật)",
    ])


# ============== TÀI LIỆU THAM KHẢO ==============
def add_ieee_reference(doc, ref_num, author, title, source_type, year, url=None, accessed=None):
    """
    Thêm một tài liệu tham khảo theo chuẩn IEEE

    Args:
        doc: Document object
        ref_num: Số thứ tự tài liệu [1], [2]...
        author: Tên tác giả
        title: Tiêu đề tài liệu (sẽ được in nghiêng)
        source_type: Loại nguồn (Online, Book, Journal...)
        year: Năm xuất bản
        url: URL (cho tài liệu online)
        accessed: Ngày truy cập (cho tài liệu online)
    """
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.63)
    p.paragraph_format.first_line_indent = Cm(-0.63)
    p.paragraph_format.space_after = Pt(6)

    # [1]
    run = p.add_run(f"[{ref_num}] ")
    set_font(run, FONT_SIZE_NORMAL)

    # Author,
    run = p.add_run(f"{author}, ")
    set_font(run, FONT_SIZE_NORMAL)

    # "Title" (in nghiêng)
    run = p.add_run(f'"{title}," ')
    set_font(run, FONT_SIZE_NORMAL, italic=True)

    # Year
    run = p.add_run(f"{year}. ")
    set_font(run, FONT_SIZE_NORMAL)

    # [Online]. Available: URL
    if source_type.lower() == "online" and url:
        run = p.add_run("[Online]. Available: ")
        set_font(run, FONT_SIZE_NORMAL)

        run = p.add_run(url)
        set_font(run, FONT_SIZE_NORMAL)
        run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color for URL

        if accessed:
            run = p.add_run(f". [Accessed: {accessed}].")
            set_font(run, FONT_SIZE_NORMAL)
    else:
        run = p.add_run(f"{source_type}.")
        set_font(run, FONT_SIZE_NORMAL)


def add_references(doc):
    """
    Tài liệu tham khảo theo chuẩn IEEE
    Sử dụng Heading 1 để có thể thêm vào mục lục tự động
    """
    doc.add_page_break()

    # Sử dụng Heading 1 để có thể tạo mục lục
    p = doc.add_paragraph(style='Heading 1')
    run = p.add_run("TÀI LIỆU THAM KHẢO")
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    # IEEE format references
    add_ieee_reference(doc,
        ref_num=1,
        author="Oracle Corporation",
        title="Oracle Database Documentation",
        source_type="Online",
        year="2024",
        url="https://docs.oracle.com/en/database/",
        accessed="Jan. 2026"
    )

    add_ieee_reference(doc,
        ref_num=2,
        author="VMware",
        title="Spring Batch Reference Documentation",
        source_type="Online",
        year="2024",
        url="https://docs.spring.io/spring-batch/",
        accessed="Jan. 2026"
    )

    add_ieee_reference(doc,
        ref_num=3,
        author="RESTfulAPI.net",
        title="RESTful API Design Guidelines",
        source_type="Online",
        year="2024",
        url="https://restfulapi.net/",
        accessed="Jan. 2026"
    )

    add_ieee_reference(doc,
        ref_num=4,
        author="Anthropic",
        title="Claude Documentation",
        source_type="Online",
        year="2024",
        url="https://docs.anthropic.com/",
        accessed="Jan. 2026"
    )

    # Tài liệu nội bộ
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.63)
    p.paragraph_format.first_line_indent = Cm(-0.63)
    run = p.add_run("[5] ")
    set_font(run, FONT_SIZE_NORMAL)
    run = p.add_run("Tài liệu nội bộ công ty SY PARTNERS., JSC (không công khai).")
    set_font(run, FONT_SIZE_NORMAL)


# ============== PHỤ LỤC ==============
def add_appendix(doc):
    """Phụ lục - sử dụng Heading 1 để có thể thêm vào mục lục"""
    doc.add_page_break()

    # Sử dụng Heading 1 để có thể tạo mục lục
    p = doc.add_paragraph(style='Heading 1')
    run = p.add_run("PHỤ LỤC")
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Phụ lục A: Nhật ký thực tập
    add_section_title(doc, "Phụ lục A: Nhật ký thực tập")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Đính kèm nhật ký thực tập chi tiết]")
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_NORMAL
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Phụ lục B: Hình ảnh, tài liệu minh chứng
    add_section_title(doc, "Phụ lục B: Hình ảnh, tài liệu minh chứng")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Đính kèm hình ảnh minh chứng quá trình thực tập]")
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_NORMAL
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Phụ lục C: Sản phẩm thực tập
    add_section_title(doc, "Phụ lục C: Sản phẩm thực tập")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Mô tả hoặc đính kèm các sản phẩm thiết kế đã hoàn thành]")
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_NORMAL
    run.font.color.rgb = RGBColor(128, 128, 128)


# ============== HÀM CHÍNH ==============
def create_report():
    """Hàm chính tạo báo cáo"""
    print("Đang tạo báo cáo thực tập theo mẫu mới...")

    doc = Document()

    # Thiết lập document
    set_document_margins(doc)
    setup_styles(doc)

    # 1. Trang bìa chính
    add_cover_page(doc)

    # 2. Trang bìa phụ
    add_secondary_cover_page(doc)

    # 3. Bản nhận xét của cơ sở thực tập
    add_company_review_page(doc)

    # 4. Lời cảm ơn
    add_acknowledgment_page(doc)

    # 5. Mục lục + Danh mục hình vẽ + Danh mục bảng biểu
    add_toc_page(doc)
    add_list_of_figures(doc)
    add_list_of_tables(doc)

    # 6. Danh mục từ viết tắt
    add_abbreviations(doc)

    # 7. Nội dung chính - 4 chương
    add_chapter1(doc)  # Giới thiệu chung về đơn vị thực tập
    add_chapter2(doc)  # Nội dung thực tập
    add_chapter3(doc)  # Kết quả và đánh giá
    add_chapter4(doc)  # Nhận xét và định hướng

    # 8. Tài liệu tham khảo
    add_references(doc)

    # 9. Phụ lục
    add_appendix(doc)

    # Thêm số trang
    add_page_number_header(doc)

    # Lưu file
    output_path = "BAO_CAO_THUC_TAP.docx"
    doc.save(output_path)

    print(f"Da tao file: {output_path}")
    print(f"Cau truc bao cao theo mau moi:")
    print(f"  1. Bia chinh (co bang thong tin SV, logo, chu CU NHAN mau vang)")
    print(f"  2. Bia phu (them truong Don vi thuc tap)")
    print(f"  3. Ban nhan xet cua co so thuc tap")
    print(f"  4. Loi cam on")
    print(f"  5. Muc luc + Danh muc hinh ve + Danh muc bang bieu")
    print(f"  6. Danh muc tu viet tat")
    print(f"  7. 4 Chuong noi dung chinh")
    print(f"  8. Tai lieu tham khao (IEEE)")
    print(f"  9. Phu luc")

    return output_path


if __name__ == "__main__":
    create_report()
