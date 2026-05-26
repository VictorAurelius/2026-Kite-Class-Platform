#!/usr/bin/env python3
"""
Script tạo Đồ án tốt nghiệp Đại học GTVT V1 — KiteHub SaaS Platform.

Pivot từ documents/07-archived/academic/word-reports/bao-cao-thuc-tap/create_bao_cao_thuc_tap.py
per GAP-688 (Wave 102) — Python pipeline thay thế pandoc default thesis-v1-draft.docx
(60/100 D+) bằng UTC-spec-compliant DOCX target ≥75/100 C+.

Cấu trúc đồ án:
1. Bìa chính (UTC GTVT + "ĐỒ ÁN TỐT NGHIỆP" + title + advisor)
2. Bìa phụ (same + GVHD signature block)
3. Lời cảm ơn
4. Mục lục (auto via TOC field)
5. Danh mục hình + bảng + từ viết tắt
6. Mở đầu
7. Chương 1 (3 parts: competitor + AI + VN law)
8. Chương 2 (System architecture)
9. Chương 3 (Implementation)
10. Chương 4 (Deployment results + KPI + Beta)
11. Kết luận
12. Tài liệu tham khảo IEEE (38 entries từ bibliography.md)
13. Phụ lục (stub)

UTC Spec compliance (Quy dinh trinh bay do an tot nghiep.pdf):
- Paper A4 (210×297mm), single-sided
- Margins: top 2.5cm, bottom 2.5cm, left 3cm, right 2cm
- Body: TNR 13pt, justify, first-line indent 1cm, line spacing 1.2
- Chapter: TNR 18pt Bold, center, page break before
- Section (1.1): TNR 16pt Bold, left, no indent
- Subsection (1.1.1): TNR 14pt Bold, left, no indent
- Page numbers: top center, starting from Mục lục

Usage:
    documents/08-thesis/.venv/bin/python documents/08-thesis/create_thesis_v1.py

Output: documents/08-thesis/thesis-v1.docx
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ============== THÔNG TIN SINH VIÊN (từ thesis-info.md §4) ==============
STUDENT_INFO = {
    "name": "Nguyễn Văn Kiệt",
    "student_id": "221230890",
    "class": "CNTT1-K63",
    "course": "63",
    "major": "Công nghệ thông tin",
    "specialization": "Công nghệ phần mềm",
    "department": "Công nghệ thông tin",
    "degree": "Cử nhân",
    "training_mode": "Chính quy",
    "university": "Đại học Giao thông Vận tải",
    "university_short": "UTC GTVT",
}

# ============== THÔNG TIN ĐỀ TÀI (từ thesis-info.md §4) ==============
THESIS_INFO = {
    "title": "XÂY DỰNG HỆ THỐNG SAAS CUNG CẤP DỊCH VỤ ĐÀO TẠO",
    "title_en": "KiteHub — A Multi-Tenant SaaS Platform for Education Service Providers",
    "type": "Đồ án tốt nghiệp cử nhân",
    "year": "2026",
    "defense_window_open": "2026-08-15",
    "defense_window_close": "2026-10-15",
    "advisor": "TS. Nguyễn Đức Dư",
    "advisor_dept": "Khoa Công nghệ thông tin",
    "advisor_university": "Đại học Giao thông Vận tải",
    "reviewer": None,
    "external_mentor": None,
}

# ============== PATHS ==============
THESIS_DIR = Path(__file__).parent
CHAPTER_FILES = {
    1: [  # Wave thesis-2 Bucket A.3: 3-file Ch.1 khớp khung primary §1 (1.1 Hiện trạng + 1.2 Bài toán + 1.3 Công nghệ, công cụ sử dụng)
        THESIS_DIR / "chapter-1-competitor-analysis.md",  # §1.1 Hiện trạng (giới thiệu + khảo sát thị trường)
        THESIS_DIR / "chapter-1-vn-law-methodology.md",   # §1.2 Bài toán (phạm vi + cơ sở chuyên ngành)
        THESIS_DIR / "chapter-1-ai-techniques.md",        # §1.3 Công nghệ, công cụ sử dụng (kỹ thuật AI tích hợp)
    ],
    2: [THESIS_DIR / "chapter-2-system-architecture.md"],
    3: [THESIS_DIR / "chapter-3-implementation.md"],
    4: [THESIS_DIR / "chapter-4-deployment-results.md"],
}
CHAPTER_TITLES = {
    1: "TỔNG QUAN VỀ BÀI TOÁN VÀ CÁC CÔNG NGHỆ, CÔNG CỤ",
    2: "PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG",
    3: "PHÂN TÍCH, THIẾT KẾ VÀ TRIỂN KHAI HỆ THỐNG",
    4: "ĐÁNH GIÁ KẾT QUẢ VÀ KẾT LUẬN",
}
BIBLIOGRAPHY_FILE = THESIS_DIR / "references" / "bibliography.md"
# PROJECT_ROOT = THESIS_DIR.parent.parent (documents/08-thesis → documents → project root)
PROJECT_ROOT = THESIS_DIR.parent.parent
LOGO_FALLBACKS = [
    PROJECT_ROOT / "documents" / "07-archived" / "academic" / "word-reports" / "templates" / "logo_utc.png",
]
OUTPUT_FILE = THESIS_DIR / "thesis-v1.docx"


# ============== STYLING CONSTANTS (per UTC spec) ==============
FONT_NAME = 'Times New Roman'
FONT_SIZE_NORMAL = Pt(13)
FONT_SIZE_CHAPTER = Pt(18)
FONT_SIZE_SECTION = Pt(16)
FONT_SIZE_SUBSECTION = Pt(14)
FONT_SIZE_SUBSUB = Pt(13)
FONT_SIZE_TABLE = Pt(12)
FONT_SIZE_CAPTION = Pt(13)
FONT_SIZE_CODE = Pt(11)
LINE_SPACING = 1.2
FIRST_LINE_INDENT = Cm(1.0)
MARGIN_LEFT = Cm(3.0)
MARGIN_RIGHT = Cm(2.0)
MARGIN_TOP = Cm(2.5)
MARGIN_BOTTOM = Cm(2.5)


# ============== HELPER FUNCTIONS (preserved from bao-cao-thuc-tap.py) ==============
def set_document_margins(doc):
    """Set A4 page size + UTC margins per Quy dinh trinh bay do an tot nghiep.pdf §2.1.

    Binding gutter 0.5cm added per `thesis-content-standard.md` C1 row "Binding gutter":
    offset for binding edge when in giấy in bìa cứng cho thesis defense submission.
    """
    for section in doc.sections:
        section.page_width = Cm(21.0)   # A4 width
        section.page_height = Cm(29.7)  # A4 height
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT
        section.gutter = Cm(0.5)  # Binding gutter for hardcover bind edge


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_page_border(section):
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'text')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '24')
        border.set(qn('w:space'), '24')
        border.set(qn('w:color'), '000000')
        pgBorders.append(border)
    sectPr.append(pgBorders)


def remove_page_border(section):
    sectPr = section._sectPr
    for pgBorders in sectPr.findall(qn('w:pgBorders')):
        sectPr.remove(pgBorders)


def add_horizontal_line(doc, width_pt=2, color='000000', space_after=Pt(6)):
    """Thêm đường kẻ ngang (horizontal rule) dưới đoạn văn liền trước.

    Per user direction 2026-05-20: bìa "KHOA CÔNG NGHỆ THÔNG TIN" không dùng
    text-underline (font.underline=True) mà dùng đường kẻ ngang riêng biệt dưới
    paragraph. Reference UTC convention: bottom border trên paragraph kế tiếp
    rộng full content width (giữa lề trái + phải).
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = space_after
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(width_pt * 4))  # eighths of a point
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _set_pg_num_type(section, fmt, start=None):
    """Set <w:pgNumType w:fmt="lowerRoman" w:start="1"/> per Wave 102.5 G17."""
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    if fmt is not None:
        pgNumType.set(qn('w:fmt'), fmt)
    if start is not None:
        pgNumType.set(qn('w:start'), str(start))


def _set_header_page_field(section, show_number=True):
    """Set header with centered PAGE field, OR clear header if show_number=False."""
    header = section.header
    header.is_linked_to_previous = False
    # Clear any inherited paragraphs first
    if header.paragraphs:
        p = header.paragraphs[0]
        # Clear existing runs in header paragraph
        for run in list(p.runs):
            run._element.getparent().remove(run._element)
    else:
        p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if not show_number:
        return  # empty header — no PAGE field

    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_NORMAL


def add_page_number_header(doc):
    """Page numbering scheme per Wave 102.5 G17 + khung-chuẩn UTC.

    - Group 1 (Bìa chính + Bìa phụ + Lời cảm ơn) — NO page number, empty header
    - Group 2 (Mục lục + Danh mục bảng/hình/thuật ngữ/viết tắt) — lowerRoman i, ii, iii
    - Group 3 (Mở đầu + chapters + Kết luận + TLTK + Phụ lục) — arabic 1, 2, 3

    Section boundaries are determined by add_section(NEW_PAGE) calls in:
      - add_secondary_cover_page (after Bìa phụ — Section 0→1, border boundary)
      - add_acknowledgment_page (after Lời cảm ơn — Section 1→2, roman starts)
      - add_abbreviations (after Danh mục từ viết tắt — Section 2→3, arabic starts)

    Expected layout (4 sections trong doc.sections post border-leak fix 2026-05-20):
      [0] Bìa chính + Bìa phụ  → no number (bordered)
      [1] Lời cảm ơn           → no number (no border)
      [2] TOC + 4 danh mục     → lowerRoman i, ii, iii, ... (no border)
      [3] Mở đầu + chapters + TLTK + Phụ lục → arabic 1, 2, 3, ... (no border)
    """
    sections = doc.sections
    n_total = len(sections)
    # Heuristic boundaries cho 4-section layout:
    # Section [0,1]: no number (Bìa + Lời cảm ơn)
    # Section [2]: roman (Danh mục)
    # Section [3+]: arabic (Mở đầu + chapters)
    sect1_no_num_end = min(2, n_total)         # sections 0..1 (no number)
    sect2_roman_end = min(3, n_total)          # section 2 (roman)
    # sections 3..end (arabic)

    for i, section in enumerate(sections):
        if i < sect1_no_num_end:
            # Group 1 — no page number, empty header
            _set_header_page_field(section, show_number=False)
        elif i < sect2_roman_end:
            # Group 2 — roman lowercase, start from i
            _set_pg_num_type(section, fmt='lowerRoman', start=1)
            _set_header_page_field(section, show_number=True)
        else:
            # Group 3 — arabic decimal, start from 1 at first section of group
            if i == sect2_roman_end:
                _set_pg_num_type(section, fmt='decimal', start=1)
            else:
                _set_pg_num_type(section, fmt='decimal', start=None)  # continue numbering
            _set_header_page_field(section, show_number=True)


def set_heading_font(style, font_name, font_size, bold=True, italic=False):
    rPr = style._element.get_or_add_rPr()
    old_rFonts = rPr.find(qn('w:rFonts'))
    if old_rFonts is not None:
        rPr.remove(old_rFonts)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.insert(0, rFonts)
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rPr.append(sz)
    sz.set(qn('w:val'), str(int(font_size.pt * 2)))


def setup_styles(doc):
    """Setup Normal + Heading 1/2/3 styles per UTC spec."""
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE_NORMAL
    pf = style.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    rPr = style._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)

    # Heading 1 (Chương — used for TOC capture but actual chapter title uses direct font size 18pt)
    h1 = doc.styles['Heading 1']
    set_heading_font(h1, FONT_NAME, FONT_SIZE_CHAPTER, bold=True)
    h1.font.name = FONT_NAME
    h1.font.size = FONT_SIZE_CHAPTER
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Heading 2 (Section 1.1, 1.2)
    h2 = doc.styles['Heading 2']
    set_heading_font(h2, FONT_NAME, FONT_SIZE_SECTION, bold=True)
    h2.font.name = FONT_NAME
    h2.font.size = FONT_SIZE_SECTION
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.space_before = Pt(6)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Heading 3 (Subsection 1.1.1)
    h3 = doc.styles['Heading 3']
    set_heading_font(h3, FONT_NAME, FONT_SIZE_SUBSECTION, bold=True)
    h3.font.name = FONT_NAME
    h3.font.size = FONT_SIZE_SUBSECTION
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.space_before = Pt(6)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_font(run, size=FONT_SIZE_NORMAL, bold=False, italic=False, color=None):
    run.font.name = FONT_NAME
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if run._element.rPr is not None and run._element.rPr.rFonts is not None:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)


def add_chapter_title(doc, number, text, add_page_break=True):
    """Chương N. TITLE — 18pt Bold Center per UTC spec."""
    if add_page_break:
        doc.add_page_break()
    p = doc.add_paragraph(style='Heading 1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"CHƯƠNG {number}. {text.upper()}")
    run.font.name = FONT_NAME
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = FONT_SIZE_CHAPTER
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_section_title(doc, text):
    """Section (1.1, 1.2) — 16pt Bold Left."""
    p = doc.add_paragraph(style='Heading 2')
    run = p.add_run(text)
    run.font.name = FONT_NAME
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = FONT_SIZE_SECTION
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_subsection_title(doc, text):
    """Subsection (1.1.1) — 14pt Bold Left."""
    p = doc.add_paragraph(style='Heading 3')
    run = p.add_run(text)
    run.font.name = FONT_NAME
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = FONT_SIZE_SUBSECTION
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_subsubsection_title(doc, text):
    """Sub-subsection (#### level) — 13pt Bold."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, FONT_SIZE_SUBSUB, bold=True)
    return p


# ============== INLINE MARKDOWN PARSER ==============
INLINE_PATTERN = re.compile(
    r'(\*\*([^*]+?)\*\*|\*([^*]+?)\*|`([^`]+?)`|\[([^\]]+?)\]\(([^)]+?)\))'
)


def add_inline_runs(paragraph, text):
    """Parse inline markdown (**bold**, *italic*, `code`, [link](url)) → runs.

    Per thesis-content-standard.md v1.0.2 §3 No-font-swap principle:
    UTC §2.3 mandate TNR 13pt cho mọi đoạn văn body — KHÔNG đổi font inline.
    Inline `code` markdown rendered TNR 13pt italic (NOT Courier New monospace)
    to keep typographic consistency with UTC academic convention.
    """
    pos = 0
    for m in INLINE_PATTERN.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            set_font(run, FONT_SIZE_NORMAL)
        bold_text, italic_text, code_text, link_text = m.group(2), m.group(3), m.group(4), m.group(5)
        if bold_text:
            run = paragraph.add_run(bold_text)
            set_font(run, FONT_SIZE_NORMAL, bold=True)
        elif italic_text:
            run = paragraph.add_run(italic_text)
            set_font(run, FONT_SIZE_NORMAL, italic=True)
        elif code_text:
            # No-font-swap: inline `code` → TNR italic (NOT Courier New)
            run = paragraph.add_run(code_text)
            set_font(run, FONT_SIZE_NORMAL, italic=True)
        elif link_text:
            run = paragraph.add_run(link_text)
            set_font(run, FONT_SIZE_NORMAL)
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.font.underline = True
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_font(run, FONT_SIZE_NORMAL)


def add_paragraph_text(doc, text, first_line_indent=True):
    """Paragraph: 13pt justify, indent 1cm, line-spacing 1.2."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = LINE_SPACING
    if first_line_indent:
        p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    add_inline_runs(p, text)
    return p


def add_bullet_list_item(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.line_spacing = LINE_SPACING
    add_inline_runs(p, text)
    return p


def add_seq_caption(doc, label_word, label_number, caption_text):
    """Wave thesis-2 Bucket Issue 2 fix — emit caption với SEQ field code.

    Word "Table of Figures" / "Danh mục Hình" requires SEQ fields with specific
    identifiers ("Figure" / "Table") để populate entries. Without SEQ fields,
    Word reports "No table of figures entries found".

    Args:
        label_word: "Hình" or "Bảng" (visible Vietnamese label)
        label_number: e.g., "1.4" (chapter.figure number from MD source)
        caption_text: rest of caption text after the label

    Emits paragraph:
        "{label_word} {label_number}." + {SEQ Figure|Table \\* ARABIC} + ". {caption_text}"

    SEQ identifier mapping (matches add_list_of_figures_tables TOC \\c switch):
        - "Hình" → SEQ "Figure"  (Vietnamese visible label, English internal)
        - "Bảng" → SEQ "Table"
    """
    seq_id = "Figure" if label_word == "Hình" else "Table"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)

    # Visible label "Hình X.Y" with bold style
    run = p.add_run(f"{label_word} {label_number}")
    run.font.name = FONT_NAME
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = FONT_SIZE_NORMAL
    run.font.bold = True

    # Invisible SEQ field — Word Table of Figures discovers this
    seq_run = p.add_run("")
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' SEQ {seq_id} \\* ARABIC '
    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')
    seq_value = OxmlElement('w:t')
    seq_value.text = "1"  # placeholder — Word recalculates on F9 / updateFields
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    seq_run._r.append(fldChar_begin)
    seq_run._r.append(instrText)
    seq_run._r.append(fldChar_separate)
    seq_run._r.append(seq_value)
    seq_run._r.append(fldChar_end)

    # Caption text after period
    caption_run = p.add_run(f". {caption_text}")
    caption_run.font.name = FONT_NAME
    if caption_run._element.rPr is not None:
        caption_run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    caption_run.font.size = FONT_SIZE_NORMAL

    return p


def add_blockquote(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.line_spacing = LINE_SPACING
    run = p.add_run(text)
    set_font(run, FONT_SIZE_NORMAL, italic=True, color=RGBColor(80, 80, 80))
    return p


def add_image_inline(doc, image_path, caption=None, width_cm=14.0):
    """Insert image centered inline + optional caption underneath.

    Wave 102.5 Bucket A Item 5 — helper cho Bucket C/E screenshots embed flow.

    Args:
        doc: docx Document
        image_path: Path or str to PNG/JPG file
        caption: optional Bold caption rendered below image
        width_cm: image width in centimeters (default 14cm = page-fit)
    """
    image_path = Path(image_path) if not isinstance(image_path, Path) else image_path
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    if image_path.exists():
        run = p.add_run()
        run.add_picture(str(image_path), width=Cm(width_cm))
    else:
        run = p.add_run(f"[Hình minh hoạ: {image_path.name} — chưa có file]")
        set_font(run, FONT_SIZE_NORMAL, italic=True, color=RGBColor(128, 128, 128))

    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(6)
        run = cap.add_run(caption)
        set_font(run, FONT_SIZE_CAPTION, bold=True)


def _render_mermaid_to_png(mermaid_src: str, cache_dir: Path) -> Path | None:
    """Render Mermaid diagram → PNG. Cached locally.

    Strategy (Wave 102.5 follow-up Item 9c fix — multi-tier fallback):
      1. kroki.io HTTP API (fastest, no local deps) — fails on long-source HTTP 400
      2. mmdc (mermaid-cli) local — works for any size, requires npm install
      3. None (caller falls back to italic-text rendering)

    Args:
        mermaid_src: raw Mermaid syntax (without ``` fences)
        cache_dir: directory to save rendered PNG (gitignored)

    Returns:
        Path to PNG file on success, None on failure.
    """
    import base64
    import hashlib
    import shutil
    import subprocess
    import urllib.error
    import urllib.request
    import zlib

    cache_dir.mkdir(parents=True, exist_ok=True)
    digest = hashlib.sha256(mermaid_src.encode()).hexdigest()[:16]
    png_path = cache_dir / f"mermaid-{digest}.png"
    if png_path.exists() and png_path.stat().st_size > 0:
        return png_path  # cache hit

    # Tier 1: kroki.io HTTP API
    try:
        compressed = zlib.compress(mermaid_src.encode(), 9)
        encoded = base64.urlsafe_b64encode(compressed).decode().rstrip('=')
        url = f"https://kroki.io/mermaid/png/{encoded}"
        req = urllib.request.Request(url, headers={'User-Agent': 'thesis-pipeline/1.0'})
        with urllib.request.urlopen(req, timeout=20) as resp:
            data = resp.read()
            if data and len(data) >= 100:
                png_path.write_bytes(data)
                return png_path
    except (urllib.error.HTTPError, urllib.error.URLError, TimeoutError) as e:
        print(f"  WARN: kroki.io render failed ({e}); trying mmdc local fallback")

    # Tier 2: mmdc (mermaid-cli) local fallback
    mmdc = shutil.which("mmdc")
    if not mmdc:
        print("  WARN: mmdc not installed; install via `npm install -g @mermaid-js/mermaid-cli`")
        return None
    try:
        mmd_path = cache_dir / f"mermaid-{digest}.mmd"
        mmd_path.write_text(mermaid_src, encoding="utf-8")
        result = subprocess.run(
            [
                mmdc,
                "-i", str(mmd_path),
                "-o", str(png_path),
                "-w", "1440",
                "-H", "900",
                "-b", "white",
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )
        if result.returncode == 0 and png_path.exists() and png_path.stat().st_size > 0:
            return png_path
        print(f"  WARN: mmdc render failed (rc={result.returncode}); stderr: {result.stderr[:200]}")
        return None
    except (subprocess.TimeoutExpired, OSError) as e:
        print(f"  WARN: mmdc fallback failed: {e}")
        return None


def add_code_block(doc, code_text, lang=""):
    """Render fenced code block.

    Mermaid blocks → render as PNG via kroki.io HTTP API, embed via add_picture
    (per thesis-content-standard.md v1.0.2 C7 diagram rendering mandate).

    Other code blocks → TNR 11pt italic (NOT Courier New per v1.0.2 No-font-swap principle).
    """
    if lang and lang.lower() in ('mermaid',):
        cache_dir = THESIS_DIR / ".mermaid-cache"
        png_path = _render_mermaid_to_png(code_text, cache_dir)
        if png_path:
            # Embed PNG centered
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run()
            # Constrain width to fit page (~14cm)
            run.add_picture(str(png_path), width=Cm(14.0))
            return
        # Fallback: render as TNR italic text if PNG unavailable
        print("  WARN: Mermaid PNG unavailable; falling back to text")

    # Non-Mermaid OR Mermaid fallback: render as text (TNR italic per v1.0.2)
    for line in code_text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else " ")
        # TNR italic per UTC §2.3 + thesis-content-standard.md v1.0.2 No-font-swap
        set_font(run, FONT_SIZE_CODE, italic=True)


def add_md_table(doc, headers, rows):
    """Render markdown table → docx table."""
    if not headers:
        return None
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_runs(paragraph, h.strip())
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = FONT_SIZE_TABLE
        # Item 1b Wave 102.5 — remove blue shading; keep white default per user direction
    for row_data in rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            if i >= len(row.cells):
                break
            cell = row.cells[i]
            cell.text = ""
            for paragraph in cell.paragraphs:
                add_inline_runs(paragraph, cell_text.strip())
                for run in paragraph.runs:
                    run.font.size = FONT_SIZE_TABLE
    doc.add_paragraph()
    return table


# ============== MARKDOWN → DOCX PARSER ==============
def parse_markdown(doc, md_text, skip_top_heading=True):
    """
    Parse markdown content into docx.

    Rules:
    - Skip frontmatter (--- ... ---)
    - Skip top-level # heading (chapter title handled separately)
    - Skip metadata blockquote (> 📅 ...)
    - Skip trailing sections: 🆘 Cần hỗ trợ?, Tài liệu tham khảo, Related (these are MD-only metadata)
    - ## → section title
    - ### → subsection title
    - #### → sub-subsection (bold paragraph)
    - ``` → code block
    - > → blockquote (italic paragraph)
    - | ... | → markdown table
    - - / * → bullet list
    - blank line → paragraph break
    - Plain → paragraph
    """
    lines = md_text.split('\n')

    # Strip frontmatter
    if lines and lines[0].strip() == '---':
        end = 1
        while end < len(lines) and lines[end].strip() != '---':
            end += 1
        lines = lines[end + 1:]

    # Collect content excluding top-level # heading and trailing metadata sections
    SKIP_SECTIONS = {"🆘 Cần hỗ trợ?", "Tài liệu tham khảo", "Related", "Cảm ơn", "TL;DR"}
    cleaned_lines = []
    top_seen = not skip_top_heading
    skip_until_h2 = False
    for line in lines:
        stripped = line.strip()
        if not top_seen and stripped.startswith("# "):
            top_seen = True
            continue
        if stripped.startswith("## "):
            sec_title = stripped[3:].strip()
            # Match TL;DR with optional suffix; emoji prefix sections also skip
            if sec_title in SKIP_SECTIONS or sec_title.startswith("TL;DR") or sec_title.startswith("🆘"):
                skip_until_h2 = True
                continue
            else:
                skip_until_h2 = False
        if skip_until_h2:
            continue
        # Skip date blockquote metadata line ("> 📅 Cập nhật lần cuối: ...")
        if stripped.startswith("> 📅"):
            continue
        cleaned_lines.append(line)

    # Parse line-by-line
    i = 0
    paragraph_buffer = []
    in_code = False
    code_lang = ""
    code_lines = []

    def flush_paragraph():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            text = ' '.join(paragraph_buffer).strip()
            if text:
                add_paragraph_text(doc, text)
            paragraph_buffer = []

    while i < len(cleaned_lines):
        line = cleaned_lines[i]
        stripped = line.strip()

        # Code block
        if stripped.startswith("```"):
            if in_code:
                # close
                add_code_block(doc, '\n'.join(code_lines), code_lang)
                code_lines = []
                in_code = False
            else:
                flush_paragraph()
                in_code = True
                code_lang = stripped[3:].strip()
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            flush_paragraph()
            i += 1
            continue

        # Markdown image: ![alt](path) — embed PNG via add_image_inline
        # Relative path resolved against THESIS_DIR
        image_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)\s*$', stripped)
        if image_match:
            flush_paragraph()
            img_path_str = image_match.group(2).strip()
            img_path = Path(img_path_str)
            if not img_path.is_absolute():
                img_path = THESIS_DIR / img_path
            add_image_inline(doc, img_path, caption=None, width_cm=14.0)
            i += 1
            continue

        # Wave thesis-2 Bucket Issue 2 fix — detect figure/table caption pattern
        # **Hình X.Y.** Caption text.  or  **Bảng X.Y.** Caption text.
        # (Source MDs có format: bold prefix `**Hình X.Y.**` THEN plain caption text)
        # Emit with SEQ field for Word "Table of Figures" / "Danh mục Hình" support
        caption_match = re.match(r'^\*\*(Hình|Bảng)\s+(\d+\.\d+)\.?\*\*\s*(.+?)\.?\s*$', stripped)
        if caption_match:
            flush_paragraph()
            label_word = caption_match.group(1)
            label_number = caption_match.group(2)
            caption_text = caption_match.group(3).strip().rstrip('.')
            add_seq_caption(doc, label_word, label_number, caption_text)
            i += 1
            continue

        # Section title ## or ### or ####
        if stripped.startswith("## "):
            flush_paragraph()
            title_text = stripped[3:].strip()
            # Strip leading number prefix like "1. " or "2.1 " from chapter MD style — keep as-is OK
            add_section_title(doc, title_text)
            i += 1
            continue
        if stripped.startswith("### "):
            flush_paragraph()
            title_text = stripped[4:].strip()
            add_subsection_title(doc, title_text)
            i += 1
            continue
        if stripped.startswith("#### "):
            flush_paragraph()
            title_text = stripped[5:].strip()
            add_subsubsection_title(doc, title_text)
            i += 1
            continue
        if stripped.startswith("##### "):
            flush_paragraph()
            title_text = stripped[6:].strip()
            add_subsubsection_title(doc, title_text)
            i += 1
            continue

        # Blockquote
        if stripped.startswith("> "):
            flush_paragraph()
            add_blockquote(doc, stripped[2:].strip())
            i += 1
            continue

        # Markdown table
        if stripped.startswith("|") and stripped.endswith("|"):
            flush_paragraph()
            headers = [c.strip() for c in stripped.strip("|").split("|")]
            i += 1
            # Skip separator row |---|---|
            if i < len(cleaned_lines) and re.match(r'\s*\|[\s\-:|]+\|\s*$', cleaned_lines[i]):
                i += 1
            rows = []
            while i < len(cleaned_lines):
                row_line = cleaned_lines[i].strip()
                if not (row_line.startswith("|") and row_line.endswith("|")):
                    break
                rows.append([c.strip() for c in row_line.strip("|").split("|")])
                i += 1
            add_md_table(doc, headers, rows)
            continue

        # Bullet list
        if stripped.startswith("- ") or stripped.startswith("* "):
            flush_paragraph()
            add_bullet_list_item(doc, stripped[2:].strip())
            i += 1
            continue

        # Numbered list
        m = re.match(r'^\d+\.\s+(.+)$', stripped)
        if m:
            flush_paragraph()
            add_bullet_list_item(doc, m.group(1))
            i += 1
            continue

        # Blank line → flush
        if not stripped:
            flush_paragraph()
            i += 1
            continue

        # Otherwise: collect into paragraph
        paragraph_buffer.append(stripped)
        i += 1

    flush_paragraph()


# ============== TRANG BÌA CHÍNH ==============
def add_cover_page(doc):
    """Bìa chính: UTC + KHOA CNTT + LOGO + ĐỒ ÁN TỐT NGHIỆP + title + 9-field info + year.

    Per BAO_CAO_THUC_TAP.pdf reference + user direction 2026-05-20:
    - NO "BỘ GIÁO DỤC VÀ ĐÀO TẠO" line (báo cáo thực tập không có)
    - "ĐỒ ÁN TỐT NGHIỆP" plain black bold (NO gold + NO underline)
    """
    # TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI (header line 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI")
    set_font(run, Pt(14), bold=True)

    # KHOA CÔNG NGHỆ THÔNG TIN — Wave 102.7.1 Bucket P Fix 1:
    # Đổi text-underline (font.underline=True) sang đường kẻ ngang riêng biệt
    # (bottom border) dưới đoạn văn — đúng convention UTC bìa cứng đồ án.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("KHOA CÔNG NGHỆ THÔNG TIN")
    set_font(run, Pt(14), bold=True)
    add_horizontal_line(doc, width_pt=2, color='000000', space_after=Pt(0))

    # Logo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(36)
    logo_found = None
    for candidate in LOGO_FALLBACKS:
        if candidate.exists():
            logo_found = candidate
            break
    if logo_found:
        run = p.add_run()
        run.add_picture(str(logo_found), width=Cm(3.5))
    else:
        run = p.add_run("[LOGO UTC]")
        set_font(run, Pt(12), italic=True, color=RGBColor(128, 128, 128))

    # ĐỒ ÁN TỐT NGHIỆP (plain black, no underline per user direction 2026-05-20)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("ĐỒ ÁN TỐT NGHIỆP")
    set_font(run, Pt(24), bold=True)

    # CỬ NHÂN
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("CỬ NHÂN")
    set_font(run, Pt(18), bold=True)

    # Tên đề tài
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Đề tài:")
    set_font(run, Pt(14), italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run(THESIS_INFO["title"])
    set_font(run, Pt(20), bold=True)

    # Bảng thông tin 7-field — Wave 102.7.1 Bucket P Fix 2: bỏ "Chuyên ngành" + "Năm bảo vệ"
    # (per user direction 2026-05-20). Năm bảo vệ vẫn hiển thị qua dòng "Hà Nội – 2026" cuối bìa.
    info_rows = [
        ("Sinh viên thực hiện", STUDENT_INFO["name"]),
        ("Mã số sinh viên", STUDENT_INFO["student_id"]),
        ("Lớp", STUDENT_INFO["class"]),
        ("Khóa", STUDENT_INFO["course"]),
        ("Ngành đào tạo", STUDENT_INFO["major"]),
        ("Hệ đào tạo", STUDENT_INFO["training_mode"]),
        ("Giảng viên hướng dẫn", THESIS_INFO["advisor"]),
    ]
    table = doc.add_table(rows=len(info_rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, (label, value) in enumerate(info_rows):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].width = Cm(5.5)
        for paragraph in row.cells[0].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_font(run, Pt(13), bold=True)
        row.cells[1].text = value
        row.cells[1].width = Cm(9.0)
        for paragraph in row.cells[1].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_font(run, Pt(13))

    # Spacer + Hà Nội – 2026
    for _ in range(2):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Hà Nội – {THESIS_INFO['year']}")
    set_font(run, Pt(14), bold=True, italic=True)

    # Page break (NOT section break) — Bìa chính + Bìa phụ + Lời cảm ơn cùng Section 1
    doc.add_page_break()


# ============== TRANG BÌA PHỤ ==============
def add_secondary_cover_page(doc):
    """Bìa phụ: layout y hệt bìa chính trừ logo (thay bằng khoảng trắng cùng kích thước).

    Per user direction 2026-05-20:
    - NO "BỘ GIÁO DỤC VÀ ĐÀO TẠO" line (báo cáo thực tập không có)
    - "ĐỒ ÁN TỐT NGHIỆP" plain black bold (NO gold + NO underline)
    - Bảng info 9 fields y hệt bìa chính (NO Giáo viên phản biện)
    - Khoảng trắng vị trí logo same as bìa chính (logo width Cm(3.5) + spacing Pt(36)×2)
    """
    # TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI")
    set_font(run, Pt(14), bold=True)

    # KHOA CÔNG NGHỆ THÔNG TIN — Wave 102.7.1 Bucket P Fix 1 (bìa phụ tương tự bìa chính)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("KHOA CÔNG NGHỆ THÔNG TIN")
    set_font(run, Pt(14), bold=True)
    add_horizontal_line(doc, width_pt=2, color='000000', space_after=Pt(0))

    # Khoảng trắng tương đương logo position trên bìa chính
    # Logo bìa chính: space_before=Pt(36), height ~Cm(3.5)≈Pt(99), space_after=Pt(36)
    # → tổng vertical ~Pt(171) cho logo block. Replicate via 3 empty paragraphs với spacing.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(36)
    run = p.add_run("")
    set_font(run, Pt(99))  # invisible spacer ≈ logo height Cm(3.5)

    # ĐỒ ÁN TỐT NGHIỆP (plain black, no underline)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("ĐỒ ÁN TỐT NGHIỆP")
    set_font(run, Pt(24), bold=True)

    # CỬ NHÂN
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("CỬ NHÂN")
    set_font(run, Pt(18), bold=True)

    # Tên đề tài
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Đề tài:")
    set_font(run, Pt(14), italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run(THESIS_INFO["title"])
    set_font(run, Pt(20), bold=True)

    # Bảng 7-field Y HỆT bìa chính — Wave 102.7.1 Bucket P Fix 2:
    # bỏ "Chuyên ngành" + "Năm bảo vệ" (per user direction 2026-05-20)
    info_rows = [
        ("Sinh viên thực hiện", STUDENT_INFO["name"]),
        ("Mã số sinh viên", STUDENT_INFO["student_id"]),
        ("Lớp", STUDENT_INFO["class"]),
        ("Khóa", STUDENT_INFO["course"]),
        ("Ngành đào tạo", STUDENT_INFO["major"]),
        ("Hệ đào tạo", STUDENT_INFO["training_mode"]),
        ("Giảng viên hướng dẫn", THESIS_INFO["advisor"]),
    ]
    table = doc.add_table(rows=len(info_rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, (label, value) in enumerate(info_rows):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].width = Cm(5.5)
        for paragraph in row.cells[0].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_font(run, Pt(13), bold=True)
        row.cells[1].text = value
        row.cells[1].width = Cm(9.0)
        for paragraph in row.cells[1].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_font(run, Pt(13))

    # Spacer + Hà Nội – Năm
    for _ in range(2):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Hà Nội – {THESIS_INFO['year']}")
    set_font(run, Pt(14), bold=True, italic=True)

    # Section break Section 0 (Bìa chính + Bìa phụ — bordered) → Section 1 (Lời cảm ơn — no border)
    # Per user direction 2026-05-20: border của bìa KHÔNG leak sang Lời cảm ơn + Danh mục
    doc.add_section(WD_SECTION.NEW_PAGE)


# ============== LỜI CẢM ƠN ==============
def add_acknowledgment_page(doc):
    # NOTE: page break đã được xử lý bởi add_secondary_cover_page() khi exit
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("LỜI CẢM ƠN")
    set_font(run, Pt(16), bold=True)

    # Wave 102.5 follow-up 2026-05-20 — expanded 5-paragraph structure ~450 từ
    # Wave 102.7.1 Bucket P Fix 3: đổi đại từ "tôi" → "em" cho phù hợp khóa luận
    # tốt nghiệp (sinh viên perspective formal-respectful trước GVHD + hội đồng).
    # Phần 1 — Mở đầu
    add_paragraph_text(doc,
        "Trong suốt quá trình nghiên cứu, thiết kế và thực hiện đồ án tốt nghiệp này, em đã nhận "
        "được sự quan tâm, hướng dẫn và giúp đỡ quý báu từ nhiều tập thể và cá nhân. Đây là nguồn "
        "động viên to lớn, giúp em có thêm động lực và kiến thức để hoàn thành tốt đề tài. Em "
        "xin được gửi những lời cảm ơn chân thành nhất đến tất cả những người đã đồng hành cùng "
        "em trên hành trình này.")

    # Phần 2 — GVHD chi tiết
    add_paragraph_text(doc,
        f"Trước hết, em xin bày tỏ lòng biết ơn sâu sắc đến {THESIS_INFO['advisor']}, "
        f"giảng viên hướng dẫn thuộc {THESIS_INFO['advisor_dept']}, {THESIS_INFO['advisor_university']}. "
        "Thầy đã tận tình hướng dẫn em từ giai đoạn xác định đề tài, định hướng phạm vi nghiên cứu "
        "và phương pháp luận, đến việc đóng góp ý kiến chuyên môn quan trọng trong suốt quá trình "
        "thực hiện. Những kiến thức chuyên sâu, kinh nghiệm thực tiễn cũng như tư duy nghiêm túc "
        "trong học thuật và tinh thần phản biện mà thầy truyền đạt đã giúp em nâng cao chất lượng "
        "đồ án một cách rõ rệt, đồng thời rèn luyện cho em tác phong nghiên cứu khoa học cần thiết.")

    # Phần 3 — Khoa + Trường
    add_paragraph_text(doc,
        f"Em xin chân thành cảm ơn Khoa {STUDENT_INFO['department']} và "
        f"{STUDENT_INFO['university']} đã tạo điều kiện thuận lợi để em được tiếp cận với các kiến "
        "thức nền tảng về Công nghệ phần mềm, Kiến trúc hệ thống phân tán, Cơ sở dữ liệu, An toàn "
        "thông tin và các công nghệ thực tiễn trong ngành công nghiệp phần mềm. Môi trường học tập "
        "chuyên nghiệp cùng với chương trình đào tạo bài bản là nền tảng quan trọng giúp em có "
        "đủ năng lực và sự tự tin thực hiện đề tài này.")

    # Phần 4 — Quý thầy cô bộ môn
    add_paragraph_text(doc,
        "Bên cạnh đó, em xin gửi lời cảm ơn chân thành đến quý thầy cô trong Bộ môn Công nghệ "
        "phần mềm và toàn thể giảng viên Khoa Công nghệ thông tin đã nhiệt tình giảng dạy, chia "
        "sẻ kinh nghiệm chuyên môn trong suốt bốn năm học, qua đó giúp em xây dựng được tư duy "
        "kỹ thuật vững vàng và phương pháp tiếp cận vấn đề có hệ thống — những phẩm chất thiết "
        "yếu cho hành trình phát triển nghề nghiệp sau này.")

    # Phần 5 — Gia đình + bạn bè + đóng kết
    add_paragraph_text(doc,
        "Cuối cùng, em xin gửi lời cảm ơn sâu sắc tới gia đình, bạn bè và những người thân đã "
        "luôn quan tâm, động viên và hỗ trợ em cả về tinh thần lẫn vật chất trong suốt thời gian "
        "học tập và thực hiện đồ án. Mặc dù đã rất cố gắng, song do thời gian và kinh nghiệm thực "
        "tiễn còn hạn chế, đồ án không tránh khỏi những thiếu sót; em rất mong nhận được sự đóng "
        "góp ý kiến từ quý thầy cô để đồ án được hoàn thiện hơn. Em xin chân thành cảm ơn!")

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"Hà Nội, ngày … tháng … năm {THESIS_INFO['year']}")
    set_font(run, Pt(13), italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Sinh viên thực hiện")
    set_font(run, Pt(13), bold=True)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(STUDENT_INFO["name"])
    set_font(run, Pt(13), bold=True)

    # Wave 102.5 G17 — end of Section 1 (no page number group)
    # Next section (TOC + danh mục) starts với roman lowercase numbering
    doc.add_section(WD_SECTION.NEW_PAGE)


# ============== MỤC LỤC ==============
def add_toc_page(doc):
    # NOTE: section break (Section 1→2) đã advance page automatically — KHÔNG cần page break dư
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("MỤC LỤC")
    set_font(run, Pt(16), bold=True)

    # TOC field
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = r'TOC \o "1-3" \h \z \u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    # Wave thesis-2 Issue 1 fix — Remove placeholder note. With updateFields=true
    # flag in settings.xml (per auto_populate_fields), Word/LibreOffice tự update
    # TOC on open without manual F9 — placeholder note thừa.


def add_list_of_figures_tables(doc):
    """Add 'Danh mục hình vẽ' + 'Danh mục bảng biểu' on the same page."""
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("DANH MỤC HÌNH VẼ")
    set_font(run, Pt(16), bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\h \\z \\c "Figure"'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    # Wave thesis-2 Issue 1 fix — Remove placeholder note. SEQ Figure fields
    # injected via add_seq_caption (Issue 2 fix) — TOC \\c "Figure" sẽ populate
    # khi Word open + updateFields=true trigger F9.

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("DANH MỤC BẢNG BIỂU")
    set_font(run, Pt(16), bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\h \\z \\c "Table"'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _add_table_2col(doc, rows_data, col0_width_cm=4.0, col1_width_cm=12.0,
                    header_row=("Mục", "Giải thích")):
    """Helper — render 2-column table with header shading."""
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_widths = [Cm(col0_width_cm), Cm(col1_width_cm)]

    for i, h in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.width = col_widths[i]
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_font(run, FONT_SIZE_TABLE, bold=True)
        # Item 1b Wave 102.5 — remove blue shading; keep white default per user direction

    for col0_val, col1_val in rows_data:
        row = table.add_row()
        row.cells[0].text = col0_val
        row.cells[0].width = col_widths[0]
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                set_font(run, FONT_SIZE_TABLE, bold=True)
        row.cells[1].text = col1_val
        row.cells[1].width = col_widths[1]
        for paragraph in row.cells[1].paragraphs:
            for run in paragraph.runs:
                set_font(run, FONT_SIZE_TABLE)


def add_abbreviations(doc):
    """Danh mục THUẬT NGỮ + Danh mục TỪ VIẾT TẮT — 2 H1 danh mục riêng biệt per UTC §2 mandate."""
    # ============ Danh mục 1: THUẬT NGỮ (H1 riêng biệt) ============
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("DANH MỤC THUẬT NGỮ")
    set_font(run, Pt(16), bold=True)

    # Wave 102.7.5 Bucket B Item 10 — sorted ABC (VN-aware: Đ sau D)
    terms = [
        ("Continuous Deployment", "Quy trình triển khai tự động từ code lên môi trường production sau khi pass CI"),
        ("Continuous Integration", "Quy trình tích hợp code thường xuyên vào nhánh chung kèm automated build + test"),
        ("Defense-in-depth", "Chiến lược bảo mật nhiều lớp — mỗi lớp độc lập kiểm tra tránh single-point failure"),
        ("Domain-Driven Design", "Phương pháp thiết kế phần mềm hướng theo miền nghiệp vụ — chia hệ thống theo bounded contexts"),
        ("Multi-tenant", "Kiến trúc phần mềm cho phép nhiều tổ chức (tenant) dùng chung một hệ thống với dữ liệu cách ly"),
        ("Outbox Pattern", "Mẫu thiết kế đảm bảo tính nhất quán giữa lưu DB + phát message qua message broker"),
        ("Pool model", "Mô hình multi-tenant chia sẻ tài nguyên hạ tầng dùng RLS để cách ly dữ liệu"),
        ("Row-Level Security", "Cơ chế cách ly dữ liệu cấp dòng trên PostgreSQL — DB enforces filtering theo tenant context"),
        ("Software as a Service", "Mô hình triển khai phần mềm dạng dịch vụ điện toán đám mây, người dùng truy cập qua web"),
        ("Test-Driven Development", "Phương pháp phát triển dựa trên kiểm thử — viết test trước khi viết code (Red-Green-Refactor)"),
    ]
    _add_table_2col(doc, terms, col0_width_cm=4.5, col1_width_cm=11.5,
                    header_row=("Thuật ngữ", "Giải thích"))

    # ============ Danh mục 2: TỪ VIẾT TẮT (H1 riêng biệt, page break) ============
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("DANH MỤC TỪ VIẾT TẮT")
    set_font(run, Pt(16), bold=True)

    # Wave 102.7.5 Bucket B Item 10 — sorted ABC (VN-aware: Đ sau D)
    abbrevs = [
        ("ALB", "Application Load Balancer — bộ cân bằng tải ứng dụng AWS"),
        ("API", "Application Programming Interface"),
        ("CI/CD", "Continuous Integration / Continuous Deployment"),
        ("DDD", "Domain-Driven Design"),
        ("DPIA", "Data Protection Impact Assessment — Đánh giá Tác động Bảo vệ Dữ liệu"),
        ("DPO", "Data Protection Officer — Cán bộ Bảo vệ Dữ liệu theo PDPL"),
        ("ECR", "Elastic Container Registry — kho chứa Docker images trên AWS"),
        ("ECS", "Elastic Container Service — dịch vụ điều phối container AWS"),
        ("IEEE", "Institute of Electrical and Electronics Engineers"),
        ("ISO", "International Organization for Standardization"),
        ("JWT", "JSON Web Token"),
        ("KMS", "Key Management Service — dịch vụ quản lý khóa mã hóa"),
        ("KPI", "Key Performance Indicator — chỉ số hiệu suất chính"),
        ("LMS", "Learning Management System — hệ quản lý học tập"),
        ("MVP", "Minimum Viable Product — sản phẩm tối thiểu khả dụng"),
        ("OIDC", "OpenID Connect — chuẩn xác thực OAuth 2.0 mở rộng"),
        ("OWASP", "Open Worldwide Application Security Project"),
        ("PDPL", "Personal Data Protection Law — Luật Bảo vệ Dữ liệu Cá nhân 2023 (Số 49/2023/QH15)"),
        ("RDS", "Relational Database Service — dịch vụ cơ sở dữ liệu quan hệ AWS"),
        ("REST", "Representational State Transfer"),
        ("RLS", "Row-Level Security"),
        ("SaaS", "Software as a Service"),
        ("SES", "Simple Email Service — dịch vụ gửi email AWS"),
        ("TDD", "Test-Driven Development"),
        ("UTC GTVT", "Trường Đại học Giao thông Vận tải — University of Transport and Communications"),
        ("VPC", "Virtual Private Cloud — mạng riêng ảo trên cloud"),
    ]
    _add_table_2col(doc, abbrevs, col0_width_cm=3.5, col1_width_cm=12.5,
                    header_row=("Từ viết tắt", "Nghĩa đầy đủ"))

    # Wave 102.5 G17 — end of Section 2 (roman group)
    # Next section (Mở đầu + chapters) starts với arabic numbering từ 1
    doc.add_section(WD_SECTION.NEW_PAGE)


# ============== MỞ ĐẦU ==============
def add_introduction(doc):
    """Wave 102.7.1 Bucket P Fix 7 — wrap 6 mục dưới MỞ ĐẦU như H1 chapter-style
    title (Heading 1 style, 18pt bold center, page break trước) để 6 mục
    1. Lý do / 2. Mục tiêu / 3. Phạm vi / 4. Phương pháp / 5. Tóm tắt / 6. Cấu trúc
    trở thành H2 children dưới MỞ ĐẦU H1, đúng UTC structure.
    """
    # NOTE: section break (Section 2→3) đã advance page automatically — KHÔNG cần page break dư
    # Tuy nhiên thêm explicit page break để chắc chắn MỞ ĐẦU bắt đầu trang mới
    doc.add_page_break()

    p = doc.add_paragraph(style='Heading 1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("MỞ ĐẦU")
    run.font.name = FONT_NAME
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = FONT_SIZE_CHAPTER
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    add_section_title(doc, "1. Lý do chọn đề tài")
    add_paragraph_text(doc,
        "Thị trường phần mềm quản lý trung tâm giáo dục Việt Nam đang tăng trưởng mạnh sau khi "
        "Thông tư 29/2024/TT-BGDĐT chính thức hóa hoạt động dạy thêm có thu phí. Tuy nhiên, đa "
        "số trung tâm nhỏ và vừa (1-10 chi nhánh, 100-2000 học viên) vẫn dùng Excel hoặc các "
        "phần mềm enterprise không phù hợp về giá và độ phức tạp. Khoảng trống này tạo cơ hội "
        "cho một giải pháp SaaS multi-tenant gốc, Vietnamese-first UX, và tự động hóa các tác "
        "vụ branding bằng AI — đây chính là động lực để em chọn đề tài \"" + THESIS_INFO["title"] + "\".")

    add_section_title(doc, "2. Mục tiêu nghiên cứu")
    add_bullet_list_item(doc, "Xây dựng nền tảng SaaS multi-tenant cho trung tâm giáo dục Việt Nam, hỗ trợ scale từ 1 chi nhánh lên 100+ chi nhánh không cần re-architect.")
    add_bullet_list_item(doc, "Tích hợp AI Branding tự động sinh logo + banner + hero image, giảm thời gian go-live của trung tâm từ tuần xuống ngày.")
    add_bullet_list_item(doc, "Đảm bảo tuân thủ pháp luật Việt Nam: PDPL 2023, Luật An ninh mạng 2018, Thông tư 78/2021/TT-BTC về hóa đơn điện tử.")
    add_bullet_list_item(doc, "Áp dụng phương pháp luận audit-driven development để duy trì chất lượng code + docs trong quá trình phát triển dài hạn.")

    add_section_title(doc, "3. Phạm vi nghiên cứu")
    add_paragraph_text(doc,
        "Đồ án tập trung vào giai đoạn thử nghiệm tenant của nền tảng KiteHub, được giới hạn theo ba "
        "chiều: không gian, thời gian và đối tượng. Về không gian, hệ thống triển khai cloud trên "
        "AWS Singapore (ap-southeast-1) Free Tier theo ADR-025, phục vụ nhóm trung tâm giáo dục "
        "thương mại tại Việt Nam (Hà Nội, TP. Hồ Chí Minh và các tỉnh lân cận). Về thời gian, "
        "phạm vi triển khai kéo dài từ tháng 5/2026 đến hết giai đoạn thử nghiệm tenant (~tháng 9/2026), "
        "với cohort 7-10 trung tâm dùng miễn phí trong 9 tuần đầu để thu thập phản hồi.")
    add_paragraph_text(doc,
        "Về đối tượng, đồ án phục vụ ba persona chính: Solo Teacher (giáo viên độc lập, 1-50 học "
        "viên), Center Owner (chủ trung tâm, 1-10 chi nhánh, 100-2.000 học viên), và Center "
        "Manager (quản lý vận hành trung tâm). Persona K-12 Parent + Student được hoãn sang giai "
        "đoạn mở rộng K-12 do yêu cầu bổ sung DPO + DPIA theo Điều 26 Luật Bảo vệ Dữ liệu Cá "
        "nhân 2023. Kiến trúc hệ thống cấu thành từ ba lớp dịch vụ. Lớp nền tảng KiteHub gồm "
        "sáu dịch vụ độc lập đảm nhận các trách nhiệm khác nhau: quản trị (kitehub-admin), "
        "nhận diện thương hiệu (kitehub-branding), thư điện tử (kitehub-email), điều phối yêu "
        "cầu (kitehub-gateway), thư viện dùng chung (kitehub-platform) và quản lý đăng ký "
        "(kitehub-subscription). Lớp nghiệp vụ tenant KiteClass tập trung tại dịch vụ "
        "kiteclass-core. Lớp giao diện gồm hai ứng dụng Next.js phục vụ tập người dùng khác "
        "nhau.")

    add_section_title(doc, "4. Phương pháp nghiên cứu")
    add_paragraph_text(doc,
        "Đồ án kết hợp phương pháp nghiên cứu lý thuyết (literature review) với phương pháp thực "
        "nghiệm (experimental design + implementation):")
    add_bullet_list_item(doc,
        "Phân tích thị trường VN edu SaaS qua so sánh có hệ thống 4 hệ thống tương tự "
        "(MISA AMIS Trường Học, Mona eLMS, Easy Edu, DotB) trên các tiêu chí giá, tính năng "
        "multi-tenant, mức tích hợp AI và mức tuân thủ pháp luật Việt Nam.")
    add_bullet_list_item(doc,
        "Thiết kế kiến trúc multi-tenant theo các pattern industry-standard (single-bucket "
        "Row-Level Security, defense-in-depth 5 lớp) đối chiếu với AWS SaaS Lens và Azure "
        "multi-tenant whitepaper.")
    add_bullet_list_item(doc,
        "Áp dụng phương pháp luận Quality-Driven Development bốn trụ cột (TDD per Beck 2002, "
        "DDD per Evans 2003, PDCA per Deming 1986, Lean per Poppendieck 2003) — mỗi miss "
        "được chuyển thành rule + cơ chế enforcement trong cùng pull request.")
    add_bullet_list_item(doc,
        "Đánh giá chất lượng hệ thống qua bộ audit bảy chiều (Quality, UI, Security, "
        "Performance, API Contract, Business Logic, Ops Readiness) chấm điểm /100 hoặc /128 "
        "định kỳ sau mỗi wave.")

    add_section_title(doc, "5. Tóm tắt nội dung")
    add_paragraph_text(doc,
        "Đồ án trình bày bốn đóng góp chính tương ứng bốn chương nội dung. Thứ nhất, đồ án "
        "phân tích thị trường EdTech Việt Nam và khung pháp lý liên quan (PDPL 2023, Luật An "
        "ninh mạng 2018, Thông tư 29/2024/TT-BGDĐT) làm cơ sở xác định khoảng trống và nhu cầu "
        "của nhóm trung tâm vừa và nhỏ. Thứ hai, đồ án thiết kế và mô tả kiến trúc hệ thống "
        "multi-tenant SaaS với mô hình C4 bốn cấp, kết hợp Use Case + Class + ERD và sơ đồ "
        "tuần tự cho các luồng quan trọng (đăng ký, kích hoạt subscription).")
    add_paragraph_text(doc,
        "Thứ ba, đồ án triển khai và kiểm thử hệ thống trên AWS Singapore Free Tier, với các "
        "pattern cốt lõi (Row-Level Security NULL force-fail, Outbox Pattern, JWT propagation, "
        "REST 3-tier) được kiểm chứng qua unit test, integration test và E2E test. Thứ tư, đồ "
        "án trình bày kết quả triển khai thực tế trong giai đoạn thử nghiệm tenant kèm các KPI đo "
        "lường, đánh giá độ trưởng thành và đề xuất hướng phát triển tiếp theo. Kết quả của "
        "đồ án vừa là sản phẩm phần mềm hoạt động được, vừa là tài liệu tham chiếu cho các "
        "công trình nghiên cứu kế tiếp về EdTech multi-tenant tại Việt Nam.")

    add_section_title(doc, "6. Cấu trúc đồ án")
    add_paragraph_text(doc, "Đồ án gồm bốn chương nội dung chính:")
    add_bullet_list_item(doc, "Chương 1 — Tổng quan: phân tích đối tượng tham khảo trên thị trường, kỹ thuật AI tích hợp, và khung pháp lý Việt Nam tác động đến nền tảng.")
    add_bullet_list_item(doc, "Chương 2 — Kiến trúc hệ thống: yêu cầu chức năng + phi chức năng, mô hình hóa C4 + Use Case + Class + ERD, thiết kế cơ sở dữ liệu, multi-tenant single-bucket, defense-in-depth 5 lớp.")
    add_bullet_list_item(doc, "Chương 3 — Triển khai: kết quả triển khai giao diện sản phẩm và bộ kiểm thử ba lớp (unit / integration / E2E) với các sample test case cụ thể.")
    add_bullet_list_item(doc, "Chương 4 — Kết quả triển khai: cloud AWS giai đoạn thử nghiệm tenant, user onboarding flow, KPI metrics, scope thử nghiệm.")


# ============== CHAPTER LOADER (MD parser) ==============
def add_chapter_from_md(doc, chapter_num, chapter_title, md_paths):
    """Load 1 or more MD files → inject as chapter."""
    add_chapter_title(doc, chapter_num, chapter_title)

    for md_path in md_paths:
        if not md_path.exists():
            print(f"WARN: chapter MD not found: {md_path}")
            add_paragraph_text(doc, f"[Chapter source missing: {md_path.name}]")
            continue
        print(f"  Loading: {md_path.name}")
        md_text = md_path.read_text(encoding='utf-8')
        parse_markdown(doc, md_text, skip_top_heading=True)


# ============== KẾT LUẬN ==============
def add_conclusion(doc):
    """Wave 102.7.1 Bucket P Fix 8 — wrap 5 mục (Tổng kết / Hạn chế / Hướng
    phát triển / Đóng góp khoa học / Kiến nghị) dưới KẾT LUẬN VÀ KIẾN NGHỊ
    H1 chapter-style title (Heading 1, 18pt bold center, page break trước).
    5 mục trở thành H2 children dưới KẾT LUẬN H1, đúng UTC structure.
    """
    doc.add_page_break()

    p = doc.add_paragraph(style='Heading 1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("KẾT LUẬN VÀ KIẾN NGHỊ")
    run.font.name = FONT_NAME
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = FONT_SIZE_CHAPTER
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    add_section_title(doc, "1. Tổng kết kết quả đạt được")
    add_paragraph_text(doc,
        "Đồ án đã hoàn thành các mục tiêu đặt ra ban đầu trong phạm vi giai đoạn thử nghiệm tenant của "
        "nền tảng KiteHub. Cụ thể, hệ thống được triển khai trên AWS Singapore Free Tier với "
        "kiến trúc multi-tenant single-bucket RLS-protected, 7 microservice cùng 2 ứng dụng "
        "frontend, đảm bảo tuân thủ các yêu cầu pháp lý Việt Nam (PDPL 2023, Luật An ninh mạng "
        "2018, Thông tư 78/2021/TT-BTC).")

    add_paragraph_text(doc,
        "Các pattern kiến trúc cốt lõi đã được hiện thực hóa qua 5 đoạn code đại diện (JWT "
        "authentication tại Gateway, Row-Level Security cho multi-tenant query, Outbox "
        "Pattern cho email dispatch, REST API 3-tier cho beta access, Next.js App Router "
        "page) — tất cả đều được kiểm thử qua bộ unit test + integration test, đáp ứng tỉ lệ "
        "coverage tối thiểu yêu cầu cho production-ready.")

    # Wave 102.7.5 Bucket B Item 8 — Phụ lục B reframed as inline GitHub link
    add_paragraph_text(doc,
        "Toàn bộ mã nguồn dự án được công bố tại kho lưu trữ công khai trên GitHub: "
        "https://github.com/VictorAurelius/2026-Kite-Class-Platform — bao gồm các thành phần "
        "kitehub/ (6 microservice nền tảng + frontend), kiteclass/ (dịch vụ nghiệp vụ tenant), "
        "infrastructure/ (cấu hình hạ tầng Terraform), và documents/ (tài liệu thiết kế + audit). "
        "Người đọc có thể tham khảo trực tiếp mã nguồn để đối chiếu với các pattern và đoạn code "
        "trình bày trong báo cáo.")

    add_section_title(doc, "2. Hạn chế")
    add_bullet_list_item(doc, "Phạm vi giai đoạn thử nghiệm tenant chỉ phục vụ 3 persona tenant (Solo Teacher, Center Owner, Center Manager); persona K-12 Parent + Student được hoãn sang giai đoạn K-12 expansion do yêu cầu DPO + DPIA bổ sung theo PDPL.")
    add_bullet_list_item(doc, "Một số KPI thực tế (Time to First Value, Daily Active Users, Monthly Recurring Revenue) chưa có số liệu thực tế do beta cohort 7-10 tenant đang trong giai đoạn triển khai 9 tuần.")
    add_bullet_list_item(doc, "AI Branding mới được tích hợp ở mức MVP với 1 nhà cung cấp (Replicate Stable Diffusion XL); các phương án multi-vendor failover sẽ được triển khai trong giai đoạn production.")

    add_section_title(doc, "3. Hướng phát triển tiếp theo")
    add_bullet_list_item(doc,
        "Giai đoạn trả phí (paid tier): tích hợp thanh toán (VNPay, MoMo), partnership MISA MeInvoice cho "
        "hóa đơn điện tử theo Thông tư 78/2021/TT-BTC, mở rộng tenant cohort beta lên 30-50 trung tâm.")
    add_bullet_list_item(doc,
        "Giai đoạn vận hành chính thức: kiến trúc đa-region (Singapore + Hà Nội data localization "
        "Việt Nam theo Nghị định 53/2022), AI Quality Gate phiên bản nâng cao với multi-vendor failover, "
        "mobile app native iOS + Android.")
    add_bullet_list_item(doc,
        "Giai đoạn mở rộng K-12: mở rộng sang persona trường công lập với DPO chính thức + DPIA cho dữ "
        "liệu trẻ em theo Luật Bảo vệ Dữ liệu Cá nhân Điều 26.")

    # Wave thesis-2 Bucket Issue 7 fix — Remove "4. Đóng góp khoa học" + "5. Kiến nghị" separate sections per khung primary §1 Ch.4 §4.4
    # (khung mandate "Kết luận, kiến nghị + Phương hướng phát triển" gộp 1 mục §4.4; KHÔNG separate "Đóng góp khoa học" section).
    # Audit map item 7: "Bỏ Đóng góp khoa học + Kiến nghị" — confirmed.


# ============== TÀI LIỆU THAM KHẢO ==============
def parse_bibliography_md(md_text):
    """
    Parse bibliography.md → list of dicts {num, raw_text}.
    Each entry starts with `[N] ` and may span multiple lines.
    """
    entries = []
    lines = md_text.split('\n')
    current = None
    current_text = []
    pattern = re.compile(r'^\[(\d+)\]\s+(.*)$')

    for line in lines:
        # Skip frontmatter, headers, and metadata
        if line.startswith('---'):
            continue
        if line.startswith('#'):
            # Save current entry if any
            if current is not None:
                entries.append((current, ' '.join(current_text).strip()))
                current = None
                current_text = []
            continue
        m = pattern.match(line.strip())
        if m:
            # Save previous
            if current is not None:
                entries.append((current, ' '.join(current_text).strip()))
            current = int(m.group(1))
            current_text = [m.group(2)]
        elif current is not None and line.strip():
            current_text.append(line.strip())

    if current is not None:
        entries.append((current, ' '.join(current_text).strip()))

    # Sort by num
    entries.sort(key=lambda x: x[0])
    return entries


def add_references_from_md(doc):
    """Load bibliography.md → render IEEE-formatted paragraphs."""
    doc.add_page_break()

    p = doc.add_paragraph(style='Heading 1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TÀI LIỆU THAM KHẢO")
    run.font.name = FONT_NAME
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = FONT_SIZE_CHAPTER
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    if not BIBLIOGRAPHY_FILE.exists():
        add_paragraph_text(doc, f"[Bibliography missing: {BIBLIOGRAPHY_FILE}]")
        return

    md_text = BIBLIOGRAPHY_FILE.read_text(encoding='utf-8')
    entries = parse_bibliography_md(md_text)
    print(f"  Loaded {len(entries)} bibliography entries")

    for num, text in entries:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.63)
        p.paragraph_format.first_line_indent = Cm(-0.63)
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.15

        # Add [N] prefix
        run = p.add_run(f"[{num}] ")
        set_font(run, FONT_SIZE_NORMAL)

        # Render rest with inline markdown
        add_inline_runs(p, text)


# Wave 102.7.5 Bucket B Items 8+9 — `add_appendix` function removed entirely.
# Phụ lục A (cấu hình triển khai): removed — nội dung quan trọng đã có trong Chương 4 deployment results.
# Phụ lục B (mã nguồn): reframed as inline GitHub link trong KẾT LUẬN §1 "Tổng kết kết quả đạt được".
# Phụ lục C (audit chất lượng): removed entirely.


# ============== AUTO-POPULATE TOC + SEQ FIELDS ==============
def auto_populate_fields(docx_path: Path) -> bool:
    """
    Auto-populate Word field codes (TOC, SEQ, etc).

    Wave thesis-2 Bucket A.4.1 fix: previous impl chỉ `libreoffice --convert-to
    docx` (re-encode without field calculation) — fields stayed as placeholder
    "Bấm Ctrl+A...F9". User-visible bug.

    Fix: 2-step approach
    1. Set <w:updateFields w:val="true"/> trong word/settings.xml (python-docx) —
       Word / LibreOffice tự auto-update fields khi mở file lần đầu.
    2. Run LibreOffice headless macro để force update NOW (optional — nếu
       LibreOffice available, save docx với fields đã populated; otherwise
       Step 1 đủ để user thấy populated fields khi mở docx local).
    """
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Step 1: Set updateFields=true flag trong settings.xml
    try:
        doc = Document(str(docx_path))
        settings_element = doc.settings.element
        # Remove existing updateFields if any
        for existing in settings_element.findall(qn('w:updateFields')):
            settings_element.remove(existing)
        # Add new updateFields=true
        update_fields = OxmlElement('w:updateFields')
        update_fields.set(qn('w:val'), 'true')
        settings_element.append(update_fields)
        doc.save(str(docx_path))
        print(f"✅ Set <w:updateFields w:val='true'/> trong settings.xml — "
              f"Word/LibreOffice tự auto-update fields khi mở file (TOC + Danh mục bảng + Danh mục hình)")
    except Exception as e:
        print(f"⚠️  Could not set updateFields flag: {e}")
        print("   Fallback: manual Ctrl+A + F9 trong Word per thesis-pre-defense-checklist.md §1")
        return False

    # Step 2: Try LibreOffice headless macro để force update NOW (best-effort)
    import subprocess
    macro_inline = f'''
import uno
def update_fields_now():
    ctx = uno.getComponentContext()
    smgr = ctx.ServiceManager
    desktop = smgr.createInstanceWithContext("com.sun.star.frame.Desktop", ctx)
    url = "file://{docx_path}"
    doc = desktop.loadComponentFromURL(url, "_blank", 0, tuple())
    if doc is not None:
        # Update TOC + indexes + fields
        if hasattr(doc, "DocumentIndexes"):
            for idx in doc.DocumentIndexes:
                idx.update()
        if hasattr(doc, "TextFields"):
            for field in doc.TextFields:
                if hasattr(field, "update"):
                    field.update()
        doc.store()
        doc.close(True)
update_fields_now()
'''
    macro_path = docx_path.parent / "_update_fields_macro.py"
    try:
        macro_path.write_text(macro_inline, encoding='utf-8')
        # LibreOffice Python bridge - use UnoCommand via cli isn't easy.
        # Skip this step — Step 1 updateFields flag đủ để user thấy fields
        # populated khi mở file. Manual F9 trong Word vẫn workflow per
        # thesis-pre-defense-checklist.md §1 nếu fields chưa update ngay.
        macro_path.unlink(missing_ok=True)
        print("   ℹ️  Step 2 macro skipped — Step 1 flag đủ cho user view trong Word/LibreOffice")
    except Exception as e:
        print(f"   ⚠️  Step 2 macro setup failed (non-fatal): {e}")

    return True


# ============== MAIN ENTRY POINT ==============
def create_thesis():
    print("=" * 60)
    print(f"Đang tạo đồ án tốt nghiệp: {THESIS_INFO['title'][:60]}")
    print("=" * 60)

    doc = Document()
    setup_styles(doc)

    # 1. Bìa chính
    print("[1/8] Trang bìa chính")
    add_cover_page(doc)

    # 2. Bìa phụ
    print("[2/8] Trang bìa phụ")
    add_secondary_cover_page(doc)

    # Wave thesis-2 Bucket A.2 — REMOVED non-khung frontmatter pages
    # (NHẬN XÉT GVHD / LỜI CAM ĐOAN / TÓM TẮT / ABSTRACT) per khung-bao-cao-do-an.png
    # primary source. Sequence now: Bìa → Bìa phụ → Lời cảm ơn → Mục lục → ...
    # Per thesis-content-standard.md v2.0.0 §3 banned non-khung sections.

    # 3. Lời cảm ơn
    print("[3/8] Lời cảm ơn")
    add_acknowledgment_page(doc)

    # 4. Mục lục + Danh mục
    print("[4/8] Mục lục + Danh mục hình/bảng + Danh mục từ viết tắt")
    add_toc_page(doc)
    add_list_of_figures_tables(doc)
    add_abbreviations(doc)

    # 5. Mở đầu
    print("[5/8] Mở đầu")
    add_introduction(doc)

    # 6. 4 Chương nội dung
    for ch_num in [1, 2, 3, 4]:
        title = CHAPTER_TITLES[ch_num]
        md_files = CHAPTER_FILES[ch_num]
        print(f"[6/8] CHƯƠNG {ch_num}. {title} ({len(md_files)} MD source{'s' if len(md_files) > 1 else ''})")
        add_chapter_from_md(doc, ch_num, title, md_files)

    # 7. Kết luận
    print("[7/8] Kết luận")
    add_conclusion(doc)

    # 8. Tài liệu tham khảo
    # Wave 102.7.5 Bucket B: Phụ lục A+C removed; Phụ lục B inline trong KẾT LUẬN §1.
    print("[8/8] Tài liệu tham khảo IEEE")
    add_references_from_md(doc)

    # Apply borders: chỉ section 0 = Bìa chính + Bìa phụ (border)
    # Wave 102.5 follow-up 2026-05-20 border-leak fix: bìa + bìa phụ giờ cùng Section 0
    # (NEW_PAGE break giữa bìa chính+bìa phụ chỉ là page break, không phải section break);
    # section 1+ = Lời cảm ơn / Danh mục / Mở đầu — KHÔNG có border
    print(f"DEBUG: Total sections = {len(doc.sections)}")
    if len(doc.sections) >= 2:
        add_page_border(doc.sections[0])
        for i in range(1, len(doc.sections)):
            remove_page_border(doc.sections[i])

    # Apply margins to all sections
    set_document_margins(doc)

    # Add page numbers (skip 2 covers)
    add_page_number_header(doc)

    # Save
    doc.save(str(OUTPUT_FILE))
    print()
    print("=" * 60)
    print(f"✅ Đã tạo file: {OUTPUT_FILE}")
    print(f"   Số sections: {len(doc.sections)}")
    print(f"   Số paragraphs: {len(doc.paragraphs)}")
    print("=" * 60)

    # Auto-populate TOC + SEQ fields if LibreOffice available (graceful fallback to Word F9)
    auto_populate_fields(OUTPUT_FILE)

    return OUTPUT_FILE


# ============== RUBRIC VALIDATION (heuristic post-bake) ==============
def validate_rubric(docx_path: Path) -> dict:
    """Heuristic check theo `thesis-content-standard.md` v1.1.0 9-category rubric /100.

    Output: dict {category_id, score, max, notes} + total /100 + verdict (PASS ≥75 / PARTIAL 60-74 / FAIL <60).

    Scope: subset of rubric mà có thể tự verify từ source MD + docx — NOT replace human reviewer.
    Categories đo được heuristic: C1 Format (margin/font/section count), C2 Content (page count proxy),
    C3 Bibliography (entries count + utilization), C5 Project-internal scrub, C6 Draft-marker scrub.
    Categories KHÔNG đo được heuristic (defer human reviewer): C4 Academic tone, C7 Figure rendering nuances,
    C8 Examiner readiness, C9 Compliance phrasing.
    """
    import re

    from docx import Document as _Document

    notes = []
    scores = {}

    # Load docx
    if not docx_path.exists():
        return {"error": f"docx not found: {docx_path}", "total": 0, "verdict": "FAIL"}
    doc = _Document(str(docx_path))

    # ============== C1 — Format compliance (15 pts) — partial heuristic ==============
    c1 = 0
    section = doc.sections[0] if doc.sections else None
    if section:
        # A4 page size check (210x297mm)
        a4_width_cm = 21.0
        a4_height_cm = 29.7
        actual_w_cm = section.page_width.cm if section.page_width else 0
        actual_h_cm = section.page_height.cm if section.page_height else 0
        if abs(actual_w_cm - a4_width_cm) < 0.5 and abs(actual_h_cm - a4_height_cm) < 0.5:
            c1 += 2
            notes.append(f"C1: A4 page size verified ({actual_w_cm:.1f}x{actual_h_cm:.1f}cm)")
        else:
            notes.append(f"C1: page size NOT A4 ({actual_w_cm:.1f}x{actual_h_cm:.1f}cm) -- expected 21.0x29.7cm")

        # Margins check (T=2.5 B=2.5 L=3.0 R=2.0cm)
        margins_ok = (
            abs((section.top_margin.cm if section.top_margin else 0) - 2.5) < 0.3
            and abs((section.bottom_margin.cm if section.bottom_margin else 0) - 2.5) < 0.3
            and abs((section.left_margin.cm if section.left_margin else 0) - 3.0) < 0.3
            and abs((section.right_margin.cm if section.right_margin else 0) - 2.0) < 0.3
        )
        if margins_ok:
            c1 += 2
            notes.append("C1: margins verified T=2.5 B=2.5 L=3.0 R=2.0cm")
        else:
            notes.append("C1: margins NOT match UTC spec -- check section.{top,bottom,left,right}_margin")

    # TNR + heading font heuristic (count Normal style references)
    normal_style = doc.styles["Normal"] if "Normal" in [s.name for s in doc.styles] else None
    if normal_style and normal_style.font and normal_style.font.name == "Times New Roman":
        c1 += 3
        notes.append("C1: Normal style font = Times New Roman")
    else:
        notes.append(f"C1: Normal style font = {normal_style.font.name if normal_style and normal_style.font else 'unknown'} (expected TNR)")

    # Sub-section numbering heuristic — count strict CHƯƠNG N. + N.M
    para_texts = [p.text for p in doc.paragraphs]
    chapter_count = sum(1 for t in para_texts if re.match(r"^\s*CHƯƠNG\s+[0-9]+\.", t))
    if chapter_count >= 4:
        c1 += 1
        notes.append(f"C1: detected {chapter_count} CHƯƠNG N. headings (>=4 expected)")
    scores["C1"] = {"score": c1, "max": 15}

    # ============== C2 — Content + page count (15 pts) — heuristic via paragraph count ==============
    c2 = 0
    para_count = len(doc.paragraphs)
    # Char-based estimate ~2500 chars per page body TNR 13pt 1.2 line-spacing.
    # Tables: ~10% screen real estate each; Figures (inline shapes): ~30% screen real estate each.
    para_chars = sum(len(p.text) for p in doc.paragraphs)
    table_chars = sum(len(c.text) for tbl in doc.tables for r in tbl.rows for c in r.cells)
    n_tables = len(doc.tables)
    n_figures = len(doc.inline_shapes)
    # Text body pages
    text_pages = (para_chars + table_chars) / 2500
    # Tables: add 0.5 page each (since some are large)
    table_pages = n_tables * 0.5
    # Figures: add 0.4 page each (caption + image)
    figure_pages = n_figures * 0.4
    # Frontmatter overhead (covers + TOC + danh mục) ~10 pages
    frontmatter_pages = 10
    est_pages = int(text_pages + table_pages + figure_pages + frontmatter_pages)
    if 60 <= est_pages <= 80:
        c2 += 6
        notes.append(f"C2: estimated pages ~{est_pages} (60-80 target hit)")
    elif 81 <= est_pages <= 90:
        c2 += 4
        notes.append(f"C2: estimated pages ~{est_pages} (81-90 soft window)")
    elif est_pages <= 100:
        c2 += 2
        notes.append(f"C2: estimated pages ~{est_pages} (over 90 cap soft-deduct)")
    elif est_pages <= 120:
        c2 += 0
        notes.append(f"C2: estimated pages ~{est_pages} -- over 100, near target +30, deduct")
    else:
        notes.append(f"C2: estimated pages ~{est_pages} -- significantly over target")
    # Chapter content present
    if chapter_count >= 4:
        c2 += 4
    # Conclusion section heuristic
    if any("KẾT LUẬN" in t for t in para_texts):
        c2 += 2
        notes.append("C2: KẾT LUẬN section present")
    scores["C2"] = {"score": c2, "max": 15}

    # ============== C3 — Bibliography IEEE format (15 pts) — heuristic ==============
    c3 = 0
    # Count [N] references in bibliography section
    bib_pattern = re.compile(r"^\s*\[[0-9]+\]\s")
    bib_entries = sum(1 for t in para_texts if bib_pattern.match(t))
    if bib_entries >= 30:
        c3 += 3
        notes.append(f"C3: {bib_entries} bibliography entries (>=30 cử nhân target)")
    elif bib_entries >= 20:
        c3 += 2
        notes.append(f"C3: {bib_entries} bibliography entries (between 20-29)")
    else:
        notes.append(f"C3: only {bib_entries} bibliography entries (<20)")

    if any("TÀI LIỆU THAM KHẢO" in t for t in para_texts):
        c3 += 1
    # Cite utilization heuristic — count distinct [N] in body
    body_text = "\n".join(para_texts)
    cite_matches = set(re.findall(r"\[([0-9]+)\]", body_text))
    if len(cite_matches) >= bib_entries * 0.9 and bib_entries > 0:
        c3 += 4
        notes.append(f"C3: cite utilization {len(cite_matches)}/{bib_entries} (>=90%)")
    elif bib_entries > 0:
        c3 += 2
        notes.append(f"C3: cite utilization {len(cite_matches)}/{bib_entries} (<90%)")
    c3 = min(c3, 15)
    scores["C3"] = {"score": c3, "max": 15}

    # ============== C4 — Academic tone (15 pts) — partial heuristic ==============
    c4 = 0
    # Narrative-only banned words (technical OK / HTTP 200 OK = legitimate technical context).
    # Refined heuristic: only count standalone words trong narrative, not technical tokens.
    narrative_banned = {
        "đối thủ": "competitor business jargon",
        "🎉": "emoji",
        "📅": "emoji",
        "🆘": "emoji",
    }
    # Refined: only flag 'bạn' as informal pronoun if NOT preceded by other Vietnamese
    # compound (bạn bè / bạn đọc / quý bạn = legitimate); detect via word-boundary + lookahead.
    informal_pronouns = {
        r"\bbạn\s+(đã|sẽ|là|có|đang|được|cần)\b": "informal 'bạn' as pronoun-subject",
        r"\bchúng ta\b": "informal 'chúng ta' pronoun",
    }
    banned_hits = 0
    for w, label in narrative_banned.items():
        body_hits = body_text.count(w)
        if body_hits > 0:
            notes.append(f"C4: banned '{w}' ({label}) found {body_hits} times")
            banned_hits += body_hits
    for pat, label in informal_pronouns.items():
        body_hits = len(re.findall(pat, body_text))
        if body_hits > 0:
            notes.append(f"C4: informal pronoun '{pat}' ({label}) found {body_hits} times")
            banned_hits += body_hits
    if banned_hits == 0:
        c4 += 12
        notes.append("C4: no banned narrative tokens (đối thủ / emoji / informal pronoun)")
    elif banned_hits < 3:
        c4 += 8
    elif banned_hits < 10:
        c4 += 4
    scores["C4"] = {"score": c4, "max": 15}

    # ============== C5 — Project-internal scrub (10 pts) ==============
    c5 = 0
    banned_internal = ["Claude", "Wave [0-9]", "Phase 1 BETA", "GAP-[0-9]", r"\.claude/"]
    internal_hits = 0
    for pat in banned_internal:
        hits = len(re.findall(pat, body_text))
        if hits > 0:
            notes.append(f"C5: banned internal ref '{pat}' found {hits} times")
            internal_hits += hits
    if internal_hits == 0:
        c5 += 10
        notes.append("C5: no project-internal refs in body")
    elif internal_hits < 3:
        c5 += 6
    elif internal_hits < 10:
        c5 += 3
    scores["C5"] = {"score": c5, "max": 10}

    # ============== C6 — Draft-marker scrub (5 pts) ==============
    c6 = 0
    draft_patterns = ["TL;DR", "TODO", "FIXME", "placeholder", "[stub]", "v0.9.0-beta", "Cập nhật lần cuối"]
    draft_hits = 0
    for pat in draft_patterns:
        hits = body_text.count(pat)
        if hits > 0:
            notes.append(f"C6: draft marker '{pat}' found {hits} times")
            draft_hits += hits
    if draft_hits == 0:
        c6 += 5
        notes.append("C6: no draft markers")
    elif draft_hits < 3:
        c6 += 3
    elif draft_hits < 10:
        c6 += 1
    scores["C6"] = {"score": c6, "max": 5}

    # ============== C7 — Figure rendering (10 pts) — heuristic via inline_shapes ==============
    c7 = 0
    fig_count = len(doc.inline_shapes)
    if fig_count >= 5:
        c7 += 6
        notes.append(f"C7: {fig_count} inline shapes (figures) embedded")
    elif fig_count >= 1:
        c7 += 3
        notes.append(f"C7: only {fig_count} inline shapes (target >=5)")
    else:
        notes.append("C7: NO inline shapes -- mermaid blocks may not be rendered as images")

    # No raw mermaid text leak
    if "```mermaid" not in body_text and "flowchart " not in body_text[:500]:
        c7 += 2
        notes.append("C7: no raw mermaid code in body")
    # Hình caption present
    fig_caption_count = len(re.findall(r"Hình\s+[0-9]+\.[0-9]+", body_text))
    if fig_caption_count >= 3:
        c7 += 2
    scores["C7"] = {"score": c7, "max": 10}

    # ============== C8 — Examiner readiness (10 pts) — partial heuristic ==============
    c8 = 0
    if any("LỜI CẢM ƠN" in t for t in para_texts):
        c8 += 2
    if any("MỞ ĐẦU" in t for t in para_texts):
        c8 += 2
    if any("KẾT LUẬN" in t for t in para_texts):
        c8 += 2
    if c3 >= 10:
        c8 += 2
    if fig_count >= 5:
        c8 += 2
    scores["C8"] = {"score": c8, "max": 10}

    # ============== C9 — Compliance + legal (5 pts) — defer human ==============
    # Heuristic: refined -- "vi phạm dữ liệu" / "vi phạm pháp luật" trong context của VN law citation
    # legitimate (data breach notification = legal term); ONLY flag if pattern is "thừa nhận vi phạm" / "vi phạm Decree".
    c9 = 0
    self_admit_violation = bool(re.search(r"thừa nhận vi phạm|vi phạm (Decree|Nghị định) 53/2022|compliance debt được chấp nhận", body_text))
    if not self_admit_violation:
        c9 += 3
        notes.append("C9: no explicit self-admit violation phrasing")
    else:
        notes.append("C9: self-admit violation phrasing found -- review per thesis-content-standard.md §C9")
    if "DPO" in body_text or "DPIA" in body_text:
        c9 += 2
    scores["C9"] = {"score": c9, "max": 5}

    # ============== Total ==============
    total = sum(s["score"] for s in scores.values())
    max_total = sum(s["max"] for s in scores.values())

    if total >= 85:
        verdict = "PASS A (≥85)"
    elif total >= 75:
        verdict = "PASS C+ (≥75 minimum per thesis-content-standard.md §1)"
    elif total >= 60:
        verdict = "PARTIAL C (60-74)"
    else:
        verdict = "FAIL (<60)"

    return {
        "total": total,
        "max_total": max_total,
        "verdict": verdict,
        "scores": scores,
        "notes": notes,
        "para_count": para_count,
        "est_pages": est_pages,
        "bib_entries": bib_entries,
        "cite_utilization": f"{len(cite_matches)}/{bib_entries}" if bib_entries else "0/0",
        "fig_count": fig_count,
    }


def print_rubric_report(result: dict) -> None:
    """Print rubric validation result in readable format."""
    print()
    print("=" * 60)
    print("RUBRIC VALIDATION (heuristic per thesis-content-standard.md v1.1.0)")
    print("=" * 60)
    if "error" in result:
        print(f"ERROR: {result['error']}")
        return
    print(f"Total: {result['total']}/{result['max_total']} -- {result['verdict']}")
    print()
    print("Per-category scores:")
    for cat_id, s in result["scores"].items():
        marker = "PASS" if s["score"] >= s["max"] * 0.7 else "PARTIAL" if s["score"] >= s["max"] * 0.5 else "FAIL"
        print(f"  {cat_id}: {s['score']}/{s['max']} ({marker})")
    print()
    print(f"Metadata: {result['para_count']} paragraphs / ~{result['est_pages']} pages / "
          f"{result['bib_entries']} bib entries / cite {result['cite_utilization']} / {result['fig_count']} figures")
    print()
    print("Notes:")
    for n in result["notes"]:
        print(f"  - {n}")
    print("=" * 60)


# ============== ENTRY POINT (argparse) ==============
def main():
    import argparse
    parser = argparse.ArgumentParser(
        description="Thesis V1 DOCX pipeline -- generate khóa luận tốt nghiệp from chapter MDs."
    )
    parser.add_argument(
        "--execute",
        action="store_true",
        help="Production mode: full bake -- save docx + auto-populate fields. Default behavior (kept for backward compat).",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Dry-run mode: parse chapter MDs + report stats WITHOUT saving docx (useful for CI lint).",
    )
    parser.add_argument(
        "--validate-rubric",
        action="store_true",
        help="Post-bake: run heuristic rubric validation per thesis-content-standard.md v1.1.0 9-category /100.",
    )
    args = parser.parse_args()

    # --dry-run + --execute mutually exclusive: --execute wins (production)
    if args.dry_run and not args.execute:
        # Dry-run only: parse + report but don't save docx
        print("=" * 60)
        print("DRY-RUN MODE -- parse chapter MDs + report stats (no docx save)")
        print("=" * 60)
        total_chars = 0
        total_lines = 0
        for ch_num, md_paths in CHAPTER_FILES.items():
            for md_path in md_paths:
                if md_path.exists():
                    content = md_path.read_text(encoding="utf-8")
                    total_chars += len(content)
                    total_lines += content.count("\n")
                    print(f"  Ch.{ch_num}: {md_path.name} -- {len(content)} chars / {content.count(chr(10))} lines")
                else:
                    print(f"  Ch.{ch_num}: {md_path.name} -- MISSING")
        bib_path = BIBLIOGRAPHY_FILE
        if bib_path.exists():
            bib_content = bib_path.read_text(encoding="utf-8")
            bib_entries = bib_content.count("\n[")
            print(f"  Bibliography: {bib_entries} entries / {len(bib_content)} chars")
        print(f"Total source: {total_chars} chars / {total_lines} lines")
        print("=" * 60)
        return

    # Default = --execute (production mode)
    output_path = create_thesis()

    if args.validate_rubric:
        result = validate_rubric(output_path)
        print_rubric_report(result)


if __name__ == "__main__":
    main()
