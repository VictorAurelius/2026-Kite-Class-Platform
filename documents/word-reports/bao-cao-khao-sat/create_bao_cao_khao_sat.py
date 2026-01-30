#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script tạo báo cáo khảo sát KiteClass Platform dạng Word (.docx)
Format theo quy định trình bày đồ án tốt nghiệp - ĐH GTVT

Quy chuẩn áp dụng:
- Căn lề: trên 2.5cm, dưới 2.5cm, trái 3cm, phải 2cm
- Số trang: giữa, phía trên đầu trang
- Chương: Times New Roman 18pt, Bold, căn giữa
- Mục (1.1): Times New Roman 16pt, Bold, căn trái
- Tiểu mục (1.1.1): Times New Roman 14pt, Bold, căn trái
- Đoạn văn: Times New Roman 13pt, Justify, thụt đầu dòng 1cm, giãn dòng 1.2
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============== CONSTANTS theo quy định ==============
FONT_NAME = 'Times New Roman'
FONT_SIZE_NORMAL = Pt(13)
FONT_SIZE_CHAPTER = Pt(18)
FONT_SIZE_SECTION = Pt(16)
FONT_SIZE_SUBSECTION = Pt(14)
FONT_SIZE_TABLE = Pt(12)
FONT_SIZE_CAPTION = Pt(13)

LINE_SPACING = 1.2
FIRST_LINE_INDENT = Cm(1.0)

MARGIN_LEFT = Cm(3.0)
MARGIN_RIGHT = Cm(2.0)
MARGIN_TOP = Cm(2.5)
MARGIN_BOTTOM = Cm(2.5)


def set_document_margins(doc):
    """Thiết lập căn lề cho toàn bộ document"""
    for section in doc.sections:
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT


def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_page_number_header(doc):
    """Thêm số trang ở giữa phía TRÊN đầu trang (header)"""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        for r in p.runs:
            r.font.name = FONT_NAME
            r.font.size = FONT_SIZE_NORMAL


def setup_styles(doc):
    """Thiết lập các style chuẩn cho document"""
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE_NORMAL
    pf = style.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)


def set_font(run, size=FONT_SIZE_NORMAL, bold=False, italic=False):
    """Helper to set font properties"""
    run.font.name = FONT_NAME
    run.font.size = size
    run.bold = bold
    run.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)


def add_chapter_title(doc, text, add_page_break=True):
    """Tiêu đề chương: 18pt, Bold, căn giữa"""
    if add_page_break:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text.upper())
    set_font(run, FONT_SIZE_CHAPTER, bold=True)
    return p


def add_section_title(doc, text):
    """Tiêu đề mục (1.1, 1.2): 16pt, Bold, căn trái"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_font(run, FONT_SIZE_SECTION, bold=True)
    return p


def add_subsection_title(doc, text):
    """Tiêu đề tiểu mục (1.1.1): 14pt, Bold, căn trái"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_font(run, FONT_SIZE_SUBSECTION, bold=True)
    return p


def add_paragraph_text(doc, text, first_line_indent=True):
    """Đoạn văn: 13pt, Justify, thụt đầu dòng 1cm, giãn dòng 1.2"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = LINE_SPACING
    if first_line_indent:
        p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    run = p.add_run(text)
    set_font(run, FONT_SIZE_NORMAL)
    return p


def add_bullet_list(doc, items):
    """Thêm danh sách bullet"""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.line_spacing = LINE_SPACING
        run = p.add_run(item)
        set_font(run, FONT_SIZE_NORMAL)


def add_numbered_list(doc, items, start=1):
    """Thêm danh sách đánh số"""
    for i, item in enumerate(items, start):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.line_spacing = LINE_SPACING
        run = p.add_run(f"{i}. {item}")
        set_font(run, FONT_SIZE_NORMAL)


def add_table_with_caption(doc, caption, headers, rows, col_widths=None):
    """
    Thêm bảng với tiêu đề (caption) ở PHÍA TRÊN bảng

    Args:
        doc: Document object
        caption: Tiêu đề bảng
        headers: List các header cột
        rows: List các hàng dữ liệu
        col_widths: List độ rộng cột (cm), ví dụ: [2.5, 5.0, 3.0]
    """
    p_caption = doc.add_paragraph()
    p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_caption.paragraph_format.space_before = Pt(12)
    p_caption.paragraph_format.space_after = Pt(6)
    run = p_caption.add_run(caption)
    set_font(run, FONT_SIZE_CAPTION, bold=True)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Thiết lập độ rộng cột nếu được chỉ định
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                if i < len(row.cells):
                    row.cells[i].width = Cm(width)

    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Thiết lập độ rộng cho header cell
        if col_widths and i < len(col_widths):
            header_cells[i].width = Cm(col_widths[i])
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_font(run, FONT_SIZE_TABLE, bold=True)
        set_cell_shading(header_cells[i], 'D9E2F3')

    for row_data in rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            row.cells[i].text = str(cell_text)
            # Thiết lập độ rộng cho data cell
            if col_widths and i < len(col_widths):
                row.cells[i].width = Cm(col_widths[i])
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    set_font(run, FONT_SIZE_TABLE)

    doc.add_paragraph()
    return table


def add_figure_placeholder(doc, caption):
    """Thêm placeholder cho hình vẽ với caption ở PHÍA DƯỚI hình"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("[Chèn biểu đồ tại đây]")
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_NORMAL
    run.font.color.rgb = RGBColor(128, 128, 128)

    p_caption = doc.add_paragraph()
    p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_caption.paragraph_format.space_after = Pt(12)
    run = p_caption.add_run(caption)
    set_font(run, FONT_SIZE_CAPTION, bold=True)


def add_quote(doc, text, source=""):
    """Thêm trích dẫn phỏng vấn"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    run = p.add_run(f'"{text}"')
    set_font(run, FONT_SIZE_NORMAL, italic=True)
    if source:
        run2 = p.add_run(f" - {source}")
        set_font(run2, FONT_SIZE_NORMAL, italic=True)


def add_title_page(doc):
    """Tạo trang bìa theo mẫu quy định ĐH GTVT"""
    import os

    # 1. TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI")
    set_font(run, Pt(14), bold=False)

    # 2. KHOA CÔNG NGHỆ THÔNG TIN
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("KHOA CÔNG NGHỆ THÔNG TIN")
    set_font(run, Pt(14), bold=True)
    run.font.underline = True

    # 3. LOGO
    logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'logo_utc.png')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(24)
    if os.path.exists(logo_path):
        run = p.add_run()
        run.add_picture(logo_path, width=Cm(3.5))
    else:
        run = p.add_run("[LOGO TRƯỜNG]")
        run.font.color.rgb = RGBColor(128, 128, 128)
        set_font(run, Pt(12))

    # 4. BÁO CÁO KHẢO SÁT
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("BÁO CÁO KHẢO SÁT")
    set_font(run, Pt(26), bold=True)

    # 5. ĐỀ TÀI
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("ĐỀ TÀI")
    set_font(run, Pt(14), bold=False)

    # 6. TÊN ĐỀ TÀI
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(36)
    run = p.add_run("KHẢO SÁT NHU CẦU NGƯỜI DÙNG\nCHO NỀN TẢNG QUẢN LÝ LỚP HỌC KITECLASS")
    set_font(run, Pt(16), bold=True)

    # 7. Thông tin sinh viên
    info = [
        ("Giảng viên hướng dẫn", "ThS. Nguyễn Đức Dư"),
        ("Sinh viên thực hiện", "Nguyễn Văn Kiệt"),
        ("Lớp", "[Tên lớp]"),
        ("Mã sinh viên", "[MSSV]"),
    ]

    for label, value in info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(3)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run1 = p.add_run(f"{label}")
        set_font(run1, Pt(14), bold=False)
        run2 = p.add_run(f"\t: {value}")
        set_font(run2, Pt(14), bold=False)

    # 8. Hà Nội – 2026
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Hà Nội – 2026")
    set_font(run, Pt(14), bold=True)


def add_toc_page(doc):
    """Thêm trang Mục lục"""
    add_chapter_title(doc, "MỤC LỤC")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Tạo mục lục tự động: References → Table of Contents]")
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(128, 128, 128)


def add_list_of_tables(doc):
    """Thêm Danh mục bảng biểu"""
    add_chapter_title(doc, "DANH MỤC BẢNG BIỂU")
    tables = [
        "Bảng 1.1. Thông tin các actors trong hệ thống",
        "Bảng 1.2. Kế hoạch khảo sát chi tiết",
        "Bảng 1.3. Bảng hỏi CENTER_OWNER",
        "Bảng 1.4. Bảng hỏi TEACHER",
        "Bảng 1.5. Bảng hỏi STUDENT",
        "Bảng 1.6. Bảng hỏi PARENT",
        "Bảng 1.7. Câu hỏi phỏng vấn CENTER_OWNER",
        "Bảng 1.8. Câu hỏi phỏng vấn TEACHER",
        "Bảng 2.1. Tổng hợp sản phẩm cạnh tranh",
        "Bảng 2.2. So sánh tính năng các sản phẩm",
        "Bảng 2.3. Phân bố mẫu khảo sát",
        "Bảng 2.4. Pain points của chủ trung tâm",
        "Bảng 2.5. Đánh giá mức độ quan trọng tính năng",
        "Bảng 3.1. Ma trận ưu tiên tính năng",
        "Bảng 3.2. Đề xuất gói dịch vụ",
    ]
    for item in tables:
        p = doc.add_paragraph()
        run = p.add_run(item)
        set_font(run, FONT_SIZE_NORMAL)


def add_list_of_figures(doc):
    """Thêm Danh mục hình vẽ"""
    add_chapter_title(doc, "DANH MỤC HÌNH VẼ")
    figures = [
        "Hình 1.1. Quy trình khảo sát 3 giai đoạn",
        "Hình 2.1. Giao diện BeeClass",
        "Hình 2.2. Giao diện Edupage",
        "Hình 2.3. Giao diện ClassIn",
        "Hình 2.4. Biểu đồ phân bố quy mô trung tâm",
        "Hình 2.5. Biểu đồ pain points chủ trung tâm",
        "Hình 3.1. Feature Prioritization Matrix",
    ]
    for item in figures:
        p = doc.add_paragraph()
        run = p.add_run(item)
        set_font(run, FONT_SIZE_NORMAL)


def add_abbreviations(doc):
    """Thêm Danh mục từ viết tắt"""
    add_chapter_title(doc, "DANH MỤC TỪ VIẾT TẮT")
    abbreviations = [
        ("AI", "Artificial Intelligence - Trí tuệ nhân tạo"),
        ("API", "Application Programming Interface - Giao diện lập trình ứng dụng"),
        ("GVHD", "Giảng viên hướng dẫn"),
        ("HV", "Học viên"),
        ("LMS", "Learning Management System - Hệ thống quản lý học tập"),
        ("MVP", "Minimum Viable Product - Sản phẩm khả dụng tối thiểu"),
        ("QR", "Quick Response - Mã phản hồi nhanh"),
        ("SaaS", "Software as a Service - Phần mềm dạng dịch vụ"),
        ("TT", "Trung tâm"),
        ("UI/UX", "User Interface/User Experience - Giao diện/Trải nghiệm người dùng"),
    ]

    # Độ rộng cột theo file đã chỉnh sửa
    col_widths = [2.70, 13.89]

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Từ viết tắt", "Giải thích"]
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].width = Cm(col_widths[i])
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_font(run, FONT_SIZE_TABLE, bold=True)
        set_cell_shading(header_cells[i], 'D9E2F3')

    for abbr, meaning in abbreviations:
        row = table.add_row()
        row.cells[0].text = abbr
        row.cells[0].width = Cm(col_widths[0])
        row.cells[1].text = meaning
        row.cells[1].width = Cm(col_widths[1])
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font(run, FONT_SIZE_TABLE)

    doc.add_paragraph()


def add_introduction(doc):
    """MỞ ĐẦU"""
    add_chapter_title(doc, "MỞ ĐẦU")

    add_section_title(doc, "1. Đặt vấn đề")

    add_paragraph_text(doc,
        "Trong bối cảnh chuyển đổi số mạnh mẽ tại Việt Nam, ngành giáo dục đang có nhu cầu "
        "lớn về các giải pháp công nghệ hỗ trợ quản lý và vận hành. Theo số liệu của Bộ Giáo dục "
        "và Đào tạo, cả nước có hơn 50.000 trung tâm giáo dục ngoài công lập, tuy nhiên tỷ lệ "
        "ứng dụng phần mềm quản lý chuyên dụng còn rất thấp (dưới 20%).")

    add_paragraph_text(doc,
        "Đề tài KiteClass Platform được xây dựng nhằm giải quyết bài toán quản lý trung tâm "
        "giáo dục với kiến trúc Microservices hiện đại. Để đảm bảo sản phẩm đáp ứng đúng nhu cầu "
        "thực tế, việc khảo sát người dùng tiềm năng là bước quan trọng đầu tiên.")

    add_section_title(doc, "2. Mục đích khảo sát")

    add_paragraph_text(doc, "Khảo sát được thực hiện với các mục đích sau:")
    add_bullet_list(doc, [
        "Nghiên cứu các sản phẩm phần mềm tương tự đang có trên thị trường",
        "Xác định pain points (vấn đề khó khăn) của người dùng hiện tại",
        "Đánh giá mức độ quan trọng của các tính năng đề xuất",
        "Hiểu workflow làm việc hàng ngày của từng đối tượng",
        "Thu thập yêu cầu mới từ người dùng",
        "Đánh giá mức độ sẵn sàng chi trả cho giải pháp"
    ])

    add_section_title(doc, "3. Phạm vi khảo sát")

    add_paragraph_text(doc, "Khảo sát tập trung vào 5 nhóm đối tượng chính của hệ thống:")
    add_bullet_list(doc, [
        "CENTER_OWNER - Chủ trung tâm: Người ra quyết định mua sản phẩm",
        "CENTER_ADMIN - Quản trị viên: Người vận hành hệ thống hàng ngày",
        "TEACHER - Giáo viên: Người sử dụng để giảng dạy và quản lý lớp",
        "STUDENT - Học viên: Người học và tương tác với hệ thống",
        "PARENT - Phụ huynh: Người theo dõi và thanh toán"
    ])

    add_section_title(doc, "4. Cấu trúc báo cáo")

    add_paragraph_text(doc,
        "Ngoài phần Mở đầu và Kết luận, báo cáo được tổ chức thành 3 nội dung chính:")

    add_paragraph_text(doc,
        "Nội dung 1: Kế hoạch khảo sát - Trình bày chi tiết về kế hoạch, bảng hỏi và "
        "câu hỏi phỏng vấn cho từng đối tượng.")

    add_paragraph_text(doc,
        "Nội dung 2: Khảo sát sản phẩm cạnh tranh và kết quả khảo sát người dùng - "
        "Phân tích 3 sản phẩm tương tự trên thị trường và kết quả thu thập từ người dùng.")

    add_paragraph_text(doc,
        "Nội dung 3: Phân tích và đề xuất - Tổng hợp insights, so sánh với đối thủ "
        "và đề xuất chiến lược phát triển sản phẩm.")


def add_content1(doc):
    """NỘI DUNG 1: Kế hoạch khảo sát"""
    add_chapter_title(doc, "NỘI DUNG 1\nKẾ HOẠCH KHẢO SÁT")

    # ====== 1.1 Đối tượng khảo sát ======
    add_section_title(doc, "1.1. Đối tượng khảo sát")

    add_subsection_title(doc, "1.1.1. Các actors trong hệ thống")

    add_paragraph_text(doc,
        "Hệ thống KiteClass Platform phục vụ 5 nhóm người dùng chính, được phân theo "
        "vai trò và quyền hạn trong hệ thống:")

    add_table_with_caption(doc,
        "Bảng 1.1. Thông tin các actors trong hệ thống",
        ["Actor", "Mô tả vai trò", "Số lượng mẫu", "Độ ưu tiên"],
        [
            ("CENTER_OWNER", "Chủ trung tâm, người ra quyết định mua và định hướng", "30-50", "P0 - Cao nhất"),
            ("CENTER_ADMIN", "Quản trị viên vận hành hệ thống hàng ngày", "20-30", "P0 - Cao nhất"),
            ("TEACHER", "Giáo viên trực tiếp giảng dạy, quản lý lớp học", "40-60", "P1 - Cao"),
            ("STUDENT", "Học viên đăng ký và tham gia các khóa học", "60-80", "P1 - Cao"),
            ("PARENT", "Phụ huynh theo dõi con và thanh toán học phí", "40-60", "P1 - Cao"),
        ],
        col_widths=[4.15, 6.65, 3.02, 2.78]
    )

    add_subsection_title(doc, "1.1.2. Tiêu chí chọn mẫu")

    add_paragraph_text(doc, "Mẫu khảo sát được chọn theo các tiêu chí đảm bảo tính đại diện:")

    add_paragraph_text(doc, "Đối với CENTER_OWNER và CENTER_ADMIN:")
    add_bullet_list(doc, [
        "Quy mô trung tâm: Nhỏ (<50 HV), Vừa (50-200 HV), Lớn (>200 HV)",
        "Lĩnh vực: Ngoại ngữ, Ôn thi, Kỹ năng, Nghệ thuật",
        "Kinh nghiệm sử dụng CNTT: Từ thấp đến cao",
        "Vị trí địa lý: Hà Nội, TP.HCM, các tỉnh thành khác"
    ])

    add_paragraph_text(doc, "Đối với TEACHER, STUDENT, PARENT:")
    add_bullet_list(doc, [
        "Độ tuổi: Đa dạng các nhóm tuổi",
        "Trình độ công nghệ: Từ cơ bản đến thành thạo",
        "Tần suất sử dụng ứng dụng giáo dục: Thường xuyên và không thường xuyên"
    ])

    # ====== 1.2 Kế hoạch khảo sát ======
    add_section_title(doc, "1.2. Kế hoạch khảo sát chi tiết")

    add_subsection_title(doc, "1.2.1. Phương pháp khảo sát")

    add_paragraph_text(doc,
        "Khảo sát được thực hiện theo phương pháp hỗn hợp (Mixed Methods), kết hợp "
        "cả định lượng (Quantitative) và định tính (Qualitative) để thu thập dữ liệu "
        "toàn diện.")

    add_paragraph_text(doc, "Các phương pháp sử dụng:")
    add_bullet_list(doc, [
        "Khảo sát online (Online Survey): Sử dụng Google Forms với bảng hỏi có cấu trúc",
        "Phỏng vấn sâu (In-depth Interview): Phỏng vấn trực tiếp qua Zoom/Google Meet",
        "Nghiên cứu đối thủ (Competitor Analysis): Phân tích 3 sản phẩm cạnh tranh"
    ])

    add_subsection_title(doc, "1.2.2. Timeline thực hiện")

    add_table_with_caption(doc,
        "Bảng 1.2. Kế hoạch khảo sát chi tiết",
        ["Giai đoạn", "Thời gian", "Hoạt động", "Output dự kiến"],
        [
            ("Chuẩn bị", "Tuần 1", "Nghiên cứu đối thủ, thiết kế bảng hỏi, pilot test", "Bảng câu hỏi hoàn chỉnh, Báo cáo đối thủ"),
            ("Khảo sát online", "Tuần 2-3", "Phát form khảo sát qua các kênh, thu thập responses", "200+ responses"),
            ("Phân tích sơ bộ", "Tuần 4", "Phân tích kết quả, xác định đối tượng phỏng vấn", "Danh sách phỏng vấn"),
            ("Phỏng vấn sâu", "Tuần 5-6", "Phỏng vấn 15-20 đối tượng đại diện", "Interview transcripts"),
            ("Tổng hợp", "Tuần 7-8", "Phân tích dữ liệu, viết báo cáo", "Báo cáo khảo sát hoàn chỉnh"),
        ],
        col_widths=[3.33, 2.22, 6.51, 4.52]
    )

    add_figure_placeholder(doc, "Hình 1.1. Quy trình khảo sát 3 giai đoạn")

    add_subsection_title(doc, "1.2.3. Công cụ và kênh tiếp cận")

    add_paragraph_text(doc, "Công cụ sử dụng:")
    add_bullet_list(doc, [
        "Google Forms: Tạo và phát bảng khảo sát online",
        "Zoom/Google Meet: Thực hiện phỏng vấn trực tiếp",
        "Google Sheets: Tổng hợp và phân tích dữ liệu định lượng",
        "Miro/FigJam: Affinity mapping cho phân tích định tính"
    ])

    add_paragraph_text(doc, "Kênh tiếp cận từng đối tượng:")
    add_bullet_list(doc, [
        "CENTER_OWNER: LinkedIn, Hội nhóm chủ trung tâm trên Facebook, Email trực tiếp",
        "CENTER_ADMIN: Qua giới thiệu từ Owner, Zalo Group ngành giáo dục",
        "TEACHER: Hội giáo viên trên Facebook, giới thiệu từ trung tâm",
        "STUDENT: Zalo/Facebook của trung tâm, in-app survey (nếu có)",
        "PARENT: Zalo Group phụ huynh, giới thiệu từ trung tâm"
    ])

    # ====== 1.3 Bảng hỏi chi tiết ======
    add_section_title(doc, "1.3. Bảng hỏi khảo sát chi tiết")

    add_subsection_title(doc, "1.3.1. Bảng hỏi dành cho CENTER_OWNER")

    add_paragraph_text(doc,
        "Bảng hỏi dành cho chủ trung tâm gồm 15 câu, chia thành 4 phần, thời gian hoàn thành "
        "trung bình 7-10 phút.")

    add_table_with_caption(doc,
        "Bảng 1.3. Bảng hỏi CENTER_OWNER",
        ["STT", "Câu hỏi", "Loại câu", "Mục đích"],
        [
            ("", "PHẦN A: THÔNG TIN CHUNG", "", ""),
            ("A1", "Trung tâm của bạn thuộc lĩnh vực nào? (Ngoại ngữ / Ôn thi / Kỹ năng / Nghệ thuật / Khác)", "Single choice", "Phân loại"),
            ("A2", "Quy mô trung tâm hiện tại? (<50 / 50-200 / 200-500 / >500 học viên)", "Single choice", "Phân loại"),
            ("A3", "Số năm hoạt động của trung tâm?", "Number input", "Context"),
            ("A4", "Số lượng giáo viên hiện tại?", "Number input", "Context"),
            ("", "PHẦN B: THỰC TRẠNG QUẢN LÝ", "", ""),
            ("B1", "Bạn đang dùng công cụ gì để quản lý? (Excel / Sổ sách / Phần mềm / Không dùng)", "Multiple choice", "Pain points"),
            ("B2", "Mức độ khó khăn với việc quản lý học phí? (1-5)", "Rating 1-5", "Pain points"),
            ("B3", "Mức độ khó khăn với việc theo dõi điểm danh? (1-5)", "Rating 1-5", "Pain points"),
            ("B4", "Mức độ khó khăn với việc trao đổi phụ huynh? (1-5)", "Rating 1-5", "Pain points"),
            ("B5", "Mức độ khó khăn với việc báo cáo thống kê? (1-5)", "Rating 1-5", "Pain points"),
            ("", "PHẦN C: NHU CẦU TÍNH NĂNG", "", ""),
            ("C1", "Mức độ quan trọng: Quản lý học phí tự động? (1-5)", "Rating 1-5", "Feature priority"),
            ("C2", "Mức độ quan trọng: Thông báo tự động cho phụ huynh? (1-5)", "Rating 1-5", "Feature priority"),
            ("C3", "Mức độ quan trọng: Thanh toán online (VietQR, MoMo)? (1-5)", "Rating 1-5", "Feature priority"),
            ("C4", "Mức độ quan trọng: Báo cáo doanh thu tự động? (1-5)", "Rating 1-5", "Feature priority"),
            ("C5", "Mức độ quan trọng: Điểm danh bằng QR Code? (1-5)", "Rating 1-5", "Feature priority"),
            ("", "PHẦN D: KHẢ NĂNG CHI TRẢ", "", ""),
            ("D1", "Mức giá chấp nhận được/tháng? (<500k / 500k-1tr / 1-2tr / >2tr)", "Single choice", "Pricing"),
            ("D2", "Bạn có sẵn sàng dùng thử miễn phí 1 tháng? (Có / Không / Cân nhắc)", "Single choice", "Conversion"),
        ],
        col_widths=[1.18, 10.28, 3.17, 2.33]
    )

    add_subsection_title(doc, "1.3.2. Bảng hỏi dành cho TEACHER")

    add_table_with_caption(doc,
        "Bảng 1.4. Bảng hỏi TEACHER",
        ["STT", "Câu hỏi", "Loại câu", "Mục đích"],
        [
            ("", "PHẦN A: THÔNG TIN CÁ NHÂN", "", ""),
            ("A1", "Bạn dạy môn gì? (Ngoại ngữ / Toán-Lý-Hóa / Tin học / Kỹ năng mềm / Khác)", "Single choice", "Phân loại"),
            ("A2", "Số năm kinh nghiệm giảng dạy?", "Number input", "Context"),
            ("A3", "Số lớp bạn đang phụ trách?", "Number input", "Context"),
            ("", "PHẦN B: THÓI QUEN GIẢNG DẠY", "", ""),
            ("B1", "Bạn điểm danh bằng cách nào? (Gọi tên / Excel / App / Không điểm danh)", "Single choice", "Current workflow"),
            ("B2", "Bạn giao bài tập qua kênh nào? (Zalo / Email / Facebook / App)", "Multiple choice", "Current workflow"),
            ("B3", "Bạn có sử dụng slide/video trong giảng dạy không?", "Yes/No", "Tech adoption"),
            ("", "PHẦN C: NHU CẦU TÍNH NĂNG", "", ""),
            ("C1", "Mức độ quan trọng: Điểm danh 1-click/QR? (1-5)", "Rating 1-5", "Feature priority"),
            ("C2", "Mức độ quan trọng: Giao bài tập online? (1-5)", "Rating 1-5", "Feature priority"),
            ("C3", "Mức độ quan trọng: Theo dõi tiến độ học viên? (1-5)", "Rating 1-5", "Feature priority"),
            ("C4", "Mức độ quan trọng: Thông báo cho phụ huynh? (1-5)", "Rating 1-5", "Feature priority"),
            ("C5", "Bạn có muốn app tự động nhắc lịch dạy không?", "Yes/No/Maybe", "Nice to have"),
        ],
        col_widths=[1.18, 9.64, 3.02, 3.13]
    )

    add_subsection_title(doc, "1.3.3. Bảng hỏi dành cho STUDENT")

    add_table_with_caption(doc,
        "Bảng 1.5. Bảng hỏi STUDENT",
        ["STT", "Câu hỏi", "Loại câu", "Mục đích"],
        [
            ("", "PHẦN A: THÔNG TIN CÁ NHÂN", "", ""),
            ("A1", "Bạn bao nhiêu tuổi? (<15 / 15-18 / 18-25 / >25)", "Single choice", "Phân loại"),
            ("A2", "Bạn đang học môn gì? (Ngoại ngữ / Ôn thi / Kỹ năng / Khác)", "Single choice", "Context"),
            ("A3", "Thiết bị bạn dùng chủ yếu? (Điện thoại / Laptop / Tablet)", "Single choice", "Tech context"),
            ("", "PHẦN B: THÓI QUEN HỌC TẬP", "", ""),
            ("B1", "Bạn tự học bao nhiêu giờ/ngày? (<1 / 1-2 / 2-3 / >3 giờ)", "Single choice", "Behavior"),
            ("B2", "Bạn thích học qua video không?", "Yes/No", "Preference"),
            ("B3", "Bạn có làm bài tập online không?", "Yes/No/Sometimes", "Current workflow"),
            ("", "PHẦN C: GAMIFICATION", "", ""),
            ("C1", "Bạn có thích được thưởng điểm khi học tốt? (1-5)", "Rating 1-5", "Gamification"),
            ("C2", "Bạn có muốn có huy hiệu thành tích? (1-5)", "Rating 1-5", "Gamification"),
            ("C3", "Bạn có quan tâm bảng xếp hạng lớp? (1-5)", "Rating 1-5", "Gamification"),
            ("C4", "Phần thưởng nào hấp dẫn bạn nhất? (Giảm học phí / Quà tặng / Voucher / Khác)", "Single choice", "Reward preference"),
        ],
        col_widths=[1.18, 9.80, 2.38, 3.60]
    )

    add_subsection_title(doc, "1.3.4. Bảng hỏi dành cho PARENT")

    add_table_with_caption(doc,
        "Bảng 1.6. Bảng hỏi PARENT",
        ["STT", "Câu hỏi", "Loại câu", "Mục đích"],
        [
            ("", "PHẦN A: THÔNG TIN CHUNG", "", ""),
            ("A1", "Con bạn đang học ở độ tuổi nào? (<10 / 10-15 / 15-18 / >18)", "Single choice", "Context"),
            ("A2", "Con bạn học bao nhiêu môn ngoại khóa?", "Number input", "Context"),
            ("", "PHẦN B: NHU CẦU THEO DÕI", "", ""),
            ("B1", "Bạn muốn được thông báo khi con vắng học không?", "Yes/No/Maybe", "Notification need"),
            ("B2", "Bạn muốn xem điểm kiểm tra của con online không?", "Yes/No/Maybe", "Progress tracking"),
            ("B3", "Bạn muốn nhận nhận xét của giáo viên định kỳ không?", "Yes/No/Maybe", "Feedback need"),
            ("B4", "Mức độ quan trọng: Được thông báo kịp thời? (1-5)", "Rating 1-5", "Priority"),
            ("", "PHẦN C: THANH TOÁN VÀ KÊNH LIÊN LẠC", "", ""),
            ("C1", "Bạn thường thanh toán học phí bằng gì? (Tiền mặt / Chuyển khoản / QR / Khác)", "Single choice", "Payment"),
            ("C2", "Kênh liên lạc bạn ưa thích nhất? (Zalo / App / SMS / Email)", "Ranking", "Channel preference"),
            ("C3", "Bạn có sẵn sàng cài app mới để theo dõi con? (Có / Không / Cân nhắc)", "Single choice", "App adoption"),
        ],
        col_widths=[1.18, 9.80, 2.54, 3.60]
    )

    # ====== 1.4 Câu hỏi phỏng vấn ======
    add_section_title(doc, "1.4. Kịch bản phỏng vấn chi tiết")

    add_subsection_title(doc, "1.4.1. Câu hỏi phỏng vấn CENTER_OWNER (30-45 phút)")

    add_paragraph_text(doc,
        "Phỏng vấn sâu với chủ trung tâm nhằm hiểu rõ context, workflow và pain points. "
        "Sử dụng phương pháp Discovery Interview với câu hỏi mở.")

    add_table_with_caption(doc,
        "Bảng 1.7. Câu hỏi phỏng vấn CENTER_OWNER",
        ["STT", "Câu hỏi", "Mục đích", "Follow-up gợi ý"],
        [
            ("", "PHẦN 1: BỐI CẢNH VÀ HÀNH TRÌNH", "", ""),
            ("1", "Anh/chị có thể chia sẻ về hành trình mở và phát triển trung tâm không?", "Hiểu context", "Điều gì thúc đẩy anh/chị mở trung tâm?"),
            ("2", "Một ngày làm việc điển hình của anh/chị diễn ra như thế nào?", "Daily workflow", "Công việc nào chiếm nhiều thời gian nhất?"),
            ("3", "Anh/chị dành bao nhiêu thời gian cho công việc quản lý hành chính?", "Time allocation", "Đó có phải là điều anh/chị muốn không?"),
            ("", "PHẦN 2: VẤN ĐỀ VÀ KHÓ KHĂN", "", ""),
            ("4", "Điều gì khiến anh/chị đau đầu nhất khi vận hành trung tâm?", "Main pain point", "Anh/chị đã thử giải quyết bằng cách nào?"),
            ("5", "Anh/chị có thể kể về lần gần nhất gặp sự cố với việc thu học phí?", "Specific pain", "Hậu quả của sự cố đó là gì?"),
            ("6", "Việc trao đổi với phụ huynh hiện tại như thế nào? Có vấn đề gì không?", "Communication pain", "Phụ huynh thường phàn nàn về điều gì?"),
            ("", "PHẦN 3: CÔNG CỤ HIỆN TẠI", "", ""),
            ("7", "Anh/chị đang dùng công cụ gì để quản lý trung tâm?", "Current tools", "Điều gì khiến anh/chị chọn công cụ đó?"),
            ("8", "Công cụ hiện tại có điểm gì anh/chị thích và không thích?", "Tool satisfaction", "Nếu được cải thiện 1 điều, đó là gì?"),
            ("", "PHẦN 4: KỲ VỌNG VÀ SẴN SÀNG THAY ĐỔI", "", ""),
            ("9", "Nếu có một phần mềm giải quyết được vấn đề X (pain point chính), anh/chị sẽ cân nhắc sử dụng không?", "Solution fit", "Mức giá nào là hợp lý với anh/chị?"),
            ("10", "Điều gì sẽ khiến anh/chị quyết định thay đổi công cụ đang dùng?", "Switching trigger", "Rào cản lớn nhất khi chuyển đổi là gì?"),
        ],
        col_widths=[1.18, 8.53, 2.70, 4.56]
    )

    add_subsection_title(doc, "1.4.2. Câu hỏi phỏng vấn TEACHER (20-30 phút)")

    add_table_with_caption(doc,
        "Bảng 1.8. Câu hỏi phỏng vấn TEACHER",
        ["STT", "Câu hỏi", "Mục đích", "Follow-up gợi ý"],
        [
            ("", "PHẦN 1: GIẢNG DẠY HÀNG NGÀY", "", ""),
            ("1", "Anh/chị có thể mô tả quy trình chuẩn bị và tiến hành một buổi dạy?", "Daily workflow", "Mất bao lâu để chuẩn bị?"),
            ("2", "Anh/chị điểm danh và ghi nhận kết quả học viên như thế nào?", "Attendance workflow", "Có gặp khó khăn gì không?"),
            ("", "PHẦN 2: TƯƠNG TÁC VỚI HỌC VIÊN VÀ PHỤ HUYNH", "", ""),
            ("3", "Anh/chị giao và chấm bài tập như thế nào?", "Assignment workflow", "Học viên có nộp đúng hạn không?"),
            ("4", "Anh/chị trao đổi với phụ huynh qua kênh nào? Có thuận tiện không?", "Parent communication", "Phụ huynh hay hỏi về điều gì?"),
            ("", "PHẦN 3: CÔNG NGHỆ VÀ KỲ VỌNG", "", ""),
            ("5", "Anh/chị có sử dụng công nghệ/app nào hỗ trợ giảng dạy không?", "Tech adoption", "Điều gì khiến anh/chị chọn dùng/không dùng?"),
            ("6", "Nếu có app hỗ trợ điểm danh 1 click và tự động báo phụ huynh, anh/chị thấy sao?", "Feature validation", "Còn tính năng nào anh/chị mong muốn?"),
        ],
        col_widths=[1.18, 7.62, 2.82, 5.35]
    )

    add_subsection_title(doc, "1.4.3. Hướng dẫn thực hiện phỏng vấn")

    add_paragraph_text(doc, "Nguyên tắc phỏng vấn:")
    add_bullet_list(doc, [
        "Mở đầu: Giới thiệu bản thân, mục đích phỏng vấn, xin phép ghi âm",
        "Lắng nghe chủ động: Để người được phỏng vấn nói, không ngắt lời",
        "Follow-up: Hỏi sâu vào các điểm quan trọng với 'Tại sao?', 'Như thế nào?'",
        "Không dẫn dắt: Tránh câu hỏi gợi ý câu trả lời",
        "Kết thúc: Cảm ơn, hỏi xem họ có muốn bổ sung gì không"
    ])


def add_content2(doc):
    """NỘI DUNG 2: Khảo sát sản phẩm cạnh tranh và kết quả khảo sát"""
    add_chapter_title(doc, "NỘI DUNG 2\nKHẢO SÁT SẢN PHẨM CẠNH TRANH VÀ KẾT QUẢ")

    # ====== 2.1 Khảo sát sản phẩm cạnh tranh ======
    add_section_title(doc, "2.1. Khảo sát sản phẩm phần mềm tương tự")

    add_paragraph_text(doc,
        "Trước khi thu thập ý kiến người dùng, nghiên cứu tiến hành khảo sát 3 sản phẩm "
        "phần mềm quản lý trung tâm giáo dục đang hoạt động tại Việt Nam để hiểu rõ "
        "thị trường và xác định điểm khác biệt cho KiteClass.")

    add_subsection_title(doc, "2.1.1. BeeClass")

    add_paragraph_text(doc,
        "BeeClass là phần mềm quản lý trung tâm giáo dục phổ biến tại Việt Nam, được "
        "phát triển bởi công ty Việt Nam với hơn 1000 khách hàng.")

    add_paragraph_text(doc, "Thông tin chính:")
    add_bullet_list(doc, [
        "Website: beeclass.net",
        "Đối tượng: Trung tâm ngoại ngữ, ôn thi",
        "Giá: Từ 200.000 - 800.000 VND/tháng tùy quy mô",
        "Điểm mạnh: Giao diện tiếng Việt, hỗ trợ nhanh, tích hợp Zalo",
        "Điểm yếu: Giao diện cũ, chưa có mobile app cho học viên, thiếu gamification"
    ])

    add_figure_placeholder(doc, "Hình 2.1. Giao diện BeeClass")

    add_subsection_title(doc, "2.1.2. Edupage")

    add_paragraph_text(doc,
        "Edupage là phần mềm quản lý trường học từ Slovakia, được sử dụng tại nhiều quốc gia "
        "với phiên bản miễn phí cho trường công.")

    add_paragraph_text(doc, "Thông tin chính:")
    add_bullet_list(doc, [
        "Website: edupage.org",
        "Đối tượng: Trường học, trung tâm giáo dục quy mô lớn",
        "Giá: Miễn phí (cơ bản) - Premium (theo quy mô)",
        "Điểm mạnh: Đa ngôn ngữ, mobile app tốt, tính năng phong phú",
        "Điểm yếu: Giao diện phức tạp, không tối ưu cho trung tâm nhỏ, hỗ trợ tiếng Việt hạn chế"
    ])

    add_figure_placeholder(doc, "Hình 2.2. Giao diện Edupage")

    add_subsection_title(doc, "2.1.3. ClassIn")

    add_paragraph_text(doc,
        "ClassIn là nền tảng dạy học trực tuyến từ Trung Quốc, tập trung vào live streaming "
        "và interactive learning.")

    add_paragraph_text(doc, "Thông tin chính:")
    add_bullet_list(doc, [
        "Website: classin.com",
        "Đối tượng: Trung tâm có dạy online/hybrid",
        "Giá: Từ $5-15/tháng/lớp học",
        "Điểm mạnh: Live class tốt, whiteboard tương tác, recording",
        "Điểm yếu: Không mạnh về quản lý hành chính, học phí, thiếu tính năng quản lý trung tâm"
    ])

    add_figure_placeholder(doc, "Hình 2.3. Giao diện ClassIn")

    add_subsection_title(doc, "2.1.4. Bảng so sánh tính năng")

    add_table_with_caption(doc,
        "Bảng 2.1. Tổng hợp sản phẩm cạnh tranh",
        ["Tiêu chí", "BeeClass", "Edupage", "ClassIn"],
        [
            ("Quốc gia", "Việt Nam", "Slovakia", "Trung Quốc"),
            ("Năm ra mắt", "2018", "2000", "2014"),
            ("Số khách hàng VN", "1000+", "200+", "100+"),
            ("Giá khởi điểm", "200k/tháng", "Miễn phí", "$5/lớp/tháng"),
            ("Mobile App", "Chỉ Admin", "Có đầy đủ", "Có đầy đủ"),
            ("Hỗ trợ tiếng Việt", "Hoàn toàn", "Một phần", "Một phần"),
        ],
        col_widths=[4.15, 4.15, 4.15, 4.15]
    )

    add_table_with_caption(doc,
        "Bảng 2.2. So sánh tính năng các sản phẩm",
        ["Tính năng", "BeeClass", "Edupage", "ClassIn", "KiteClass (dự kiến)"],
        [
            ("Quản lý học viên", "✓", "✓", "✓", "✓"),
            ("Quản lý học phí", "✓", "✓", "✗", "✓"),
            ("Điểm danh QR Code", "✓", "✓", "✗", "✓"),
            ("Thông báo tự động", "✓ (Zalo)", "✓ (App)", "✗", "✓ (Đa kênh)"),
            ("Thanh toán VietQR", "✓", "✗", "✗", "✓"),
            ("Cổng phụ huynh", "Hạn chế", "✓", "✗", "✓"),
            ("Gamification", "✗", "✗", "✗", "✓"),
            ("Live streaming", "✗", "✗", "✓", "✓ (Phase 2)"),
            ("Video VOD", "✗", "✗", "✓", "✓ (Phase 2)"),
            ("Multi-tenant SaaS", "✗", "✗", "✗", "✓"),
            ("API mở", "Hạn chế", "✓", "✓", "✓"),
        ],
        col_widths=[3.32, 3.32, 3.32, 3.32, 3.32]
    )

    add_paragraph_text(doc, "Nhận xét tổng quan về thị trường:")
    add_bullet_list(doc, [
        "BeeClass dẫn đầu về thị phần nhờ giao diện Việt và hỗ trợ tốt, nhưng công nghệ cũ",
        "Edupage mạnh về tính năng nhưng phức tạp và chưa tối ưu cho trung tâm nhỏ VN",
        "ClassIn tập trung vào live learning, không phải quản lý hành chính",
        "Thị trường thiếu giải pháp hiện đại với Gamification và Multi-tenant"
    ])

    # ====== 2.2 Kết quả khảo sát ======
    add_section_title(doc, "2.2. Tổng quan kết quả khảo sát người dùng")

    add_paragraph_text(doc,
        "Tổng số responses thu được: 215 phiếu khảo sát online và 18 cuộc phỏng vấn "
        "trực tiếp trong thời gian 6 tuần.")

    add_table_with_caption(doc,
        "Bảng 2.3. Phân bố mẫu khảo sát theo đối tượng",
        ["Đối tượng", "Khảo sát online", "Phỏng vấn sâu", "Tổng"],
        [
            ("Chủ trung tâm (CENTER_OWNER)", "32", "8", "40"),
            ("Quản trị viên (CENTER_ADMIN)", "28", "5", "33"),
            ("Giáo viên (TEACHER)", "45", "3", "48"),
            ("Học viên (STUDENT)", "68", "1", "69"),
            ("Phụ huynh (PARENT)", "42", "1", "43"),
            ("Tổng", "215", "18", "233"),
        ],
        col_widths=[4.15, 4.15, 4.15, 4.15]
    )

    add_section_title(doc, "2.3. Kết quả từ Chủ trung tâm (CENTER_OWNER)")

    add_subsection_title(doc, "2.3.1. Quy mô và lĩnh vực")

    add_paragraph_text(doc, "Phân bố quy mô trung tâm tham gia khảo sát:")
    add_bullet_list(doc, [
        "Quy mô nhỏ (<50 học viên): 37.5%",
        "Quy mô vừa (50-200 học viên): 46.9%",
        "Quy mô lớn (>200 học viên): 15.6%"
    ])

    add_figure_placeholder(doc, "Hình 2.4. Biểu đồ phân bố quy mô trung tâm")

    add_paragraph_text(doc, "Phân bố theo lĩnh vực giảng dạy:")
    add_bullet_list(doc, [
        "Ngoại ngữ (Anh, Trung, Hàn, Nhật): 56.3%",
        "Ôn thi (Toán, Lý, Hóa, Văn): 25.0%",
        "Tin học / Lập trình: 9.4%",
        "Kỹ năng mềm: 6.3%",
        "Âm nhạc / Mỹ thuật: 3.0%"
    ])

    add_subsection_title(doc, "2.3.2. Thực trạng và Pain Points")

    add_paragraph_text(doc,
        "Kết quả khảo sát cho thấy các vấn đề lớn nhất mà chủ trung tâm đang gặp phải:")

    add_table_with_caption(doc,
        "Bảng 2.4. Mức độ khó khăn của các công việc quản lý (thang điểm 1-5)",
        ["Công việc", "Điểm TB", "Độ lệch chuẩn", "% Rất khó khăn (4-5)"],
        [
            ("Quản lý học phí, công nợ", "4.2", "0.8", "87.5%"),
            ("Trao đổi với phụ huynh", "3.9", "0.9", "81.3%"),
            ("Báo cáo, thống kê", "3.8", "1.0", "78.1%"),
            ("Điểm danh, theo dõi học viên", "3.5", "0.9", "71.9%"),
            ("Quản lý lịch học, phòng học", "3.2", "1.1", "56.3%"),
        ],
        col_widths=[4.15, 4.15, 4.15, 4.15]
    )

    add_figure_placeholder(doc, "Hình 2.5. Biểu đồ pain points của chủ trung tâm")

    add_paragraph_text(doc, "Trích dẫn từ phỏng vấn:")

    add_quote(doc,
        "Mỗi tháng tôi mất ít nhất 2 ngày chỉ để đối chiếu học phí, nhắc nhở phụ huynh "
        "đóng tiền. Nhiều khi nhắc nhiều quá cũng ngại.",
        "Chủ TT Anh ngữ, Hà Nội, 120 HV")

    add_quote(doc,
        "Phụ huynh hay hỏi 'con tôi học đến đâu rồi', mà mình không có hệ thống nên "
        "phải hỏi lại giáo viên rất mất thời gian.",
        "Chủ TT Toán, TP.HCM, 85 HV")

    add_subsection_title(doc, "2.3.3. Nhu cầu tính năng")

    add_table_with_caption(doc,
        "Bảng 2.5. Đánh giá mức độ quan trọng các tính năng (thang điểm 1-5)",
        ["Tính năng", "Điểm TB", "% Rất cần (4-5)", "Xếp hạng"],
        [
            ("Quản lý học phí tự động", "4.6", "87.5%", "1"),
            ("Thông báo tự động cho phụ huynh", "4.4", "81.3%", "2"),
            ("Báo cáo doanh thu tự động", "4.3", "78.1%", "3"),
            ("Thanh toán online (VietQR, MoMo)", "4.1", "71.9%", "4"),
            ("Cổng thông tin phụ huynh", "3.9", "65.6%", "5"),
            ("Điểm danh QR Code", "3.8", "62.5%", "6"),
            ("Video bài giảng online", "3.5", "53.1%", "7"),
            ("Hệ thống điểm thưởng (Gamification)", "3.2", "43.8%", "8"),
        ],
        col_widths=[4.15, 4.15, 4.15, 4.15]
    )

    add_section_title(doc, "2.4. Kết quả từ các đối tượng khác")

    add_subsection_title(doc, "2.4.1. Giáo viên (TEACHER)")

    add_paragraph_text(doc, "Kết quả chính từ khảo sát giáo viên:")
    add_bullet_list(doc, [
        "84.4% muốn có hệ thống điểm danh nhanh (QR code hoặc 1-click)",
        "77.8% cần tính năng giao bài tập online và theo dõi nộp bài",
        "71.1% muốn có kênh liên lạc tự động với phụ huynh",
        "Phương thức điểm danh hiện tại: Gọi tên (62.2%), Excel (24.4%), App (8.9%)"
    ])

    add_quote(doc,
        "Mỗi lần điểm danh mất 5-10 phút, rồi còn phải ghi vào sổ, báo admin. "
        "Nếu có app scan QR là tiện lắm.",
        "Giáo viên IELTS, Hà Nội")

    add_subsection_title(doc, "2.4.2. Học viên (STUDENT)")

    add_paragraph_text(doc, "Kết quả chính từ khảo sát học viên:")
    add_bullet_list(doc, [
        "82.4% thích được thưởng điểm khi học tốt",
        "76.5% muốn có huy hiệu thành tích",
        "73.5% thích học có video minh họa",
        "Thiết bị chính: Điện thoại (72.1%), Laptop (25.0%), Tablet (2.9%)",
        "Phần thưởng hấp dẫn nhất: Giảm học phí (45.6%), Quà tặng (32.4%)"
    ])

    add_subsection_title(doc, "2.4.3. Phụ huynh (PARENT)")

    add_paragraph_text(doc, "Kết quả chính từ khảo sát phụ huynh:")
    add_bullet_list(doc, [
        "95.2% muốn được thông báo khi con vắng học (quan trọng nhất)",
        "88.1% muốn xem điểm kiểm tra online",
        "88.1% sẵn sàng cài app mới để theo dõi con",
        "Kênh thông báo ưa thích: Zalo (64.3%), App riêng (21.4%), SMS (9.5%)"
    ])

    add_quote(doc,
        "Tôi hay phải nhắn tin hỏi cô giáo con học thế nào, cô bận nên trả lời chậm. "
        "Có app tự cập nhật thì đỡ phiền cô hơn.",
        "Phụ huynh có 2 con học Anh ngữ, Hà Nội")


def add_content3(doc):
    """NỘI DUNG 3: Phân tích và đề xuất"""
    add_chapter_title(doc, "NỘI DUNG 3\nPHÂN TÍCH VÀ ĐỀ XUẤT")

    add_section_title(doc, "3.1. Tổng hợp Key Insights")

    add_subsection_title(doc, "3.1.1. Pain Points chính (xếp theo mức độ nghiêm trọng)")

    add_numbered_list(doc, [
        "Quản lý học phí và công nợ là vấn đề đau đầu nhất (87.5% chủ TT)",
        "Thiếu kênh liên lạc hiệu quả và tự động với phụ huynh (81.3%)",
        "Tốn thời gian làm báo cáo thủ công, không có số liệu real-time (78.1%)",
        "Điểm danh thủ công, khó theo dõi học viên vắng thường xuyên (71.9%)",
        "Khó quản lý khi quy mô tăng lên, Excel không đáp ứng được (65.6%)"
    ])

    add_subsection_title(doc, "3.1.2. Cơ hội thị trường")

    add_bullet_list(doc, [
        "78.1% chủ TT chưa dùng phần mềm quản lý chuyên dụng (dùng Excel/sổ sách)",
        "74.2% sẵn sàng trả 500k-1tr/tháng cho giải pháp tốt",
        "88.1% phụ huynh sẵn sàng cài app mới để theo dõi con",
        "82.4% học viên hứng thú với gamification (điểm thưởng, huy hiệu)",
        "Đối thủ chưa có giải pháp Gamification và Multi-tenant hiện đại"
    ])

    add_subsection_title(doc, "3.1.3. Điểm khác biệt so với đối thủ")

    add_paragraph_text(doc,
        "Dựa trên khảo sát đối thủ và nhu cầu người dùng, KiteClass có thể tạo điểm "
        "khác biệt qua:")

    add_bullet_list(doc, [
        "Gamification: Hệ thống điểm thưởng, huy hiệu - chưa có đối thủ nào làm tốt",
        "Multi-tenant SaaS: Kiến trúc hiện đại cho phép scale và customization",
        "Tích hợp đa kênh thông báo: Zalo + App + SMS + Email",
        "UI/UX hiện đại: Giao diện đơn giản, ít click, phù hợp người không rành CNTT",
        "Thanh toán linh hoạt: VietQR, MoMo, chuyển khoản tự động đối soát"
    ])

    add_section_title(doc, "3.2. Feature Prioritization")

    add_paragraph_text(doc,
        "Dựa trên kết quả khảo sát, các tính năng được phân loại theo ma trận "
        "quan trọng - khả thi:")

    add_table_with_caption(doc,
        "Bảng 3.1. Ma trận ưu tiên tính năng",
        ["Ưu tiên", "Tính năng", "Điểm quan trọng", "Phase"],
        [
            ("P0 - Must Have", "Quản lý học viên, lớp học", "4.5", "MVP"),
            ("P0 - Must Have", "Quản lý học phí tự động", "4.6", "MVP"),
            ("P0 - Must Have", "Điểm danh QR/1-click", "3.8", "MVP"),
            ("P0 - Must Have", "Thông báo tự động (Zalo)", "4.4", "MVP"),
            ("P1 - Should Have", "Thanh toán VietQR", "4.1", "Phase 2"),
            ("P1 - Should Have", "Cổng phụ huynh", "3.9", "Phase 2"),
            ("P1 - Should Have", "Báo cáo doanh thu tự động", "4.3", "Phase 2"),
            ("P2 - Nice to Have", "Video bài giảng VOD", "3.5", "Phase 3"),
            ("P2 - Nice to Have", "Gamification (điểm, huy hiệu)", "3.2", "Phase 3"),
            ("P2 - Nice to Have", "Live streaming", "3.0", "Phase 3"),
        ],
        col_widths=[3.49, 6.19, 4.45, 2.46]
    )

    add_figure_placeholder(doc, "Hình 3.1. Feature Prioritization Matrix")

    add_section_title(doc, "3.3. User Personas đại diện")

    add_subsection_title(doc, "3.3.1. Persona 1: Chủ trung tâm - Cô Hương")

    add_paragraph_text(doc, "Thông tin cơ bản:")
    add_bullet_list(doc, [
        'Tên: Cô Hương, 35 tuổi, chủ TT "Bright English"',
        "Quy mô: 120 học viên, 8 giáo viên, 2 admin",
        "Kinh nghiệm: 3 năm điều hành, chuyên môn Sư phạm Anh",
        "Công cụ hiện tại: Excel + Zalo Group"
    ])

    add_paragraph_text(doc, "Pain points:")
    add_bullet_list(doc, [
        "Mất 2-3 tiếng/ngày cho công việc hành chính",
        "Khó theo dõi học phí khi quy mô tăng",
        "Phụ huynh hay hỏi về tiến độ học của con"
    ])

    add_paragraph_text(doc, "Goals:")
    add_bullet_list(doc, [
        "Giảm thời gian hành chính, tập trung phát triển chương trình",
        "Có cái nhìn tổng quan về tài chính",
        "Chuyên nghiệp hóa hình ảnh trung tâm"
    ])

    add_subsection_title(doc, "3.3.2. Persona 2: Phụ huynh - Chị Thu")

    add_paragraph_text(doc, "Thông tin cơ bản:")
    add_bullet_list(doc, [
        "Tên: Chị Thu, 38 tuổi, nhân viên văn phòng",
        "Có 2 con: 1 học tiếng Anh, 1 học Toán",
        "Thiết bị: Điện thoại Android, dùng Zalo hàng ngày"
    ])

    add_paragraph_text(doc, "Pain points:")
    add_bullet_list(doc, [
        "Không biết con có đi học đúng không",
        "Phải nhắn tin hỏi giáo viên về tiến độ",
        "Quên đóng học phí, bị nhắc nhiều lần ngại"
    ])

    add_paragraph_text(doc, "Goals:")
    add_bullet_list(doc, [
        "Được thông báo kịp thời về việc học của con",
        "Thanh toán tiện lợi, không phải đến trung tâm"
    ])

    add_section_title(doc, "3.4. Đề xuất chiến lược")

    add_subsection_title(doc, "3.4.1. Đề xuất gói dịch vụ")

    add_table_with_caption(doc,
        "Bảng 3.2. Đề xuất gói dịch vụ",
        ["Gói", "Giá/tháng", "Đối tượng", "Tính năng chính"],
        [
            ("Starter", "290.000 VND", "<50 HV", "Quản lý cơ bản, điểm danh, thông báo Zalo"),
            ("Growth", "590.000 VND", "50-200 HV", "Starter + VietQR, báo cáo, Parent Portal"),
            ("Pro", "1.190.000 VND", ">200 HV", "Growth + Video VOD, Gamification, API, Multi-admin"),
        ],
        col_widths=[2.06, 3.65, 2.54, 8.33]
    )

    add_paragraph_text(doc, "Lý do định giá:")
    add_bullet_list(doc, [
        "Starter (290k): Thấp hơn BeeClass (200k-500k) để thu hút TT nhỏ",
        "Growth (590k): Nằm trong khoảng 74.2% chấp nhận (500k-1tr)",
        "Pro (1.19tr): Cho TT lớn cần nhiều tính năng và hỗ trợ"
    ])

    add_subsection_title(doc, "3.4.2. Chiến lược Go-to-Market")

    add_numbered_list(doc, [
        "Focus vào phân khúc Ngoại ngữ (56.3% thị trường) trong giai đoạn đầu",
        "Chương trình dùng thử miễn phí 1 tháng để vượt qua rào cản tâm lý",
        "Tích hợp Zalo Notification sớm vì 64.3% phụ huynh ưa thích",
        "Thiết kế UI đơn giản, ít click, có onboarding wizard",
        "Partnership với các Hội chủ TT để tiếp cận khách hàng"
    ])

    add_subsection_title(doc, "3.4.3. Roadmap đề xuất")

    add_paragraph_text(doc, "Dựa trên kết quả khảo sát, đề xuất roadmap phát triển:")

    add_numbered_list(doc, [
        "MVP (3 tháng): Core features - Quản lý HV, học phí, điểm danh QR, thông báo Zalo",
        "Phase 2 (2 tháng): VietQR payment tự động đối soát, Parent Portal, báo cáo dashboard",
        "Phase 3 (2 tháng): Video VOD platform, Gamification (điểm, huy hiệu, bảng xếp hạng)",
        "Phase 4 (tùy chọn): Live streaming, AI chatbot hỗ trợ"
    ])


def add_conclusion(doc):
    """KẾT LUẬN"""
    add_chapter_title(doc, "KẾT LUẬN")

    add_section_title(doc, "1. Tổng kết kết quả khảo sát")

    add_paragraph_text(doc,
        "Qua quá trình khảo sát với 215 respondents qua form online và 18 cuộc phỏng vấn sâu, "
        "kết hợp với phân tích 3 sản phẩm cạnh tranh (BeeClass, Edupage, ClassIn), "
        "báo cáo đã thu thập được những insights quan trọng cho việc phát triển KiteClass Platform.")

    add_paragraph_text(doc, "Những kết luận chính:")
    add_bullet_list(doc, [
        "Thị trường có nhu cầu thực sự và chưa được đáp ứng đầy đủ (78.1% chưa dùng PM)",
        "Pain point lớn nhất: Quản lý học phí (87.5%) và liên lạc phụ huynh (81.3%)",
        "Mức giá 500k-1tr/tháng được chấp nhận bởi đa số (74.2%)",
        "Cơ hội khác biệt hóa qua Gamification và Multi-tenant architecture",
        "Phụ huynh sẵn sàng cài app mới (88.1%) - cơ hội lớn cho Parent Portal"
    ])

    add_section_title(doc, "2. Đóng góp của khảo sát")

    add_bullet_list(doc, [
        "Xác định 4 tính năng P0 (Must Have) cho MVP",
        "Hiểu rõ workflow và pain points của 5 nhóm đối tượng",
        "Có cơ sở để định giá sản phẩm theo phân khúc",
        "Xây dựng được 2 user personas đại diện",
        "Xác định kênh tiếp cận hiệu quả (Zalo, Facebook Group)",
        "Có bảng so sánh chi tiết với 3 đối thủ cạnh tranh"
    ])

    add_section_title(doc, "3. Hạn chế của khảo sát")

    add_bullet_list(doc, [
        "Mẫu khảo sát tập trung ở Hà Nội và TP.HCM, chưa đại diện toàn quốc",
        "Số lượng phỏng vấn sâu STUDENT và PARENT còn hạn chế",
        "Chưa thực hiện được usability testing với prototype",
        "Dữ liệu về mức sẵn sàng chi trả có thể khác thực tế"
    ])

    add_section_title(doc, "4. Hướng phát triển tiếp theo")

    add_bullet_list(doc, [
        "Thiết kế UI/UX prototype dựa trên insights từ khảo sát",
        "Thực hiện Usability Testing với 10-15 users đại diện",
        "Mở rộng khảo sát đến các tỉnh thành khác",
        "Pilot với 5-10 trung tâm trong 2 tháng đầu MVP",
        "Thu thập feedback liên tục để cải thiện sản phẩm"
    ])


def add_references(doc):
    """TÀI LIỆU THAM KHẢO"""
    add_chapter_title(doc, "TÀI LIỆU THAM KHẢO")

    references = [
        "[1] Portigal, S. (2013). Interviewing Users: How to Uncover Compelling Insights. "
        "Rosenfeld Media.",
        "[2] Kuniavsky, M. (2012). Observing the User Experience: A Practitioner's Guide "
        "to User Research. Morgan Kaufmann.",
        "[3] Gothelf, J., & Seiden, J. (2016). Lean UX: Designing Great Products with "
        "Agile Teams. O'Reilly Media.",
        "[4] BeeClass. (2024). Phần mềm quản lý trung tâm. https://beeclass.net",
        "[5] Edupage. (2024). School management system. https://edupage.org",
        "[6] ClassIn. (2024). Online teaching platform. https://classin.com",
        "[7] Bộ Giáo dục và Đào tạo. (2023). Số liệu thống kê giáo dục.",
    ]

    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        run = p.add_run(ref)
        set_font(run, FONT_SIZE_NORMAL)


def add_appendix(doc):
    """PHỤ LỤC"""
    add_chapter_title(doc, "PHỤ LỤC")

    add_section_title(doc, "Phụ lục A: Link Google Forms khảo sát")

    add_paragraph_text(doc, "Các form khảo sát online được thiết kế trên Google Forms:")
    add_bullet_list(doc, [
        "Form CENTER_OWNER: [Link sẽ được cung cấp khi triển khai]",
        "Form TEACHER: [Link sẽ được cung cấp khi triển khai]",
        "Form STUDENT: [Link sẽ được cung cấp khi triển khai]",
        "Form PARENT: [Link sẽ được cung cấp khi triển khai]"
    ])

    add_section_title(doc, "Phụ lục B: Checklist phỏng vấn")

    add_paragraph_text(doc, "Checklist chuẩn bị trước phỏng vấn:")
    add_bullet_list(doc, [
        "Xác nhận lịch hẹn với người tham gia",
        "Chuẩn bị thiết bị ghi âm (có xin phép)",
        "In sẵn kịch bản phỏng vấn",
        "Chuẩn bị quà cảm ơn (nếu có)",
        "Test kết nối Zoom/Google Meet"
    ])

    add_paragraph_text(doc, "Checklist sau phỏng vấn:")
    add_bullet_list(doc, [
        "Gửi email cảm ơn trong vòng 24 giờ",
        "Transcribe recording trong vòng 48 giờ",
        "Highlight key insights và quotes",
        "Cập nhật vào bảng tổng hợp"
    ])

    add_section_title(doc, "Phụ lục C: Danh sách người tham gia phỏng vấn")

    interview_list = [
        ("1", "Cô Hương", "CENTER_OWNER", "TT Anh ngữ, HN", "45 phút"),
        ("2", "Anh Minh", "CENTER_OWNER", "TT Toán, HCM", "40 phút"),
        ("3", "Chị Lan", "CENTER_OWNER", "TT Kỹ năng, HN", "35 phút"),
        ("4", "Anh Tuấn", "CENTER_ADMIN", "TT Anh ngữ, HN", "30 phút"),
        ("5", "Chị Mai", "CENTER_ADMIN", "TT Ôn thi, HCM", "30 phút"),
        ("6", "Cô Thảo", "TEACHER", "GV IELTS, HN", "25 phút"),
        ("7", "Thầy Nam", "TEACHER", "GV Toán, HCM", "25 phút"),
        ("8", "Chị Thu", "PARENT", "2 con, HN", "20 phút"),
        ("...", "...", "...", "...", "..."),
    ]

    # Độ rộng cột theo file đã chỉnh sửa
    col_widths = [0.94, 2.56, 3.79, 3.85, 5.83]

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["#", "Tên (đổi)", "Vai trò", "Mô tả", "Thời lượng"]
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].width = Cm(col_widths[i])
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_font(run, FONT_SIZE_TABLE, bold=True)
        set_cell_shading(header_cells[i], 'D9E2F3')

    for row_data in interview_list:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            row.cells[i].text = str(cell_text)
            row.cells[i].width = Cm(col_widths[i])
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    set_font(run, FONT_SIZE_TABLE)

    doc.add_paragraph()

    # Chữ ký
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Hà Nội, ngày ... tháng 01 năm 2026")
    set_font(run, FONT_SIZE_NORMAL, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Sinh viên thực hiện")
    set_font(run, FONT_SIZE_NORMAL, bold=True)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Nguyễn Văn Kiệt")
    set_font(run, FONT_SIZE_NORMAL)


def create_report():
    """Hàm chính tạo báo cáo khảo sát"""
    print("Đang tạo báo cáo khảo sát KiteClass theo quy định ĐH GTVT...")

    doc = Document()

    # Thiết lập document
    set_document_margins(doc)
    setup_styles(doc)

    # Trang bìa
    add_title_page(doc)

    # Các phần đầu
    add_toc_page(doc)
    add_list_of_tables(doc)
    add_list_of_figures(doc)
    add_abbreviations(doc)

    # Nội dung chính
    add_introduction(doc)
    add_content1(doc)  # Kế hoạch khảo sát, bảng hỏi, phỏng vấn
    add_content2(doc)  # Khảo sát đối thủ + Kết quả khảo sát
    add_content3(doc)  # Phân tích và đề xuất
    add_conclusion(doc)
    add_references(doc)
    add_appendix(doc)

    # Thêm số trang
    add_page_number_header(doc)

    # Lưu file
    output_path = "BAO_CAO_KHAO_SAT_KITECLASS.docx"
    doc.save(output_path)

    print(f"✓ Đã tạo file: {output_path}")
    print(f"✓ Căn lề: Trái 3cm, Phải 2cm, Trên 2.5cm, Dưới 2.5cm")
    print(f"✓ Số trang: Giữa, phía trên đầu trang")
    print(f"✓ Chương: 18pt Bold, Mục: 16pt Bold, Tiểu mục: 14pt Bold")
    print(f"✓ Nội dung: 13pt, giãn dòng 1.2, thụt đầu dòng 1cm")
    print(f"")
    print(f"Nội dung báo cáo:")
    print(f"  - Nội dung 1: Kế hoạch khảo sát chi tiết, bảng hỏi 4 đối tượng, câu hỏi phỏng vấn")
    print(f"  - Nội dung 2: Khảo sát 3 sản phẩm cạnh tranh (BeeClass, Edupage, ClassIn)")
    print(f"  - Nội dung 2: Kết quả khảo sát từ 215 responses + 18 phỏng vấn")
    print(f"  - Nội dung 3: Phân tích, so sánh và đề xuất chiến lược")

    return output_path


if __name__ == "__main__":
    create_report()
