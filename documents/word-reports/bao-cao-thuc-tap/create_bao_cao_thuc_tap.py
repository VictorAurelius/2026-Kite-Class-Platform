#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script tạo báo cáo thực tập tốt nghiệp dạng Word (.docx)
Format theo mẫu "Huong dan trinh bay bao cao TTTN.pdf" - ĐH GTVT

Cấu trúc báo cáo:
1. Bìa chính (có bảng thông tin sinh viên, logo, chữ CỬ NHÂN màu vàng)
2. Bìa phụ (thêm trường Đơn vị thực tập)
3. Bản nhận xét của cơ sở thực tập
4. Lời cảm ơn
5. Mục lục + Danh mục hình vẽ + Danh mục bảng biểu
6. Danh mục từ viết tắt
7. 4 Chương nội dung chính
8. Tài liệu tham khảo (IEEE)
9. Phụ lục
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============== THÔNG TIN SINH VIÊN ==============
STUDENT_INFO = {
    "name": "Nguyễn Văn Kiệt",
    "student_id": "221230890",
    "class": "CNTT1-K63",
    "course": "63",
    "major": "Công nghệ thông tin",
    "department": "Công nghệ thông tin",
    "degree": "Cử nhân",
    "university": "Đại học Giao thông Vận tải",
}

INTERNSHIP_INFO = {
    "company": "SY PARTNERS., JSC",
    "address": "Tầng 3, Tòa nhà Luxury, Số 99 Võ Chí Công, Quận Tây Hồ, Hà Nội",
    "position": "Software Engineer",
    "advisor": "TS. Nguyễn Đức Dư",
    "company_mentor": "[Tên CBHD]",
    "start_date": "26/06/2025",
    "end_date": "26/09/2025",
}

# ============== CONSTANTS ==============
FONT_NAME = 'Times New Roman'
FONT_SIZE_NORMAL = Pt(13)  # Đoạn văn
FONT_SIZE_CHAPTER = Pt(18)  # Chương - Theo quy định UTC
FONT_SIZE_SECTION = Pt(16)  # Mục 1.1 - Theo quy định UTC
FONT_SIZE_SUBSECTION = Pt(14)  # Tiểu mục 1.1.1 - Theo quy định UTC
FONT_SIZE_TABLE = Pt(12)
FONT_SIZE_CAPTION = Pt(13)

LINE_SPACING = 1.2  # Theo quy định UTC
FIRST_LINE_INDENT = Cm(1.0)  # Theo quy định UTC

MARGIN_LEFT = Cm(3.0)
MARGIN_RIGHT = Cm(2.0)
MARGIN_TOP = Cm(2.5)  # Theo quy định UTC
MARGIN_BOTTOM = Cm(2.5)  # Theo quy định UTC


def set_document_margins(doc):
    """Thiết lập căn lề cho toàn bộ document"""
    for section in doc.sections:
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT


def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def remove_cell_borders(cell):
    """Remove all borders from a cell"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_page_border(section):
    """Thêm khung viền cho trang bìa"""
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'text')

    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '24')  # Border width (1/8 point)
        border.set(qn('w:space'), '24')  # Space from text (points)
        border.set(qn('w:color'), '000000')  # Black
        pgBorders.append(border)

    sectPr.append(pgBorders)


def remove_page_border(section):
    """Xóa khung viền khỏi section (nếu có)"""
    sectPr = section._sectPr
    # Tìm và xóa tất cả pgBorders elements
    pgBorders_list = sectPr.findall(qn('w:pgBorders'))
    for pgBorders in pgBorders_list:
        sectPr.remove(pgBorders)


def add_page_number_header(doc):
    """Thêm số trang ở giữa phía TRÊN đầu trang (header)"""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False

        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        for r in p.runs:
            r.font.name = FONT_NAME
            r.font.size = FONT_SIZE_NORMAL


def set_heading_font(style, font_name, font_size, bold=True, italic=False):
    """
    Thiết lập font cho Heading style một cách đầy đủ.

    Word Heading styles mặc định sử dụng theme fonts (Calibri Headings)
    và theme colors (màu xanh). Cần:
    - XÓA các thuộc tính theme (asciiTheme, hAnsiTheme, themeColor)
    - Thiết lập font và color trực tiếp
    """
    # Đảm bảo rPr element tồn tại
    rPr = style._element.get_or_add_rPr()

    # Xóa rFonts cũ (có thể chứa theme fonts) và tạo mới
    old_rFonts = rPr.find(qn('w:rFonts'))
    if old_rFonts is not None:
        rPr.remove(old_rFonts)

    # Tạo rFonts mới với font name trực tiếp (không dùng theme)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    # KHÔNG đặt asciiTheme, hAnsiTheme để tránh bị override bởi theme
    rPr.insert(0, rFonts)

    # Thiết lập font size
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rPr.append(sz)
    sz.set(qn('w:val'), str(int(font_size.pt * 2)))  # Word uses half-points

    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = OxmlElement('w:szCs')
        rPr.append(szCs)
    szCs.set(qn('w:val'), str(int(font_size.pt * 2)))

    # Thiết lập bold
    b = rPr.find(qn('w:b'))
    if bold:
        if b is None:
            b = OxmlElement('w:b')
            rPr.append(b)
    else:
        if b is not None:
            rPr.remove(b)

    # Thiết lập italic
    i = rPr.find(qn('w:i'))
    if italic:
        if i is None:
            i = OxmlElement('w:i')
            rPr.append(i)
    else:
        if i is not None:
            rPr.remove(i)

    # Xóa color cũ (có thể chứa themeColor) và tạo mới
    old_color = rPr.find(qn('w:color'))
    if old_color is not None:
        rPr.remove(old_color)

    # Tạo color mới với màu đen trực tiếp (không dùng theme)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '000000')
    # KHÔNG đặt themeColor để tránh bị override
    rPr.append(color)


def setup_styles(doc):
    """Thiết lập các style chuẩn cho document bao gồm Heading styles"""
    # Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE_NORMAL

    pf = style.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Thiết lập font cho Normal style qua XML
    rPr = style._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)

    # Heading 1 style (Chương) - 14pt Bold, Left
    h1_style = doc.styles['Heading 1']
    # Thiết lập qua XML để xóa theme
    set_heading_font(h1_style, FONT_NAME, Pt(14), bold=True, italic=False)
    # Thiết lập qua API để đảm bảo
    h1_style.font.name = FONT_NAME
    h1_style.font.size = Pt(14)
    h1_style.font.bold = True
    h1_style.font.italic = False
    h1_style.font.color.rgb = RGBColor(0, 0, 0)
    h1_style.paragraph_format.space_before = Pt(12)
    h1_style.paragraph_format.space_after = Pt(6)
    h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Heading 2 style (Mục 1.1, 1.2) - 13pt Bold, Left
    h2_style = doc.styles['Heading 2']
    set_heading_font(h2_style, FONT_NAME, Pt(13), bold=True, italic=False)
    h2_style.font.name = FONT_NAME
    h2_style.font.size = Pt(13)
    h2_style.font.bold = True
    h2_style.font.italic = False
    h2_style.font.color.rgb = RGBColor(0, 0, 0)
    h2_style.paragraph_format.space_before = Pt(12)
    h2_style.paragraph_format.space_after = Pt(6)
    h2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Heading 3 style (Tiểu mục 1.1.1) - 13pt Bold + Italic, Left
    h3_style = doc.styles['Heading 3']
    set_heading_font(h3_style, FONT_NAME, Pt(13), bold=True, italic=True)
    h3_style.font.name = FONT_NAME
    h3_style.font.size = Pt(13)
    h3_style.font.bold = True
    h3_style.font.italic = True
    h3_style.font.color.rgb = RGBColor(0, 0, 0)
    h3_style.paragraph_format.space_before = Pt(6)
    h3_style.paragraph_format.space_after = Pt(6)
    h3_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Caption style cho bảng và hình
    try:
        caption_style = doc.styles['Caption']
    except KeyError:
        caption_style = doc.styles.add_style('Caption', 1)  # 1 = paragraph style
    set_heading_font(caption_style, FONT_NAME, Pt(12), bold=True, italic=True)
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_style.paragraph_format.space_before = Pt(6)
    caption_style.paragraph_format.space_after = Pt(6)


def set_font(run, size=FONT_SIZE_NORMAL, bold=False, italic=False, color=None):
    """Helper to set font properties"""
    run.font.name = FONT_NAME
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)


def add_chapter_title(doc, number, text, add_page_break=True):
    """
    Tiêu đề chương: Sử dụng Heading 1 style để tạo mục lục tự động
    Format: "1. GIỚI THIỆU CHUNG VỀ ĐƠN VỊ THỰC TẬP"
    - Heading 1: 18pt Bold, Center, Times New Roman (theo quy định UTC)
    """
    if add_page_break:
        doc.add_page_break()

    # Sử dụng Heading 1 style để Word có thể tạo mục lục tự động
    p = doc.add_paragraph(style='Heading 1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Căn giữa theo quy định UTC

    run = p.add_run(f"{number}. {text.upper()}")

    # Thiết lập font trực tiếp cho run để đảm bảo không bị theme override
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = FONT_SIZE_CHAPTER  # 18pt theo quy định UTC
    run.font.bold = True
    run.font.italic = False
    run.font.color.rgb = RGBColor(0, 0, 0)

    return p


def add_section_title(doc, text):
    """
    Tiêu đề mục (1.1, 1.2): Sử dụng Heading 2 style để tạo mục lục tự động
    - Heading 2: 16pt Bold, Left, Times New Roman (theo quy định UTC)
    """
    p = doc.add_paragraph(style='Heading 2')
    run = p.add_run(text)

    # Thiết lập font trực tiếp cho run
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = FONT_SIZE_SECTION  # 16pt theo quy định UTC
    run.font.bold = True
    run.font.italic = False
    run.font.color.rgb = RGBColor(0, 0, 0)

    return p


def add_subsection_title(doc, text):
    """
    Tiêu đề tiểu mục (1.1.1, 1.1.2): Sử dụng Heading 3 style để tạo mục lục tự động
    - Heading 3: 14pt Bold, Left, Times New Roman (theo quy định UTC)
    """
    p = doc.add_paragraph(style='Heading 3')
    run = p.add_run(text)

    # Thiết lập font trực tiếp cho run
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = FONT_SIZE_SUBSECTION  # 14pt theo quy định UTC
    run.font.bold = True
    run.font.italic = False  # Không italic theo quy định UTC
    run.font.color.rgb = RGBColor(0, 0, 0)

    return p


def add_seq_field(paragraph, seq_name, prefix="", reset_chapter=False):
    """
    Thêm SEQ field vào paragraph để đánh số tự động
    SEQ field giống như khi dùng Insert > Caption trong Word

    Args:
        paragraph: Paragraph object
        seq_name: Tên sequence (ví dụ: "Table", "Figure") - KHÔNG thêm số chapter
        prefix: Text đứng trước số (ví dụ: "Bảng ", "Hình ")
        reset_chapter: True nếu cần reset về 1 ở đầu chapter mới
    """
    if prefix:
        run = paragraph.add_run(prefix)
        set_font(run, Pt(12), bold=True, italic=True)

    # Tạo SEQ field
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    # Thêm \r để reset numbering nếu cần (ở đầu chapter mới)
    if reset_chapter:
        instrText.text = f' SEQ {seq_name} \\* ARABIC \\r 1 '
    else:
        instrText.text = f' SEQ {seq_name} \\* ARABIC '

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    # Placeholder text (sẽ được cập nhật khi F9 trong Word)
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    run2 = paragraph.add_run("1")  # Placeholder number
    set_font(run2, Pt(12), bold=True, italic=True)

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run2._r.append(fldChar3)

    return paragraph


def add_table_caption(doc, chapter_num, caption_text, reset=False):
    """
    Thêm caption cho bảng với SEQ field tự động đánh số
    Format: "Bảng X.Y. Caption text" (với X là số chương, Y là SEQ tự động)

    Args:
        doc: Document object
        chapter_num: Số chương (1, 2, 3, 4)
        caption_text: Nội dung caption
        reset: True nếu đây là bảng đầu tiên trong chapter (reset về 1)

    Lưu ý: Trong Word, bấm Ctrl+A rồi F9 để cập nhật tất cả fields
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    # Thêm "Bảng X." với X là số chương
    run = p.add_run(f"Bảng {chapter_num}.")
    set_font(run, Pt(12), bold=True, italic=True)

    # Thêm SEQ field cho số thứ tự - QUAN TRỌNG: Dùng "Table" chung cho tất cả chapter
    # Để TOC \c "Table" có thể tìm thấy
    add_seq_field(p, "Table", "", reset_chapter=reset)

    # Thêm phần caption text
    run = p.add_run(f". {caption_text}")
    set_font(run, Pt(12), bold=True, italic=True)

    return p


def add_figure_caption(doc, chapter_num, caption_text, reset=False):
    """
    Thêm caption cho hình với SEQ field tự động đánh số
    Format: "Hình X.Y. Caption text" (với X là số chương, Y là SEQ tự động)

    Args:
        doc: Document object
        chapter_num: Số chương (1, 2, 3, 4)
        caption_text: Nội dung caption
        reset: True nếu đây là hình đầu tiên trong chapter (reset về 1)

    Lưu ý: Trong Word, bấm Ctrl+A rồi F9 để cập nhật tất cả fields
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)

    # Thêm "Hình X." với X là số chương
    run = p.add_run(f"Hình {chapter_num}.")
    set_font(run, Pt(12), bold=True, italic=True)

    # Thêm SEQ field cho số thứ tự - QUAN TRỌNG: Dùng "Figure" chung cho tất cả chapter
    # Để TOC \c "Figure" có thể tìm thấy
    add_seq_field(p, "Figure", "", reset_chapter=reset)

    # Thêm phần caption text
    run = p.add_run(f". {caption_text}")
    set_font(run, Pt(12), bold=True, italic=True)

    return p


def add_paragraph_text(doc, text, first_line_indent=True):
    """
    Đoạn văn: 13pt, Justify, thụt đầu dòng 1.27cm, giãn dòng 1.5
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = LINE_SPACING

    if first_line_indent:
        p.paragraph_format.first_line_indent = FIRST_LINE_INDENT

    run = p.add_run(text)
    set_font(run, FONT_SIZE_NORMAL)

    return p


def add_bullet_list(doc, items):
    """
    Thêm danh sách bullet

    Format: Bullet nhô ra ngoài, nội dung thẳng hàng khi xuống dòng
    - left_indent: Vị trí nội dung text (1.5cm)
    - first_line_indent: Bullet nhô ra (-0.5cm)
    """
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        # Text nằm ở 1.5cm từ lề trái
        p.paragraph_format.left_indent = Cm(1.5)
        # Bullet nhô ra ngoài 0.5cm (hanging indent)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.line_spacing = LINE_SPACING

        run = p.add_run(item)
        set_font(run, FONT_SIZE_NORMAL)


def add_table_with_caption(doc, chapter_num, caption_text, headers, rows, col_widths=None, reset=False):
    """
    Thêm bảng với caption sử dụng SEQ field để đánh số tự động
    Caption ở PHÍA TRÊN bảng theo đúng chuẩn Word

    Args:
        doc: Document object
        chapter_num: Số chương (1, 2, 3, 4)
        caption_text: Nội dung caption
        headers: List các header cột
        rows: List các hàng dữ liệu
        col_widths: List độ rộng cột (cm)
        reset: True nếu đây là bảng đầu tiên trong chapter
    """
    # Thêm caption với SEQ field
    add_table_caption(doc, chapter_num, caption_text, reset=reset)

    # Tạo bảng
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        if col_widths and i < len(col_widths):
            header_cells[i].width = Cm(col_widths[i])
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_font(run, FONT_SIZE_TABLE, bold=True)
        set_cell_shading(header_cells[i], 'D9E2F3')

    for row_data in rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            row.cells[i].text = str(cell_text)
            if col_widths and i < len(col_widths):
                row.cells[i].width = Cm(col_widths[i])
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    set_font(run, FONT_SIZE_TABLE)

    doc.add_paragraph()
    return table


def add_figure_placeholder(doc, chapter_num, caption_text):
    """
    Thêm placeholder cho hình vẽ với caption sử dụng SEQ field
    Caption ở PHÍA DƯỚI hình theo đúng chuẩn Word

    Args:
        doc: Document object
        chapter_num: Số chương (1, 2, 3, 4)
        caption_text: Nội dung caption
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)

    run = p.add_run("[Chèn hình vẽ tại đây]")
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_NORMAL
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Thêm caption với SEQ field
    add_figure_caption(doc, chapter_num, caption_text)


# ============== TRANG BÌA CHÍNH ==============
def add_cover_page(doc):
    """Tạo trang bìa chính theo mẫu PDF (trang 1)"""
    import os

    # TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI")
    set_font(run, Pt(14), bold=False)

    # KHOA CÔNG NGHỆ THÔNG TIN (đậm, gạch chân)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("KHOA CÔNG NGHỆ THÔNG TIN")
    set_font(run, Pt(14), bold=True)
    run.font.underline = True

    # Logo
    logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'templates', 'logo_utc.png')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(24)

    if os.path.exists(logo_path):
        run = p.add_run()
        run.add_picture(logo_path, width=Cm(3.5))
    else:
        run = p.add_run("[LOGO TRƯỜNG]")
        run.font.color.rgb = RGBColor(128, 128, 128)
        set_font(run, Pt(12))

    # BÁO CÁO (gạch chân)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("BÁO CÁO")
    set_font(run, Pt(14), bold=True)
    run.font.underline = True

    # THỰC TẬP TỐT NGHIỆP
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("THỰC TẬP TỐT NGHIỆP")
    set_font(run, Pt(22), bold=True)

    # CỬ NHÂN (không màu vàng, không gạch chân)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("CỬ NHÂN")
    set_font(run, Pt(22), bold=True)

    # Bảng thông tin sinh viên (CÓ viền)
    table = doc.add_table(rows=9, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'  # Thêm style có viền

    info_rows = [
        ("Sinh viên thực hiện", STUDENT_INFO["name"]),
        ("Mã sinh viên", STUDENT_INFO["student_id"]),
        ("Lớp", STUDENT_INFO["class"]),
        ("Khóa", STUDENT_INFO["course"]),
        ("Ngành đào tạo", STUDENT_INFO["major"]),
        ("Đơn vị thực tập", INTERNSHIP_INFO["company"]),
        ("Giảng viên hướng dẫn", INTERNSHIP_INFO["advisor"]),
        ("CBHD tại đơn vị TT", INTERNSHIP_INFO["company_mentor"]),
        ("Thời gian thực tập", f"Từ ngày {INTERNSHIP_INFO['start_date']} đến ngày {INTERNSHIP_INFO['end_date']}"),
    ]

    for i, (label, value) in enumerate(info_rows):
        row = table.rows[i]
        # Label cell
        row.cells[0].text = label
        row.cells[0].width = Cm(5.0)
        for paragraph in row.cells[0].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_font(run, Pt(13))

        # Value cell
        row.cells[1].text = f": {value}"
        row.cells[1].width = Cm(9.0)
        for paragraph in row.cells[1].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_font(run, Pt(13))

    # Khoảng trống
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # Hà Nội – 2026
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Hà Nội – 2026")
    set_font(run, Pt(14), bold=True, italic=True)

    # QUAN TRỌNG: Tạo section break để tách bìa chính thành section riêng
    doc.add_section(WD_SECTION.NEW_PAGE)


# ============== TRANG BÌA PHỤ ==============
def add_secondary_cover_page(doc):
    """Tạo trang bìa phụ theo mẫu PDF (trang 2) - thêm Đơn vị thực tập

    Trang này sẽ nằm trong section 1 (nhờ section break ở cuối add_cover_page)
    """
    import os

    # Section break đã được tạo trong add_cover_page, nội dung này sẽ ở section 1

    # TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI")
    set_font(run, Pt(14), bold=False)

    # KHOA CÔNG NGHỆ THÔNG TIN (đậm, gạch chân)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(36)
    run = p.add_run("KHOA CÔNG NGHỆ THÔNG TIN")
    set_font(run, Pt(14), bold=True)
    run.font.underline = True

    # BÁO CÁO (gạch chân)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("BÁO CÁO")
    set_font(run, Pt(14), bold=True)
    run.font.underline = True

    # THỰC TẬP TỐT NGHIỆP
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("THỰC TẬP TỐT NGHIỆP")
    set_font(run, Pt(22), bold=True)

    # CỬ NHÂN (không màu vàng, không gạch chân)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("CỬ NHÂN")
    set_font(run, Pt(22), bold=True)

    # Bảng thông tin sinh viên (CÓ viền, có thêm Đơn vị thực tập)
    table = doc.add_table(rows=9, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'  # Thêm style có viền

    info_rows = [
        ("Sinh viên thực hiện", STUDENT_INFO["name"]),
        ("Mã sinh viên", STUDENT_INFO["student_id"]),
        ("Lớp", STUDENT_INFO["class"]),
        ("Khóa", STUDENT_INFO["course"]),
        ("Ngành đào tạo", STUDENT_INFO["major"]),
        ("Đơn vị thực tập", INTERNSHIP_INFO["company"]),
        ("Giảng viên hướng dẫn", INTERNSHIP_INFO["advisor"]),
        ("CBHD tại đơn vị TT", INTERNSHIP_INFO["company_mentor"]),
        ("Thời gian thực tập", f"Từ ngày {INTERNSHIP_INFO['start_date']} đến ngày {INTERNSHIP_INFO['end_date']}"),
    ]

    for i, (label, value) in enumerate(info_rows):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].width = Cm(5.0)
        for paragraph in row.cells[0].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_font(run, Pt(13))

        row.cells[1].text = f": {value}"
        row.cells[1].width = Cm(9.0)
        for paragraph in row.cells[1].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_font(run, Pt(13))

    # Khoảng trống
    for _ in range(2):
        p = doc.add_paragraph()

    # Hà Nội – 2026
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Hà Nội – 2026")
    set_font(run, Pt(14), bold=True, italic=True)

    # QUAN TRỌNG: Tạo section break để tách bìa phụ thành section riêng
    # Sau section break này, các trang còn lại sẽ KHÔNG có border
    doc.add_section(WD_SECTION.NEW_PAGE)


# ============== BẢN NHẬN XÉT CỦA CƠ SỞ THỰC TẬP ==============
def add_company_review_page(doc):
    """Tạo trang Bản nhận xét của cơ sở thực tập (trang 3)"""
    doc.add_page_break()

    # CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    set_font(run, Pt(13), bold=True)

    # Độc lập - Tự do - Hạnh phúc
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Độc lập - Tự do - Hạnh phúc")
    set_font(run, Pt(13), bold=True)
    run.font.underline = True

    # Ngày tháng năm
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Hà Nội, ngày …. tháng … năm 2026")
    set_font(run, Pt(13), italic=True)

    # Tiêu đề
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("BẢN NHẬN XÉT CỦA CƠ SỞ THỰC TẬP")
    set_font(run, Pt(14), bold=True)

    # Kính gửi
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Kính gửi: ")
    set_font(run, Pt(13), bold=True, italic=True)
    run = p.add_run("Khoa Công nghệ thông tin, Trường Đại học Giao thông vận tải")
    set_font(run, Pt(13), bold=True, italic=True)

    # Thông tin cơ sở thực tập
    info_lines = [
        f"Cơ sở thực tập: {INTERNSHIP_INFO['company']}",
        "Người đại diện: ",
        "Chức vụ: ",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_font(run, Pt(13))

    # Xác nhận sinh viên
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Xác nhận sinh viên:")
    set_font(run, Pt(13), bold=True)

    # Thông tin sinh viên
    p = doc.add_paragraph()
    run = p.add_run(f"Họ tên: {STUDENT_INFO['name']}")
    set_font(run, Pt(13))
    run = p.add_run(f"\t\t\tMã sinh viên: {STUDENT_INFO['student_id']}")
    set_font(run, Pt(13))

    p = doc.add_paragraph()
    run = p.add_run(f"Lớp: {STUDENT_INFO['class']}")
    set_font(run, Pt(13))

    p = doc.add_paragraph()
    run = p.add_run(f"Đã thực tập tốt nghiệp tại cơ sở trong thời gian từ: {INTERNSHIP_INFO['start_date']} đến: {INTERNSHIP_INFO['end_date']}")
    set_font(run, Pt(13))

    # Các mục nhận xét
    review_items = [
        "Nội dung thực tập:",
        "Về tinh thần, ý thức, thái độ đối với công việc được giao:",
        "Về trình độ, kỹ năng làm việc/ khả năng thực hành:",
        "Ưu điểm nổi bật:",
        "Hạn chế cần khắc phục:",
        "Các nhận xét khác (nếu có):",
    ]

    for item in review_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run(f"- {item}")
        set_font(run, Pt(13))
        run.font.underline = True

    # Điểm thực tập
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run("Điểm thực tập (thang điểm 10): …… điểm.")
    set_font(run, Pt(13))

    # Chữ ký
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("ĐẠI DIỆN CƠ SỞ THỰC TẬP")
    set_font(run, Pt(13), bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("(Họ tên, chữ ký và đóng dấu nếu có)")
    set_font(run, Pt(13), italic=True)


# ============== LỜI CẢM ƠN ==============
def add_acknowledgment_page(doc):
    """Tạo trang Lời cảm ơn (trang 4) - theo mẫu tham khảo"""
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("LỜI CẢM ƠN")
    set_font(run, Pt(14), bold=True, italic=True)

    # Nội dung lời cảm ơn - chi tiết hơn theo mẫu tham khảo
    add_paragraph_text(doc,
        "Trong suốt quá trình thực tập và hoàn thành báo cáo này, em đã nhận được sự "
        "quan tâm, hướng dẫn và giúp đỡ quý báu từ nhiều tập thể và cá nhân. Đây là nguồn động "
        "viên to lớn, giúp em có thêm động lực và kiến thức để hoàn thành tốt đợt thực tập của mình.")

    add_paragraph_text(doc,
        f"Trước hết, em xin bày tỏ lòng biết ơn sâu sắc đến {INTERNSHIP_INFO['advisor']}, giảng "
        "viên hướng dẫn thuộc Trường Đại học Giao thông Vận tải. Thầy đã tận tình hướng dẫn, "
        "định hướng nội dung thực tập, đóng góp nhiều ý kiến chuyên môn quan trọng và luôn "
        "theo sát, hỗ trợ em trong suốt quá trình thực hiện đề tài. Những kiến thức chuyên sâu, "
        "kinh nghiệm thực tiễn cũng như sự nghiêm túc trong học thuật mà thầy truyền đạt đã "
        "giúp em nâng cao tư duy chuyên môn và hoàn thiện báo cáo một cách tốt hơn.")

    add_paragraph_text(doc,
        "Em xin chân thành cảm ơn Khoa Công nghệ thông tin, Trường Đại học Giao "
        "thông Vận tải đã tạo điều kiện thuận lợi để em được tham gia thực tập, tiếp cận với môi "
        "trường làm việc thực tế và vận dụng những kiến thức đã học vào thực tiễn. Sự hỗ trợ "
        "của Khoa là nền tảng quan trọng giúp sinh viên có cơ hội học hỏi, rèn luyện kỹ năng và "
        "tích lũy kinh nghiệm thực tế.")

    add_paragraph_text(doc,
        f"Bên cạnh đó, em xin gửi lời cảm ơn chân thành đến {INTERNSHIP_INFO['company']} cùng các "
        "anh/chị trong đơn vị đã nhiệt tình hướng dẫn, chia sẻ kinh nghiệm chuyên môn, tạo điều "
        "kiện thuận lợi để em được tham gia vào các công việc thực tế, qua đó giúp em hiểu rõ "
        "hơn về quy trình làm việc cũng như yêu cầu của môi trường nghề nghiệp sau này.")

    add_paragraph_text(doc,
        "Cuối cùng, em xin cảm ơn gia đình, bạn bè và những người thân đã luôn quan "
        "tâm, động viên, hỗ trợ em cả về tinh thần lẫn vật chất trong suốt thời gian thực tập và học tập.")

    add_paragraph_text(doc,
        "Mặc dù đã rất cố gắng, song do thời gian thực tập và kinh nghiệm thực tiễn còn "
        "hạn chế, báo cáo không tránh khỏi những thiếu sót. Em rất mong nhận được sự đóng "
        "góp ý kiến từ quý thầy cô để báo cáo được hoàn thiện hơn.")

    add_paragraph_text(doc, "Em xin chân thành cảm ơn!")

    # Chữ ký
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Hà Nội, ngày ... tháng 01 năm 2026")
    set_font(run, Pt(13), italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Sinh viên thực hiện")
    set_font(run, Pt(13), bold=True)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(STUDENT_INFO["name"])
    set_font(run, Pt(13))


# ============== MỤC LỤC ==============
def add_toc_page(doc):
    """Thêm trang Mục lục"""
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("MỤC LỤC")
    set_font(run, Pt(14), bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Tạo mục lục tự động: References → Table of Contents]")
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(128, 128, 128)


def add_list_of_figures(doc):
    """
    Thêm Danh mục hình vẽ với TOC field tự động

    Sau khi mở file Word:
    1. Bấm Ctrl+A (chọn tất cả)
    2. Bấm F9 (cập nhật tất cả fields)
    3. Danh mục hình vẽ sẽ tự động hiển thị
    """
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("DANH MỤC HÌNH VẼ")
    set_font(run, Pt(14), bold=True)

    # Thêm TOC field tự động cho hình vẽ
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_toc_field(p, toc_type="Figure")

    # Thêm hướng dẫn
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("(Bấm Ctrl+A rồi F9 trong Word để cập nhật danh mục)")
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(128, 128, 128)


def add_toc_field(paragraph, toc_type="Table"):
    """
    Thêm TOC field cho danh mục bảng biểu hoặc hình vẽ

    Args:
        paragraph: Paragraph object
        toc_type: "Table" hoặc "Figure"
    """
    run = paragraph.add_run()

    # Begin field character
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    # Instruction text
    # TOC \h \z \c "Table" - creates table of figures for "Table" label
    # \h: Include hyperlinks
    # \z: Hide page numbers in web layout view
    # \c "Table": Only include items with "Table" label
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f'TOC \\h \\z \\c "{toc_type}"'

    # End field character
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    # Add all elements to run
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    return run


def add_list_of_tables(doc):
    """
    Thêm Danh mục bảng biểu với TOC field tự động

    Sau khi mở file Word:
    1. Bấm Ctrl+A (chọn tất cả)
    2. Bấm F9 (cập nhật tất cả fields)
    3. Danh mục bảng biểu sẽ tự động hiển thị
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("DANH MỤC BẢNG BIỂU")
    set_font(run, Pt(14), bold=True)

    # Thêm TOC field tự động cho bảng biểu
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_toc_field(p, toc_type="Table")

    # Thêm hướng dẫn
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("(Bấm Ctrl+A rồi F9 trong Word để cập nhật danh mục)")
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(128, 128, 128)


def add_abbreviations(doc):
    """Thêm Danh mục từ viết tắt - 3 cột theo mẫu tham khảo"""
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("DANH MỤC TỪ VIẾT TẮT")
    set_font(run, Pt(14), bold=True)

    # 3 cột: Từ viết tắt | Diễn giải | Ghi chú trong báo cáo
    abbreviations = [
        ("AI", "Artificial Intelligence", "Trí tuệ nhân tạo, công cụ hỗ trợ kiểm tra thiết kế"),
        ("API", "Application Programming Interface", "Giao diện lập trình ứng dụng"),
        ("CSDL", "Cơ sở dữ liệu", "Database, nơi lưu trữ dữ liệu hệ thống"),
        ("DB", "Database", "Cơ sở dữ liệu Oracle"),
        ("IDE", "Integrated Development Environment", "Môi trường phát triển tích hợp"),
        ("RESTful", "Representational State Transfer", "Kiến trúc thiết kế API"),
        ("Batch", "Batch Processing", "Xử lý hàng loạt theo lịch trình"),
        ("Shiteki", "指摘 (tiếng Nhật)", "Phản hồi/góp ý từ review"),
        ("BrSE", "Bridge System Engineer", "Kỹ sư cầu nối Việt-Nhật"),
        ("QA", "Quality Assurance", "Đảm bảo chất lượng"),
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header 3 cột
    headers = ["Từ viết tắt", "Diễn giải", "Ghi chú trong báo cáo"]
    col_widths = [Cm(3.0), Cm(5.5), Cm(7.5)]
    header_cells = table.rows[0].cells
    for i, (header, width) in enumerate(zip(headers, col_widths)):
        header_cells[i].text = header
        header_cells[i].width = width
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_font(run, FONT_SIZE_TABLE, bold=True)
        set_cell_shading(header_cells[i], 'D9E2F3')

    for abbr, meaning, note in abbreviations:
        row = table.add_row()
        row.cells[0].text = abbr
        row.cells[0].width = col_widths[0]
        row.cells[1].text = meaning
        row.cells[1].width = col_widths[1]
        row.cells[2].text = note
        row.cells[2].width = col_widths[2]
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font(run, FONT_SIZE_TABLE)


# ============== CHƯƠNG 1: GIỚI THIỆU CHUNG VỀ ĐƠN VỊ THỰC TẬP ==============
def add_chapter1(doc):
    """Chương 1: Giới thiệu chung về đơn vị thực tập - theo mẫu tham khảo"""
    add_chapter_title(doc, "1", "GIỚI THIỆU CHUNG VỀ ĐƠN VỊ THỰC TẬP")

    # 1.1 Thông tin chung về đơn vị thực tập
    add_section_title(doc, "1.1. Thông tin chung về đơn vị thực tập")

    add_paragraph_text(doc,
        f"Tên đơn vị: {INTERNSHIP_INFO['company']} (viết tắt: SYP)")

    add_paragraph_text(doc, "Lĩnh vực hoạt động:")

    add_paragraph_text(doc,
        f"{INTERNSHIP_INFO['company']} là công ty công nghệ thông tin chuyên về "
        "phát triển phần mềm gia công (offshore development) cho thị trường Nhật Bản và quốc tế. "
        "Công ty được thành lập tại Hà Nội, Việt Nam vào năm 2022, với đội ngũ lãnh đạo có hơn 20 năm "
        "kinh nghiệm trong lĩnh vực phát triển phần mềm offshore trên toàn cầu. Công ty tập trung vào "
        "các lĩnh vực:")

    add_bullet_list(doc, [
        "Phát triển phần mềm gia công cho khách hàng Nhật Bản",
        "Thiết kế hệ thống và cơ sở dữ liệu",
        "Phát triển ứng dụng web và mobile",
        "Tích hợp hệ thống và DevOps",
        "Tư vấn và chuyển giao công nghệ",
    ])

    add_paragraph_text(doc, "Thông tin liên hệ:")

    add_bullet_list(doc, [
        f"Địa chỉ: {INTERNSHIP_INFO['address']}",
        "Website: https://syp.vn",
        "Quy mô: Hơn 95 nhân viên (tính đến tháng 6/2024)",
    ])

    add_paragraph_text(doc, "Cơ cấu tổ chức:")

    add_paragraph_text(doc,
        "Công ty được tổ chức theo mô hình chuẩn của doanh nghiệp phần mềm, với các bộ phận chính:")

    add_bullet_list(doc, [
        "Ban Giám đốc: Điều hành và quản lý chiến lược công ty",
        "Phòng Phát triển phần mềm (Development): Thiết kế và lập trình các dự án",
        "Phòng Kiểm thử (QA/QC): Đảm bảo chất lượng sản phẩm",
        "Phòng BrSE (Bridge System Engineer): Kết nối với khách hàng Nhật Bản",
        "Phòng Nhân sự và Hành chính: Quản lý nguồn nhân lực",
    ])

    add_paragraph_text(doc, "Quy mô và định hướng phát triển:")

    add_paragraph_text(doc,
        "Công ty hướng tới trở thành đối tác chiến lược tin cậy cho các doanh nghiệp Nhật Bản "
        "trong lĩnh vực phát triển phần mềm. Các mục tiêu chính bao gồm:")

    add_bullet_list(doc, [
        "Mở rộng quy mô lên 200+ nhân viên vào năm 2027",
        "Tăng cường hợp tác với khách hàng Nhật Bản lớn",
        "Phát triển năng lực AI và Machine Learning trong các dự án",
        "Xây dựng môi trường làm việc đạt chuẩn quốc tế",
    ])

    # 1.2 Chức năng, nhiệm vụ của bộ phận thực tập
    add_section_title(doc, "1.2. Chức năng, nhiệm vụ của bộ phận thực tập")

    add_paragraph_text(doc,
        "Sinh viên được phân công vào bộ phận Development với vai trò Software Engineer. "
        "Bộ phận thực tập tại công ty có các chức năng và nhiệm vụ chính sau đây:")

    add_paragraph_text(doc, "Tổ chức và quản lý chương trình thực tập:")

    add_paragraph_text(doc,
        "Phối hợp với giảng viên hướng dẫn để lập kế hoạch, phân công nhiệm vụ thực tập "
        "cho sinh viên. Đảm bảo thực tập phù hợp với chương trình đào tạo, đáp ứng yêu cầu "
        "về thời lượng, nội dung chuyên môn và mục tiêu phát triển kỹ năng thực tiễn.")

    add_paragraph_text(doc, "Hướng dẫn và giám sát sinh viên thực tập:")

    add_paragraph_text(doc,
        "Hướng dẫn sinh viên thực hiện các công việc thiết kế hệ thống, viết tài liệu kỹ thuật, "
        "tham gia vào quy trình review và cải tiến chất lượng. Tổ chức các buổi training định kỳ, "
        "kiểm tra tiến độ và góp ý chuyên môn để đảm bảo chất lượng công việc.")

    add_paragraph_text(doc, "Cung cấp điều kiện cơ sở vật chất và tài nguyên hỗ trợ:")

    add_paragraph_text(doc,
        "Cung cấp máy tính, phần mềm chuyên dụng (IntelliJ IDEA, VS Code, Oracle Database, "
        "các công cụ thiết kế), tài liệu tham khảo và môi trường phát triển. Hỗ trợ sinh viên "
        "tiếp cận các dự án thực tế của công ty.")

    add_paragraph_text(doc, "Đánh giá và phản hồi:")

    add_paragraph_text(doc,
        "Tham gia đánh giá kết quả thực tập dựa trên các tiêu chí: chất lượng thiết kế, "
        "tinh thần học hỏi, khả năng làm việc nhóm và thái độ làm việc. Phối hợp với "
        "giảng viên hướng dẫn để đưa ra nhận xét cuối cùng.")

    # 1.3 Môi trường làm việc và quy trình công tác
    add_section_title(doc, "1.3. Môi trường làm việc và quy trình công tác")

    add_paragraph_text(doc, "Môi trường làm việc:")

    add_paragraph_text(doc,
        f"Công ty có văn phòng hiện đại tại {INTERNSHIP_INFO['address']}. "
        "Môi trường làm việc được trang bị đầy đủ:")

    add_bullet_list(doc, [
        "Hệ thống máy tính cấu hình cao, phòng họp với thiết bị video conference",
        "Không gian làm việc thoáng đãng, có Wi-Fi tốc độ cao",
        "Môi trường thân thiện, chuyên nghiệp với đội ngũ nhân viên trẻ trung",
        "Văn hóa chia sẻ kiến thức qua các buổi training nội bộ hàng tuần",
    ])

    add_paragraph_text(doc, "Quy trình công tác:")

    add_paragraph_text(doc,
        "Công ty áp dụng mô hình làm việc hybrid, kết hợp giữa làm việc tại văn phòng "
        "và làm việc từ xa. Quy trình phát triển phần mềm theo mô hình Agile/Scrum "
        "với các đặc điểm:")

    add_bullet_list(doc, [
        "Sprint 2 tuần với daily standup meeting hàng ngày",
        "Quy trình review nhiều cấp: Leader Review → Customer Review → End-user Review",
        "Hệ thống quản lý task bằng Backlog (công cụ quản lý dự án của Nhật)",
        "Sử dụng Git cho quản lý mã nguồn và tài liệu",
        "Thời gian làm việc linh hoạt, tập trung vào kết quả",
    ])


# ============== CHƯƠNG 2: NỘI DUNG THỰC TẬP ==============
def add_chapter2(doc):
    """Chương 2: Nội dung thực tập - theo mẫu tham khảo"""
    add_chapter_title(doc, "2", "NỘI DUNG THỰC TẬP")

    # 2.1 Mục tiêu và yêu cầu của đợt thực tập
    add_section_title(doc, "2.1. Mục tiêu và yêu cầu của đợt thực tập")

    add_paragraph_text(doc,
        "Mục tiêu của đợt thực tập nhằm giúp sinh viên củng cố và vận dụng những kiến "
        "thức lý thuyết đã học vào môi trường làm việc thực tế, qua đó nâng cao năng lực chuyên "
        "môn và kỹ năng nghề nghiệp trong lĩnh vực công nghệ thông tin.")

    add_paragraph_text(doc, "Cụ thể, đợt thực tập hướng tới các mục tiêu sau:")

    add_bullet_list(doc, [
        "Giúp sinh viên hiểu rõ hơn về môi trường làm việc, quy trình công tác và yêu cầu "
        "chuyên môn trong lĩnh vực phát triển phần mềm offshore",
        "Rèn luyện kỹ năng thiết kế hệ thống: cơ sở dữ liệu, màn hình, API và batch processing",
        "Nâng cao kỹ năng làm việc độc lập, làm việc nhóm, quản lý thời gian và báo cáo công việc",
        "Tạo điều kiện cho sinh viên tiếp cận với các công nghệ, công cụ và phương pháp "
        "thiết kế phần mềm hiện đại theo tiêu chuẩn Nhật Bản",
        "Học cách sử dụng AI (Claude AI) [4] hỗ trợ kiểm tra chất lượng thiết kế",
        "Chuẩn bị nền tảng kiến thức và kỹ năng cần thiết cho đồ án tốt nghiệp cũng như "
        "công việc sau khi ra trường",
    ])

    add_paragraph_text(doc,
        "Yêu cầu của đợt thực tập là sinh viên phải thực hiện đầy đủ kế hoạch thực tập đã đề ra, "
        "chấp hành nghiêm túc nội quy của công ty, hoàn thành các nhiệm vụ được giao, "
        "báo cáo tiến độ đúng thời hạn và tổng hợp kết quả thực tập thành báo cáo theo đúng "
        "quy định của Nhà trường.")

    # 2.2 Kế hoạch thực tập
    add_section_title(doc, "2.2. Kế hoạch thực tập")

    add_paragraph_text(doc,
        "Bảng kế hoạch thực tập được xây dựng theo từng tuần nhằm đảm bảo tiến độ và "
        "chất lượng công việc trong suốt thời gian thực tập.")

    # Bảng đầu tiên trong Chapter 2 - reset=True để đánh số từ 1
    add_table_with_caption(doc, 2, "Kế hoạch thực tập chi tiết",
        ["Tuần", "Thời gian", "Nội dung công việc chính"],
        [
            ("1", "26/06 – 02/07", "Làm quen môi trường, tìm hiểu dự án, ôn tập kiến thức"),
            ("2", "03/07 – 09/07", "Training thiết kế cơ sở dữ liệu (DB Design)"),
            ("3", "10/07 – 16/07", "Thực hành thiết kế bảng, index, constraints"),
            ("4", "17/07 – 23/07", "Training và thực hành thiết kế màn hình (Screen Design)"),
            ("5", "24/07 – 30/07", "Hoàn thiện Screen Design, bắt đầu API Design"),
            ("6", "31/07 – 06/08", "Training và thực hành thiết kế API RESTful"),
            ("7", "07/08 – 13/08", "Giới thiệu AI Checker, training thiết kế Batch"),
            ("8", "14/08 – 20/08", "Thực hành thiết kế Batch Processing"),
            ("9", "21/08 – 27/08", "Thiết kế độc lập, xử lý Shiteki (feedback) đợt 1"),
            ("10", "28/08 – 03/09", "Xử lý Shiteki đợt 2, hoàn thiện thiết kế"),
            ("11", "04/09 – 10/09", "Tổng hợp sản phẩm, kiểm tra chất lượng"),
            ("12", "11/09 – 17/09", "Viết báo cáo thực tập, chuẩn bị tài liệu"),
            ("13", "18/09 – 26/09", "Hoàn thiện báo cáo, nộp sản phẩm cuối cùng"),
        ],
        col_widths=[2.0, 3.5, 10.5],
        reset=True
    )

    # 2.3 Các công việc đã thực hiện
    add_section_title(doc, "2.3. Các công việc đã thực hiện")

    add_paragraph_text(doc,
        "Trong thời gian thực tập tại công ty, em đã thực hiện các công việc "
        "theo sự phân công và hướng dẫn trực tiếp của cán bộ hướng dẫn tại đơn vị.")

    add_paragraph_text(doc, "Mô tả các công việc đã thực hiện:")

    add_bullet_list(doc, [
        "Tìm hiểu đề tài thực tập và các yêu cầu kỹ thuật do cán bộ hướng dẫn giao",
        "Nghiên cứu tài liệu liên quan đến thiết kế hệ thống phần mềm theo chuẩn Nhật Bản",
        "Thực hiện thiết kế cơ sở dữ liệu, màn hình, API và batch processing",
        "Ghi chép tiến độ công việc và báo cáo kết quả định kỳ cho cán bộ hướng dẫn",
        "Hoàn thiện sản phẩm và tổng hợp nội dung báo cáo thực tập",
    ])

    add_paragraph_text(doc, "Vai trò và trách nhiệm của sinh viên:")

    add_paragraph_text(doc,
        "Sinh viên là người trực tiếp thực hiện toàn bộ các công việc được giao, chủ động "
        "nghiên cứu, học hỏi và áp dụng kiến thức chuyên môn để hoàn thành nhiệm vụ. "
        "Đồng thời, sinh viên có trách nhiệm tuân thủ kế hoạch thực tập, đảm bảo tiến độ "
        "công việc và chất lượng kết quả theo yêu cầu của cán bộ hướng dẫn.")

    add_subsection_title(doc, "2.3.1. Thiết kế cơ sở dữ liệu")

    add_paragraph_text(doc,
        "Trong quá trình thực tập, sinh viên được giao nhiệm vụ thiết kế cơ sở dữ liệu "
        "cho các module của hệ thống. Công việc bắt đầu với việc nghiên cứu cấu trúc "
        "database hiện có bao gồm: Entity Info (thông tin cơ bản về entity như tên logic, "
        "tên vật lý, hệ thống), Column Info (chi tiết các cột bao gồm tên logic, tên vật lý, "
        "kiểu dữ liệu, ràng buộc), và Index Info (thông tin về Primary Key, Foreign Key, các index).")

    add_paragraph_text(doc,
        "Các kiến thức về Oracle Database [1] được áp dụng trong quá trình thiết kế bao gồm:")

    add_bullet_list(doc, [
        "Kiểu dữ liệu: CHAR, VARCHAR2, NUMBER, DATE, CLOB, BLOB, JSON",
        "Constraints: PRIMARY KEY, FOREIGN KEY, NOT NULL, UNIQUE, CHECK",
        "Index: B-tree index, Bitmap index, Function-based index",
        "Naming convention: Quy tắc đặt tên chuẩn cho table, column, index",
    ])

    add_paragraph_text(doc,
        "Sinh viên học được cách định nghĩa bảng theo chuẩn với các thành phần: "
        "tên vật lý (sử dụng PascalCase hoặc snake_case), kiểu dữ liệu theo chuẩn Oracle, "
        "các ràng buộc (NOT NULL, UNIQUE, DEFAULT), và Primary Key (định danh duy nhất).")

    add_subsection_title(doc, "2.3.2. Thiết kế màn hình")

    add_paragraph_text(doc,
        "Thiết kế màn hình là quá trình xác định giao diện người dùng, bao gồm bố cục, "
        "các thành phần tương tác và luồng xử lý. Đây là loại thiết kế phức tạp nhất "
        "vì cần tham chiếu đến nhiều tài liệu liên quan.")

    add_paragraph_text(doc, "Các thành phần chính trong thiết kế màn hình bao gồm:")

    add_bullet_list(doc, [
        "Layout: Bố cục tổng thể của màn hình",
        "Items: Định nghĩa từng phần tử trên màn hình (ID, tên, loại, I/O, bắt buộc)",
        "Validation: Quy tắc kiểm tra dữ liệu nhập (đơn lẻ và tương quan)",
        "Messages: Các thông báo hiển thị cho người dùng",
        "Item Control: Điều khiển trạng thái các phần tử (hiển thị/ẩn, enable/disable)",
    ])

    add_paragraph_text(doc,
        "Validation được chia thành hai loại: Validate đơn lẻ (kiểm tra từng trường độc lập "
        "như bắt buộc nhập, độ dài, format, kiểu dữ liệu) và Validate tương quan (kiểm tra "
        "logic phụ thuộc giữa các trường như ngày kết thúc phải sau ngày bắt đầu).")

    add_subsection_title(doc, "2.3.3. Thiết kế API RESTful")

    add_paragraph_text(doc,
        "Thiết kế API là quá trình định nghĩa các endpoint để frontend và các hệ thống "
        "khác giao tiếp với backend. API trong dự án tuân theo kiến trúc RESTful [3], [5] với "
        "các nguyên tắc: Resource-based URL, sử dụng HTTP Methods chuẩn, Status Codes "
        "rõ ràng, và Content-Type là application/json.")

    add_paragraph_text(doc, "Cấu trúc thiết kế API bao gồm các thành phần:")

    add_bullet_list(doc, [
        "API ID: Mã định danh duy nhất cho mỗi API",
        "HTTP Method: GET (đọc), POST (tạo), PUT (cập nhật toàn bộ), PATCH (cập nhật một phần), DELETE (xóa)",
        "Endpoint: URL path của API",
        "Request Parameters và Request Body: Các tham số đầu vào",
        "Response: Định dạng phản hồi với status, data, message",
        "Error Handling: Xử lý các trường hợp lỗi với status code phù hợp",
    ])

    add_subsection_title(doc, "2.3.4. Thiết kế Batch Processing")

    add_paragraph_text(doc,
        "Batch processing (xử lý hàng loạt) là phương pháp xử lý khối lượng lớn dữ liệu "
        "theo lịch trình định sẵn, không cần tương tác người dùng. Công việc thiết kế batch "
        "áp dụng kiến trúc Spring Batch [2] với các thành phần: Job (đơn vị công việc cao nhất), "
        "Step (các bước trong một Job), ItemReader, ItemProcessor, và ItemWriter.")

    add_paragraph_text(doc, "Một thiết kế batch tiêu chuẩn bao gồm các phần:")

    add_bullet_list(doc, [
        "Tổng quan chức năng: Mục đích, đối tượng batch, cách khởi động",
        "Shell Script: Định nghĩa tham số đầu vào, mã kết thúc",
        "Xử lý Java: Logic xử lý chính với 5 block cơ bản (Chuẩn bị, Khởi tạo, Kiểm tra, Xử lý chính, Kết thúc)",
        "Yêu cầu tìm kiếm: Các câu SQL SELECT",
        "Yêu cầu cập nhật: Các câu SQL INSERT/UPDATE/DELETE",
    ])

    add_subsection_title(doc, "2.3.5. Quy trình xử lý Review (Shiteki)")

    add_paragraph_text(doc,
        "Sau khi hoàn thành thiết kế, sản phẩm được đưa qua quy trình review nhiều cấp: "
        "Leader Review → Customer Review → End-user Review. Mỗi cấp review sẽ đưa ra "
        "các phản hồi (shiteki) cần được xử lý trước khi chuyển sang cấp tiếp theo.")

    add_paragraph_text(doc, "Các loại shiteki thường gặp và cách xử lý:")

    add_bullet_list(doc, [
        "Lỗi chính tả (sai tên bảng, cột, API): Kiểm tra lại tài liệu tham khảo",
        "Lỗi logic (sai điều kiện, thiếu trường hợp): Phân tích lại nghiệp vụ",
        "Thiếu thông tin (thiếu validation, message): Bổ sung theo yêu cầu",
        "Không nhất quán (khác biệt giữa các phần): Đồng bộ toàn bộ thiết kế",
    ])

    # 2.4 Công nghệ, công cụ và kỹ thuật sử dụng
    add_section_title(doc, "2.4. Công nghệ, công cụ và kỹ thuật sử dụng")

    add_table_with_caption(doc, 2, "Công nghệ và công cụ sử dụng",
        ["Loại", "Tên", "Mục đích"],
        [
            ("Ngôn ngữ", "Java, SQL", "Lập trình backend, truy vấn CSDL"),
            ("Framework", "Spring Boot, Spring Batch", "Phát triển ứng dụng"),
            ("CSDL", "Oracle Database", "Lưu trữ dữ liệu"),
            ("IDE", "IntelliJ IDEA, VS Code", "Viết code"),
            ("Quản lý mã nguồn", "Git, GitHub", "Version control"),
            ("AI Tools", "Claude AI", "Kiểm tra chất lượng thiết kế"),
        ],
        col_widths=[4.0, 5.0, 7.0]
    )


# ============== CHƯƠNG 3: KẾT QUẢ VÀ ĐÁNH GIÁ ==============
def add_chapter3(doc):
    """Chương 3: Kết quả và đánh giá - theo mẫu tham khảo"""
    add_chapter_title(doc, "3", "KẾT QUẢ VÀ ĐÁNH GIÁ")

    # 3.1 Kết quả đạt được
    add_section_title(doc, "3.1. Kết quả đạt được trong quá trình thực tập")

    add_paragraph_text(doc,
        "Trong suốt thời gian thực tập tại công ty, em đã hoàn thành "
        "đầy đủ các nội dung và nhiệm vụ theo kế hoạch thực tập đã đề ra dưới sự hướng dẫn của "
        "cán bộ phụ trách. Các công việc được giao đều được thực hiện nghiêm túc, đúng "
        "tiến độ và đảm bảo yêu cầu về chất lượng.")

    add_paragraph_text(doc,
        "Thông qua quá trình thực tập, em đã từng bước tiếp cận với các công việc chuyên "
        "môn trong lĩnh vực thiết kế hệ thống phần mềm, từ việc nghiên cứu tài liệu, phân tích yêu cầu "
        "đến triển khai và hoàn thiện các nội dung liên quan đến đề tài thực tập. Kết quả đạt được "
        "không chỉ thể hiện qua sản phẩm hoặc nội dung công việc đã hoàn thành mà còn ở sự "
        "tiến bộ rõ rệt về tư duy, kỹ năng và thái độ làm việc.")

    add_paragraph_text(doc, "Kết quả đạt được:")

    add_bullet_list(doc, [
        "Hoàn thành các nội dung công việc theo kế hoạch đề ra",
        "Hoàn thành thiết kế cơ sở dữ liệu cho các module được giao đạt tiêu chuẩn",
        "Hoàn thành thiết kế màn hình với đầy đủ validation và message",
        "Hoàn thành thiết kế API RESTful theo chuẩn công ty",
        "Tham gia sử dụng và góp ý cải tiến hệ thống AI Checker",
        "Nâng cao kỹ năng thiết kế, tư duy logic và khả năng giải quyết vấn đề",
        "Hoàn thành báo cáo thực tập đúng quy định, phản ánh trung thực quá trình thực hiện",
    ])

    # 3.2 Kiến thức và kỹ năng tích lũy được
    add_section_title(doc, "3.2. Kiến thức và kỹ năng tích lũy được")

    add_subsection_title(doc, "3.2.1. Kiến thức chuyên môn")

    add_paragraph_text(doc,
        "Trong quá trình thực tập, em đã củng cố và mở rộng các kiến thức chuyên môn đã "
        "được học trên giảng đường. Đặc biệt là kiến thức về thiết kế hệ thống, phân tích yêu cầu, "
        "cơ sở dữ liệu và quy trình phát triển phần mềm. Việc áp dụng lý thuyết vào các "
        "bài toán thực tế giúp em hiểu sâu hơn bản chất vấn đề, đồng thời nâng cao khả năng vận "
        "dụng kiến thức vào thực tiễn.")

    add_bullet_list(doc, [
        "Nắm vững quy trình thiết kế hệ thống phần mềm chuyên nghiệp theo chuẩn Nhật Bản",
        "Hiểu sâu về thiết kế cơ sở dữ liệu với Oracle Database [1]",
        "Biết cách thiết kế API RESTful [3], [5] theo chuẩn",
        "Hiểu về kiến trúc Spring Batch [2] và thiết kế batch processing",
        "Biết cách sử dụng AI (Claude) [4] hỗ trợ kiểm tra chất lượng thiết kế",
    ])

    add_subsection_title(doc, "3.2.2. Kỹ năng làm việc nhóm")

    add_paragraph_text(doc,
        "Quá trình thực tập giúp em rèn luyện kỹ năng làm việc nhóm thông qua việc trao "
        "đổi, thảo luận và phối hợp với cán bộ hướng dẫn và các đồng nghiệp. Em học được "
        "cách lắng nghe ý kiến đóng góp, chia sẻ công việc hợp lý, hỗ trợ lẫn nhau để hoàn thành "
        "nhiệm vụ chung.")

    add_bullet_list(doc, [
        "Kỹ năng giao tiếp và trao đổi với team member trong môi trường chuyên nghiệp",
        "Kỹ năng review thiết kế và tiếp nhận feedback (shiteki)",
        "Kỹ năng làm việc với khách hàng nước ngoài (thông qua BrSE)",
        "Kỹ năng giao tiếp trong môi trường học thuật và chuyên môn",
    ])

    add_subsection_title(doc, "3.2.3. Kỹ năng phân tích và giải quyết vấn đề")

    add_paragraph_text(doc,
        "Thông qua các công việc được giao, em đã rèn luyện khả năng phân tích yêu cầu, "
        "xác định vấn đề và đề xuất hướng giải quyết phù hợp. Khi gặp khó khăn trong quá trình "
        "thực hiện, em học được cách chủ động tìm kiếm tài liệu, tham khảo ý kiến mentor "
        "và tự đánh giá, điều chỉnh phương án thực hiện.")

    add_bullet_list(doc, [
        "Kỹ năng đọc hiểu và phân tích yêu cầu nghiệp vụ từ tài liệu khách hàng",
        "Kỹ năng viết QA (Question & Answer) để xác nhận yêu cầu không rõ ràng",
        "Kỹ năng tư duy logic, xử lý tình huống và khắc phục lỗi trong thiết kế",
        "Kỹ năng sử dụng AI hỗ trợ công việc một cách hiệu quả",
    ])

    # 3.3 Thuận lợi và khó khăn
    add_section_title(doc, "3.3. Thuận lợi và khó khăn")

    add_paragraph_text(doc, "Thuận lợi:")

    add_bullet_list(doc, [
        "Nhận được sự quan tâm, hướng dẫn tận tình của cán bộ hướng dẫn trong suốt quá trình thực tập",
        "Môi trường học tập và làm việc tại công ty thân thiện, chuyên nghiệp, "
        "tạo điều kiện thuận lợi cho việc học hỏi và nghiên cứu",
        "Được tiếp cận với các tài liệu chuyên môn, cơ sở vật chất và công cụ hỗ trợ "
        "phục vụ cho quá trình thực tập",
        "Kiến thức nền tảng đã được trang bị trong quá trình học tập tại trường giúp "
        "sinh viên dễ dàng tiếp cận nội dung thực tập",
    ])

    add_paragraph_text(doc, "Khó khăn:")

    add_bullet_list(doc, [
        "Một số kiến thức và công nghệ còn mới (Oracle Database [1], Spring Batch [2], Java SE [6]), "
        "đòi hỏi em phải tự nghiên cứu và học hỏi thêm trong thời gian ngắn",
        "Kinh nghiệm thực tiễn còn hạn chế nên trong giai đoạn đầu gặp một số khó khăn "
        "khi triển khai công việc",
        "Thời gian thực tập có hạn, trong khi khối lượng công việc và yêu cầu chuyên môn "
        "tương đối nhiều",
        "Kiến thức tiếng Nhật còn hạn chế trong việc đọc hiểu một số tài liệu",
    ])

    add_paragraph_text(doc,
        "Tuy nhiên, nhờ sự hướng dẫn của cán bộ hướng dẫn và sự nỗ lực của bản thân, em đã từng bước "
        "khắc phục được những khó khăn trên và hoàn thành tốt đợt thực tập.")


# ============== CHƯƠNG 4: NHẬN XÉT VÀ ĐỊNH HƯỚNG ==============
def add_chapter4(doc):
    """Chương 4: Nhận xét và định hướng - theo mẫu tham khảo"""
    add_chapter_title(doc, "4", "NHẬN XÉT VÀ ĐỊNH HƯỚNG")

    # 4.1 Nhận xét chung về đợt thực tập
    add_section_title(doc, "4.1. Nhận xét chung về đợt thực tập")

    add_paragraph_text(doc,
        f"Đợt thực tập tại {INTERNSHIP_INFO['company']} "
        "là một trải nghiệm học tập có ý nghĩa và mang lại nhiều giá trị thiết thực đối với em.")

    add_paragraph_text(doc,
        "Thông qua quá trình thực tập, em đã có cơ hội tiếp cận với môi trường làm việc mang "
        "tính chuyên nghiệp cao, từ đó hiểu rõ hơn về yêu cầu và tính chất công việc "
        "trong lĩnh vực phát triển phần mềm offshore cho thị trường Nhật Bản.")

    add_paragraph_text(doc,
        "Các nội dung thực tập được xây dựng phù hợp với chương trình đào tạo, gắn liền "
        "giữa lý thuyết và thực tiễn, giúp em từng bước làm quen với quy trình làm việc, phương "
        "pháp thiết kế và triển khai các nhiệm vụ chuyên môn. Sự hướng dẫn tận tình của "
        "cán bộ hướng dẫn đã giúp em định hướng đúng đắn, kịp thời khắc phục những hạn chế trong "
        "quá trình thực hiện.")

    add_paragraph_text(doc,
        "Nhìn chung, đợt thực tập đã đạt được các mục tiêu đề ra, góp phần nâng cao kiến "
        "thức, kỹ năng và ý thức nghề nghiệp của em, đồng thời tạo nền tảng quan trọng cho quá trình "
        "học tập và làm việc sau này.")

    # 4.2 Bài học kinh nghiệm rút ra
    add_section_title(doc, "4.2. Bài học kinh nghiệm rút ra")

    add_paragraph_text(doc,
        "Từ quá trình thực tập, em đã rút ra được nhiều bài học kinh nghiệm quý báu. "
        "Trước hết là bài học về tinh thần tự giác và chủ động học tập. Trong môi trường thực "
        "tế, việc tự tìm hiểu tài liệu, chủ động đặt câu hỏi và đề xuất giải pháp là yếu tố quan "
        "trọng giúp nâng cao hiệu quả công việc.")

    add_paragraph_text(doc,
        "Bên cạnh đó, em nhận thức rõ hơn về tầm quan trọng của việc nắm vững kiến "
        "thức nền tảng và khả năng vận dụng linh hoạt kiến thức đã học vào các tình huống cụ "
        "thể. Việc làm việc theo kế hoạch, tuân thủ quy trình và đảm bảo tiến độ cũng là những "
        "kinh nghiệm cần thiết được rút ra trong suốt quá trình thực tập.")

    add_paragraph_text(doc, "Các bài học kinh nghiệm cụ thể:")

    add_bullet_list(doc, [
        "Cần chủ động trong việc học hỏi và đặt câu hỏi (viết QA) khi gặp khó khăn",
        "Tầm quan trọng của việc đọc kỹ tài liệu và hiểu rõ yêu cầu trước khi thực hiện",
        "Cần kiểm tra kỹ lưỡng (self-review) trước khi gửi sản phẩm cho leader review",
        "Kỹ năng mềm (giao tiếp, làm việc nhóm) quan trọng không kém kỹ năng chuyên môn",
        "Học cách tiếp nhận ý kiến góp ý (shiteki) một cách nghiêm túc và cải thiện",
    ])

    # 4.3 Định hướng nghề nghiệp và học tập sau thực tập
    add_section_title(doc, "4.3. Định hướng nghề nghiệp và học tập sau thực tập")

    add_paragraph_text(doc,
        "Sau đợt thực tập, em đã có cái nhìn rõ ràng hơn về định hướng nghề nghiệp trong "
        "tương lai. Trên cơ sở những kiến thức và kỹ năng đã tích lũy được, em định hướng tiếp "
        "tục nâng cao trình độ chuyên môn trong lĩnh vực công nghệ thông tin, đặc biệt là các "
        "mảng liên quan đến thiết kế hệ thống và phát triển phần mềm.")

    add_paragraph_text(doc,
        "Trong thời gian tới, em sẽ tập trung củng cố kiến thức chuyên ngành, học hỏi "
        "thêm các công nghệ mới, nâng cao kỹ năng lập trình, kỹ năng làm việc nhóm và kỹ năng "
        "nghiên cứu. Đồng thời, em cũng sẽ chuẩn bị tốt cho đồ án tốt nghiệp và sẵn sàng tham "
        "gia vào môi trường làm việc chuyên nghiệp sau khi ra trường.")

    add_paragraph_text(doc, "Các mục tiêu cụ thể:")

    add_bullet_list(doc, [
        "Hoàn thành tốt đồ án tốt nghiệp với đề tài KiteClass Platform - áp dụng kiến thức đã tích lũy",
        "Tiếp tục học hỏi và nâng cao kỹ năng lập trình (Java, TypeScript, Python)",
        "Tìm hiểu sâu hơn về AI và ứng dụng trong phát triển phần mềm",
        "Cải thiện kỹ năng ngoại ngữ (tiếng Anh, tiếng Nhật) để làm việc trong môi trường quốc tế",
        "Xây dựng portfolio cá nhân với các dự án thực tế",
    ])

    add_paragraph_text(doc,
        "Đợt thực tập là bước đệm quan trọng giúp em xác định rõ mục tiêu học tập và "
        "nghề nghiệp, tạo động lực để không ngừng rèn luyện và phát triển bản thân trong tương lai.")

    # 4.4 Những đóng góp của đề tài
    add_section_title(doc, "4.4. Những đóng góp của đề tài")

    add_subsection_title(doc, "4.4.1. Đóng góp về mặt sản phẩm")

    add_bullet_list(doc, [
        "Hoàn thành các thiết kế cơ sở dữ liệu, màn hình, API theo chuẩn doanh nghiệp Nhật Bản",
        "Tham gia vào quy trình thiết kế Batch Processing theo kiến trúc Spring Batch",
        "Đóng góp vào việc cải tiến chất lượng thiết kế thông qua quy trình review nhiều cấp",
        "Xây dựng sản phẩm hoặc mô hình phục vụ học tập và nghiên cứu",
    ])

    add_subsection_title(doc, "4.4.2. Đóng góp về mặt quy trình")

    add_bullet_list(doc, [
        "Hiểu và áp dụng quy trình làm việc thực tế trong doanh nghiệp offshore Nhật Bản",
        "Nắm vững quy trình xử lý Shiteki (review feedback) từ nhiều cấp độ: "
        "Leader → Customer → End-user",
        "Tích lũy kinh nghiệm sử dụng AI (Claude) hỗ trợ kiểm tra chất lượng thiết kế",
        "Hình thành phong cách làm việc chủ động và khoa học",
    ])

    add_subsection_title(doc, "4.4.3. Đóng góp về mặt kiến thức")

    add_bullet_list(doc, [
        "Tổng hợp kiến thức về thiết kế hệ thống phần mềm theo chuẩn quốc tế",
        "Chuẩn bị nền tảng kiến thức vững chắc cho đồ án tốt nghiệp KiteClass Platform",
        "Tài liệu báo cáo có thể làm tham khảo cho các sinh viên khóa sau về quy trình "
        "thực tập tại doanh nghiệp offshore",
    ])


# ============== TÀI LIỆU THAM KHẢO ==============
def add_ieee_reference(doc, ref_num, author, title, source_type, year, url=None, accessed=None, publisher=None, pages=None):
    """
    Thêm một tài liệu tham khảo theo chuẩn IEEE Citation Style

    IEEE Standard Format:
    - Online: [1] A. Author, "Title," Website/Source, Month Day, Year. [Online]. Available: URL. (accessed Month Day, Year).
    - Book: [1] A. Author, Title. City: Publisher, Year.
    - Journal: [1] A. Author, "Article title," Journal Name, vol. X, no. Y, pp. Z, Year.

    Args:
        doc: Document object
        ref_num: Số thứ tự tài liệu [1], [2]...
        author: Tên tác giả (e.g., "J. Smith" hoặc "Smith, J.")
        title: Tiêu đề tài liệu
        source_type: Loại nguồn ("online", "book", "journal", "conference")
        year: Năm xuất bản (e.g., "2024")
        url: URL (cho tài liệu online)
        accessed: Ngày truy cập (e.g., "Jan. 15, 2026")
        publisher: Nhà xuất bản (cho sách)
        pages: Số trang (cho journal/conference)
    """
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.63)
    p.paragraph_format.first_line_indent = Cm(-0.63)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # [1]
    run = p.add_run(f"[{ref_num}] ")
    set_font(run, FONT_SIZE_NORMAL)

    # Author,
    run = p.add_run(f"{author}, ")
    set_font(run, FONT_SIZE_NORMAL)

    # "Title" (in nghiêng cho article/online, không nghiêng cho book)
    if source_type.lower() in ["online", "journal", "conference"]:
        run = p.add_run(f'"{title}," ')
        set_font(run, FONT_SIZE_NORMAL, italic=True)
    else:  # book
        run = p.add_run(f"{title}. ")
        set_font(run, FONT_SIZE_NORMAL, italic=True)

    # Format theo loại nguồn
    if source_type.lower() == "online" and url:
        # Online format: Website, Year. [Online]. Available: URL. (accessed Date).
        run = p.add_run(f"{year}. ")
        set_font(run, FONT_SIZE_NORMAL)

        run = p.add_run("[Online]. Available: ")
        set_font(run, FONT_SIZE_NORMAL)

        run = p.add_run(url)
        set_font(run, FONT_SIZE_NORMAL)
        run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color for URL
        run.font.underline = True

        if accessed:
            run = p.add_run(f" (accessed {accessed})")
            set_font(run, FONT_SIZE_NORMAL)

        run = p.add_run(".")
        set_font(run, FONT_SIZE_NORMAL)

    elif source_type.lower() == "book":
        # Book format: City: Publisher, Year.
        if publisher:
            run = p.add_run(f"{publisher}, ")
            set_font(run, FONT_SIZE_NORMAL)
        run = p.add_run(f"{year}.")
        set_font(run, FONT_SIZE_NORMAL)

    elif source_type.lower() == "journal":
        # Journal format: Journal Name, vol. X, no. Y, pp. Z, Year.
        if pages:
            run = p.add_run(f"pp. {pages}, ")
            set_font(run, FONT_SIZE_NORMAL)
        run = p.add_run(f"{year}.")
        set_font(run, FONT_SIZE_NORMAL)

    else:
        # Generic format
        run = p.add_run(f"{year}.")
        set_font(run, FONT_SIZE_NORMAL)


def add_references(doc):
    """
    Tài liệu tham khảo theo chuẩn IEEE
    Sử dụng Heading 1 để có thể thêm vào mục lục tự động
    """
    doc.add_page_break()

    # Sử dụng Heading 1 để có thể tạo mục lục
    p = doc.add_paragraph(style='Heading 1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Căn giữa theo quy định UTC

    run = p.add_run("TÀI LIỆU THAM KHẢO")
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = FONT_SIZE_CHAPTER  # 18pt theo quy định UTC
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    # IEEE format references - Ví dụ chuẩn IEEE Citation Style

    # [1] Online resource - Database documentation
    add_ieee_reference(doc,
        ref_num=1,
        author="Oracle Corporation",
        title="Oracle Database 19c Documentation",
        source_type="online",
        year="2024",
        url="https://docs.oracle.com/en/database/oracle/oracle-database/19/",
        accessed="Jan. 15, 2026"
    )

    # [2] Online resource - Framework documentation
    add_ieee_reference(doc,
        ref_num=2,
        author="VMware Inc.",
        title="Spring Batch Reference Documentation v5.0",
        source_type="online",
        year="2024",
        url="https://docs.spring.io/spring-batch/docs/current/reference/html/",
        accessed="Jan. 18, 2026"
    )

    # [3] Online resource - API design guide
    add_ieee_reference(doc,
        ref_num=3,
        author="M. Masse",
        title="REST API Design Rulebook",
        source_type="book",
        year="2011",
        publisher="O'Reilly Media"
    )

    # [4] Online resource - API documentation
    add_ieee_reference(doc,
        ref_num=4,
        author="Anthropic PBC",
        title="Claude API Reference Documentation",
        source_type="online",
        year="2025",
        url="https://docs.anthropic.com/en/api/",
        accessed="Jan. 20, 2026"
    )

    # [5] Technical standard
    add_ieee_reference(doc,
        ref_num=5,
        author="R. T. Fielding",
        title="Architectural Styles and the Design of Network-based Software Architectures",
        source_type="online",
        year="2000",
        url="https://www.ics.uci.edu/~fielding/pubs/dissertation/top.htm",
        accessed="Jan. 10, 2026"
    )

    # [6] Java documentation
    add_ieee_reference(doc,
        ref_num=6,
        author="Oracle Corporation",
        title="Java SE 17 Documentation",
        source_type="online",
        year="2024",
        url="https://docs.oracle.com/en/java/javase/17/",
        accessed="Jan. 12, 2026"
    )

    # [7] Tài liệu nội bộ
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.63)
    p.paragraph_format.first_line_indent = Cm(-0.63)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("[7] ")
    set_font(run, FONT_SIZE_NORMAL)
    run = p.add_run("Tài liệu thiết kế nội bộ dự án SORA STEP4, SY PARTNERS., JSC (không công khai), 2025.")
    set_font(run, FONT_SIZE_NORMAL)


# ============== PHỤ LỤC ==============
def add_appendix(doc):
    """Phụ lục - sử dụng Heading 1 để có thể thêm vào mục lục"""
    doc.add_page_break()

    # Sử dụng Heading 1 để có thể tạo mục lục
    p = doc.add_paragraph(style='Heading 1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Căn giữa theo quy định UTC

    run = p.add_run("PHỤ LỤC")
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = FONT_SIZE_CHAPTER  # 18pt theo quy định UTC
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Phụ lục A: Nhật ký thực tập chi tiết (theo mẫu báo cáo tham khảo)
    add_section_title(doc, "Phụ lục A: Nhật ký thực tập")

    add_paragraph_text(doc,
        f"Thời gian thực tập: Từ ngày {INTERNSHIP_INFO['start_date']} đến ngày {INTERNSHIP_INFO['end_date']}",
        first_line_indent=False)

    # Bảng nhật ký thực tập theo tuần
    diary_headers = ["Tuần", "Thời gian", "Mục tiêu tuần", "Công việc thực hiện", "Kết quả", "Ghi chú"]
    diary_col_widths = [Cm(1.2), Cm(2.0), Cm(2.8), Cm(4.5), Cm(3.0), Cm(2.5)]

    diary_data = [
        ("1-2", "26/06 – 09/07", "Làm quen môi trường", "Tìm hiểu dự án, training thiết kế DB", "Nắm cấu trúc DB", ""),
        ("3-4", "10/07 – 23/07", "Thiết kế CSDL", "Thực hành thiết kế bảng, index, constraints", "Hoàn thành thiết kế DB", ""),
        ("5-6", "24/07 – 06/08", "Thiết kế màn hình", "Training và thực hành thiết kế Screen", "Hoàn thành thiết kế Screen", ""),
        ("7-8", "07/08 – 20/08", "Thiết kế API", "Training và thực hành thiết kế RESTful API", "Hoàn thành thiết kế API", ""),
        ("9-10", "21/08 – 03/09", "Thiết kế Batch", "Giới thiệu AI Checker, training thiết kế Batch", "Hoàn thành thiết kế Batch", ""),
        ("11-13", "04/09 – 26/09", "Hoàn thiện", "Thiết kế độc lập, hoàn thành báo cáo", "Báo cáo hoàn chỉnh", ""),
    ]

    table = doc.add_table(rows=1, cols=len(diary_headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    header_cells = table.rows[0].cells
    for i, (header, width) in enumerate(zip(diary_headers, diary_col_widths)):
        header_cells[i].text = header
        header_cells[i].width = width
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_font(run, Pt(11), bold=True)
        set_cell_shading(header_cells[i], 'D9E2F3')

    # Data rows
    for row_data in diary_data:
        row = table.add_row()
        for i, (cell_text, width) in enumerate(zip(row_data, diary_col_widths)):
            row.cells[i].text = cell_text
            row.cells[i].width = width
            for paragraph in row.cells[i].paragraphs:
                if i == 0:  # Tuần - căn giữa
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_font(run, Pt(11))

    doc.add_paragraph()  # Spacing

    # Phụ lục B: Hình ảnh, tài liệu minh chứng
    add_section_title(doc, "Phụ lục B: Hình ảnh, tài liệu minh chứng")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Đính kèm hình ảnh minh chứng quá trình thực tập]")
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_NORMAL
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Phụ lục C: Sản phẩm thực tập
    add_section_title(doc, "Phụ lục C: Sản phẩm thực tập")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Mô tả hoặc đính kèm các sản phẩm thiết kế đã hoàn thành]")
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_NORMAL
    run.font.color.rgb = RGBColor(128, 128, 128)


# ============== HÀM CHÍNH ==============
def create_report():
    """Hàm chính tạo báo cáo"""
    print("Đang tạo báo cáo thực tập theo mẫu mới...")

    doc = Document()

    # Thiết lập document
    # Lưu ý: margins sẽ được apply lại sau khi tạo xong tất cả sections
    setup_styles(doc)

    # 1. Trang bìa chính
    add_cover_page(doc)

    # 2. Trang bìa phụ
    add_secondary_cover_page(doc)

    # 3. Bản nhận xét của cơ sở thực tập - BỎ (sẽ xin vật lý)
    # add_company_review_page(doc)

    # 4. Lời cảm ơn
    add_acknowledgment_page(doc)

    # 5. Mục lục + Danh mục bảng biểu (KHÔNG có danh mục hình vẽ vì không có hình)
    add_toc_page(doc)
    add_list_of_tables(doc)

    # 6. Danh mục từ viết tắt
    add_abbreviations(doc)

    # 7. Nội dung chính - 4 chương
    add_chapter1(doc)  # Giới thiệu chung về đơn vị thực tập
    add_chapter2(doc)  # Nội dung thực tập
    add_chapter3(doc)  # Kết quả và đánh giá
    add_chapter4(doc)  # Nhận xét và định hướng

    # 8. Tài liệu tham khảo
    add_references(doc)

    # 9. Phụ lục
    add_appendix(doc)

    # DEBUG: Kiểm tra số sections
    print(f"DEBUG: Total sections = {len(doc.sections)}")

    # QUAN TRỌNG: Thêm khung viền TRƯỚC KHI thêm số trang và apply margins
    # để tránh border properties bị propagate
    if len(doc.sections) >= 3:
        # Bước 1: Thêm border cho sections 0 và 1
        add_page_border(doc.sections[0])  # Trang bìa chính
        add_page_border(doc.sections[1])  # Trang bìa phụ

        # Bước 2: Xóa EXPLICITLY borders từ section 2 trở đi
        for i in range(2, len(doc.sections)):
            remove_page_border(doc.sections[i])

        print(f"DEBUG: Applied border to sections 0 and 1")
        print(f"DEBUG: Explicitly removed borders from sections 2-{len(doc.sections)-1}")
    else:
        print(f"WARNING: Không đủ 3 sections! Chỉ có {len(doc.sections)} sections")

    # Apply margins cho TẤT CẢ sections SAU KHI đã apply/remove borders
    set_document_margins(doc)

    # Thêm số trang SAU CÙNG
    add_page_number_header(doc)

    # Lưu file
    output_path = "BAO_CAO_THUC_TAP.docx"
    doc.save(output_path)

    print(f"Da tao file: {output_path}")
    print(f"Cau truc bao cao:")
    print(f"  1. Bia chinh (co khung vien, bang thong tin SV co vien, logo UTC)")
    print(f"  2. Bia phu")
    print(f"  3. Loi cam on")
    print(f"  4. Muc luc + Danh muc bang bieu")
    print(f"  5. Danh muc tu viet tat")
    print(f"  6. 4 Chuong noi dung chinh")
    print(f"  7. Tai lieu tham khao (IEEE citations)")
    print(f"  8. Phu luc")
    print(f"")
    print(f"Luu y: Ban nhan xet cua co so thuc tap se xin vat ly (khong co trong file Word)")

    return output_path


if __name__ == "__main__":
    create_report()
