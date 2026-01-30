#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script tạo Đề cương Đồ án Tốt nghiệp Cử nhân dạng Word (.docx)
Format theo mẫu "Mau-Decuong DATN-Cử nhân.pdf" - ĐH GTVT

Cấu trúc đề cương:
1. Header: Trường + Khoa | Quốc hiệu
2. Ngày tháng
3. Tiêu đề: ĐỀ CƯƠNG ĐỒ ÁN TỐT NGHIỆP CỬ NHÂN
4. Thông tin sinh viên
5. Thông tin giảng viên hướng dẫn
6. Tên đề tài
7. 4 mục nội dung chính
8. Chữ ký các bên
9. Logo UTC
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============== THÔNG TIN SINH VIÊN ==============
STUDENT_INFO = {
    "name": "Nguyễn Văn Kiệt",
    "student_id": "221230890",
    "class": "CNTT1-K63",
    "course": "63",
    "phone": "...",
    "email": "...",
    "major": "Công nghệ thông tin",
    "system": "Chính quy",
}

# ============== THÔNG TIN GIẢNG VIÊN ==============
ADVISOR_INFO = {
    "name": "ThS. Nguyễn Đức Dư",
    "department": "Khoa Công nghệ thông tin - Trường ĐH GTVT",
    "phone": "...",
    "email": "...",
}

# ============== THÔNG TIN ĐỀ TÀI ==============
THESIS_INFO = {
    "title": "Xây dựng hệ thống quản lý lớp học trực tuyến theo kiến trúc Microservices - KiteClass Platform",
}

# ============== CONSTANTS ==============
FONT_NAME = 'Times New Roman'
FONT_SIZE_NORMAL = Pt(13)
FONT_SIZE_TITLE = Pt(14)

MARGIN_LEFT = Cm(2.5)
MARGIN_RIGHT = Cm(2.0)
MARGIN_TOP = Cm(2.0)
MARGIN_BOTTOM = Cm(2.0)


def set_document_margins(doc):
    """Thiết lập căn lề cho toàn bộ document"""
    for section in doc.sections:
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT


def set_font(run, size=FONT_SIZE_NORMAL, bold=False, italic=False, underline=False, color=None):
    """Helper to set font properties"""
    run.font.name = FONT_NAME
    run.font.size = size
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if color:
        run.font.color.rgb = color
    # Đảm bảo font cho tiếng Việt
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)


def set_cell_borders(cell, border_size=4):
    """Set cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(border_size))
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def remove_cell_borders(cell):
    """Remove all borders from a cell"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_header_section(doc):
    """Thêm phần header: Trường + Khoa | Quốc hiệu"""
    # Tạo bảng 2 cột cho header
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Cột trái: Trường + Khoa
    cell_left_1 = table.rows[0].cells[0]
    cell_left_1.width = Cm(8)
    p = cell_left_1.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TRƯỜNG ĐẠI HỌC GTVT")
    set_font(run, Pt(12), bold=True)
    remove_cell_borders(cell_left_1)

    cell_left_2 = table.rows[1].cells[0]
    p = cell_left_2.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("KHOA CÔNG NGHỆ THÔNG TIN")
    set_font(run, Pt(12), bold=True)
    remove_cell_borders(cell_left_2)

    # Cột phải: Quốc hiệu
    cell_right_1 = table.rows[0].cells[1]
    cell_right_1.width = Cm(9)
    p = cell_right_1.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    set_font(run, Pt(12), bold=True)
    remove_cell_borders(cell_right_1)

    cell_right_2 = table.rows[1].cells[1]
    p = cell_right_2.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Độc lập - Tự do - Hạnh phúc")
    set_font(run, Pt(12), bold=True, underline=True)
    remove_cell_borders(cell_right_2)

    # Ngày tháng
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Hà Nội, ngày ..... tháng ..... năm 2026")
    set_font(run, Pt(13), italic=True)


def add_title(doc):
    """Thêm tiêu đề: ĐỀ CƯƠNG ĐỒ ÁN TỐT NGHIỆP CỬ NHÂN"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)

    run = p.add_run("ĐỀ CƯƠNG ĐỒ ÁN TỐT NGHIỆP ")
    set_font(run, Pt(14), bold=True)

    # CỬ NHÂN với highlight màu vàng
    run = p.add_run("CỬ NHÂN")
    set_font(run, Pt(14), bold=True, underline=True)
    # Thêm highlight màu vàng
    rPr = run._element.get_or_add_rPr()
    highlight = OxmlElement('w:highlight')
    highlight.set(qn('w:val'), 'yellow')
    rPr.append(highlight)


def add_student_info(doc):
    """Thêm thông tin sinh viên"""
    # Họ và tên sinh viên
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Họ và tên sinh viên: ")
    set_font(run, Pt(13), bold=True)
    run = p.add_run(STUDENT_INFO["name"])
    set_font(run, Pt(13))

    # Mã SV, Lớp, Khóa (trên cùng một dòng)
    p = doc.add_paragraph()
    run = p.add_run("Mã SV")
    set_font(run, Pt(13))
    run = p.add_run(f"\t\t: {STUDENT_INFO['student_id']}")
    set_font(run, Pt(13))
    run = p.add_run(f"\t\tLớp: {STUDENT_INFO['class']}")
    set_font(run, Pt(13))
    run = p.add_run(f"\t\tKhóa: {STUDENT_INFO['course']}")
    set_font(run, Pt(13))

    # Số điện thoại, Email
    p = doc.add_paragraph()
    run = p.add_run("Số điện thoại")
    set_font(run, Pt(13))
    run = p.add_run(f"\t: {STUDENT_INFO['phone']}")
    set_font(run, Pt(13))
    run = p.add_run(f"\t\tEmail: {STUDENT_INFO['email']}")
    set_font(run, Pt(13))

    # Ngành
    p = doc.add_paragraph()
    run = p.add_run("Ngành")
    set_font(run, Pt(13))
    run = p.add_run(f"\t\t: ")
    set_font(run, Pt(13))
    run = p.add_run("Công nghệ thông tin")
    set_font(run, Pt(13), underline=True)
    run = p.add_run("/Khoa học máy tính")
    set_font(run, Pt(13))

    # Hệ
    p = doc.add_paragraph()
    run = p.add_run(f"Hệ: {STUDENT_INFO['system']}")
    set_font(run, Pt(13))


def add_advisor_info(doc):
    """Thêm thông tin giảng viên hướng dẫn"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Giảng viên (cán bộ) hướng dẫn: ")
    set_font(run, Pt(13), bold=True)
    run = p.add_run(ADVISOR_INFO["name"])
    set_font(run, Pt(13))

    p = doc.add_paragraph()
    run = p.add_run("Đơn vị công tác")
    set_font(run, Pt(13))
    run = p.add_run(f"\t: {ADVISOR_INFO['department']}")
    set_font(run, Pt(13))

    p = doc.add_paragraph()
    run = p.add_run("Số điện thoại")
    set_font(run, Pt(13))
    run = p.add_run(f"\t: {ADVISOR_INFO['phone']}")
    set_font(run, Pt(13))
    run = p.add_run(f"\t\tEmail: {ADVISOR_INFO['email']}")
    set_font(run, Pt(13))


def add_thesis_title(doc):
    """Thêm tên đề tài"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Tên đề tài: ")
    set_font(run, Pt(13), bold=True)
    run = p.add_run(THESIS_INFO["title"])
    set_font(run, Pt(13))


def add_section_title(doc, number, title):
    """Thêm tiêu đề mục"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"{number}. {title}")
    set_font(run, Pt(13), bold=True)


def add_section_content(doc, text, indent=True):
    """Thêm nội dung mục"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_font(run, Pt(13))


def add_bullet_item(doc, text):
    """Thêm bullet item"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(f"- {text}")
    set_font(run, Pt(13))


def add_subsection_title(doc, title):
    """Thêm tiêu đề tiểu mục (2.1, 2.2...)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    set_font(run, Pt(13), bold=True)


def add_content_sections(doc):
    """Thêm 4 mục nội dung chính - theo mẫu tham khảo"""

    # 1. Nội dung, phạm vi của đề tài
    add_section_title(doc, "1", "Nội dung, phạm vi của đề tài")

    # Nội dung của đề tài
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Nội dung của đề tài")
    set_font(run, Pt(13), bold=True)

    add_section_content(doc,
        "Đề tài \"Xây dựng hệ thống quản lý lớp học trực tuyến theo kiến trúc Microservices - "
        "KiteClass Platform\" tập trung vào việc nghiên cứu, phân tích và xây dựng một nền tảng "
        "học tập trực tuyến có khả năng đáp ứng nhu cầu quản lý lớp học khác nhau của từng tổ chức "
        "giáo dục. Hệ thống cho phép khách hàng nhanh chóng khởi tạo và vận hành hệ thống lớp học "
        "trực tuyến riêng với giao diện và thương hiệu cá nhân hóa, được hỗ trợ bởi trí tuệ nhân tạo.")

    add_section_content(doc, "Các nội dung chính của đề tài bao gồm:", indent=True)
    add_bullet_item(doc, "Khảo sát và phân tích yêu cầu của hệ thống quản lý lớp học trực tuyến")
    add_bullet_item(doc, "Thiết kế kiến trúc tổng thể của hệ thống theo mô hình Microservices, "
        "bao gồm KiteClass Core Services và KiteHub Platform")
    add_bullet_item(doc, "Xây dựng các chức năng cơ bản: quản lý người dùng, quản lý lớp học, "
        "quản lý nội dung học tập, điểm danh, giao bài tập")
    add_bullet_item(doc, "Nghiên cứu và áp dụng AI Agent để tự động tạo branding assets "
        "(logo, banner, color palette) cho từng instance")
    add_bullet_item(doc, "Xây dựng hệ thống auto-provisioning để tự động triển khai instance mới")
    add_bullet_item(doc, "Triển khai thử nghiệm hệ thống và đánh giá kết quả đạt được")

    # Phạm vi của đề tài
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Phạm vi của đề tài")
    set_font(run, Pt(13), bold=True)

    add_section_content(doc,
        "Trong phạm vi của đề tài, hệ thống KiteClass Platform được xây dựng "
        "ở mức độ MVP (Minimum Viable Product), tập trung vào các chức năng cốt lõi phục vụ quá trình "
        "quản lý lớp học trực tuyến cho các tổ chức giáo dục nhỏ và vừa.")

    add_bullet_item(doc, "Đối tượng: Giáo viên dạy thêm THCS/THPT, gia sư chứng chỉ (IELTS, JLPT), "
        "trung tâm ngoại ngữ, tổ chức đào tạo doanh nghiệp")
    add_bullet_item(doc, "Phạm vi chức năng: Quản lý lớp học, người dùng, nội dung học tập, "
        "điểm danh, bài tập và thanh toán cơ bản")
    add_bullet_item(doc, "Kiến trúc: KiteHub Platform (Modular Monolith) + KiteClass Instances (Microservices)")
    add_bullet_item(doc, "Giới hạn: Chưa triển khai quy mô lớn, các vấn đề bảo mật nâng cao và "
        "tối ưu hiệu năng chỉ xem xét ở mức cơ bản")

    # 2. Công nghệ, công cụ và ngôn ngữ lập trình
    add_section_title(doc, "2", "Công nghệ, công cụ và ngôn ngữ lập trình")

    add_section_content(doc,
        "Để xây dựng hệ thống KiteClass Platform, đề tài sử dụng bộ công "
        "nghệ full-stack hiện đại kết hợp trí tuệ nhân tạo (AI), đảm bảo tính scalable, hiệu suất "
        "cao, dễ bảo trì và phù hợp với xu hướng phát triển phần mềm năm 2025-2026.")

    add_subsection_title(doc, "2.1. Ngôn ngữ lập trình")
    add_bullet_item(doc, "Java 21 LTS: Ngôn ngữ chính cho backend services")
    add_bullet_item(doc, "TypeScript/JavaScript: Frontend và scripting")
    add_bullet_item(doc, "Python: AI Agent và automation scripts")

    add_subsection_title(doc, "2.2. Framework và nền tảng chính")
    add_bullet_item(doc, "Backend: Spring Boot 3.2, Spring Security, Spring Data JPA")
    add_bullet_item(doc, "Frontend: Next.js 14 (App Router) + React + TypeScript")
    add_bullet_item(doc, "Database: PostgreSQL 15 (database chính), Redis 7.x (caching)")
    add_bullet_item(doc, "Message Queue: RabbitMQ 3.12 cho async communication")

    add_subsection_title(doc, "2.3. Công cụ hỗ trợ phát triển và triển khai")
    add_bullet_item(doc, "IDE và Editor: IntelliJ IDEA, VS Code")
    add_bullet_item(doc, "Quản lý mã nguồn: Git + GitHub (branch workflow: main/dev/feature)")
    add_bullet_item(doc, "Container: Docker, Kubernetes (EKS)")
    add_bullet_item(doc, "CI/CD: GitHub Actions, Terraform (Infrastructure as Code)")
    add_bullet_item(doc, "AI Services: OpenAI GPT-4 (text), Stability AI SDXL (images), Remove.bg")
    add_bullet_item(doc, "Monitoring: Prometheus, Grafana")

    add_subsection_title(doc, "2.4. Lý do lựa chọn bộ công nghệ này")
    add_bullet_item(doc, "Phù hợp xu hướng 2025-2026: Java Spring + Next.js là combo phổ biến cho enterprise apps, "
        "Python AI cho generative AI applications")
    add_bullet_item(doc, "Dễ triển khai cho sinh viên: Tài liệu phong phú, cộng đồng lớn, free tier dồi dào")
    add_bullet_item(doc, "Tính ứng dụng thực tiễn: Hệ thống scalable, có thể mở rộng thành sản phẩm thương mại")
    add_bullet_item(doc, "Hạn chế: Ưu tiên open-source để tránh phụ thuộc API trả phí, không dùng cloud enterprise full")

    # 3. Các kết quả chính dự kiến đạt được
    add_section_title(doc, "3", "Các kết quả chính dự kiến đạt được")

    add_section_content(doc,
        "Qua quá trình nghiên cứu và thực hiện đề tài, các kết quả chính dự kiến đạt được bao gồm:")

    # Sản phẩm phần mềm
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Sản phẩm phần mềm hoàn chỉnh (MVP):")
    set_font(run, Pt(13), bold=True)

    add_bullet_item(doc, "KiteHub Platform: Hệ thống quản lý trung tâm với Sale Module, AI Agent Module, "
        "Maintaining Module")
    add_bullet_item(doc, "KiteClass Instance: Hệ thống quản lý lớp học với User+Gateway Service, Core Service, "
        "và Engagement Service (optional)")
    add_bullet_item(doc, "AI Agent tự động tạo branding: Logo, banner, slogan, color palette (~$0.19/instance)")
    add_bullet_item(doc, "Auto-provisioning: Tự động triển khai instance mới trong 3-5 phút")
    add_bullet_item(doc, "Frontend dashboard: Giao diện quản trị trực quan với Next.js")

    # Kết quả kỹ thuật
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Kết quả kỹ thuật và đánh giá:")
    set_font(run, Pt(13), bold=True)

    add_bullet_item(doc, "Thời gian phản hồi API < 200ms cho 95% requests")
    add_bullet_item(doc, "Hỗ trợ 1000+ người dùng đồng thời trên mỗi instance")
    add_bullet_item(doc, "Uptime 99.9% với automatic failover")
    add_bullet_item(doc, "Báo cáo kiểm thử đầy đủ (unit test, integration test, load test)")

    # Tài liệu
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Tài liệu và sản phẩm phụ:")
    set_font(run, Pt(13), bold=True)

    add_bullet_item(doc, "Báo cáo đồ án tốt nghiệp hoàn chỉnh với UML diagrams, ERD, API documentation")
    add_bullet_item(doc, "Slide bảo vệ đồ án và video demo")
    add_bullet_item(doc, "Mã nguồn trên GitHub với README chi tiết")
    add_bullet_item(doc, "Tài liệu hướng dẫn triển khai và sử dụng")

    # Giá trị ứng dụng
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Giá trị ứng dụng và đóng góp:")
    set_font(run, Pt(13), bold=True)

    add_bullet_item(doc, "Đề tài mang tính ứng dụng cao: Hỗ trợ các tổ chức giáo dục nhỏ có nền tảng "
        "quản lý lớp học riêng với chi phí thấp")
    add_bullet_item(doc, "Đóng góp vào nghiên cứu ứng dụng AI trong giáo dục và kiến trúc Microservices")
    add_bullet_item(doc, "Kết quả có thể làm nền tảng để mở rộng thành sản phẩm thương mại thực sự")

    # 4. Kế hoạch thực hiện đề tài
    add_section_title(doc, "4", "Kế hoạch thực hiện đề tài")

    add_section_content(doc,
        "Kế hoạch thực hiện đề tài được lập theo các giai đoạn chính, với thời gian tổng cộng "
        "khoảng 4 tháng (từ tháng 2/2026 đến tháng 5/2026). Kế hoạch linh hoạt, có thể điều chỉnh "
        "theo góp ý của giảng viên hướng dẫn và tiến độ thực tế.")

    # Bảng kế hoạch với ngày cụ thể
    table = doc.add_table(rows=8, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    headers = ["STT", "Nội dung công việc", "Thời gian dự kiến", "Ghi chú"]
    header_widths = [Cm(1.2), Cm(7.5), Cm(4.5), Cm(3)]

    for i, (header, width) in enumerate(zip(headers, header_widths)):
        cell = table.rows[0].cells[i]
        cell.width = width
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_font(run, Pt(12), bold=True)
        set_cell_shading(cell, 'D9E2F3')

    # Data rows với ngày cụ thể
    plan_data = [
        ("1", "Nghiên cứu tài liệu, tổng quan đề tài", "01/02 – 21/02/2026", ""),
        ("2", "Phân tích yêu cầu, thiết kế hệ thống", "22/02 – 07/03/2026", "UML, ERD"),
        ("3", "Xây dựng KiteHub Platform", "08/03 – 28/03/2026", "Phase 1"),
        ("4", "Xây dựng KiteClass Core Services", "29/03 – 25/04/2026", "Phase 2"),
        ("5", "Phát triển AI Agent & Auto-provisioning", "26/04 – 10/05/2026", "Phase 3"),
        ("6", "Kiểm thử toàn hệ thống, fix bug, deploy", "11/05 – 25/05/2026", ""),
        ("7", "Hoàn thiện báo cáo, slide, video demo", "26/05 – 31/05/2026", ""),
    ]

    for row_idx, (stt, content, time, note) in enumerate(plan_data):
        row = table.rows[row_idx + 1]

        # STT
        cell = row.cells[0]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(stt)
        set_font(run, Pt(12))

        # Nội dung
        cell = row.cells[1]
        p = cell.paragraphs[0]
        run = p.add_run(content)
        set_font(run, Pt(12))

        # Thời gian
        cell = row.cells[2]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(time)
        set_font(run, Pt(12))

        # Ghi chú
        cell = row.cells[3]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(note)
        set_font(run, Pt(12))


def add_signatures(doc):
    """Thêm phần chữ ký"""
    # Khoảng trống
    for _ in range(2):
        doc.add_paragraph()

    # Bảng chữ ký 4 cột
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    titles = ["Trưởng Khoa", "Trưởng Bộ môn", "Giảng viên hướng dẫn", "Sinh viên thực hiện"]
    subtitles = ["(Ký và ghi rõ họ tên)", "(Ký và ghi rõ họ tên)", "(Ký và ghi rõ họ tên)", "(Ký và ghi rõ họ tên)"]

    for i, (title, subtitle) in enumerate(zip(titles, subtitles)):
        # Title
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        set_font(run, Pt(12), bold=True)
        remove_cell_borders(cell)

        # Subtitle
        cell = table.rows[1].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        set_font(run, Pt(11), italic=True)
        remove_cell_borders(cell)


def add_logo(doc):
    """Thêm logo UTC ở cuối trang"""
    import os

    # Khoảng trống cho chữ ký
    for _ in range(4):
        doc.add_paragraph()

    logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'logo_utc.png')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if os.path.exists(logo_path):
        run = p.add_run()
        run.add_picture(logo_path, width=Cm(3.0))
    else:
        run = p.add_run("[LOGO UTC]")
        run.font.color.rgb = RGBColor(128, 128, 128)
        set_font(run, Pt(12))


def create_de_cuong():
    """Hàm chính tạo đề cương"""
    print("Đang tạo Đề cương Đồ án Tốt nghiệp...")

    doc = Document()

    # Thiết lập document
    set_document_margins(doc)

    # 1. Header
    add_header_section(doc)

    # 2. Tiêu đề
    add_title(doc)

    # 3. Thông tin sinh viên
    add_student_info(doc)

    # 4. Thông tin giảng viên
    add_advisor_info(doc)

    # 5. Tên đề tài
    add_thesis_title(doc)

    # 6. Nội dung chính (4 mục)
    add_content_sections(doc)

    # 7. Chữ ký
    add_signatures(doc)

    # 8. Logo
    add_logo(doc)

    # Lưu file
    output_path = "DE_CUONG_DATN.docx"
    doc.save(output_path)

    print(f"Đã tạo file: {output_path}")
    print(f"Cấu trúc đề cương:")
    print(f"  - Thông tin sinh viên: {STUDENT_INFO['name']} - {STUDENT_INFO['student_id']}")
    print(f"  - Giảng viên hướng dẫn: {ADVISOR_INFO['name']}")
    print(f"  - Tên đề tài: {THESIS_INFO['title'][:50]}...")
    print(f"  - 4 mục nội dung chính + Bảng kế hoạch")
    print(f"  - Chữ ký 4 bên + Logo UTC")

    return output_path


if __name__ == "__main__":
    create_de_cuong()
